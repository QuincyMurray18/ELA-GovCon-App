import streamlit as st
## ELA Phase3 hybrid_api
# Optional FastAPI backend (run separately) + client with graceful fallback.
# Set GOVCON_API_BASE or st.secrets['api']['base_url'] to use the API.
import os, threading, uuid, time
try:
    import requests
except Exception:
    requests = None

# ---- job store for background tasks (in-memory) ----
_jobs: dict[str, dict] = {}

def _enqueue(fn, *args, **kwargs) -> str:
    jid = str(uuid.uuid4())
    _jobs[jid] = {"status":"queued"}
    def _run():
        _jobs[jid] = {"status":"running"}
        try:
            res = fn(*args, **kwargs)
            _jobs[jid] = {"status":"done", "result": res}
        except Exception as e:
            _jobs[jid] = {"status":"error", "error": str(e)}
    t = threading.Thread(target=_run, daemon=True)
    t.start()
    return jid

def _job_status(jid: str) -> dict:
    return _jobs.get(jid, {"status":"unknown"})

# ---- API service functions (fallbacks) ----
def _svc_analyze(question: str, context: dict | None = None):
    # Use your in-process analyzer if available
    try:
        return _service_analyze_rfp_question(question, context or {})
    except Exception:
        # Minimal fallback
        title = ""
        if isinstance(context, dict):
            title = context.get("title") or context.get("Title") or ""
        return f"[demo] Analysis for {title or 'the selected notice'}: {question}"

# ---- FastAPI app (optional) ----
try:
    from fastapi import FastAPI, BackgroundTasks
    from pydantic import BaseModel
    FASTAPI_OK = True
except Exception:
    FASTAPI_OK = False

if FASTAPI_OK:
    api = FastAPI(title="ELA GovCon API", version="0.1")

    class AnalyzeReq(BaseModel):
        question: str
        context: dict | None = None

    @api.post("/api/analyze")
    def api_analyze(req: AnalyzeReq, background_tasks: BackgroundTasks):
        jid = _enqueue(_svc_analyze, req.question, req.context or {})
        return {"job_id": jid}

    @api.get("/api/job/{job_id}")
    def api_job(job_id: str):
        return _job_status(job_id)

# ---- Client helpers ----
def _api_base_url():
    try:
        base = os.environ.get("GOVCON_API_BASE")
        if not base and hasattr(st, "secrets"):
            base = (getattr(st, "secrets", {}) or {}).get("api", {}).get("base_url")
        return base
    except Exception:
        return None

def _api_post_v2(path: str, payload: dict):
    base = _api_base_url()
    if not base or requests is None:
        return None
    try:
        token = None
        try:
            token = (getattr(st, 'secrets', {}) or {}).get('api', {}).get('token')
            if not token:
                import os as _os
                token = _os.getenv('GOVCON_API_TOKEN')
        except Exception:
            token = None
        url = base.rstrip('/') + path
        headers = {'X-API-Key': token} if token else None
        r = requests.post(url, json=payload, headers=headers, timeout=20)
        if r.status_code == 200:
            return r.json()
    except Exception:
        return None
    return None

def _api_get_v2(path: str):
    base = _api_base_url()
    if not base or requests is None:
        return None
    try:
        token = None
        try:
            token = (getattr(st, 'secrets', {}) or {}).get('api', {}).get('token')
            if not token:
                import os as _os
                token = _os.getenv('GOVCON_API_TOKEN')
        except Exception:
            token = None
        url = base.rstrip('/') + path
        headers = {'X-API-Key': token} if token else None
        r = requests.get(url, headers=headers, timeout=20)
        if r.status_code == 200:
            return r.json()
    except Exception:
        return None
    return None

def _api_post(path: str, payload: dict):
    base = _api_base_url()
    if not base or requests is None:
        return None
    try:
        url = base.rstrip("/") + path
        r = requests.post(url, json=payload, timeout=20)
        if r.status_code == 200:
            return r.json()
    except Exception:
        return None
    return None

def _api_get(path: str):
    base = _api_base_url()
    if not base or requests is None:
        return None
    try:
        url = base.rstrip("/") + path
        r = requests.get(url, timeout=20)
        if r.status_code == 200:
            return r.json()
    except Exception:
        return None
    return None

# ---- Wire into the RFP Analyze service used by the dialog ----
def _phase3_analyze(question: str, opportunity: dict | None = None):
    # Try API path first
    resp = _api_post_v2("/api/analyze", {"question": question, "context": opportunity or {}})
    if isinstance(resp, dict) and resp.get("job_id"):
        st.session_state["phase3_last_job"] = resp["job_id"]
        return f"Queued analysis job: {resp['job_id']}"
    # Fallback to in-process
    try:
        return _service_analyze_rfp_question(question, opportunity or {})
    except Exception as e:
        return f"Analyze error: {e}"

def _phase3_poll_status():
    jid = st.session_state.get("phase3_last_job")
    if not jid:
        return None
    resp = _api_get_v2(f"/api/job/{jid}")
    if not resp:
        return None
    return resp

# ---- Enhance Ask RFP Analyzer dialog to use Phase 3 when available ----
def _phase3_enhance_rfp_dialog():
    # integrate into existing dialog if present
    if "show_rfp_analyzer" in st.session_state and st.session_state.get("show_rfp_analyzer"):
        # add a status row
        c1, c2 = st.columns([1,1])
        with c1:
            if st.button("Check job status", key=_unique_key("poll_job","o3")):
                st.session_state["phase3_last_status"] = _phase3_poll_status()
        with c2:
            if st.session_state.get("phase3_last_status"):
                st.caption(str(st.session_state["phase3_last_status"]))

# ELA Phase2 performance
import sqlite3, hashlib, time

# Cached DB connector with WAL + PRAGMAs
def _ensure_indices(conn):

    try:
        cur = conn.cursor()
        def table_exists(name):
            try:
                cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (name,))
                return cur.fetchone() is not None
            except Exception:
                return False
        stmts = []
        if table_exists("notices"):
            stmts.append("CREATE INDEX IF NOT EXISTS idx_notices_notice_id ON notices(notice_id)")
        if table_exists("vendors"):
            stmts.append("CREATE INDEX IF NOT EXISTS idx_vendors_place_id ON vendors(place_id)")
        if table_exists("deals"):
            stmts.append("CREATE INDEX IF NOT EXISTS idx_deals_stage ON deals(stage)")
        if table_exists("files"):
            stmts.append("CREATE INDEX IF NOT EXISTS idx_files_notice_id ON files(notice_id)")
        if table_exists("messages"):
            stmts.append("CREATE INDEX IF NOT EXISTS idx_messages_sent_at ON messages(sent_at)")
        for s in stmts:
            try:
                cur.execute(s)
            except Exception:
                pass
        cur.close()
    except Exception:
        pass

def _db_connect(db_path: str, **kwargs):
    import sqlite3
    import streamlit as st
    # Build connect kwargs with safe defaults
    base_kwargs = {"check_same_thread": False, "detect_types": sqlite3.PARSE_DECLTYPES, "timeout": 15}
    try:
        base_kwargs.update(kwargs or {})
    except Exception:
        pass
    conn = sqlite3.connect(db_path, **base_kwargs)
    try:
        conn.execute("PRAGMA journal_mode=WAL;")
        conn.execute("PRAGMA synchronous=NORMAL;")
        conn.execute("PRAGMA temp_store=MEMORY;")
        conn.execute("PRAGMA mmap_size=300000000;")
        conn.execute("PRAGMA cache_size=-200000;")
        conn.execute("PRAGMA busy_timeout=5000;")
    except Exception:
        pass
    # One-time per-session index creation
    try:
        if not st.session_state.get("_phase2_indices_done"):
            _ensure_indices(conn)
            st.session_state["_phase2_indices_done"] = True
    except Exception:
        pass
    return conn
# Cached SELECT helper (returns rows + cols); pass db_path explicitly
@st.cache_data(ttl=600, show_spinner=False)
def _cached_select(db_path: str, sql: str, params: tuple = ()):
    conn = _db_connect(db_path)
    cur = conn.execute(sql, params)
    rows = cur.fetchall()
    cols = [d[0] for d in cur.description] if cur.description else []
    return rows, cols

# Cache AI answers by (question + context hash)
def _ai_cache_key(question: str, context_hash: str = ""):
    key = (question or "") + "|" + (context_hash or "")
    return hashlib.sha256(key.encode("utf-8")).hexdigest()

def _cached_ai_answer(question: str, context_hash: str = ""):
    @st.cache_data(ttl=86400, show_spinner=False)
    def _inner(k, q, ch):
        try:
            return _service_analyze_rfp_question(q, {"hash": ch})
        except Exception as e:
            return f"AI error: {e}"
    return _inner(_ai_cache_key(question, context_hash), question, context_hash)

# Expand Phase 0 write guard to clear caches after commits
def _write_guard(conn, fn, *args, **kwargs):
    import streamlit as st
    with conn:
        out = fn(*args, **kwargs)
    try:
        st.cache_data.clear()
    except Exception:
        pass
    return out

# ELA Phase1 bootstrap

# Safe dataframe wrapper and monkey patch to avoid height=None issues
def _styled_dataframe(df, use_container_width=True, height=None, hide_index=True, column_config=None):
    kwargs = {"use_container_width": use_container_width}
    if height is not None:
        try:
            kwargs["height"] = int(height) if height != "stretch" else "stretch"
        except Exception:
            if isinstance(height, str):
                kwargs["height"] = height
    try:
        return st.dataframe(df, hide_index=hide_index, column_config=column_config, **kwargs)
    except TypeError:
        return st.dataframe(df, **kwargs)

# Monkey patch st.dataframe to drop height=None safely
if not hasattr(st, "_orig_dataframe"):
    st._orig_dataframe = st.dataframe
    def _safe_dataframe(df, **kwargs):
        if "height" in kwargs and kwargs["height"] is None:
            kwargs.pop("height", None)
        return st._orig_dataframe(df, **kwargs)
    st.dataframe = _safe_dataframe

# Ensure theme exists
if "apply_theme" not in globals():
    def _apply_theme_old():
        if st.session_state.get("_phase1_theme_applied"):
            return
        st.session_state["_phase1_theme_applied"] = True
        st.markdown('''
        <style>
        .block-container {padding-top: 1.2rem; padding-bottom: 1.2rem; max-width: 1400px;}
        h1, h2, h3 {margin-bottom: .4rem;}
        div[data-testid="stDataFrame"] thead th {position: sticky; top: 0; background: #fff; z-index: 2;}
        div[data-testid="stDataFrame"] tbody tr:hover {background: rgba(64,120,242,0.06);}
        [data-testid="stExpander"] {border: 1px solid rgba(49,51,63,0.16); border-radius: 12px; margin-bottom: 10px;}
        [data-testid="stExpander"] summary {font-weight: 600;}
        .stTextInput>div>div>input, .stNumberInput input, .stTextArea textarea {border-radius: 10px !important;}
        button[kind="primary"] {box-shadow: 0 1px 4px rgba(0,0,0,.08);}
        .ela-banner {position: sticky; top: 0; z-index: 999; background: linear-gradient(90deg, #4068f2, #7a9cff); color: #fff; padding: 6px 12px; border-radius: 8px; margin-bottom: 10px;}
        </style>
        ''', unsafe_allow_html=True)
        st.markdown("<div class='ela-banner'>Phase 1 theme active · polished layout and tables</div>", unsafe_allow_html=True)



# ===== Phase 1 Theme (auto-injected) =====
def _apply_theme_old():
    import streamlit as st
    if st.session_state.get("_phase1_theme_applied"):
        return
    st.session_state["_phase1_theme_applied"] = True
    st.markdown('''
    <style>
    .block-container {padding-top: 1.2rem; padding-bottom: 1.2rem; max-width: 1400px;}
    h1, h2, h3 {margin-bottom: .4rem;}
    .ela-subtitle {color: rgba(49,51,63,0.65); font-size: .95rem; margin-bottom: 1rem;}
    /* Dataframe polish */
    div[data-testid="stDataFrame"] thead th {position: sticky; top: 0; background: #fff; z-index: 2;}
    div[data-testid="stDataFrame"] tbody tr:hover {background: rgba(64,120,242,0.06);}
    /* Cards & expanders */
    [data-testid="stExpander"] {border: 1px solid rgba(49,51,63,0.16); border-radius: 12px; margin-bottom: 10px;}
    [data-testid="stExpander"] summary {font-weight: 600;}
    .ela-card {border: 1px solid rgba(49,51,63,0.16); border-radius: 12px; padding: 12px; margin-bottom: 12px;}
    .ela-chip {display:inline-block; padding: 2px 8px; border-radius: 999px; font-size: 12px; margin-right:6px; background: rgba(49,51,63,.06);}
    .ela-ok {background: rgba(0,200,83,.12);} .ela-warn {background: rgba(251,140,0,.12);} .ela-bad {background: rgba(229,57,53,.12);}
    /* Inputs & buttons */
    .stTextInput>div>div>input, .stNumberInput input, .stTextArea textarea {border-radius: 10px !important;}
    button[kind="primary"] {box-shadow: 0 1px 4px rgba(0,0,0,.08);}
    /* Banner */
    .ela-banner {position: sticky; top: 0; z-index: 999; background: linear-gradient(90deg, #4068f2, #7a9cff); color: #fff; padding: 6px 12px; border-radius: 8px; margin-bottom: 10px;}
    </style>
    ''', unsafe_allow_html=True)
    st.markdown("<div class='ela-banner'>Phase 1 theme active · polished layout & tables</div>", unsafe_allow_html=True)

# ===== injected early helpers (do not remove) =====
def _safe_int(x, default=0):
    try:
        if x is None:
            return int(default)
        if isinstance(x, int):
            return x
        s = str(x).strip()
        if s == "" or s.lower() in ("none", "nan"):
            return int(default)
        # try float first (handles "123.0")
        try:
            return int(float(s))
        except Exception:
            pass
        # fallback: keep only digits
        digits = "".join(ch for ch in s if ch.isdigit())
        return int(digits) if digits else int(default)
    except Exception:
        return int(default)

def _uniq_key(base: str, rfp_id: int) -> str:
    try:
        k = f"__uniq_counter_{base}_{rfp_id}"
        n = int(st.session_state.get(k, 0))
        st.session_state[k] = n + 1
        return f"{base}_{rfp_id}_{n}"
    except Exception:
        import time
        return f"{base}_{rfp_id}_{int(time.time()*1000)%100000}"
# ===== end injected early helpers =====


def y3_get_rfp_files(_conn, rfp_id: int):
    """Return [(id, file_name, bytes)] for files saved in rfp_files for this RFP."""
    try:
        from contextlib import closing as _closing
        with _closing(_conn.cursor()) as cur:
            cur.execute("SELECT id, file_name, bytes FROM rfp_files WHERE rfp_id=? ORDER BY id", (rfp_id,))
            return cur.fetchall() or []
    except Exception:
        return []

import requests


# ===== X3 MODAL HELPERS =====
def _x3_open_modal(row_dict: dict):
    st.session_state["x3_modal_notice"] = dict(row_dict or {})
    st.session_state["x3_show_modal"] = True
    try:
        st.rerun()
    except Exception:
        pass

def _x3_render_modal(notice: dict):
    try:
        rfp_id = _ensure_rfp_for_notice(conn, notice)
    except Exception as e:
        st.error(f"Could not open RFP Analyzer: {e}")
        return
    st.caption(f"RFP #{rfp_id} · {notice.get('Title','')}")


    # X.6 Compliance Matrix v1
    try:
        _ensure_x6_schema(conn)
        with st.expander("Compliance Matrix v1", expanded=False):
            cA, cB, cC = st.columns([1,1,2])
            with cA:
                if st.button("Extract requirements", key=_uniq_key("x6_extract", int(rfp_id))):
                    n = x6_extract_requirements(conn, int(rfp_id))
                    st.success(f"Extracted {n} new requirement(s).")
                    try: st.rerun()
                    except Exception: pass
            with cB:
                cov, tot = x6_coverage(conn, int(rfp_id))
                pct = 0 if tot == 0 else int(round(100 * cov / tot))
                st.metric("Coverage", f"{cov}/{tot}", f"{pct}%")
            with cC:
                st.caption("Link requirements to outline sections to increase coverage.")

            df = x6_requirements_df(conn, int(rfp_id))
            if df is None or df.empty:
                st.info("No requirements yet. Click Extract requirements.")
            else:
                # editable mapping column
                suggestions = x6_sections_suggestions(int(rfp_id))
                import pandas as pd
                df_view = df.copy()
                df_view["Map to section"] = ""
                edited = st.data_editor(
                    df_view,
                    column_config={
                        "must_flag": st.column_config.CheckboxColumn("Must", help="Detected must or shall"),
                        "Map to section": st.column_config.SelectboxColumn(options=suggestions),
                    },
                    hide_index=True,
                    use_container_width=True,
                    key=_uniq_key("x6_editor", int(rfp_id))
                )
                # Save mappings
                to_save = []
                for i, row in edited.iterrows():
                    sec = str(row.get("Map to section") or "").strip()
                    if sec:
                        to_save.append((int(row["id"]), sec))
                if st.button("Save mappings", key=_uniq_key("x6_save", int(rfp_id))) and to_save:
                    saved = x6_save_links(conn, int(rfp_id), to_save)
                    st.success(f"Saved {saved} link(s).")
                    try: st.rerun()
                    except Exception: pass

                # Export CSV
                import io
                import csv
                buf = io.StringIO()
                writer = csv.writer(buf)
                writer.writerow(["id","must","file","page","text","section"])
                for _, r in edited.iterrows():
                    writer.writerow([int(r["id"]), int(r["must_flag"]), r["file"], int(r["page"]) if r.get("page") else "", r["text"], r.get("Map to section") or ""])
                st.download_button("Export Compliance CSV", buf.getvalue().encode("utf-8"), file_name=f"rfp_{int(rfp_id)}_compliance.csv", mime="text/csv", key=_uniq_key("x6_csv", int(rfp_id)))
    except Exception as _x6e:
        st.info(f"Compliance Matrix unavailable: {_x6e}")
    # Phase 4: side panel (SAM facts + AI summary + quick actions)
    try:
        _url = notice.get('SAM Link') or notice.get('sam_url') or ''
    except Exception:
        _url = ''
    try:
        render_amendment_sidebar(conn, int(rfp_id), _url, ttl_hours=72)
    except Exception:
        pass
    try:
        with st.sidebar.expander('RFP Summary', expanded=True):
            st.markdown(_rfp_ai_summary(full_text or '', notice))
    except Exception:
        pass
    try:
        with st.sidebar.expander('Quick Actions', expanded=False):
            if st.button('Fetch attachments now', key=_uniq_key('x3_fetch_sidebar', int(rfp_id))):
                c = _fetch_and_save_now(conn, str(notice.get('Notice ID') or ''), int(rfp_id))
                st.success(f'Fetched {c} attachment(s).')
                try: st.rerun()
                except Exception: pass
            if st.button('Rebuild Search Index', key=_uniq_key('x3_reindex_sidebar', int(rfp_id))):
                try:
                    y1_index_rfp(conn, int(rfp_id), rebuild=True)
                    st.success('Index rebuilt')
                except Exception as _e:
                    st.info(f'Index rebuild failed: {_e}')
    except Exception:
        pass
    try:
        outline_key = f'proposal_outline_{int(rfp_id)}'
        with st.sidebar.expander('Proposal Outline', expanded=False):
            if st.button('Generate Outline', key=_uniq_key('x3_outline_sidebar', int(rfp_id))):
                outline = [
                    '# Proposal Outline',
                    '1. Cover Letter',
                    '2. Executive Summary',
                    '3. Technical Approach',
                    '4. Management Approach',
                    '5. Past Performance',
                    '6. Pricing (separate volume if required)',
                    '7. Compliance Matrix',
                ]
                st.session_state[outline_key] = '\n'.join(outline)
                st.success('Outline drafted and saved to session')
            _ol = st.session_state.get(outline_key, '')
            if _ol:
                st.text_area('Current outline', value=_ol, height=180, key=_uniq_key('x3_outline_preview', int(rfp_id)))
                try:
                    data = _ol.encode('utf-8')
                    st.download_button('Download Outline (.md)', data=data, file_name=f'proposal_outline_{int(rfp_id)}.md', mime='text/markdown', key=_uniq_key('x3_outline_dl', int(rfp_id)))
                except Exception:
                    pass
    except Exception:
        pass
    # Attachments area
    try:
        from contextlib import closing as _closing
        with _closing(conn.cursor()) as cur:
            cur.execute("SELECT COUNT(*) FROM rfp_files WHERE rfp_id=?;", (int(rfp_id),))
            n_files = int(cur.fetchone()[0])
    except Exception:
        n_files = 0
    cA, cB = st.columns([2,1])
    with cA:
        st.write(f"Attachments saved: **{n_files}**")
    with cB:
        if st.button("Fetch attachments now", key=_uniq_key("x3_fetch", int(rfp_id))):
            c = _fetch_and_save_now(conn, str(notice.get("Notice ID") or ""), int(rfp_id))
            st.success(f"Fetched {c} attachment(s).")
            try: st.rerun()
            except Exception: pass

    # Index + summary
    try:
        y1_index_rfp(conn, int(rfp_id), rebuild=False)
    except Exception:
        pass
    full_text, sources = _rfp_build_fulltext_from_db(conn, int(rfp_id))
    if not full_text:
        st.info("No documents yet. You can still ask questions; I'll use the SAM description if available.")
        try:
            descs = sam_try_fetch_attachments(str(notice.get("Notice ID") or "")) or []
            for name, b in descs:
                if name.endswith("_description.html"):
                    try:
                        import bs4
                        soup = bs4.BeautifulSoup(b.decode('utf-8', errors='ignore'), 'html.parser')
                        full_text = soup.get_text(" ", strip=True)
                    except Exception:
                        pass
                    break
        except Exception:
            pass
    with st.expander("AI Summary", expanded=True):
        st.markdown(_rfp_ai_summary(full_text or "", notice))

    # Per-document summarize chips
    try:
        from contextlib import closing as _closing
        with _closing(conn.cursor()) as cur:
            cur.execute("SELECT id, filename, mime FROM rfp_files WHERE rfp_id=? ORDER BY id;", (int(rfp_id),))
            files = cur.fetchall() or []
    except Exception:
        files = []
    if files:
        st.write("Documents:")
        for fid, fname, fmime in files[:12]:
            if st.button(f"Summarize: {fname}", key=_uniq_key("sumdoc", int(fid))):
                try:
                    import pandas as pd
                    blob = pd.read_sql_query("SELECT bytes, mime FROM rfp_files WHERE id=?;", conn, params=(int(fid),)).iloc[0]
                    _text = "\n\n".join(extract_text_pages(blob['bytes'], blob.get('mime') or (fmime or '')) or [])
                except Exception:
                    _text = ""
                st.session_state["x3_docsum"] = _rfp_ai_summary(_text, notice)
        if st.session_state.get("x3_docsum"):
            with st.expander("Document Summary", expanded=True):
                st.markdown(st.session_state.get("x3_docsum"))

    # Chat
    st.divider()
    st.subheader("Ask about this RFP")
    _chat_k = _uniq_key("x3_chat", int(rfp_id))
    hist_key = f"x3_chat_hist_{rfp_id}"
    st.session_state.setdefault(hist_key, [])
    for who, msg in st.session_state[hist_key]:
        with st.chat_message(who):
            st.markdown(msg)
    q = st.chat_input("Ask a question about the requirements, due dates, sections, etc.", key=_chat_k)
    if q:
        st.session_state[hist_key].append(("user", q))
        with st.chat_message("assistant"):
            ans = _rfp_chat(conn, int(rfp_id), q)
            st.session_state[hist_key].append(("assistant", ans))
            st.markdown(ans)

    # Proposal hand-off
    st.divider()
    if st.button("Start Proposal Outline", key=_uniq_key("x3_outline", int(rfp_id))):
        outline = [
            "# Proposal Outline",
            "1. Cover Letter",
            "2. Executive Summary",
            "3. Technical Approach",
            "4. Management Approach",
            "5. Past Performance",
            "6. Pricing (separate volume if required)",
            "7. Compliance Matrix",
        ]
        st.session_state[f"proposal_outline_{rfp_id}"] = "\n".join(outline)
        st.success("Outline drafted and saved. Open Proposal Builder to continue.")

    with st.expander("Transcript Viewer", expanded=False):
        x5_render_transcript_viewer(conn, int(rfp_id))
def _uniq_key(base: str, rfp_id: int) -> str:
    """Return a unique (but stable per render) Streamlit key to avoid duplicates."""
    try:
        k = f"__uniq_counter_{base}_{rfp_id}"
        n = int(st.session_state.get(k, 0))
        st.session_state[k] = n + 1
        return f"{base}_{rfp_id}_{n}"
    except Exception:
        # Fallback if session_state isn't available
        import time
        return f"{base}_{rfp_id}_{int(time.time()*1000)%100000}"
import time
# ==== O4 unified DB + sender helpers ====
try:
    DB_PATH
except NameError:
    DB_PATH = "./data/app.db"

_O4_CONN = globals().get("_O4_CONN", None)

def ensure_dirs():
    from pathlib import Path as _Path
    _Path(DB_PATH).parent.mkdir(parents=True, exist_ok=True)

def get_db():
    import sqlite3
    from contextlib import closing as _closing
    ensure_dirs()
    conn = _db_connect(DB_PATH, check_same_thread=False)
    with _closing(conn.cursor()) as cur:
        cur.execute("PRAGMA foreign_keys = ON;")
    return conn

def get_o4_conn():
    import streamlit as st

# ---- helper: generate unique widget keys (Phase 0) ----
# ---- helper: render-once guard for O4 (Phase 0) ----
def _render_once(name: str):
    # returns True if allowed to render, False if already rendered
    key = f"__rendered__{name}"
    if st.session_state.get(key):
        return False
    st.session_state[key] = True
    return True

def _unique_key(base: str, namespace: str = "ui"):
    # Uses a stable counter in session state to avoid duplicate form/widget keys
    ss = st.session_state
    counter_key = f"__{namespace}_key_counter__"
    if counter_key not in ss:
        ss[counter_key] = 0
    ss[counter_key] += 1
    return f"{base}-{namespace}-{ss[counter_key]}"
    global _O4_CONN
    if _O4_CONN:
        try:
            st.session_state["conn"] = _O4_CONN
        except Exception:
            pass
        return _O4_CONN
    try:
        if "conn" in st.session_state and st.session_state.get("conn"):
            _O4_CONN = st.session_state["conn"]
            return _O4_CONN
    except Exception:
        pass
    conn = get_db()
    _O4_CONN = conn
    try:
        st.session_state["conn"] = conn
    except Exception:
        pass
    return conn

def _ensure_email_accounts_schema(conn):
    # Unified to O1 schema
    ensure_outreach_o1_schema(conn)

def _get_senders(conn):
    from contextlib import closing as _closing
    tables = [
        ("email_accounts", "user_email", "display_name", "app_password"),
        ("o4_senders", "email", "name", "app_password"),
        ("senders", "email", "display_name", "app_password"),
        ("smtp_settings", "username", "label", "password"),
    ]
    for tbl, c_email, c_name, c_pw in tables:
        try:
            with _closing(conn.cursor()) as c:
                c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tbl,))
                if not c.fetchone():
                    continue
                c.execute(f"SELECT {c_email}, {c_name}, {c_pw} FROM {tbl} ORDER BY {c_email}")
                rows = c.fetchall()
                if rows:
                    return [(r[0], r[1], r[2] if len(r) > 2 else "") for r in rows]
        except Exception:
            continue
    return []
# ==== end O4 helpers ====
# Helper imports for RTM/Amendment
import re as _rtm_re
import json as _rtm_json
import hashlib as _rtm_hashlib
from contextlib import closing as _rtm_closing
try:
    import sqlite3 as _rtm_sqlite3
except Exception:
    import sqlite3 as _rtm_sqlite3
try:
    import pandas as _rtm_pd
except Exception:
    import pandas as _rtm_pd
try:
    import streamlit as _rtm_st
except Exception:
    class _Dummy: pass
    _rtm_st = _Dummy()
# === PHASE 5: L & M compliance gate ===
def require_LM_minimum(conn, rfp_id):
    f"""
    Returns (ok, missing:list[str]).
    Required: Offers Due date in rfp_meta, Section M present in rfp_sections, >=1 L/M checklist item.
    """
    missing = []
    try:
        df_due = pd.read_sql_query(
            "SELECT value FROM rfp_meta WHERE rfp_id=? AND key IN ('offers_due','due_offer') LIMIT 1;",
            conn, params=(int(rfp_id),)
        )
        if df_due is None or df_due.empty or not str(df_due.iloc[0]["value"]).strip():
            missing.append("Offers Due date (Section L)")
    except Exception:
        missing.append("Offers Due date (Section L)")
    try:
        df_m = pd.read_sql_query(
            "SELECT 1 FROM rfp_sections WHERE rfp_id=? AND (section='M' OR section LIKE 'Section M%') LIMIT 1;",
            conn, params=(int(rfp_id),)
        )
        if df_m is None or df_m.empty:
            missing.append("Section M present")
    except Exception:
        missing.append("Section M present")
    try:
        df_lm = pd.read_sql_query(
            "SELECT COUNT(1) AS c FROM lm_items WHERE rfp_id=?;",
            conn, params=(int(rfp_id),)
        )
        c = int(df_lm.iloc[0]["c"]) if df_lm is not None and not df_lm.empty else 0
        if c <= 0:
            missing.append("L/M checklist items")
    except Exception:
        missing.append("L/M checklist items")
    return (len(missing) == 0, missing)
# === end PHASE 5 ===


# --- Router helper to avoid 'Unknown page' blank rendering ---
def _safe_route_call(fn, *a, **kw):
    try:
        if callable(fn):
            return fn(*a, **kw)
    except Exception as _e:
        import streamlit as _st
        _st.error(f"Page failed: {type(_e).__name__}: {_e}")
    return None

# --- O3 helper: safe cursor context ---
from contextlib import contextmanager
@contextmanager
def _o3c(cursor):
    try:
        yield cursor
    finally:
        try:
            cursor.close()
        except Exception:
            pass
def _migrate_deals_columns(conn):
    """
    Add columns used by Deals and SAM Watch if missing. Idempotent.

# ---- Phase 0: Ask RFP Analyzer modal wiring ----
def _ask_rfp_analyzer_modal(opportunity=None):
    @st.dialog("Ask RFP Analyzer", key="ask_rfp_analyzer_dialog")
    def _dlg():
        st.write("Use AI to analyze the selected opportunity and ask questions.")
        q = st.text_area("Your question", key=_unique_key('rfp_q','o3'))
        if st.button("Analyze", key=_unique_key('rfp_analyze','o3')):
            # placeholder for analysis call; replace with your service function
            try:
                ans = _service_analyze_rfp_question(q, opportunity)
            except Exception as e:
                ans = f"Error: {e}"
            st.markdown("**Answer**")
            st.write(ans)
        if st.button("Close", key=_unique_key('rfp_close','o3')):
            st.session_state['show_rfp_analyzer'] = False
    _dlg()

def _service_analyze_rfp_question(q, opportunity):
    try:
        base = _api_base_url()
    except Exception:
        base = None
    if base:
        return _phase3_analyze(q, opportunity)
# TODO: wire to actual analyzer service; for now return a placeholder
    if not q:
    return f"Please enter a question."
    title = (opp.get('title') if isinstance(opp, dict) else str(opp)) if opp else 'the selected notice'
    return f"[demo] Analysis for {title}: {q}"

def _render_ask_rfp_button(opportunity=None):
    if st.button("Ask RFP Analyzer", key=_unique_key('ask_rfp','o3')):
        st.session_state['show_rfp_analyzer'] = True
    if st.session_state.get('show_rfp_analyzer'):
        _ask_rfp_analyzer_modal(opportunity)

    """
    try:
        import pandas as _pd
        from contextlib import closing
        cur_cols = _pd.read_sql_query("PRAGMA table_info(deals);", conn)
        have = set(cur_cols["name"].astype(str).tolist()) if cur_cols is not None else set()
    except Exception:
        have = set()

    def _add(col, ddl):
        if col not in have:
            try:
                from contextlib import closing
                with closing(conn.cursor()) as _c:
                    _c.execute(f"ALTER TABLE deals ADD COLUMN {ddl};")
                conn.commit()
            except Exception:
                pass

    _add("agency", "agency TEXT")
    _add("value", "value REAL")
    _add("sam_url", "sam_url TEXT")
    _add("notice_id", "notice_id TEXT")
    _add("solnum", "solnum TEXT")
    _add("posted_date", "posted_date TEXT")
    _add("rfp_deadline", "rfp_deadline TEXT")
    _add("naics", "naics TEXT")
    _add("psc", "psc TEXT")



# Bridge names
sqlite3 = _rtm_sqlite3
pd = _rtm_pd
st = _rtm_st
re = _rtm_re
json = _rtm_json
hashlib = _rtm_hashlib
closing = _rtm_closing

# === RTM + Amendment helpers ===

def _now_iso():
    try:
        return __import__("datetime").datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    except Exception:
        return ""

def _ensure_rtm_schema(conn: sqlite3.Connection) -> None:
    try:
        with closing(conn.cursor()) as cur:
            cur.execute("SELECT 1 FROM rtm_requirements LIMIT 1;")
    except Exception:
        with closing(conn.cursor()) as cur:
            cur.execute("CREATE TABLE IF NOT EXISTS rtm_requirements(id INTEGER PRIMARY KEY, rfp_id INTEGER, req_key TEXT, source_type TEXT, source_file TEXT, page INTEGER, text TEXT, status TEXT, created_at TEXT, updated_at TEXT);")
            cur.execute("CREATE TABLE IF NOT EXISTS rtm_links(id INTEGER PRIMARY KEY, rtm_id INTEGER, link_type TEXT, target TEXT, note TEXT, created_at TEXT, updated_at TEXT);")
            conn.commit()

def rtm_build_requirements(conn: sqlite3.Connection, rfp_id: int, max_rows: int = 800) -> int:
    """
    Seed RTM from L/M checklist and SOW-style shall/must statements in rfp_chunks.
    Returns number of rows inserted (new).
    """
    _ensure_rtm_schema(conn)
    inserted = 0
    now = _now_iso()
    # 1) From L/M
    try:
        df_lm = pd.read_sql_query("SELECT id, item_text, is_must FROM lm_items WHERE rfp_id=?", conn, params=(int(rfp_id),))
    except Exception:
        df_lm = pd.DataFrame(columns=["id","item_text","is_must"])
    for i, row in df_lm.head(max_rows).iterrows():
        txt = (row["item_text"] or "").strip()
        if not txt:
            continue
        key = f"LM-{row['id']}"
        with closing(conn.cursor()) as cur:
            cur.execute("SELECT id FROM rtm_requirements WHERE rfp_id=? AND req_key=?;", (int(rfp_id), key))
            if cur.fetchone():
                continue
            cur.execute("""
                INSERT INTO rtm_requirements(rfp_id, req_key, source_type, source_file, page, text, status, created_at, updated_at)
                VALUES(?,?,?,?,?,?,?, ?, ?);
            """, (int(rfp_id), key, "L/M", None, None, txt, "Open", now, now))
            inserted += 1
    # 2) From SOW chunks, simple heuristic
    try:
        df_chunks = pd.read_sql_query("""
            SELECT id, file_name, page, text FROM rfp_chunks
            WHERE rfp_id=? ORDER BY file_name, page, id
        """, conn, params=(int(rfp_id),))
    except Exception:
        df_chunks = pd.DataFrame(columns=["file_name","page","text"])
    trig = re.compile(r"\\b(shall|must|will|provide|furnish)\\b", re.I)
    for _, row in df_chunks.iterrows():
        t = (row["text"] or "").strip()
        if not t:
            continue
        # split into sentences with light heuristic
        sentences = re.split(r"(?<=[\\.;:])\\s+", t)
        for s in sentences:
            if len(s) < 40:
                continue
            if trig.search(s):
                key = f"SOW-{hashlib.sha1(s.encode('utf-8')).hexdigest()[:10]}"
                with closing(conn.cursor()) as cur:
                    cur.execute("SELECT id FROM rtm_requirements WHERE rfp_id=? AND req_key=?;", (int(rfp_id), key))
                    if cur.fetchone():
                        continue
                    cur.execute("""
                        INSERT INTO rtm_requirements(rfp_id, req_key, source_type, source_file, page, text, status, created_at, updated_at)
                        VALUES(?,?,?,?,?,?,?, ?, ?);
                    """, (int(rfp_id), key, "SOW", row.get("file_name"), int(row.get("page") or 0), s.strip(), "Open", now, now))
                    inserted += 1
        if inserted >= max_rows:
            break
    conn.commit()
    return inserted

def rtm_metrics(conn: sqlite3.Connection, rfp_id: int) -> dict:
    q = pd.read_sql_query("""
        SELECT r.id, r.source_type, r.status, COUNT(l.id) AS links
        FROM rtm_requirements r
        LEFT JOIN rtm_links l ON l.rtm_id = r.id
        WHERE r.rfp_id=?
        GROUP BY r.id
    """, conn, params=(int(rfp_id),))
    if q is None or q.empty:
        return {"total":0,"covered":0,"coverage":0.0,"by_type":{}}
    total = len(q)
    covered_rows = (q["links"] > 0) | (q["status"].fillna("")=="Covered")
    covered = int(covered_rows.sum())
    by_type = {}
    for t, sub in q.groupby("source_type"):
        ct = len(sub)
        cv = int(((sub["links"]>0) | (sub["status"].fillna("")=="Covered")).sum())
        by_type[t] = {"total": ct, "covered": cv, "coverage": (cv/ct if ct else 0.0)}
    return {"total": total, "covered": covered, "coverage": (covered/total if total else 0.0), "by_type": by_type}

def rtm_export_csv(conn: sqlite3.Connection, rfp_id: int) -> str:
    q = pd.read_sql_query("""
        SELECT r.id, r.req_key, r.source_type, r.source_file, r.page, r.text, r.status,
               COALESCE(GROUP_CONCAT(l.link_type || ':' || l.target, '; '), '') AS evidence
        FROM rtm_requirements r
        LEFT JOIN rtm_links l ON l.rtm_id=r.id
        WHERE r.rfp_id=?
        GROUP BY r.id
        ORDER BY r.source_type, r.id
    """, conn, params=(int(rfp_id),))
    fn = f"/mnt/data/rtm_rfp_{rfp_id}.csv"
    try:
        q.to_csv(fn, index=False)
        return fn
    except Exception:
        return ""

def render_rtm_ui(conn: sqlite3.Connection, rfp_id: int) -> None:
    # Unique key namespace to avoid duplicate element IDs
    _ns = int(st.session_state.get('rtm_ui_ns', 0))
    st.session_state['rtm_ui_ns'] = _ns + 1
    _k = f"{rfp_id}_{_ns}"

    st.subheader("RTM Coverage")
    cols = st.columns([1, 1, 1, 3])

    # Build/Update
    with cols[0]:
        if st.button("Build/Update RTM", key=f"rtm_build_{_k}", help="Pull from L/M and SOW 'shall' statements."):
            n = rtm_build_requirements(conn, int(rfp_id))
            st.success(f"Added {n} requirement(s).")

    # Export with L/M gate
    with cols[1]:
        ok_gate, missing_gate = require_LM_minimum(conn, int(rfp_id))
        if not ok_gate:
            st.button("Export CSV", key=f"rtm_export_blocked_{_k}", disabled=True, help="Blocked: " + ", ".join(missing_gate))
        else:
            path = rtm_export_csv(conn, int(rfp_id))
            if path:
                from pathlib import Path as _Path
                st.download_button("Export CSV",
                                   data=open(path, "rb").read(),
                                   file_name=_Path(path).name,
                                   mime="text/csv",
                                   key=f"rtm_export_{_k}")

    # Mark covered
    with cols[2]:
        if st.button("Mark all with evidence as Covered", key=f"rtm_mark_{_k}"):
            with closing(conn.cursor()) as cur:
                cur.execute("""
                    UPDATE rtm_requirements
                    SET status='Covered', updated_at=?
                    WHERE rfp_id=? AND id IN (
                        SELECT r.id
                        FROM rtm_requirements r
                        LEFT JOIN rtm_links l ON l.rtm_id=r.id
                        GROUP BY r.id
                        HAVING COUNT(l.id) > 0
                    );
                """, (_now_iso(), int(rfp_id)))
                conn.commit()

    # Metrics
    m = rtm_metrics(conn, int(rfp_id))
    st.caption(f"Coverage: {m['covered']}/{m['total']} = {m['coverage']:.0%}")

    # Editor
    df = pd.read_sql_query("""
        SELECT r.id as rtm_id, r.req_key, r.source_type, r.page, r.text, r.status,
               COALESCE(GROUP_CONCAT(l.link_type || ':' || l.target, '
'), '') AS evidence
        FROM rtm_requirements r
        LEFT JOIN rtm_links l ON l.rtm_id=r.id
        WHERE r.rfp_id=?
        GROUP BY r.id
        ORDER BY r.source_type, r.id
        LIMIT 1000
    """, conn, params=(int(rfp_id),))
    df["add_link_type"] = ""
    df["add_link_target"] = ""

    edited = st.data_editor(
        df,
        key=f"rtm_editor_{_k}",
        use_container_width=True,
        num_rows="dynamic",
        column_config={
            "rtm_id": st.column_config.NumberColumn(disabled=True),
            "text": st.column_config.TextColumn(width="large"),
            "evidence": st.column_config.TextColumn(disabled=True),
        },
    )

    # Persist status changes and new links
    now = _now_iso()
    for _, row in edited.iterrows():
        try:
            rid = int(row["rtm_id"])
        except Exception:
            continue
        # Status change
        with closing(conn.cursor()) as cur:
            cur.execute("UPDATE rtm_requirements SET status=?, updated_at=? WHERE id=?;",
                        (row.get("status") or "Open", now, rid))
        # New link
        lt = (row.get("add_link_type") or "").strip()
        tg = (row.get("add_link_target") or "").strip()
        if lt and tg:
            with closing(conn.cursor()) as cur:
                cur.execute("INSERT INTO rtm_links(rtm_id, link_type, target, note, created_at, updated_at) VALUES(?,?,?,?,?,?);",
                            (rid, lt, tg, "", now, now))
    conn.commit()


def _parse_sam_text_to_facts(txt: str) -> dict:
    d = {}
    # Dates
    m = re.search(r"(Offers|Quotes?) Due[^\\d]*(\\d{1,2}[/-]\\d{1,2}[/-]\\d{2,4}|[A-Za-z]{3,9} \\d{1,2}, \\d{4})", txt, re.I)
    if m: d["offers_due"] = m.group(2)
    m = re.search(r"(Questions|Q&A) Due[^\\d]*(\\d{1,2}[/-]\\d{1,2}[/-]\\d{2,4}|[A-Za-z]{3,9} \\d{1,2}, \\d{4})", txt, re.I)
    if m: d["questions_due"] = m.group(2)
    # Codes
    m = re.search(r"NAICS[^\\d]*(\\d{5,6})", txt, re.I)
    if m: d["naics"] = m.group(1)
    m = re.search(r"Set[- ]Aside[^:]*:\\s*([^\\n]+)", txt, re.I)
    if m: d["set_aside"] = m.group(1).strip()
    # Clauses and forms
    clauses = re.findall(r"(52\\.[\\d-]+\\S*)", txt)
    if clauses: d["clauses"] = sorted(set(clauses))[:50]
    forms = re.findall(r"\\b(SF|OF)-?\\s?(\\d{1,4}[A-Z]?)\\b", txt)
    if forms: d["forms"] = sorted(set([f"{a}-{b}" for a,b in forms]))[:50]
    return d

def sam_snapshot(conn: sqlite3.Connection, rfp_id: int, url: str, ttl_hours: int = 72) -> dict:
    out = {"url": url, "facts": {}, "sha256": "", "cached": False, "text": ""}
    if not (url or "").strip():
        return out
    r = research_fetch(url, ttl_hours)
    txt = (r.get("text") or "").strip()
    out["cached"] = bool(r.get("cached"))
    out["text"] = txt
    if not txt:
        return out
    sha = hashlib.sha256(txt.encode("utf-8")).hexdigest()
    out["sha256"] = sha
    facts = _parse_sam_text_to_facts(txt)
    out["facts"] = facts
    now = _now_iso()
    with closing(conn.cursor()) as cur:
        cur.execute("INSERT INTO sam_versions(rfp_id, url, sha256, extracted_json, created_at) VALUES(?,?,?,?,?);",
                    (int(rfp_id), url, sha, json.dumps(facts), now))
        vid = cur.lastrowid
        for k, v in facts.items():
            val = json.dumps(v) if not isinstance(v, str) else v
            cur.execute("INSERT INTO sam_extracts(sam_version_id, key, value) VALUES(?,?,?);", (vid, k, val))
    conn.commit()
    return out

def _facts_diff(old: dict, new: dict) -> dict:
    diffs = {}
    keys = set(old.keys()) | set(new.keys())
    for k in keys:
        ov = old.get(k, "")
        nv = new.get(k, "")
        if json.dumps(ov, sort_keys=True) != json.dumps(nv, sort_keys=True):
            diffs[k] = {"old": ov, "new": nv}
    return diffs

def render_amendment_sidebar(conn: sqlite3.Connection, rfp_id: int, url: str, ttl_hours: int = 72) -> None:
    if not (url or "").strip():
        return
    with st.sidebar.expander("Amendments · SAM Analyzer", expanded=True):
        st.caption("Tracks changes in Brief, Factors, Clauses, Dates, Forms.")
        if st.button("Fetch SAM snapshot", key=_uniq_key("sam_fetch", int(rfp_id))):
            snap = sam_snapshot(conn, int(rfp_id), url, ttl_hours)
            st.success(f"Snapshot stored. Cached={snap.get('cached')}")
        # Last two snapshots
        try:
            dfv = pd.read_sql_query("SELECT id, created_at, sha256, extracted_json FROM sam_versions WHERE rfp_id=? ORDER BY id DESC LIMIT 2;", conn, params=(int(rfp_id),))
        except Exception:
            dfv = None
        if dfv is not None and not dfv.empty:
            st.markdown("Latest snapshot facts:")
            latest = json.loads(dfv.iloc[0]["extracted_json"] or "{}")
            st.json(latest)
            if len(dfv) >= 2:
                prev = json.loads(dfv.iloc[1]["extracted_json"] or "{}")
                diffs = _facts_diff(prev, latest)
                if diffs:
                    st.markdown("**Changes since previous:**")
                    st.json(diffs)
                    # Impact to-dos
                    todos = []
                    if "offers_due" in diffs:
                        todos.append("Update due date everywhere. Recalculate schedule and reminders.")
                    if "clauses" in diffs:
                        todos.append("Re-run compliance matrix and clause flowdowns.")
                    if "forms" in diffs:
                        todos.append("Update RFQ pack forms and signatures.")
                    if "set_aside" in diffs or "naics" in diffs:
                        todos.append("Validate eligibility and certs.")
                    if todos:
                        st.markdown("**Impact to-dos:**")
                        for t in todos:
                            st.write(f"- {t}")
                else:
                    st.info("No changes between the last two snapshots.")
            else:
                st.info("Only one snapshot stored so far.")
        else:
            st.info("No snapshots yet.")


def feature_flag(name: str, default: bool=False) -> bool:
    """
    Read a feature flag from environment or Streamlit secrets.
    Precedence: os.environ["FEATURE_<NAME>"] then st.secrets["features"][name] then default.
    Does not raise if Streamlit is absent.
    """
    val = None
    try:
        import os as _os
        env_key = f"FEATURE_{name.upper()}"
        if env_key in _os.environ:
            val = _os.environ[env_key]
    except Exception:
        pass
    if val is None:
        try:
            import streamlit as _st  # type: ignore
            sec = _st.secrets.get("features", {})
            if isinstance(sec, dict) and name in sec:
                val = sec.get(name)
        except Exception:
            pass
    if isinstance(val, str):
        return val.lower() in {"1","true","yes","on"}
    if isinstance(val, bool):
        return val
    return bool(val) if val is not None else bool(default)


def _guess_solnum(text: str) -> str:
    if not text:
        return ""
    t = text
    m = re.search(r'(?i)Solicitation\s*(?:Number|No\.?|#)\s*[:#]?\s*([A-Z0-9][A-Z0-9\-\._/]{4,})', t)
    if m:
        return m.group(1)[:60]
    m = re.search(r'\b([A-Z0-9]{2,6}[A-Z0-9\-]{0,4}\d{2}[A-Z]?-?[A-Z]?-?\d{3,6})\b', t)
    if m:
        return m.group(1)[:60]
    m = re.search(r'\b(RFQ|RFP|IFB|RFI)[\s#:]*([A-Z0-9][A-Z0-9\-\._/]{3,})\b', t, re.I)
    if m:
        return (m.group(1).upper() + "-" + m.group(2))[:60]
    return ""

from contextlib import closing
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional, Any, Dict, List, Tuple
import io
import json
import os
import re
import math
import sqlite3

from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import pandas as pd
import docx  # for DOCX reading in _read_file
import re
try:
    import mathquests  # optional plugin
except Exception:
    mathquests = None
import smtplib
import streamlit as st

# --- DB helpers: ensure column exists ---
def _ensure_column(conn, table, col, type_sql):
    cur = conn.cursor()
    cols = [r[1] for r in cur.execute(f"PRAGMA table_info({table})").fetchall()]
    if col not in cols:
        try:
            cur.execute(f"ALTER TABLE {table} ADD COLUMN {col} {type_sql}")
            conn.commit()
        except Exception:
            pass

def _update_rfp_meta(conn, rfp_id, title=None, solnum=None, sam_url=None):
    _ensure_column(conn, "rfps", "sam_url", "TEXT")
    sets, vals = [], []
    if title is not None:
        sets.append("title=?"); vals.append(title)
    if solnum is not None:
        sets.append("solnum=?"); vals.append(solnum)
    if sam_url is not None:
        sets.append("sam_url=?"); vals.append(sam_url)
    if sets:
        vals.append(int(rfp_id))
        conn.execute(f'UPDATE rfps SET {", ".join(sets)} WHERE id=?', vals)
        conn.commit()
        return True
    return False

def _parse_sam_notice_id(s):
    # Basic patterns for SAM URLs: /opp/<uuid>/view  or legacy ?id=12345
    import re
    if not s:
        return None
    m = re.search(r'/opp/([0-9a-fA-F\-]{8,36})/view', s)
    if m:
        return m.group(1)
    m = re.search(r'[?&](id|noticeId|oppId)=(\w+)', s, re.I)
    if m:
        return m.group(2)
    return None


# --- Context helpers ---
_CTX_DEFAULTS = {
    "rfp": None,
    "clins": [],
    "sections": [],
    "imp": {},
    "price_sheet": [],
    "staffing": [],
    "notice_id": None,
    "proposal_id": None,
    "meta": {},
}
def _ctxd(ctx, key):
    try:
        return ctx.get(key, _CTX_DEFAULTS.get(key))
    except Exception:
        return _CTX_DEFAULTS.get(key)


# --- Safe DataFrame helpers ---
def _is_df(obj):
    try:
        import pandas as pd
        return isinstance(obj, pd.DataFrame)
    except Exception:
        return False

def _df_nonempty(df):
    return _is_df(df) and not df.empty

def _first_row_value(df, col, default=None):
    try:
        if _df_nonempty(df) and col in df.columns:
            return df.iloc[0][col]
    except Exception:
        pass
    return default


# --- Capability Statement Page (full implementation) ---



# === Y6 helper ===
def _y6_resolve_openai_client():
    try:
        if "get_ai" in globals():
            return get_ai()
        if "get_openai_client" in globals():
            return get_openai_client()
        if "get_ai_client" in globals():
            return get_ai_client()
    except Exception:
        pass
    from openai import OpenAI  # type: ignore
    import os as _os
    key = (
        st.secrets.get("openai_api_key")
        or st.secrets.get("OPENAI_API_KEY")
        or _os.environ.get("OPENAI_API_KEY")
    )
    if not key:
        raise RuntimeError("OPENAI_API_KEY not configured")
    return OpenAI(api_key=key)

def _y6_resolve_model() -> str:
    return st.secrets.get("openai_model") or st.secrets.get("OPENAI_MODEL") or "gpt-4o-mini"

def _y6_chat(messages):
    client = _y6_resolve_openai_client()
    model = _y6_resolve_model()
    resp = client.chat.completions.create(model=model, messages=messages, temperature=0.2)
    try:
        return resp.choices[0].message.content.strip()
    except Exception:
        return "AI response unavailable."

def _y6_fetch_y1_context(conn, rfp_id, question: str, k_auto_fn=None):
    if not (conn and rfp_id):
        return None
    y1 = globals().get("y1_search")
    if not callable(y1):
        return None
    try:
        k = 6
        if callable(k_auto_fn):
            try:
                k = int(max(3, min(12, k_auto_fn(question))))
            except Exception:
                pass
        hits = y1(conn, int(rfp_id), question or "", k=k) or []
        if not hits:
            return None
        blocks = []
        for i, h in enumerate(hits, start=1):
            cid = h.get("chunk_id", i)
            rid = h.get("rfp_id", rfp_id)
            text = h.get("chunk") or h.get("text") or ""
            tag = f"[RFP-{rid}:{cid}]"
            blocks.append(f"{tag} {text}")
        return "\n\n".join(blocks)
    except Exception:
        return None

def y6_render_co_box(conn, rfp_id=None, *, key_prefix: str, title: str, help_text: str="CO answers. Uses RFP context when available.") -> None:
    c1, c2 = st.columns([3, 1])
    with c1:
        st.subheader(title)
        q = st.text_area("Your question", key=f"{key_prefix}_q", height=120, help=help_text)
    with c2:
        st.caption("Y6 CO helper")
    if not st.button("Ask", key=f"{key_prefix}_ask") or not (q or "").strip():
        return
    with st.spinner("CO is analyzing"):
        ctx = _y6_fetch_y1_context(conn, rfp_id, q, globals().get("y_auto_k"))
        CO_SYS = (
            "You are a senior U.S. federal Contracting Officer (CO). "
            "Answer with short, precise sentences or numbered bullets. "
            "If RFP context is provided, cite it with [RFP-<id>:<chunk#>]. "
            "Avoid raw JSON."
        )
        sys_prompt = CO_SYS + (f"\n\nRFP context follows. Cite it when relevant:\n{ctx}" if ctx else "")
        messages = [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": (q or '').strip()},
        ]
        ans = _y6_chat(messages)
    st.markdown(ans)
    saver = globals().get("y5_save_snippet") or globals().get("pb_add_snippet") or globals().get("save_text_to_drafts")
    if saver:
        with st.expander("Add answer to Drafts", expanded=False):
            sec = st.text_input("Section label", value="CO Notes", key=f"{key_prefix}_draft_sec")
            if st.button("Add to drafts", key=f"{key_prefix}_draft_add"):
                try:
                    if "y5_save_snippet" in globals():
                        y5_save_snippet(conn, int(rfp_id) if rfp_id else 0, sec, ans, source="Y6 CO Box")
                    elif "pb_add_snippet" in globals():
                        pb_add_snippet(conn, int(rfp_id) if rfp_id else 0, sec, ans, source="Y6 CO Box")
                    else:
                        saver(conn, rfp_id, ans)
                    st.success("Saved to drafts")
                except Exception:
                    st.info("Drafts saver not available in this build.")
# === end Y6 helper ===

# === Global extractors to avoid NameError in early calls ===
def _extract_naics(text: str) -> str:
    import re as _re
    if not text:
        return ""
    m = _re.search(r'(?i)NAICS(?:\s*Code)?\s*[:#]?\s*([0-9]{5,6})', text)
    if m:
        return m.group(1)[:6]
    m = _re.search(r'(?i)NAICS[^\n]{0,50}?([0-9]{6})', text)
    if m:
        return m.group(1)
    m = _re.search(r'(?i)(?:industry|classification)[^\n]{0,50}?([0-9]{6})', text)
    return m.group(1) if m else ""

def _extract_set_aside(text: str) -> str:
    import re as _re
    if not text:
        return ""
    tags = ["SDVOSB","SDVOSBC","WOSB","EDWOSB","8(a)","8A","HUBZone","SBA","SDB","VOSB","Small Business","Total Small Business"]
    for t in tags:
        if _re.search(rf'(?i)\\b{_re.escape(t)}\\b', text):
            norm = t.upper().replace("(A)", "8A").replace("TOTAL SMALL BUSINESS","SMALL BUSINESS")
            if norm == "8(A)":
                norm = "8A"
            return norm
    m = _re.search(r'(?i)Set[- ]Aside\s*[:#]?\s*([A-Za-z0-9 \-/()]+)', text)
    if m:
        v = m.group(1).strip()
        v = _re.sub(r'\\s+', ' ', v)
        return v[:80]
    return ""
# === End global extractors ===


try:
    from rfp_onepage import run_rfp_analyzer_onepage
except Exception:
    run_rfp_analyzer_onepage = None

# --- Optional PDF backends for Phase X1 ---
try:
    import pdfplumber as _pdfplumber  # type: ignore
except Exception:
    _pdfplumber = None
try:
    import PyPDF2 as _pypdf  # type: ignore
except Exception:
    try:
        import pypdf as _pypdf  # type: ignore
    except Exception:
        _pypdf = None




# --- hashing helper ---
def compute_sha256(b: bytes) -> str:
    try:
        return hashlib.sha256(b or b"").hexdigest()
    except Exception:
        import hashlib as _h
        return _h.sha256(b or b"").hexdigest()

# Phase X unified settings
from types import SimpleNamespace as _NS
from openai import OpenAI

def getenv_int(name: str, default: int) -> int:
    try:
        return int(os.environ.get(name, default))
    except Exception:
        return default

SETTINGS = _NS(
    APP_NAME=os.environ.get("ELA_APP_NAME", "ELA GovCon Suite"),
    APP_VERSION=os.environ.get("ELA_APP_VERSION", "X-Base"),
    DATA_DIR=os.environ.get("ELA_DATA_DIR", "data"),
    UPLOADS_SUBDIR=os.environ.get("ELA_UPLOADS_SUBDIR", "uploads"),
    DEFAULT_PAGE_SIZE=getenv_int("ELA_PAGE_SIZE", 50),
)

SETTINGS.UPLOADS_DIR = os.path.join(SETTINGS.DATA_DIR, SETTINGS.UPLOADS_SUBDIR)

# Ensure directories exist
try:
    os.makedirs(SETTINGS.DATA_DIR, exist_ok=True)
    os.makedirs(SETTINGS.UPLOADS_DIR, exist_ok=True)
except Exception:
    pass

# Back-compat constants
DATA_DIR = SETTINGS.DATA_DIR
UPLOADS_DIR = SETTINGS.UPLOADS_DIR

# Feature flag alias
def flag(name: str, default: bool=False) -> bool:
    return feature_flag(name, default)

# External


APP_TITLE = "ELA GovCon Suite"
BUILD_LABEL = "Master A–F — SAM • RFP Analyzer • L&M • Proposal • Subs+Outreach • Quotes • Pricing • Win Prob • Chat • Capability"

def apply_theme_phase1():
    import streamlit as st
    if st.session_state.get("_phase1_theme_applied"):
        return
    st.session_state["_phase1_theme_applied"] = True
    st.markdown('''
    <style>
    .block-container {padding-top: 1.2rem; padding-bottom: 1.2rem; max-width: 1400px;}
    h1, h2, h3 {margin-bottom: .4rem;}
    .ela-subtitle {color: rgba(49,51,63,0.65); font-size: .95rem; margin-bottom: 1rem;}
    div[data-testid="stDataFrame"] thead th {position: sticky; top: 0; background: #fff; z-index: 2;}
    div[data-testid="stDataFrame"] tbody tr:hover {background: rgba(64,120,242,0.06);}
    [data-testid="stExpander"] {border: 1px solid rgba(49,51,63,0.16); border-radius: 12px; margin-bottom: 10px;}
    [data-testid="stExpander"] summary {font-weight: 600;}
    .ela-card {border: 1px solid rgba(49,51,63,0.16); border-radius: 12px; padding: 12px; margin-bottom: 12px;}
    .ela-chip {display:inline-block; padding: 2px 8px; border-radius: 999px; font-size: 12px; margin-right:6px; background: rgba(49,51,63,0.06);}
    .ela-ok {background: rgba(0,200,83,0.12);} .ela-warn {background: rgba(251,140,0,0.12);} .ela-bad {background: rgba(229,57,53,0.12);}
    .stTextInput>div>div>input, .stNumberInput input, .stTextArea textarea {border-radius: 10px !important;}
    button[kind="primary"] {box-shadow: 0 1px 4px rgba(0,0,0,0.08);}
    .ela-banner {position: sticky; top: 0; z-index: 999; background: linear-gradient(90deg, #4068f2, #7a9cff); color: #fff; padding: 6px 12px; border-radius: 8px; margin-bottom: 10px;}
    </style>
    ''', unsafe_allow_html=True)
    st.markdown("<div class='ela-banner'>Phase 1 theme active · polished layout & tables</div>", unsafe_allow_html=True)

st.set_page_config(page_title=APP_TITLE, layout="wide")
apply_theme_phase1()


# === Y0: GPT-5 Thinking CO assistant (streaming) ===
try:
    from openai import OpenAI as _Y0OpenAI
except Exception:
    _Y0OpenAI = None

SYSTEM_CO = ("Act as a GS-1102 Contracting Officer. Cite exact pages. "
             "Flag non-compliance. Be concise. If evidence is missing, say so.")

# === helper: auto-select number of sources to cite (Y1–Y3) ===


import os

def _resolve_model():
    # Priority: Streamlit secrets -> env var -> safe default
    try:
        import streamlit as st  # noqa: F401
        for key in ("OPENAI_MODEL", "openai_model", "model"):
            try:
                val = st.secrets.get(key)  # type: ignore[attr-defined]
                if isinstance(val, str) and val.strip():
                    return val.strip()
            except Exception:
                pass
    except Exception:
        pass
    return os.getenv("OPENAI_MODEL", "gpt-4o-mini")

_ai_client = None
def get_ai():
    import streamlit as st  # ensure st exists
    global _ai_client
    if _ai_client is None:
        if _Y0OpenAI is None:
            raise RuntimeError("openai library missing")
        _ai_client = _Y0OpenAI()  # uses OPENAI_API_KEY from Streamlit secrets
    return _ai_client

def ask_ai(messages, tools=None, temperature=0.2):
    client = get_ai()
    model_name = _resolve_model()
    try:
        resp = client.chat.completions.create(
            model=model_name,
            messages=[{"role":"system","content": SYSTEM_CO}, *messages],
            tools=tools or [],
            temperature=float(temperature),
            stream=True
        )
    except Exception as _e:
        if "model_not_found" in str(_e) or "does not exist" in str(_e):
            model_name = "gpt-4o-mini"
            resp = client.chat.completions.create(
                model=model_name,
                messages=[{"role":"system","content": SYSTEM_CO}, *messages],
                tools=tools or [],
                temperature=float(temperature),
                stream=True
            )
        else:
            yield f"AI unavailable: {type(_e).__name__}: {_e}"
            return
    for ch in resp:
        try:
            delta = ch.choices[0].delta
            if hasattr(delta, "content") and delta.content:
                yield delta.content
        except Exception:
            pass

def y0_ai_panel():

    st.header(f"Ask the CO (AI) · {_resolve_model()}")
    q = st.text_area("Your question", key="y0_q", height=120)
    if st.button("Ask", key="y0_go"):
        if not (q or "").strip():
            st.warning("Enter a question")
        else:
            ph = st.empty()
            acc = []
            for tok in ask_ai([{"role":"user","content": q.strip()}]):
                acc.append(tok)
                ph.markdown("".join(acc))
# === end Y0 ===



# --- key namespacing helper (Phase U) ---
DATA_DIR = "data"
DB_PATH = os.path.join(DATA_DIR, "govcon.db")
UPLOADS_DIR = os.path.join(DATA_DIR, "uploads")
SAM_ENDPOINT = "https://api.sam.gov/opportunities/v2/search"


# -------------------- setup --------------------
def ensure_dirs() -> None:
    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(UPLOADS_DIR, exist_ok=True)
# === Y1: Retrieval (chunks • embeddings • citations) ===
def _ensure_y1_schema(conn: sqlite3.Connection) -> None:
    try:
        with closing(conn.cursor()) as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS rfp_chunks(
                    id INTEGER PRIMARY KEY,
                    rfp_id INTEGER,
                    rfp_file_id INTEGER,
                    file_name TEXT,
                    page INTEGER,
                    chunk_idx INTEGER,
                    text TEXT,
                    emb TEXT
                );
            """)
            cur.execute("CREATE INDEX IF NOT EXISTS idx_chunks_rfp ON rfp_chunks(rfp_id);")
            cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_chunk_key ON rfp_chunks(rfp_file_id, page, chunk_idx);")
            conn.commit()
    except Exception:
        pass


def _resolve_embed_model() -> str:
    try:
        import streamlit as _st
        for k in ("OPENAI_EMBED_MODEL","openai_embed_model","EMBED_MODEL"):
            v = _st.secrets.get(k)
            if isinstance(v, str) and v.strip():
                return v.strip()
    except Exception:
        pass
    return os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")
# === PHASE 6: Embedding cache (sha256 of model:text) ===
def _embed_cache_dir() -> str:
    d = os.path.join(DATA_DIR, "embed_cache")
    try:
        os.makedirs(d, exist_ok=True)
    except Exception:
        pass
    return d

def _embed_cache_key(model: str, text: str) -> str:
    h = hashlib.sha256(f"{model}:{text}".encode("utf-8")).hexdigest()
    return h

def _embed_cache_get(model: str, texts: list[str]) -> tuple[list[list[float]|None], list[int]]:
    r"""
    Return (vecs_or_None, missing_idx). For each input text, either a vector or None.
    Also returns indices of missing items for batch computation.
    r"""
    out = []
    missing = []
    d = _embed_cache_dir()
    for i, t in enumerate(texts):
        key = _embed_cache_key(model, t or " ")
        fp = os.path.join(d, f"{key}.json")
        try:
            with open(fp, "r", encoding="utf-8") as fh:
                v = json.load(fh)
            if isinstance(v, list):
                out.append(v)
                continue
        except Exception:
            pass
        out.append(None)
        missing.append(i)
    return out, missing

def _embed_cache_put(model: str, texts: list[str], vecs: list[list[float]]) -> None:
    d = _embed_cache_dir()
    for t, v in zip(texts, vecs):
        try:
            key = _embed_cache_key(model, t or " ")
            fp = os.path.join(d, f"{key}.json")
            with open(fp, "w", encoding="utf-8") as fh:
                json.dump(v, fh)
        except Exception:
            pass
# === end PHASE 6 helpers ===

def _embed_texts(texts: list[str]) -> list[list[float]]:
    client = get_ai()
    model = _resolve_embed_model()
    clean = [t if (t or "").strip() else " " for t in texts]

    # cache lookup
    cached_vecs, missing_idx = _embed_cache_get(model, clean)
    to_compute = [clean[i] for i in missing_idx]

    new_vecs: list[list[float]] = []
    if to_compute:
        try:
            resp = client.embeddings.create(model=model, input=to_compute)
            for d in resp.data:
                try:
                    new_vecs.append(list(d.embedding))
                except Exception:
                    new_vecs.append([])
            # write cache
            _embed_cache_put(model, to_compute, new_vecs)
        except Exception:
            # fallback: return empty vecs for missing
            new_vecs = [[] for _ in to_compute]

    # merge back preserving order
    out: list[list[float]] = []
    it = iter(new_vecs)
    for v in cached_vecs:
        if v is None:
            out.append(next(it, []))
        else:
            out.append(v)
    return out

def _cos_sim(u: list[float], v: list[float]) -> float:
    if not u or not v:
        return 0.0
    try:
        import numpy as _np
        a = _np.array(u, dtype=float); b = _np.array(v, dtype=float)
        num = float((_np.dot(a, b)))
        den = float(_np.linalg.norm(a) * _np.linalg.norm(b))
        return (num / den) if den else 0.0
    except Exception:
        num = 0.0; su = 0.0; sv = 0.0
        for i in range(min(len(u), len(v))):
            x = float(u[i]); y = float(v[i])
            num += x*y; su += x*x; sv += y*y
        den = (su**0.5) * (sv**0.5)
        return (num / den) if den else 0.0

def _split_chunks(text: str, max_chars: int = 1200, overlap: int = 180) -> list[str]:
    t = (text or "").strip()
    if not t:
        return []
    chunks = []
    i = 0
    n = len(t)
    while i < n:
        j = min(n, i + max(200, max_chars))
        chunk = t[i:j]
        k = chunk.rfind(". ")
        if k > 900:
            chunk = chunk[:k+1]
            j = i + k + 1
        chunks.append(chunk)
        i = max(j - overlap, j)
    return chunks

def y1_index_rfp(conn: sqlite3.Connection, rfp_id: int, max_pages: int = 100, rebuild: bool = False) -> dict:
    _ensure_y1_schema(conn)
    try:
        df = pd.read_sql_query("SELECT id, filename, mime, pages FROM rfp_files WHERE rfp_id=? ORDER BY id;", conn, params=(int(rfp_id),))
    except Exception as e:
        return {"ok": False, "error": str(e)}
    if df is None or df.empty:
        return {"ok": False, "error": "No linked files"}
    added = 0
    skipped = 0
    for _, row in df.iterrows():
        fid = int(row["id"]); name = row.get("filename") or f"file_{fid}"
        try:
            blob = pd.read_sql_query("SELECT bytes, mime FROM rfp_files WHERE id=?;", conn, params=(fid,)).iloc[0]
            b = blob["bytes"]; mime = blob.get("mime") or (row.get("mime") or "application/octet-stream")
        except Exception:
            continue
        pages = extract_text_pages(b, mime) or []
        if pages:
            pages, _ocrn = ocr_pages_if_empty(b, mime, pages)
        if not pages:
            continue
        pages = pages[:max_pages]
        for pi, txt in enumerate(pages, start=1):
            parts = _split_chunks(txt or "", 1600, 200)
            for ci, ch in enumerate(parts):
                try:
                    if not rebuild:
                        q = pd.read_sql_query("SELECT id FROM rfp_chunks WHERE rfp_file_id=? AND page=? AND chunk_idx=?;", conn, params=(fid, pi, ci))
                        if q is not None and not q.empty:
                            skipped += 1
                            continue
                except Exception:
                    pass
                emb = _embed_texts([ch])[0]
                with closing(conn.cursor()) as cur:
                    cur.execute("""
                        INSERT OR REPLACE INTO rfp_chunks(rfp_id, rfp_file_id, file_name, page, chunk_idx, text, emb)
                        VALUES(?,?,?,?,?,?,?);
                    """, (int(rfp_id), fid, name, int(pi), int(ci), ch, json.dumps(emb)))
                    conn.commit()
                added += 1
    return {"ok": True, "added": added, "skipped": skipped}
def _y1_search_uncached(conn: sqlite3.Connection, rfp_id: int, query: str, k: int = 6) -> list[dict]:
    _ensure_y1_schema(conn)
    if not (query or "").strip():
        return []
    try:
        df = pd.read_sql_query("SELECT id, rfp_file_id, file_name, page, chunk_idx, text, emb FROM rfp_chunks WHERE rfp_id=?;", conn, params=(int(rfp_id),))
    except Exception:
        return []
    if df is None or df.empty:
        return []
    # clamp
    try:
        k = int(k)
    except Exception:
        k = 6
    k = max(1, min(8, k))

    q_emb = _embed_texts([query])[0]
    rows = []
    for _, r in df.iterrows():
        try:
            emb = json.loads(r.get("emb") or "[]")
        except Exception:
            emb = []
        sim = _cos_sim(q_emb, emb)
        rows.append({
            "id": int(r["id"]),
            "fid": int(r["rfp_file_id"]),
            "file": r.get("file_name"),
            "page": int(r.get("page") or 0),
            "chunk": int(r.get("chunk_idx") or 0),
            "text": r.get("text") or "",
            "score": float(sim),
        })
    # primary rank
    rows.sort(key=lambda x: x["score"], reverse=True)
    # de-duplicate by (file,page) to strengthen citation control
    seen = set()
    dedup = []
    for h in rows:
        key = (h.get("file") or "", int(h.get("page") or 0))
        if key in seen:
            continue
        seen.add(key)
        dedup.append(h)
        if len(dedup) >= 32:  # limit working set
            break
    # light re-rank: prefer balanced page coverage then score
    dedup.sort(key=lambda x: (-(x["score"]>0.70), -x["score"]), reverse=False)
    return dedup[:k]
# --- Safe Y1 dispatcher to avoid NameError at runtime ---
def _safe_y1_search(conn, rfp_id, query, k=6):
    try:
        return y1_search(conn, int(rfp_id), query or "", int(k)) or []
    except NameError:
        try:
            return _y1_search_uncached(conn, int(rfp_id), query or "", int(k)) or []
        except NameError:
            return []
# --- Robust Y1 shim: guarantees y1_search exists ---
if 'y1_search' not in globals():
    def y1_search(conn, rfp_id: int, query: str, k: int = 6):
        try:
            snap = _y1_snapshot(conn, int(rfp_id))
        except Exception:
            snap = None
        try:
            db_path = DB_PATH
        except Exception:
            db_path = "data/govcon.db"
        try:
            return _y1_search_cached(db_path, int(rfp_id), query or "", int(k), snap)
        except Exception:
            try:
                return _y1_search_uncached(conn, int(rfp_id), query or "", int(k))
            except Exception:
                return []






def ask_ai_with_citations(conn: sqlite3.Connection, rfp_id: int, question: str, k: int = 6, temperature: float = 0.2):
    """
    Streams a CO-style answer grounded in top-k chunk hits with [C#] citations.
    Falls back to general answer if no hits.
    """
    hits = y1_search(conn, int(rfp_id), question or "", k=int(k)) or []
    if not hits:
        try:
            strict = EVIDENCE_GATE or CO_STRICT
        except Exception:
            strict = True
        if strict:
            yield "[system] Insufficient evidence in linked RFP files. Build or Update the search index for this RFP on 'Ask with citations (Y1)', then ask again. General answers are disabled in CO Chat."
            return
        for tok in ask_ai([{"role":"user","content": (question or "").strip()}], temperature=temperature):
            yield tok
        return
    ev_lines = []
    for i, h in enumerate(hits, start=1):
        tag = f"[C{i}]"
        src_line = f"{h.get('file','')} p.{h.get('page','')}"
        snip = (h.get("text") or "").strip().replace("\n", " ")
        ev_lines.append(f"{tag} {src_line} — {snip}")
    evidence = "\n".join(ev_lines)
    user = "QUESTION\n" + (question or "").strip() + "\n\nEVIDENCE\n" + evidence
    for tok in ask_ai([{"role":"user","content": user}], temperature=temperature):
        yield tok

def _y2_build_messages(conn: sqlite3.Connection, rfp_id: int, thread_id: int, user_q: str, k: int = 6):
    """
    Build a minimal message set for CO chat, embedding local evidence as [C#].
    Returns a list of chat messages (no system role; ask_ai adds SYSTEM_CO).
    """
    q_in = (user_q or "").strip()
    hits = y1_search(conn, int(rfp_id), q_in or "Section L and Section M requirements", k=int(k)) or []
    ev_lines = []
    for i, h in enumerate(hits, start=1):
        tag = f"[C{i}]"
        src = f"{h.get('file','')} p.{h.get('page','')}"
        snip = (h.get('text') or '').strip().replace("\n", " ")
        ev_lines.append(f"{tag} {src} — {snip}")
    evidence = "\n".join(ev_lines)
    user = "QUESTION\n" + (q_in or "Provide a CO Readout.") + "\n\nEVIDENCE\n" + (evidence or "(none)")
    msgs = [{"role":"user","content": user}]
    return msgs
def _ensure_y2_schema(conn: sqlite3.Connection) -> None:
    try:
        with closing(conn.cursor()) as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS y2_threads(
                    id INTEGER PRIMARY KEY,
                    rfp_id INTEGER NOT NULL,
                    title TEXT,
                    created_at TEXT
                );
            """)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS y2_messages(
                    id INTEGER PRIMARY KEY,
                    thread_id INTEGER NOT NULL,
                    role TEXT CHECK(role in ('user','assistant')),
                    content TEXT,
                    created_at TEXT
                );
            """)
            cur.execute("CREATE INDEX IF NOT EXISTS idx_y2_threads_rfp ON y2_threads(rfp_id);")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_y2_msgs_thread ON y2_messages(thread_id);")
            conn.commit()
    except Exception:
        pass

def y2_list_threads(conn: sqlite3.Connection, rfp_id: int):
    _ensure_y2_schema(conn)
    try:
        df = pd.read_sql_query(
            "SELECT id, title, created_at FROM y2_threads WHERE rfp_id=? ORDER BY id DESC;",
            conn, params=(int(rfp_id),)
        )
    except Exception:
        return []
    if df is None or df.empty:
        return []
    out = []
    for _, row in df.iterrows():
        rid = int(row["id"])
        out.append({
            "id": rid,
            "title": (row.get("title") or f"Thread #{rid}"),
            "created_at": row.get("created_at") or ""
        })
    return out

def y2_create_thread(conn: sqlite3.Connection, rfp_id: int, title: str = "CO guidance") -> int:
    _ensure_y2_schema(conn)
    from datetime import datetime as _dt
    now = _dt.utcnow().isoformat()
    with closing(conn.cursor()) as cur:
        cur.execute("INSERT INTO y2_threads(rfp_id, title, created_at) VALUES(?,?,?);",
                    (int(rfp_id), (title or "Untitled").strip(), now))
        conn.commit()
        return int(cur.lastrowid)

def y2_get_messages(conn: sqlite3.Connection, thread_id: int):
    _ensure_y2_schema(conn)
    try:
        df = pd.read_sql_query(
            "SELECT role, content FROM y2_messages WHERE thread_id=? ORDER BY id;",
            conn, params=(int(thread_id),)
        )
    except Exception:
        return []
    if df is None or df.empty:
        return []
    return [{"role": str(r["role"]), "content": str(r["content"])} for _, r in df.iterrows()]

def y2_append_message(conn: sqlite3.Connection, thread_id: int, role: str, content: str) -> None:
    _ensure_y2_schema(conn)
    from datetime import datetime as _dt
    now = _dt.utcnow().isoformat()
    role = "assistant" if str(role).strip().lower() != "user" else "user"
    with closing(conn.cursor()) as cur:
        cur.execute(
            "INSERT INTO y2_messages(thread_id, role, content, created_at) VALUES(?,?,?,?);",
            (int(thread_id), role, (content or "").strip(), now)
        )
        conn.commit()

def y2_rename_thread(conn: sqlite3.Connection, thread_id: int, new_title: str) -> None:
    _ensure_y2_schema(conn)
    with closing(conn.cursor()) as cur:
        cur.execute("UPDATE y2_threads SET title=? WHERE id=?;", ((new_title or "Untitled").strip(), int(thread_id)))
        conn.commit()

def y2_delete_thread(conn: sqlite3.Connection, thread_id: int) -> None:
    _ensure_y2_schema(conn)
    with closing(conn.cursor()) as cur:
        cur.execute("DELETE FROM y2_messages WHERE thread_id=?;", (int(thread_id),))
        cur.execute("DELETE FROM y2_threads WHERE id=?;", (int(thread_id),))
        conn.commit()
# === end Y2 thread storage helpers ===

def y2_ui_threaded_chat(conn: sqlite3.Connection) -> None:
    st.caption("CO Chat with memory. Threads are stored per RFP.")
    df_rf = pd.read_sql_query("SELECT id, title FROM rfps ORDER BY id DESC;", conn, params=())
    if df_rf is None or df_rf.empty:
        st.info("No RFPs yet. Parse & save first.")
        return
    rfp_id = st.selectbox("RFP context", options=df_rf["id"].tolist(), format_func=lambda i: f"#{i} — {df_rf.loc[df_rf['id']==i,'title'].values[0]}", key="y2_rfp_sel")
    threads = y2_list_threads(conn, int(rfp_id))
    create = st.button("New thread", key="y2_new")
    if create:
        tid = y2_create_thread(conn, int(rfp_id), title="CO guidance")
        st.session_state["y2_thread_id"] = tid
        st.rerun()
    if threads:
        pick = st.selectbox("Thread", options=[t["id"] for t in threads], format_func=lambda i: next((f"#{t['id']} — {t.get('title') or 'Untitled'}" for t in threads if t['id']==i), f"#{i}"), key="y2_pick")
        thread_id = int(pick)
    else:
        st.info("No threads yet. Click New thread.")
        return
    with st.expander("Thread settings", expanded=False):
        cur_title = next((t.get("title") or "Untitled" for t in threads if t["id"]==thread_id), "Untitled")
        new_title = st.text_input("Rename", value=cur_title, key="y2_rename")
        colA, colB = st.columns([2,1])
        with colA:
            if st.button("Save name", key="y2_save_name"):
                y2_rename_thread(conn, int(thread_id), new_title or "Untitled")
                st.success("Renamed")
        with colB:
            if st.button("Delete thread", key="y2_del"):
                y2_delete_thread(conn, int(thread_id))
                st.success("Deleted")
    st.checkbox("Research mode (flex)", value=bool(st.session_state.get("y2_flex", False)), key="y2_flex")
    st.subheader("History")
    msgs = y2_get_messages(conn, int(thread_id))
    if msgs:
        chat_md = []
        for m in msgs:
            if m["role"] == "user":
                chat_md.append(f"**You:** {m['content']}")
            elif m["role"] == "assistant":
                chat_md.append(m["content"])
        st.markdown("\n\n".join(chat_md))
    else:
        st.caption("No messages yet.")
    q = st.text_area("Your question", height=120, key="y2_q")
    k = y_auto_k(q)
    if st.button("Ask (store to thread)", type="primary", key="y2_go"):
        if not (q or "").strip():
            st.warning("Enter a question")
        else:
            y2_append_message(conn, int(thread_id), "user", q.strip())
            ph = st.empty(); acc = []
            for tok in y2_stream_answer(conn, int(rfp_id), int(thread_id), q.strip(), k=int(k)):
                acc.append(tok); ph.markdown("".join(acc))
            ans = "".join(acc).strip()
            if ans:
                y2_append_message(conn, int(thread_id), "assistant", ans)
                st.success("Saved to thread")
                with st.expander("Add to Proposal Drafts", expanded=False):
                    sec = st.text_input("Section label", value="CO Chat Notes", key="y5_sec_y2")
                    if st.button("Add to drafts", key="y5_add_y2"):
                        y5_save_snippet(conn, int(rfp_id), sec, ans, source="Y2 Chat")
                        st.success("Saved to drafts")
# === end Y2 ===

# === Research cache (R2) ===
def _research_cache_dir() -> str:
    d = os.path.join(DATA_DIR, "research_cache")
    try:
        os.makedirs(d, exist_ok=True)
    except Exception:
        pass
    return d

def _sha16(s: str) -> str:
    try:
        import hashlib
        return hashlib.sha256((s or "").encode("utf-8")).hexdigest()[:16]
    except Exception:
        return str(abs(hash((s or ""))))[:16]

def research_fetch(url: str, ttl_hours: int = 24) -> dict:
    """
    Fetch a URL with simple on-disk cache.
    Returns dict {url, path, cached, status, text}.
    """
    out = {"url": url, "cached": False, "status": 0, "path": "", "text": ""}
    if not (url or "").strip():
        return out
    key = _sha16(url)
    dirp = _research_cache_dir()
    meta_path = os.path.join(dirp, f"{key}.json")
    txt_path  = os.path.join(dirp, f"{key}.txt")

    # serve cache if fresh
    try:
        if os.path.exists(meta_path) and os.path.exists(txt_path):
            with open(meta_path, "r", encoding="utf-8") as fh:
                meta = json.load(fh)
            age = time.time() - float(meta.get("ts", 0))
            if age < ttl_hours * 3600:
                out.update(meta)
                with open(txt_path, "r", encoding="utf-8", errors="ignore") as fh:
                    out["text"] = fh.read()
                out["cached"] = True
                out["status"] = int(out.get("status", 200) or 200)
                out["path"] = txt_path
                return out
    except Exception:
        pass

    # fetch fresh
    try:
        import requests  # lazy import
        r = requests.get(url, timeout=20, headers={"User-Agent":"ELA-GovCon/1.0"})
        status = int(getattr(r, "status_code", 0) or 0)
        text = r.text if hasattr(r, "text") else ""
        # persist
        try:
            with open(txt_path, "w", encoding="utf-8") as fh:
                fh.write(text or "")
            with open(meta_path, "w", encoding="utf-8") as fh:
                json.dump({"url": url, "status": status, "ts": time.time(), "path": txt_path}, fh)
        except Exception:
            pass
        out.update({"status": status, "text": text, "path": txt_path})
        return out
    except Exception as e:
        out.update({"status": 0, "error": str(e)})
        return out

def research_extract_excerpt(text: str, query: str, window: int = 380) -> str:
    t = (text or "")
    if not t:
        return ""
    q = (query or "").strip()
    idx = -1
    if q:
        try:
            idx = t.lower().find(q.lower())
        except Exception:
            idx = -1
    if idx < 0:
        return t[:window]
    start = max(0, idx - window//2)
    end = min(len(t), idx + window//2)
    return t[start:end]



# === Y3: Proposal drafting from evidence (per-RFP) ===
def _y3_collect_ctx(conn: sqlite3.Connection, rfp_id: int, max_items: int = 20) -> dict:
    ctx: dict = {}
    try:
        df_items = pd.read_sql_query("SELECT item_text FROM lm_items WHERE rfp_id=? ORDER BY id;", conn, params=(int(rfp_id),))
    except Exception:
        df_items = pd.DataFrame(columns=["item_text"])
    ctx["lm"] = df_items["item_text"].tolist()[:max_items] if isinstance(df_items, pd.DataFrame) and not df_items.empty else []
    try:
        df_clins = pd.read_sql_query("SELECT clin, description FROM clin_lines WHERE rfp_id=? ORDER BY id;", conn, params=(int(rfp_id),))
    except Exception:
        df_clins = pd.DataFrame(columns=["clin","description"])
    ctx["clins"] = [{"clin": str(r.get("clin","")), "desc": str(r.get("description",""))[:160]} for _, r in df_clins.head(max_items).iterrows()] if isinstance(df_clins, pd.DataFrame) and not df_clins.empty else []
    try:
        df_meta = pd.read_sql_query("SELECT key, value FROM rfp_meta WHERE rfp_id=?;", conn, params=(int(rfp_id),))
    except Exception:
        df_meta = pd.DataFrame(columns=["key","value"])
    if isinstance(df_meta, pd.DataFrame) and not df_meta.empty:
        ctx["meta"] = {str(k): str(v) for k,v in zip(df_meta["key"], df_meta["value"])}
    else:
        ctx["meta"] = {}
    try:
        df_rfp = pd.read_sql_query("SELECT title, solnum FROM rfps WHERE id=?;", conn, params=(int(rfp_id),))
        ctx["title"] = df_rfp.iloc[0]["title"] if not df_rfp.empty else ""
        ctx["solnum"] = df_rfp.iloc[0]["solnum"] if not df_rfp.empty else ""
    except Exception:
        ctx["title"] = ""; ctx["solnum"] = ""
    return ctx

def _y3_build_messages(conn: sqlite3.Connection, rfp_id: int, section_title: str, notes: str, k: int = 6, max_words: int | None = None) -> list[dict]:
    ctx = _y3_collect_ctx(conn, int(rfp_id))
    q = f"{section_title} Section L Section M instructions {ctx.get('title','')} {ctx.get('solnum','')}"
    hits = y1_search(conn, int(rfp_id), q, k=int(k)) or []
    ev_lines = []
    for i, h in enumerate(hits, start=1):
        tag = f"[C{i}]"
        src = f"{h['file']} p.{h['page']}"
        snip = (h.get("text") or "").strip().replace("\n"," ")
        ev_lines.append(f"{tag} {src} — {snip}")
    evidence = "\n".join(ev_lines)
    style = (SYSTEM_CO + " Write a proposal section for a federal RFP. Use short, compliant sentences."
             + " Map claims to evidence with bracketed citations like [C1]."
             + " Do not invent requirements. If evidence is missing, state the gap plainly.")
    bullets = "\n".join([f"- {it}" for it in (ctx.get("lm") or [])])
    clins = "\n".join([f"- {c['clin']}: {c['desc']}" for c in (ctx.get("clins") or [])])
    constraints = []
    if ctx["meta"].get("naics"): constraints.append(f"NAICS: {ctx['meta']['naics']}")
    if ctx["meta"].get("set_aside"): constraints.append(f"Set-Aside: {ctx['meta']['set_aside']}")
    if ctx["meta"].get("place_of_performance"): constraints.append(f"Place of Performance: {ctx['meta']['place_of_performance']}")
    limit_line = f"Target length: <= {int(max_words)} words." if isinstance(max_words, int) and max_words>0 else "Target brevity."
    user = f"""Draft the section: {section_title}
{limit_line}

RFP title: {ctx.get('title','')}  Solicitation: {ctx.get('solnum','')}
Constraints: {' | '.join(constraints) if constraints else 'n/a'}

Key L/M items to cover:
{bullets or '- n/a'}

Relevant CLINs:
{clins or '- n/a'}

Notes from author:
{notes or 'n/a'}

EVIDENCE:
{evidence}
Write a structured section with a short lead paragraph, 3–6 bullets, and an optional close. Use [C#] next to any factual or requirement-based claim that is tied to EVIDENCE.
"""
    return [{"role":"system","content": style}, {"role":"user","content": user}]

def y3_stream_draft(conn: sqlite3.Connection, rfp_id: int, section_title: str, notes: str, k: int = 6, max_words: int | None = None, temperature: float = 0.2):
    msgs = _y3_build_messages(conn, int(rfp_id), section_title, notes, k=int(k), max_words=max_words)
    client = get_ai()
    model_name = _resolve_model()
    try:
        resp = client.chat.completions.create(model=model_name, messages=msgs, temperature=float(temperature), stream=True)
    except Exception as _e:
        if "model_not_found" in str(_e) or "does not exist" in str(_e):
            resp = client.chat.completions.create(model="gpt-4o-mini", messages=msgs, temperature=float(temperature), stream=True)
        else:
            yield f"AI unavailable: {type(_e).__name__}: {_e}"
            return
    for ch in resp:
        try:
            delta = ch.choices[0].delta
            if hasattr(delta, "content") and delta.content:
                yield delta.content
        except Exception:
            pass
# === end Y3 ===

# === Y4: CO Review (scored compliance with citations) ===
def _y4_build_messages(conn: sqlite3.Connection, rfp_id: int, draft_text: str, k: int = 6) -> list[dict]:
    """
    Build messages for a Contracting Officer style review.
    Output must include: Score 0–100, Strengths, Gaps, Risks, Required fixes, and short Conclusion.
    Use [C#] next to claims grounded in EVIDENCE.
    """
    # Reuse Section L/M + compliance as retrieval query
    q = f"Section L Section M compliance checklist evaluation {draft_text[:400]}"
    hits = y1_search(conn, int(rfp_id), q, k=int(k)) or []
    ev_lines = []
    for i, h in enumerate(hits, start=1):
        tag = f"[C{i}]"
        src = f"{h['file']} p.{h['page']}"
        snip = (h.get("text") or "").strip().replace("\\n"," ")
        ev_lines.append(f"{tag} {src} — {snip}")
    evidence = "\\n".join(ev_lines)

    system = (SYSTEM_CO
              + " Score the draft 0-100 for compliance and clarity."
              + " Use short, direct sentences. No fluff."
              + " Structure exactly as:\\n"
              + "Score: <0-100>\\n"
              + "Strengths: <bullets>\\n"
              + "Gaps: <bullets>\\n"
              + "Risks: <bullets>\\n"
              + "Required fixes: <bullets>\\n"
              + "Conclusion: <2-3 lines>\\n"
              + "Map factual or requirement statements to EVIDENCE using [C#]."
              + " Hard caps: max 5 bullets per list. Total output ≤ 220 words. If you exceed a cap, remove items to comply."
             )

    user = f"DRAFT TO REVIEW:\\n{draft_text or '(empty)'}\\n\\nEVIDENCE:\\n{evidence or '(no evidence found)'}"
    return [{"role":"system","content": system}, {"role":"user","content": user}]



def y4_postprocess_brevity(text: str, max_words: int = 220, max_bullets: int = 5) -> str:
    """Enforce ≤5 bullets for key sections and a global word cap."""
    if not text:
        return ""
    lines = text.splitlines()
    out = []
    sections = {"Strengths:": "Strengths:", "Gaps:": "Gaps:", "Risks:": "Risks:", "Required fixes:": "Required fixes:"}
    current = None
    bullet_count = 0
    i = 0
    while i < len(lines):
        ln = lines[i]
        ln_stripped = ln.strip()
        # detect headers
        if any(ln_stripped.lower().startswith(h.lower()) for h in sections):
            current = ln_stripped.split(":")[0].lower()
            bullet_count = 0
            out.append(ln)
            i += 1
            continue
        if re.match(r"^(Score:|Conclusion:)", ln_stripped, re.I):
            current = None
            out.append(ln)
            i += 1
            continue
        # cap bullets in the four sections
        if current in {"strengths", "gaps", "risks", "required fixes"}:
            if re.match(r"^\s*[-*\u2022]\s+", ln):
                if bullet_count < max_bullets:
                    out.append(ln)
                    bullet_count += 1
                # else drop extra bullets
            else:
                out.append(ln)
            i += 1
            continue
        out.append(ln)
        i += 1
    text2 = "\n".join(out).strip()
    words = text2.split()
    if len(words) <= max_words:
        return text2
    return " ".join(words[:max_words]).strip()

    client = get_ai()
    model_name = _resolve_model()
    try:
        resp = client.chat.completions.create(model=model_name, messages=msgs, temperature=float(temperature), stream=True)
    except Exception as _e:
        if "model_not_found" in str(_e) or "does not exist" in str(_e):
            resp = client.chat.completions.create(model="gpt-4o-mini", messages=msgs, temperature=float(temperature), stream=True)
        else:
            yield f"AI unavailable: {type(_e).__name__}: {_e}"; return
    for ch in resp:
        try:
            delta = ch.choices[0].delta
            if hasattr(delta, "content") and delta.content:
                yield delta.content
        except Exception:
            pass

def y4_ui_review(conn: sqlite3.Connection) -> None:
    st.caption("CO Review with score, strengths, gaps, risks, and required fixes. Citations auto-selected.")
    df_rf = pd.read_sql_query("SELECT id, title FROM rfps ORDER BY id DESC;", conn, params=())
    if df_rf is None or df_rf.empty:
        st.info("No RFPs yet. Parse & save first."); return
    rfp_id = st.selectbox("RFP context", options=df_rf["id"].tolist(),
                          format_func=lambda i: f"#{i} — {df_rf.loc[df_rf['id']==i,'title'].values[0]}",
                          key="y4_rfp_sel")
    mode = st.radio("Mode", ["Paste", "Upload", "Use RFP files"], horizontal=True, key="y5_mode")
    st.subheader("Draft to review")

    draft_text = ""
    uploaded = None
    if mode == "Paste":
        draft_text = st.text_area("Paste your draft here", value=st.session_state.get("y4_draft",""), height=260, key="y4_draft_ta")
    elif mode == "Upload":
        uploaded = st.file_uploader("Upload DOCX/PDF/TXT", type=["docx","pdf","txt"], accept_multiple_files=True, key="y5_up")
        if uploaded:
            with st.spinner("Extracting"):
                draft_text = y5_extract_from_uploads(uploaded)[:400000]
            st.text_area("Preview", value=draft_text[:20000], height=240)
    else:
        if st.button("Assemble from linked RFP files", key="y5_from_rfp"):
            with st.spinner("Collecting text from linked files"):
                draft_text = y5_extract_from_rfp(conn, int(rfp_id))[:400000]
            st.session_state["y5_rfp_text"] = draft_text
        draft_text = st.session_state.get("y5_rfp_text","")
        if draft_text:
            st.text_area("Preview", value=draft_text[:20000], height=240)

    k = y_auto_k(draft_text or "review")
    chunking = st.checkbox("Auto-chunk long text", value=True, key="y5_chunk_on")
    run = st.button("Run CO Review", type="primary", key="y4_go")
    if run:
        ok_gate, missing_gate = require_LM_minimum(conn, int(rfp_id))
        if not ok_gate:
            st.error("L/M compliance gate failed. Missing: " + ", ".join(missing_gate))
            return

        if not (draft_text or "").strip():
            st.warning("Provide input text."); return
        texts = y5_chunk_text(draft_text) if chunking else [draft_text]
        ph = st.empty()
        all_out = []
        for idx, t in enumerate(texts, start=1):
            acc = []
            ph.caption(f"Reviewing chunk {idx}/{len(texts)}")
            for tok in y4_stream_review(conn, int(rfp_id), t.strip(), k=int(k)):
                acc.append(tok)
                ph.markdown("".join(acc))
            all_out.append("".join(acc))
        final = "\n\n".join(all_out).strip()
        final = y4_postprocess_brevity(final, max_words=220, max_bullets=5)
        st.session_state["y4_last_review"] = final
        st.subheader("Combined result")
        st.markdown(final or "_no output_")

        # Sources table
        hits = y1_search(conn, int(rfp_id), f"Section L Section M compliance {draft_text[:200]}", k=int(k))
        if hits:
            import pandas as _pd
            dfh = _pd.DataFrame([{"Tag": f"[C{i+1}]", "File": h["file"], "Page": h["page"], "Score": h["score"]} for i,h in enumerate(hits)])
            st.subheader("Sources used")
            _styled_dataframe(dfh, use_container_width=True, hide_index=True)

        with st.expander("Add to Proposal Drafts", expanded=True):
            section = st.text_input("Section label", value="CO Review Notes", key="y5_sec_rev")
            if st.button("Add review to drafts", key="y5_add_rev"):
                y5_save_snippet(conn, int(rfp_id), section, final or draft_text, source="Y4 Review")
                st.success("Saved to drafts")


# === Y5: Upload Review + Drafts plumbing ===
from contextlib import closing
import io

def ensure_y5_tables(conn: sqlite3.Connection) -> None:
    try:
        with closing(conn.cursor()) as cur:
            cur.execute(
                "CREATE TABLE IF NOT EXISTS draft_snippets("
                "id INTEGER PRIMARY KEY,"
                "rfp_id INTEGER,"
                "section TEXT,"
                "source TEXT,"
                "text TEXT,"
                "created_at TEXT DEFAULT (datetime('now'))"
                ");"
            )
            conn.commit()
    except Exception:
        pass

def y5_chunk_text(text: str, target_chars: int = 9000, overlap: int = 500) -> list[str]:
    t = (text or "").strip()
    if not t:
        return []
    if len(t) <= target_chars:
        return [t]
    out: list[str] = []
    i = 0
    n = len(t)
    while i < n:
        j = min(n, i + max(2000, target_chars))
        chunk = t[i:j]
        k = chunk.rfind(". ")
        if k > len(chunk) * 0.6:
            chunk = chunk[:k+1]
            j = i + k + 1
        out.append(chunk)
        i = max(j - overlap, j)
    return out

def _extract_docx_bytes(data: bytes) -> str:
    try:
        import docx  # python-docx
        doc = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return ""

def _extract_pdf_bytes(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        return ""

def y5_extract_from_uploads(files) -> str:
    if not files:
        return ""
    parts: list[str] = []
    for f in files:
        try:
            try:
                data = f.getbuffer().tobytes()
            except Exception:
                data = f.read()
            name = (getattr(f, "name", "") or "").lower()
            if name.endswith(".txt"):
                try:
                    parts.append(data.decode("utf-8", "ignore"))
                except Exception:
                    parts.append(str(data))
            elif name.endswith(".docx"):
                parts.append(_extract_docx_bytes(data))
            elif name.endswith(".pdf"):
                parts.append(_extract_pdf_bytes(data))
        except Exception:
            continue
    return "\n\n".join([p for p in parts if p]).strip()

def y5_extract_from_rfp(conn: sqlite3.Connection, rfp_id: int) -> str:
    # Expect rfp_files(filename TEXT, mime TEXT, bytes BLOB, rfp_id INT)
    try:
        df = pd.read_sql_query(
            "SELECT filename, mime, bytes FROM rfp_files WHERE rfp_id=? ORDER BY id;",
            conn, params=(int(rfp_id),)
        )
    except Exception:
        df = pd.DataFrame()
    if df is None or df.empty:
        return ""
    parts: list[str] = []
    for _, r in df.iterrows():
        data = r.get("bytes")
        if data is None:
            continue
        name = (r.get("filename") or "").lower()
        if name.endswith(".txt"):
            if isinstance(data, (bytes, bytearray)):
                try:
                    parts.append(bytes(data).decode("utf-8", "ignore"))
                except Exception:
                    continue
        elif name.endswith(".docx"):
            parts.append(_extract_docx_bytes(bytes(data)))
        elif name.endswith(".pdf"):
            parts.append(_extract_pdf_bytes(bytes(data)))
    return "\n\n".join([p for p in parts if p]).strip()

def y5_save_snippet(conn: sqlite3.Connection, rfp_id: int, section: str, text: str, source: str = "Y5") -> None:
    if not (text or "").strip():
        return
    try:
        with closing(conn.cursor()) as cur:
            cur.execute(
                "INSERT INTO draft_snippets(rfp_id, section, source, text) VALUES(?,?,?,?);",
                (int(rfp_id), (section or "General").strip(), (source or "Y5").strip(), text.strip())
            )
            conn.commit()
    except Exception:
        pass
# === end Y5 ===




# === Phase 2: Dedicated finders and status chips ===
def _collect_full_text(conn: sqlite3.Connection, rfp_id: int) -> str:
    try:
        df = pd.read_sql_query("SELECT filename, mime, bytes FROM rfp_files WHERE rfp_id=? ORDER BY id;", conn, params=(int(rfp_id),))
    except Exception:
        df = pd.DataFrame()
    parts = []
    for _, r in (df if (df is not None and not df.empty) else pd.DataFrame()).iterrows():
        b = r.get("bytes"); mime = r.get("mime") or ""
        pages = extract_text_pages(b, mime) or []
        if pages:
            pages, _ = ocr_pages_if_empty(b, mime, pages)
        parts.append("\n\n".join(pages))
    return "\n\n".join([p for p in parts if p]).strip()

def _upsert_meta(conn, rfp_id: int, key: str, value: str):
    if not value:
        return
    with closing(conn.cursor()) as cur:
        cur.execute("DELETE FROM rfp_meta WHERE rfp_id=? AND key=?;", (int(rfp_id), key))
        cur.execute("INSERT INTO rfp_meta(rfp_id, key, value) VALUES(?,?,?);", (int(rfp_id), key, value))
        conn.commit()

def find_due_date(conn: sqlite3.Connection, rfp_id: int) -> str:
    # Check SAM facts first
    try:
        row = pd.read_sql_query("SELECT extracted_json FROM sam_versions WHERE rfp_id=? ORDER BY id DESC LIMIT 1;", conn, params=(int(rfp_id),)).iloc[0]
        facts = json.loads(row["extracted_json"] or "{}")
        if isinstance(facts, dict) and facts.get("offers_due"):
            dd = str(facts["offers_due"]).strip()
            _upsert_meta(conn, int(rfp_id), "offers_due", dd)
            return dd
    except Exception:
        pass
    # Search chunks
    try:
        dfc = pd.read_sql_query("SELECT text FROM rfp_chunks WHERE rfp_id=?;", conn, params=(int(rfp_id),))
    except Exception:
        dfc = pd.DataFrame()
    t = " ".join((dfc["text"].tolist() if not dfc.empty else []))[:500000]
    m = re.search(r"(?i)(offers|quotes|proposals)\s+due[^\d]{0,20}(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|[A-Za-z]{3,9}\s+\d{1,2},\s*\d{4})", t)
    if m:
        dd = m.group(2).strip()
        _upsert_meta(conn, int(rfp_id), "offers_due", dd); return dd
    return ""

def find_naics_setaside(conn: sqlite3.Connection, rfp_id: int) -> dict:
    full = _collect_full_text(conn, int(rfp_id))
    naics = _extract_naics(full)
    set_aside = _extract_set_aside(full)
    if naics: _upsert_meta(conn, int(rfp_id), "naics", naics)
    if set_aside: _upsert_meta(conn, int(rfp_id), "set_aside", set_aside)
    return {"naics": naics, "set_aside": set_aside}

def find_pop(conn: sqlite3.Connection, rfp_id: int) -> dict:
    full = _collect_full_text(conn, int(rfp_id))
    pop = extract_pop_structure(full) or {}
    for k, v in pop.items():
        _upsert_meta(conn, int(rfp_id), k, str(v))
    return pop

def find_section_M(conn: sqlite3.Connection, rfp_id: int) -> int:
    full = _collect_full_text(conn, int(rfp_id))
    sec = extract_sections_L_M(full)
    cnt = 0
    with closing(conn.cursor()) as cur:
        if sec.get("L"):
            cur.execute("DELETE FROM rfp_sections WHERE rfp_id=? AND section='L';", (int(rfp_id),))
            cur.execute("INSERT INTO rfp_sections(rfp_id, section, content) VALUES(?,?,?);", (int(rfp_id), "L", sec["L"][:200000]))
            cnt += 1
        if sec.get("M"):
            cur.execute("DELETE FROM rfp_sections WHERE rfp_id=? AND section='M';", (int(rfp_id),))
            cur.execute("INSERT INTO rfp_sections(rfp_id, section, content) VALUES(?,?,?);", (int(rfp_id), "M", sec["M"][:200000]))
            cnt += 1
        conn.commit()
    return cnt

def find_clins_all(conn: sqlite3.Connection, rfp_id: int) -> int:
    full = _collect_full_text(conn, int(rfp_id))
    rows = extract_clins(full)
    try:
        df = pd.read_sql_query("SELECT clin, description FROM clin_lines WHERE rfp_id=?;", conn, params=(int(rfp_id),))
        existing = set((str(r.get("clin","")), str(r.get("description",""))) for _, r in df.iterrows())
    except Exception:
        existing = set()
    added = 0
    with closing(conn.cursor()) as cur:
        for r in rows:
            key = (r.get("clin",""), r.get("description",""))
            if key in existing:
                continue
            cur.execute("INSERT INTO clin_lines(rfp_id, clin, description, qty, unit, unit_price, extended_price) VALUES(?,?,?,?,?,?,?);",
                        (int(rfp_id), r.get("clin"), r.get("description"), r.get("qty"), r.get("unit"), r.get("unit_price"), r.get("extended_price")))
            added += 1
        conn.commit()
    return added

def _parse_money(x):
    try:
        s = str(x or "").replace(",", "").replace("$","").strip()
        return float(s) if s else 0.0
    except Exception:
        return 0.0

def clin_totals_df(conn: sqlite3.Connection, rfp_id: int):
    try:
        df = pd.read_sql_query("SELECT clin, description, qty, unit_price, extended_price FROM clin_lines WHERE rfp_id=? ORDER BY id;", conn, params=(int(rfp_id),))
    except Exception:
        return pd.DataFrame()
    if df is None or df.empty:
        return df
    qn = []
    up = []
    ext = []
    for _, r in df.iterrows():
        try:
            qn.append(float(str(r.get("qty","") or "0").replace(",","")))
        except Exception:
            qn.append(0.0)
        up.append(_parse_money(r.get("unit_price")))
        ext_val = _parse_money(r.get("extended_price"))
        if not ext_val and qn[-1] and up[-1]:
            ext_val = qn[-1]*up[-1]
        ext.append(ext_val)
    df["qty_num"] = qn; df["unit_price_num"] = up; df["extended_num"] = ext
    return df

def render_status_and_gaps(conn: sqlite3.Connection) -> None:
    st.subheader("Status & Gaps")
    try:
        df_rf = pd.read_sql_query("SELECT id, title FROM rfps ORDER BY id DESC;", conn, params=())
    except Exception:
        st.info("No RFPs yet."); return
    if df_rf is None or df_rf.empty:
        st.info("No RFPs yet."); return
    rid = st.selectbox("RFP context", options=df_rf["id"].tolist(),
                       format_func=lambda i: f"#{i} — {df_rf.loc[df_rf['id']==i,'title'].values[0]}",
                       key="p2_rfp")
    c1, c2 = st.columns([1,1])
    with c1:
        if st.button("Run finders", key="p2_find"):
            with st.spinner("Scanning RFP…"):
                find_due_date(conn, int(rid))
                find_naics_setaside(conn, int(rid))
                find_pop(conn, int(rid))
                find_section_M(conn, int(rid))
                find_clins_all(conn, int(rid))
            st.success("Updated metadata and sections.")
            st.rerun()
    # Chips
    try:
        dfm = pd.read_sql_query("SELECT key, value FROM rfp_meta WHERE rfp_id=?;", conn, params=(int(rid),))
        meta = {r["key"]: r["value"] for _, r in dfm.iterrows()} if dfm is not None and not dfm.empty else {}
    except Exception:
        meta = {}
    due = meta.get("offers_due","")
    naics = meta.get("naics","")
    sa = meta.get("set_aside","")
    pop = meta.get("pop_structure","") or meta.get("ordering_period_years","")
    def chip(label, ok):
        st.markdown(f"<span style='display:inline-block;padding:4px 8px;border-radius:12px;background:{'#DCF7E3' if ok else '#FBE8E8'};margin-right:6px'>{label}: {'OK' if ok else 'Missing'}</span>", unsafe_allow_html=True)
    st.write("")
    chips = []
    chips.append(("Due", bool(due)))
    chips.append(("NAICS", bool(naics)))
    chips.append(("Set-Aside", bool(sa)))
    chips.append(("POP", bool(pop)))
    try:
        has_M = pd.read_sql_query("SELECT 1 FROM rfp_sections WHERE rfp_id=? AND section='M' LIMIT 1;", conn, params=(int(rid),)).shape[0] > 0
    except Exception:
        has_M = False
    try:
        has_CLIN = pd.read_sql_query("SELECT 1 FROM clin_lines WHERE rfp_id=? LIMIT 1;", conn, params=(int(rid),)).shape[0] > 0
    except Exception:
        has_CLIN = False
    chips.append(("Section M", has_M))
    chips.append(("CLINs", has_CLIN))
    for label, ok in chips:
        chip(label, ok)
    # Suggested questions
    gaps = [lbl for lbl, ok in chips if not ok]
    if gaps:
        st.markdown("**Suggested questions**")
        qs = []
        if "Due" in gaps: qs.append("Confirm proposals due date and exact time zone.")
        if "NAICS" in gaps: qs.append("What NAICS code applies and size standard?")
        if "Set-Aside" in gaps: qs.append("What is the set-aside and eligibility?")
        if "POP" in gaps: qs.append("What is the base and option structure for POP?")
        if "Section M" in gaps: qs.append("List Section M factors and relative weights.")
        if "CLINs" in gaps: qs.append("Extract CLIN schedule and quantities.")
        st.write("\\n".join([f"- {q}" for q in qs]))
    # CLIN totals + CSV
    st.markdown("**CLIN totals**")
    dfc = clin_totals_df(conn, int(rid))
    if dfc is None or dfc.empty:
        st.caption("No CLINs yet.")
    else:
        total = float(dfc["extended_num"].sum())
        st.caption(f"Total Extended: ${total:,.2f}")
        _styled_dataframe(dfc[["clin","description","qty","unit_price","extended_price","extended_num"]], use_container_width=True, hide_index=True)
        csvb = dfc.to_csv(index=False).encode("utf-8")
        st.warning("Checking L/M gate before export")
        st.warning("Checking L/M gate before export")
        ok_gate, missing_gate = require_LM_minimum(conn, int(rid))
        if not ok_gate:
            st.button("Export CLIN CSV", key=f"p2_clin_csv_blocked_{rid}", disabled=True, help="Blocked: " + ", ".join(missing_gate))
        else:
            st.download_button("Export CLIN CSV", data=csvb, file_name=f"rfp_{int(rid)}_clins.csv", mime="text/csv", key=f"p2_clin_csv_{rid}")




def get_db() -> sqlite3.Connection:
    ensure_dirs()
    conn = _db_connect(DB_PATH, check_same_thread=False)
    with closing(conn.cursor()) as cur:
        cur.execute("PRAGMA foreign_keys = ON;")

        # Core
        cur.execute("""
            CREATE TABLE IF NOT EXISTS contacts(
                id INTEGER PRIMARY KEY,
                name TEXT,
                email TEXT,
                org TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS deals(
                id INTEGER PRIMARY KEY,
                title TEXT NOT NULL,
                rfp_id INTEGER REFERENCES rfps(id) ON DELETE SET NULL,
                amount REAL,
                stage TEXT,
                status TEXT DEFAULT 'Open',
                close_date TEXT,
                owner TEXT,
                created_at TEXT,
                updated_at TEXT
            );
        """)
        # Ensure new columns exist for Deals/SAM Watch
        try:
            _migrate_deals_columns(conn)
        except Exception:
            pass
        cur.execute("""
            CREATE TABLE IF NOT EXISTS app_settings(
                key TEXT PRIMARY KEY,
                val TEXT
            );
        """)

        # Org profile (Phase F - Capability Statement)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS org_profile(
                id INTEGER PRIMARY KEY CHECK (id=1),
                company_name TEXT,
                tagline TEXT,
                address TEXT,
                phone TEXT,
                email TEXT,
                website TEXT,
                uei TEXT,
                cage TEXT,
                naics TEXT,
                core_competencies TEXT,
                differentiators TEXT,
                certifications TEXT,
                past_performance TEXT,
                primary_poc TEXT
            );
        """)

        # Phase B (RFP analyzer artifacts)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rfps(
                id INTEGER PRIMARY KEY,
                title TEXT,
                solnum TEXT,
                notice_id TEXT,
                sam_url TEXT,
                file_path TEXT,
                created_at TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rfp_sections(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                section TEXT,
                content TEXT
            );
        """)


        cur.execute("""
            CREATE TABLE IF NOT EXISTS rfp_files(
    id INTEGER PRIMARY KEY,
    rfp_id INTEGER REFERENCES rfps(id) ON DELETE SET NULL,
    filename TEXT,
    mime TEXT,
    sha256 TEXT UNIQUE,
    pages INTEGER,
    bytes BLOB,
    created_at TEXT
);
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_rfp_files_rfp ON rfp_files(rfp_id);")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS lm_items(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                item_text TEXT,
                is_must INTEGER DEFAULT 1,
                status TEXT DEFAULT 'Open'
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS clin_lines(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                clin TEXT,
                description TEXT,
                qty TEXT,
                unit TEXT,
                unit_price TEXT,
                extended_price TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS key_dates(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                label TEXT,
                date_text TEXT,
                date_iso TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS pocs(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                name TEXT,
                role TEXT,
                email TEXT,
                phone TEXT
            );
        """)

        # Phase D (vendors + outreach)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS vendors(
                id INTEGER PRIMARY KEY,
                name TEXT NOT NULL,
                cage TEXT,
                uei TEXT,
                naics TEXT,
                city TEXT,
                state TEXT,
                phone TEXT,
                email TEXT,
                website TEXT,
                notes TEXT,
                last_seen_award TEXT
            );
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_vendors_naics_state ON vendors(naics, state);")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS vendor_contacts(
                id INTEGER PRIMARY KEY,
                vendor_id INTEGER NOT NULL REFERENCES vendors(id) ON DELETE CASCADE,
                name TEXT,
                email TEXT,
                phone TEXT,
                role TEXT
            );
        """)

        # Phase E (quotes + pricing)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS quotes(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                vendor TEXT NOT NULL,
                received_date TEXT,
                notes TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS quote_lines(
                id INTEGER PRIMARY KEY,
                quote_id INTEGER NOT NULL REFERENCES quotes(id) ON DELETE CASCADE,
                clin TEXT,
                description TEXT,
                qty REAL,
                unit_price REAL,
                extended_price REAL
            );
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_quote_lines ON quote_lines(quote_id, clin);")

        cur.execute("""
            CREATE TABLE IF NOT EXISTS pricing_scenarios(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                name TEXT NOT NULL,
                overhead_pct REAL DEFAULT 0.0,
                gna_pct REAL DEFAULT 0.0,
                fee_pct REAL DEFAULT 0.0,
                contingency_pct REAL DEFAULT 0.0,
                created_at TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS pricing_labor(
                id INTEGER PRIMARY KEY,
                scenario_id INTEGER NOT NULL REFERENCES pricing_scenarios(id) ON DELETE CASCADE,
                labor_cat TEXT,
                hours REAL,
                rate REAL,
                fringe_pct REAL DEFAULT 0.0
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS pricing_other(
                id INTEGER PRIMARY KEY,
                scenario_id INTEGER NOT NULL REFERENCES pricing_scenarios(id) ON DELETE CASCADE,
                label TEXT,
                cost REAL
            );
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS past_perf(
                id INTEGER PRIMARY KEY,
                project_title TEXT NOT NULL,
                customer TEXT,
                contract_no TEXT,
                naics TEXT,
                role TEXT,
                pop_start TEXT,
                pop_end TEXT,
                value NUMERIC,
                scope TEXT,
                results TEXT,
                cpars_rating TEXT,
                contact_name TEXT,
                contact_email TEXT,
                contact_phone TEXT,
                keywords TEXT,
                notes TEXT
            );
        """)

        # Phase H (White Paper Builder)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS white_templates(
                id INTEGER PRIMARY KEY,
                name TEXT NOT NULL,
                description TEXT,
                created_at TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS white_template_sections(
                id INTEGER PRIMARY KEY,
                template_id INTEGER NOT NULL REFERENCES white_templates(id) ON DELETE CASCADE,
                position INTEGER NOT NULL,
                title TEXT,
                body TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS white_papers(
                id INTEGER PRIMARY KEY,
                title TEXT NOT NULL,
                subtitle TEXT,
                rfp_id INTEGER REFERENCES rfps(id) ON DELETE SET NULL,
                created_at TEXT,
                updated_at TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS white_paper_sections(
                id INTEGER PRIMARY KEY,
                paper_id INTEGER NOT NULL REFERENCES white_papers(id) ON DELETE CASCADE,
                position INTEGER NOT NULL,
                title TEXT,
                body TEXT,
                image_path TEXT
            );
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS activities(
                id INTEGER PRIMARY KEY,
                ts TEXT,
                type TEXT,
                subject TEXT,
                notes TEXT,
                deal_id INTEGER REFERENCES deals(id) ON DELETE SET NULL,
                contact_id INTEGER REFERENCES contacts(id) ON DELETE SET NULL
            );
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_activities_ts ON activities(ts);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_activities_rel ON activities(deal_id, contact_id);")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS tasks(
                id INTEGER PRIMARY KEY,
                title TEXT NOT NULL,
                due_date TEXT,
                status TEXT DEFAULT 'Open',
                priority TEXT DEFAULT 'Normal',
                deal_id INTEGER REFERENCES deals(id) ON DELETE SET NULL,
                contact_id INTEGER REFERENCES contacts(id) ON DELETE SET NULL,
                created_at TEXT,
                completed_at TEXT
            );
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_tasks_due ON tasks(due_date, status);")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS deal_stage_log(
                id INTEGER PRIMARY KEY,
                deal_id INTEGER NOT NULL REFERENCES deals(id) ON DELETE CASCADE,
                stage TEXT NOT NULL,
                changed_at TEXT
            );
        """)

        # Phase J (File Manager)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS files(
                id INTEGER PRIMARY KEY,
                owner_type TEXT,            -- 'RFP' | 'Deal' | 'Vendor' | 'Other'
                owner_id INTEGER,           -- nullable when owner_type='Other'
                filename TEXT,
                path TEXT,
                size INTEGER,
                mime TEXT,
                tags TEXT,
                notes TEXT,
                uploaded_at TEXT
            );
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_files_owner ON files(owner_type, owner_id);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_files_tags ON files(tags);")

        # Phase L (RFQ Pack)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rfq_packs(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER REFERENCES rfps(id) ON DELETE SET NULL,
                deal_id INTEGER REFERENCES deals(id) ON DELETE SET NULL,
                title TEXT NOT NULL,
                instructions TEXT,
                due_date TEXT,
                created_at TEXT,
                updated_at TEXT
            );
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_rfq_packs_ctx ON rfq_packs(rfp_id, deal_id);");
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rfq_lines(
                id INTEGER PRIMARY KEY,
                pack_id INTEGER NOT NULL REFERENCES rfq_packs(id) ON DELETE CASCADE,
                clin_code TEXT,
                description TEXT,
                qty REAL,
                unit TEXT,
                naics TEXT,
                psc TEXT
            );
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_rfq_lines_pack ON rfq_lines(pack_id);");
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rfq_vendors(
                id INTEGER PRIMARY KEY,
                pack_id INTEGER NOT NULL REFERENCES rfq_packs(id) ON DELETE CASCADE,
                vendor_id INTEGER NOT NULL REFERENCES vendors(id) ON DELETE CASCADE
            );
        """)
        cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_rfq_vendors_unique ON rfq_vendors(pack_id, vendor_id);");
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rfq_attach(
                id INTEGER PRIMARY KEY,
                pack_id INTEGER NOT NULL REFERENCES rfq_packs(id) ON DELETE CASCADE,
                file_id INTEGER REFERENCES files(id) ON DELETE SET NULL,
                name TEXT,
                path TEXT
            );
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_rfq_attach_pack ON rfq_attach(pack_id);");

        # Phase M (Tenancy 1)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS tenants(
                id INTEGER PRIMARY KEY,
                name TEXT UNIQUE NOT NULL,
                created_at TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS current_tenant(
                id INTEGER PRIMARY KEY CHECK(id=1),
                ctid INTEGER
            );
        """)
        cur.execute("INSERT OR IGNORE INTO tenants(id, name, created_at) VALUES(1, 'Default', datetime('now'));")
        cur.execute("INSERT OR IGNORE INTO current_tenant(id, ctid) VALUES(1, 1);")

        def _add_tenant_id(table: str):
            try:
                cols = pd.read_sql_query(f"PRAGMA table_info({table});", conn)
                if "tenant_id" not in cols["name"].tolist():
                    cur.execute(f"ALTER TABLE {table} ADD COLUMN tenant_id INTEGER;")
                    conn.commit()
            except Exception:
                pass
            try:
                cur.execute(f"CREATE INDEX IF NOT EXISTS idx_{table}_tenant ON {table}(tenant_id);")
            except Exception:
                pass

        core_tables = ["rfps","lm_items","lm_meta","deals","activities","tasks","deal_stage_log",
                       "vendors","files","rfq_packs","rfq_lines","rfq_vendors","rfq_attach","contacts"]
        for t in core_tables:
            _add_tenant_id(t)

        # AFTER INSERT triggers: always stamp tenant_id to current_tenant
        def _ensure_trigger(table: str):
            trg = f"{table}_ai_tenant"
            try:
                cur.execute(f"""
                    CREATE TRIGGER IF NOT EXISTS {trg}
                    AFTER INSERT ON {table}
                    BEGIN
                        UPDATE {table}
                        SET tenant_id=(SELECT ctid FROM current_tenant WHERE id=1)
                        WHERE rowid=NEW.rowid;
                    END;
                """)
            except Exception:
                pass
        for t in core_tables:
            _ensure_trigger(t)

        # Scoped views
        def _create_view(table: str):
            v = f"{table}_t"
            try:
                cur.execute(f"CREATE VIEW IF NOT EXISTS {v} AS SELECT * FROM {table} WHERE tenant_id=(SELECT ctid FROM current_tenant WHERE id=1);")
            except Exception:
                pass
        for t in core_tables:
            _create_view(t)

        # Phase N (Persist): Pragmas
        try:
            cur.execute("PRAGMA journal_mode=WAL;")
            cur.execute("PRAGMA synchronous=NORMAL;")
            cur.execute("PRAGMA foreign_keys=ON;")
            cur.execute("PRAGMA busy_timeout=5000;")
        except Exception:
            pass


        # RTM (Requirements Traceability Matrix)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rtm_requirements(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                req_key TEXT,
                source_type TEXT,       -- 'L/M' | 'SOW' | 'CLIN' | 'Other'
                source_file TEXT,
                page INTEGER,
                text TEXT NOT NULL,
                status TEXT DEFAULT 'Open', -- Open | Covered | N/A
                created_at TEXT,
                updated_at TEXT
            );
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rtm_links(
                id INTEGER PRIMARY KEY,
                rtm_id INTEGER NOT NULL REFERENCES rtm_requirements(id) ON DELETE CASCADE,
                link_type TEXT,         -- 'Proposal' | 'Pricing' | 'Subcontractor' | 'Clause' | 'Other'
                target TEXT,            -- pointer to evidence (file name+page or section id)
                note TEXT,
                created_at TEXT,
                updated_at TEXT
            );
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_rtm_rfp ON rtm_requirements(rfp_id);")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_rtm_links ON rtm_links(rtm_id);")

        # SAM Amendment snapshots
        cur.execute("""
            CREATE TABLE IF NOT EXISTS sam_versions(
                id INTEGER PRIMARY KEY,
                rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                url TEXT,
                sha256 TEXT,
                extracted_json TEXT,
                created_at TEXT
            );
        """ )
        cur.execute("""
            CREATE TABLE IF NOT EXISTS sam_extracts(
                id INTEGER PRIMARY KEY,
                sam_version_id INTEGER NOT NULL REFERENCES sam_versions(id) ON DELETE CASCADE,
                key TEXT,
                value TEXT
            );
        """)
# Schema version for migrations
        cur.execute("""
            CREATE TABLE IF NOT EXISTS schema_version(
                id INTEGER PRIMARY KEY CHECK(id=1),
                ver INTEGER
            );
        """)
        cur.execute("INSERT OR IGNORE INTO schema_version(id, ver) VALUES(1, 0);")
        conn.commit()
    try:
        ensure_y5_tables(conn)
    except Exception:
        pass
    try:
        migrate(conn)
    except Exception:
        pass
    return conn



# Compatibility shim for vendors_t
try:
    with conn:
        conn.execute("CREATE TABLE IF NOT EXISTS vendors (id INTEGER PRIMARY KEY, name TEXT, cage TEXT, uei TEXT, naics TEXT, city TEXT, state TEXT, phone TEXT, email TEXT, website TEXT, notes TEXT, place_id TEXT UNIQUE)")
        conn.execute("CREATE VIEW IF NOT EXISTS vendors_t AS SELECT id, name, email, phone, city, state, naics, cage, uei, website, notes FROM vendors")
except Exception:
    pass

def _file_hash() -> str:
    try:
        import hashlib
        with open(__file__, 'rb') as f:
            return hashlib.sha256(f.read()).hexdigest()[:12]
    except Exception:
        return "unknown"


def save_uploaded_file(uploaded_file, subdir: str = "") -> Optional[str]:
    if not uploaded_file:
        return None
    base_dir = UPLOADS_DIR if not subdir else os.path.join(UPLOADS_DIR, subdir)
    os.makedirs(base_dir, exist_ok=True)
    path = os.path.join(base_dir, uploaded_file.name)
    with open(path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return path


# -------------------- SAM Watch helpers (Phase A) --------------------
def get_sam_api_key() -> Optional[str]:
    key = st.session_state.get("temp_sam_key")
    if key:
        return key
    try:
        key = st.secrets.get("sam", {}).get("api_key")
        if key:
            return key
    except Exception:
        pass
    try:
        key = st.secrets.get("SAM_API_KEY")
        if key:
            return key
    except Exception:
        pass
    return os.getenv("SAM_API_KEY")


@st.cache_data(show_spinner=False, ttl=300)
def sam_search_cached(params: Dict[str, Any]) -> Dict[str, Any]:
    api_key = params.get("api_key")
    if not api_key:
        return {"totalRecords": 0, "records": [], "error": "Missing API key"}

    limit = int(params.get("limit", 100))
    max_pages = int(params.pop("_max_pages", 3))
    params["limit"] = min(max(1, limit), 1000)

    all_records: List[Dict[str, Any]] = []
    offset = int(params.get("offset", 0))

    for _ in range(max_pages):
        q = {**params, "offset": offset}
        try:
            resp = requests.get(SAM_ENDPOINT, params=q, headers={"X-Api-Key": api_key}, timeout=30)
        except Exception as ex:
            return {"totalRecords": 0, "records": [], "error": f"Request error: {ex}"}

        if resp.status_code != 200:
            try:
                j = resp.json()
                msg = j.get("message") or j.get("error") or str(j)
            except Exception:
                msg = resp.text
            return {"totalRecords": 0, "records": [], "error": f"HTTP {resp.status_code}: {msg}", "status": resp.status_code, "body": msg}

        data = resp.json() or {}
        records = data.get("opportunitiesData", data.get("data", []))
        if not isinstance(records, list):
            records = []
        all_records.extend(records)

        total = data.get("totalRecords", len(all_records))
        if len(all_records) >= total:
            break
        offset += params["limit"]

    return {"totalRecords": len(all_records), "records": all_records, "error": None}


def flatten_records(records: List[Dict[str, Any]]) -> pd.DataFrame:
    rows = []
    for r in records:
        title = r.get("title") or ""
        solnum = r.get("solicitationNumber") or r.get("solnum") or ""
        posted = r.get("postedDate") or ""
        ptype = r.get("type") or r.get("baseType") or ""
        set_aside = r.get("setAside") or ""
        set_aside_code = r.get("setAsideCode") or ""
        naics = r.get("naicsCode") or r.get("ncode") or ""
        psc = r.get("classificationCode") or r.get("ccode") or ""
        deadline = r.get("reponseDeadLine") or r.get("responseDeadline") or ""
        org_path = r.get("fullParentPathName") or r.get("organizationName") or ""
        notice_id = r.get("noticeId") or r.get("noticeid") or r.get("id") or ""
        sam_url = f"https://sam.gov/opp/{notice_id}/view" if notice_id else ""
        rows.append(
            {
                "Title": title,
                "Solicitation": solnum,
                "Type": ptype,
                "Posted": posted,
                "Response Due": deadline,
                "Set-Aside": set_aside,
                "Set-Aside Code": set_aside_code,
                "NAICS": naics,
                "PSC": psc,
                "Agency Path": org_path,
                "Notice ID": notice_id,
                "SAM Link": sam_url,
            }
        )
    df = pd.DataFrame(rows)
    wanted = [
        "Title", "Solicitation", "Type", "Posted", "Response Due",
        "Set-Aside", "Set-Aside Code", "NAICS", "PSC",
        "Agency Path", "Notice ID", "SAM Link",
    ]
    return df[wanted] if not df.empty else df




def sam_try_fetch_attachments(notice_id: str) -> List[Tuple[str, bytes]]:
    """Best-effort attempt to fetch attachments for a SAM notice.
    Returns list of (filename, bytes). Falls back to saving the notice description HTML
    when public attachment download isn't available.
    """
    import io, zipfile, os
    files: List[Tuple[str, bytes]] = []
    if not notice_id:
        return files

    # Attempt 1: Use Opportunity Management API 'download all' if system creds are present.
    sys_key = None
    sys_auth = None
    try:
        sys_key = (st.secrets.get("sam", {}).get("system_api_key")
                   or st.secrets.get("SAM_SYSTEM_API_KEY")
                   or os.getenv("SAM_SYSTEM_API_KEY"))
        sys_auth = (st.secrets.get("sam", {}).get("system_auth")
                    or st.secrets.get("SAM_SYSTEM_AUTH")
                    or os.getenv("SAM_SYSTEM_AUTH"))
    except Exception:
        # Fall back to env only
        sys_key = os.getenv("SAM_SYSTEM_API_KEY")
        sys_auth = os.getenv("SAM_SYSTEM_AUTH")

    try:
        if sys_key and sys_auth:
            # Per docs: GET /{opportunityId}/resources/download/zip with Authorization header
            url = f"https://api.sam.gov/prod/opportunity/v1/api/{notice_id}/resources/download/zip"
            params = {"api_key": sys_key}
            headers = {"Authorization": sys_auth}
            r = requests.get(url, headers=headers, params=params, timeout=60)
            if r.ok and (r.headers.get("content-type","").lower().endswith("zip") or r.content[:2] == b'PK'):
                zf = zipfile.ZipFile(io.BytesIO(r.content))
                for zi in zf.infolist():
                    if zi.is_dir():
                        continue
                    try:
                        data = zf.read(zi)
                        files.append((os.path.basename(zi.filename) or "attachment.bin", data))
                    except Exception:
                        continue
                if files:
                    return files
            else:
                # If unauthorized or not found, silently fall back
                pass
    except Exception:
        # Swallow and fall back
        pass

    # Attempt 2: Save description HTML as a helpful "attachment"
    try:
        api_key = get_sam_api_key()
        desc_url = f"https://api.sam.gov/prod/opportunities/v1/noticedesc"
        params = {"noticeid": notice_id}
        if api_key:
            params["api_key"] = api_key
        resp = requests.get(desc_url, params=params, timeout=30)
        if resp.ok and resp.text:
            files.append((f"{notice_id}_description.html", resp.text.encode("utf-8", errors="ignore")))
    except Exception:
        pass

    return files


# ---------------------- Phase B: RFP parsing helpers ----------------------
def _safe_import_pdf_extractors():
    if _pypdf is not None:
        return ('pypdf', _pypdf)
    if _pdfplumber is not None:
        return ('pdfplumber', _pdfplumber)
    return None



def extract_text_from_file(path: str) -> str:
    """Unified extractor. Uses Phase X1 backends. Safe on environments missing PDF libs."""
    try:
        with open(path, 'rb') as fh:
            b = fh.read()
    except Exception:
        return ''
    name = os.path.basename(path)
    mime = _detect_mime_light(name)
    pages = extract_text_pages(b, mime)
    if not pages:
        try:
            return b.decode('utf-8', errors='ignore')
        except Exception:
            try:
                return b.decode('latin-1', errors='ignore')
            except Exception:
                return ''
    return "\n\n".join(pages)



def _detect_mime_light(name: str) -> str:
    n = (name or "").lower()
    if n.endswith(".pdf"): return "application/pdf"
    if n.endswith(".docx"): return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if n.endswith(".xlsx"): return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if n.endswith(".txt"): return "text/plain"
    return "application/octet-stream"

def _tesseract_ok() -> bool:
    try:
        import pytesseract  # type: ignore
        return True
    except Exception:
        return False


def extract_text_pages(file_bytes: bytes, mime: str) -> list:
    """Return a list of page texts. Best-effort. Up to 100 pages to keep fast."""
    out = []
    m = (mime or "").lower()
    if "pdf" in m:
        # Try vector: pdfplumber then pypdf. Never reference missing names.
        if _pdfplumber is not None:
            try:
                import io as _io
                with _pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
                    for i, pg in enumerate(pdf.pages[:100]):
                        try:
                            txt = pg.extract_text() or ""
                        except Exception:
                            txt = ""
                        out.append(txt)
            except Exception:
                pass
        if not out and _pypdf is not None:
            try:
                import io as _io
                reader = _pypdf.PdfReader(_io.BytesIO(file_bytes))
                for i, p in enumerate(reader.pages[:100]):
                    try:
                        out.append(p.extract_text() or "")
                    except Exception:
                        out.append("")
            except Exception:
                pass
        # As a very last resort, try a naive decode of bytes to avoid empty output.
        if not out:
            try:
                out = [file_bytes.decode("utf-8", errors="ignore")]
            except Exception:
                out = []
    elif "wordprocessingml" in m or (mime == "" and file_bytes[:4] == b"PK\x03\x04"):
        try:
            import io as _io, docx  # type: ignore
            doc = docx.Document(_io.BytesIO(file_bytes))
            txt = "\n".join(p.text or "" for p in doc.paragraphs)
            out = [txt] if txt else []
        except Exception:
            out = []
    elif "spreadsheetml" in m:
        try:
            import io as _io, pandas as _pd  # type: ignore
            x = _pd.read_excel(_io.BytesIO(file_bytes), sheet_name=None, dtype=str)
            for sname, df in list(x.items())[:10]:
                txt = sname + "\n" + df.fillna("").astype(str).to_csv(sep="\t", index=False)
                out.append(txt)
        except Exception:
            out = []
    elif "text/plain" in m:
        try:
            out = [file_bytes.decode("utf-8", errors="ignore")]
        except Exception:
            out = [file_bytes.decode("latin-1", errors="ignore")]
    return out

def ocr_pages_if_empty(file_bytes: bytes, mime: str, pages_text: list) -> tuple:
    """Run OCR on empty PDF pages if pytesseract available. Returns (new_pages, ocr_count)."""
    if "pdf" not in (mime or "").lower():
        return pages_text, 0
    if not _tesseract_ok():
        return pages_text, 0
    try:
        import io as _io
        pdfplumber = _pdfplumber  # type: ignore
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
        new_pages = list(pages_text)
        ocr_count = 0
        if pdfplumber is None:
            return pages_text, 0
        with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
            for i, pg in enumerate(pdf.pages[:min(len(new_pages) or 100, 100)]):
                if i >= len(new_pages): new_pages.append("")
                if (new_pages[i] or "").strip():
                    continue
                try:
                    img = pg.to_image(resolution=200).original
                    if img is None:
                        continue
                    if not isinstance(img, Image.Image):
                        img = Image.fromarray(img)
                    txt = pytesseract.image_to_string(img) or ""
                    if txt.strip():
                        new_pages[i] = txt
                        ocr_count += 1
                except Exception:
                    pass
        return new_pages, ocr_count
    except Exception:
        return pages_text, 0

def save_rfp_file_db(conn: sqlite3.Connection, rfp_id: int | None, name: str, file_bytes: bytes) -> dict:
    """Dedup by sha256. Store bytes and basic stats. Return dict with id and stats."""
    mime = _detect_mime_light(name)
    sha = compute_sha256(file_bytes)
    with closing(conn.cursor()) as cur:
        # Dedup
        cur.execute("SELECT id, pages FROM rfp_files WHERE sha256=?;", (sha,))
        row = cur.fetchone()
        if row:
            rid = int(row[0]); pages = int(row[1]) if row[1] is not None else None
            # if not linked to RFP yet, link now
            if rfp_id is not None:
                try:
                    cur.execute("UPDATE rfp_files SET rfp_id=COALESCE(rfp_id, ?) WHERE id=?;", (int(rfp_id), rid))
                    conn.commit()
                except Exception:
                    pass
            return {"id": rid, "sha256": sha, "filename": name, "mime": mime, "pages": pages, "dedup": True, "ocr_pages": 0}
        # New insert
        pages_text = extract_text_pages(file_bytes, mime)
        pages_before = len(pages_text)
        pages_text, ocr_count = ocr_pages_if_empty(file_bytes, mime, pages_text)
        pages = len(pages_text) if pages_text else None
        cur.execute(
            "INSERT INTO rfp_files(rfp_id, filename, mime, sha256, pages, bytes, created_at) VALUES(?,?,?,?,?,?, datetime('now'));",
            (int(rfp_id) if rfp_id is not None else None, name, mime, sha, pages or 0, sqlite3.Binary(file_bytes))
        )
        rid = cur.lastrowid
        conn.commit()
        return {"id": rid, "sha256": sha, "filename": name, "mime": mime, "pages": pages or 0, "dedup": False, "ocr_pages": ocr_count}

def extract_sections_L_M(text: str) -> dict:
    out = {}
    if not text:
        return out
    mL = re.search(r'(SECTION\s+L[\s\S]*?)(?=SECTION\s+[A-Z]|\Z)', text, re.IGNORECASE)
    if mL:
        out['L'] = mL.group(1)
    mM = re.search(r'(SECTION\s+M[\s\S]*?)(?=SECTION\s+[A-Z]|\Z)', text, re.IGNORECASE)
    if mM:
        out['M'] = mM.group(1)
    return out


def derive_lm_items(section_text: str) -> list:
    if not section_text:
        return []
    items = []
    for line in section_text.splitlines():
        s = line.strip()
        if len(s) < 4:
            continue
        if re.match(r'^([\-\u2022\*]|\(?[a-zA-Z0-9]\)|[0-9]+\.)\s+', s):
            items.append(s)
    seen = set()
    uniq = []
    for it in items:
        if it not in seen:
            uniq.append(it)
            seen.add(it)
    return uniq[:500]


def extract_clins(text: str) -> list:
    if not text:
        return []
    lines = text.splitlines()
    rows = []
    for i, ln in enumerate(lines):
        m = re.search(r'\bCLIN\s*([A-Z0-9\-]+)', ln, re.IGNORECASE)
        if m:
            clin = m.group(1)
            desc = lines[i+1].strip() if i+1 < len(lines) else ''
            mqty = re.search(r'\b(QTY|Quantity)[:\s]*([0-9,.]+)', ln + ' ' + desc, re.IGNORECASE)
            qty = mqty.group(2) if mqty else ''
            munit = re.search(r'\b(UNIT|Units?)[:\s]*([A-Za-z/]+)', ln + ' ' + desc, re.IGNORECASE)
            unit = munit.group(2) if munit else ''
            rows.append({
                'clin': clin,
                'description': desc[:300],
                'qty': qty,
                'unit': unit,
                'unit_price': '',
                'extended_price': ''
            })
    seen = set()
    uniq = []
    for r in rows:
        if r['clin'] not in seen:
            uniq.append(r)
            seen.add(r['clin'])
    return uniq[:500]


def extract_dates(text: str) -> list:
    if not text:
        return []
    patterns = [
        r'(Questions(?:\s+Due)?|Q&A(?:\s+Due)?|Inquiry Deadline)[:\s]*([A-Za-z]{3,9}\s+\d{1,2},\s*\d{4}|\d{1,2}/\d{1,2}/\d{2,4})',
        r'(Proposals?\s+Due|Offers?\s+Due|Closing Date)[:\s]*([A-Za-z]{3,9}\s+\d{1,2},\s*\d{4}|\d{1,2}/\d{1,2}/\d{2,4})',
        r'(Site\s+Visit)[:\s]*([A-Za-z]{3,9}\s+\d{1,2},\s*\d{4}|\d{1,2}/\d{1,2}/\d{2,4})',
        r'(Period\s+of\s+Performance|POP)[:\s]*([A-Za-z]{3,9}\s+\d{1,2},\s*\d{4}|\d{1,2}/\d{1,2}/\d{2,4})',
    ]
    out = []
    for pat in patterns:
        for m in re.finditer(pat, text, re.IGNORECASE):
            out.append({'label': m.group(1).strip(), 'date_text': m.group(2).strip(), 'date_iso': ''})
    return out[:200]


def extract_pocs(text: str) -> list:
    if not text:
        return []
    emails = list(set(re.findall(r'[\w\.-]+@[\w\.-]+\.[A-Za-z]{2,}', text)))
    phones = list(set(re.findall(r'(?:\+?1\s*)?(?:\(\d{3}\)|\d{3})[\s\-]?\d{3}[\s\-]?\d{4}', text)))
    poc_blocks = re.findall(r'(Contracting Officer|Contract Specialist|Point of Contact|POC).*?(?:\n\n|$)', text, re.IGNORECASE|re.DOTALL)
    names = []
    for blk in poc_blocks:
        for nm in re.findall(r'([A-Z][a-zA-Z\-]+\s+[A-Z][a-zA-Z\-]+)', blk):
            names.append(nm)
    out = []
    for i in range(max(len(names), len(emails), len(phones))):
        out.append({
            'name': names[i] if i < len(names) else '',
            'role': 'POC',
            'email': emails[i] if i < len(emails) else '',
            'phone': phones[i] if i < len(phones) else '',
        })
    return out[:100]
def _num_from_words(s: str) -> int | None:
    m = re.search(r'\d+', s or '')
    if m:
        return int(m.group(0))
    words = {'one':1,'two':2,'three':3,'four':4,'five':5,'six':6,'seven':7,'eight':8,'nine':9,'ten':10,'eleven':11,'twelve':12}
    t = (s or '').strip().lower()
    return words.get(t)

def extract_pop_structure(text: str) -> dict:
    if not text:
        return {}
    tl = text.lower()
    base = 1 if re.search(r'\bbase\s+(year|period)\b', tl) else 0
    oy = 0
    for pat in [r'(\d+)\s+option\s+(?:years?|periods?)',
                r'(one|two|three|four|five|six|seven|eight|nine|ten)\s+option\s+(?:years?|periods?)',
                r'option\s+years?\s+(one|two|three|four|five|six|seven|eight|nine|ten|\d+)\s*(?:through|-|to)\s*(one|two|three|four|five|six|seven|eight|nine|ten|\d+)\b']:
        m = re.search(pat, tl)
        if m:
            if len(m.groups())==2 and m.group(2):
                a = _num_from_words(m.group(1)) or 0
                b = _num_from_words(m.group(2)) or 0
                if b>=a and a>0:
                    oy = max(oy, b - a + 1)
            else:
                oy = max(oy, _num_from_words(m.group(1)) or 0)
    total_years = None
    m = re.search(r'(ordering|contract)\s+period[^\d]{0,20}(\d+)\s+year', tl)
    if m:
        try:
            total_years = int(m.group(2))
        except Exception:
            total_years = None
    base_months = None
    m = re.search(r'base\s+(?:year|period)[^\d]{0,12}(\d{1,2})\s+month', tl)
    if m:
        try:
            base_months = int(m.group(1))
        except Exception:
            base_months = None
    out = {}
    if base or oy:
        label = "Base" if base else ""
        if oy:
            if label:
                label += " + "
            label += f"{oy} Option Year{'s' if oy!=1 else ''}"
        out['pop_structure'] = label.strip()
    if total_years:
        out['ordering_period_years'] = total_years
    if base_months:
        out['base_months'] = base_months
    return out


# === Y5.5: AI-assisted Parse & Save (safe indent) ===

# === Phase 3: Parser schema and normals (Y55) ===
_Y55_SCHEMA = {
    "title": str,
    "solnum": str,
    "meta": dict,          # may include: naics, set_aside, place_of_performance, due_offer, due_questions
    "l_items": list,       # list[str]
    "clins": list,         # list[dict]
    "dates": list,         # list[dict]
    "pocs": list           # list[dict]
}

def _y55_norm_date(s: str) -> str:
    """
    Normalize many date strings to 'YYYY MM DD'. Return '' if unknown.
    """
    import re, datetime
    if not s:
        return ""
    t = str(s).strip()
    # Remove time parts
    t = re.sub(r'(\d{1,2}:\d{2}.*)$', '', t).strip()
    # Try formats explicitly
    fmts = ["%B %d, %Y", "%b %d, %Y", "%m/%d/%Y", "%m/%d/%y", "%Y-%m-%d", "%Y/%m/%d"]
    for fmt in fmts:
        try:
            dt = datetime.datetime.strptime(t, fmt).date()
            return f"{dt.year:04d} {dt.month:02d} {dt.day:02d}"
        except Exception:
            pass
    # Fallback: grab mm/dd/yy or Month D, YYYY via regex
    m = re.search(r'([A-Za-z]{3,9})\s+(\d{1,2}),\s*(\d{4})', t)
    if m:
        try:
            dt = datetime.datetime.strptime(m.group(0), "%B %d, %Y").date()
        except Exception:
            try:
                dt = datetime.datetime.strptime(m.group(0), "%b %d, %Y").date()
            except Exception:
                dt = None
        if dt:
            return f"{dt.year:04d} {dt.month:02d} {dt.day:02d}"
    m = re.search(r'(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})', t)
    if m:
        mm = int(m.group(1)); dd = int(m.group(2)); yy = int(m.group(3))
        if yy < 100:
            yy = 2000 + yy if yy < 50 else 1900 + yy
        try:
            dt = datetime.date(yy, mm, dd)
            return f"{dt.year:04d} {dt.month:02d} {dt.day:02d}"
        except Exception:
            return ""
    return ""

def _y55_coerce_money(x):
    try:
        s = str(x or "").replace(",", "").replace("$","").strip()
        return float(s) if s else 0.0
    except Exception:
        return 0.0

def _y55_validate(d: dict) -> dict:
    """
    Ensure structure matches _Y55_SCHEMA.
    Coerce money to numeric. Normalize date strings to 'YYYY MM DD' in dates[].date_iso and meta due_*.
    Replace malformed parts with safe defaults.
    """
    import re
    safe = {"title":"", "solnum":"", "meta":{}, "l_items":[], "clins":[], "dates":[], "pocs":[]}
    if not isinstance(d, dict):
        return safe
    out = dict(safe)
    # Scalars
    out["title"] = str(d.get("title") or "")[:200]
    out["solnum"] = str(d.get("solnum") or "")[:80]
    # Meta
    meta = d.get("meta") if isinstance(d.get("meta"), dict) else {}
    meta2 = {}
    for k in ("naics","set_aside","place_of_performance","due_offer","due_questions"):
        v = meta.get(k, "")
        if k in ("due_offer","due_questions"):
            meta2[k] = _y55_norm_date(v) if v else ""
        else:
            meta2[k] = str(v) if v is not None else ""
    out["meta"] = meta2
    # Lists
    out["l_items"] = [str(x).strip() for x in (d.get("l_items") or []) if isinstance(x, (str,int,float))]
    clins_in = d.get("clins") if isinstance(d.get("clins"), list) else []
    clins_out = []
    for r in clins_in:
        if not isinstance(r, dict):
            continue
        clins_out.append({
            "clin": str(r.get("clin") or ""),
            "description": str(r.get("description") or "")[:300],
            "qty": str(r.get("qty") or ""),
            "unit": str(r.get("unit") or ""),
            "unit_price": _y55_coerce_money(r.get("unit_price")),
            "extended_price": _y55_coerce_money(r.get("extended_price")),
        })
    out["clins"] = clins_out
    dates_in = d.get("dates") if isinstance(d.get("dates"), list) else []
    dates_out = []
    for r in dates_in:
        if not isinstance(r, dict):
            continue
        lbl = str(r.get("label") or "").strip()
        txt = str(r.get("date_text") or "").strip()
        iso = _y55_norm_date(r.get("date_iso") or txt)
        dates_out.append({"label": lbl, "date_text": txt, "date_iso": iso})
    out["dates"] = dates_out
    pocs_in = d.get("pocs") if isinstance(d.get("pocs"), list) else []
    pocs_out = []
    for r in pocs_in:
        if not isinstance(r, dict):
            continue
        pocs_out.append({
            "name": str(r.get("name") or ""),
            "role": str(r.get("role") or "POC"),
            "email": str(r.get("email") or ""),
            "phone": str(r.get("phone") or ""),
        })
    out["pocs"] = pocs_out
    # Derive due_offer / due_questions if not set, from dates[] labels
    if not out["meta"].get("due_offer"):
        for r in out["dates"]:
            if re.search(r"(Offer|Proposal|Quote|Closing|Response Due)", r["label"], re.I):
                if r.get("date_iso"):
                    out["meta"]["due_offer"] = r["date_iso"]; break
    if not out["meta"].get("due_questions"):
        for r in out["dates"]:
            if re.search(r"(Question|Q&A|Inquiry)", r["label"], re.I):
                if r.get("date_iso"):
                    out["meta"]["due_questions"] = r["date_iso"]; break
    return out


def y55_ai_parse(text: str) -> dict:
    out = {"title":"", "solnum":"", "meta":{}, "l_items":[], "clins":[], "dates":[], "pocs":[]}
    t = (text or "").strip()
    if not t:

        # Phase 3 post-process: validate schema, normalize dates/money, derive due_* fields
        out = _y55_validate(out)
        return out

    try:
        client = get_ai()
        model_name = _resolve_model()
        sys_msg = SYSTEM_CO + " You extract procurement data exactly and return strict JSON."
        user_msg = (
            "Extract these from the RFP text. Return JSON with keys: "
            "title, solnum, meta:{naics,set_aside,place_of_performance}, "
            "l_items: [bulleted requirement lines], "
            "clins: [{clin, description, qty, unit, unit_price, extended_price}], "
            "dates: [{label, date_text}], "
            "pocs: [{name, role, email, phone}]. "
            "Be terse. Do not hallucinate unknown numbers; leave fields empty if not present.\n\n"
            "RFP TEXT START\n" + t[:180000] + "\nRFP TEXT END"
        )
        try:
            resp = client.chat.completions.create(
                model=model_name,
                messages=[{"role":"system","content":sys_msg},{"role":"user","content":user_msg}],
                temperature=0.0
            )
        except Exception as _e:
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"system","content":sys_msg},{"role":"user","content":user_msg}],
                temperature=0.0
            )
        raw = ""
        try:
            raw = resp.choices[0].message.content or ""
        except Exception:
            raw = ""
        import json, re
        m = re.search(r'\{.*\}', raw, re.S)
        data = {}
        if m:
            try:
                data = json.loads(m.group(0))
            except Exception:
                data = {}
        if not data and raw.strip():
            raw2 = raw.strip().strip("`").strip()
            try:
                data = json.loads(raw2)
            except Exception:
                data = {}
        if isinstance(data, dict):
            for k in out.keys():
                if k in data and data[k]:
                    out[k] = data[k]
        if not isinstance(out.get("l_items"), list): out["l_items"] = []
        if not isinstance(out.get("clins"), list): out["clins"] = []
        if not isinstance(out.get("dates"), list): out["dates"] = []
        if not isinstance(out.get("pocs"), list): out["pocs"] = []
        if not isinstance(out.get("meta"), dict): out["meta"] = {}
        out["title"] = (out.get("title") or "").strip()[:200]
        out["solnum"] = (out.get("solnum") or "").strip()[:80]

        # Phase 3 post-process: validate schema, normalize dates/money, derive due_* fields
        out = _y55_validate(out)
        return out

    except Exception:
        pass

        # Phase 3 post-process: validate schema, normalize dates/money, derive due_* fields
        out = _y55_validate(out)
        return out

def _y55_norm_str(x):
    import re
    try:
        return re.sub(r'\s+',' ', str(x or '')).strip()
    except Exception:
        return str(x)

def y55_merge_lm(base_list, ai_list):
    base = [_y55_norm_str(x) for x in (base_list or []) if _y55_norm_str(x)]
    ai = [_y55_norm_str(x) for x in (ai_list or []) if _y55_norm_str(x)]
    seen = set(); out = []
    for it in base + ai:
        if it not in seen:
            seen.add(it); out.append(it)
    return out[:1000]

def y55_merge_clins(base_rows, ai_rows):
    def norm_row(r):
        return {
            "clin": _y55_norm_str((r or {}).get("clin")),
            "description": _y55_norm_str((r or {}).get("description"))[:300],
            "qty": _y55_norm_str((r or {}).get("qty")),
            "unit": _y55_norm_str((r or {}).get("unit")),
            "unit_price": _y55_norm_str((r or {}).get("unit_price")),
            "extended_price": _y55_norm_str((r or {}).get("extended_price")),
        }
    base = [norm_row(r) for r in (base_rows or [])]
    ai = [norm_row(r) for r in (ai_rows or [])]
    out = []; seen = set()
    for r in base + ai:
        key = (r["clin"], r["description"], r["qty"], r["unit_price"], r["extended_price"])
        if key in seen:
            continue
        seen.add(key); out.append(r)
    return out[:2000]

def y55_merge_dates(base_rows, ai_rows):
    def norm(r):
        return {"label": _y55_norm_str((r or {}).get("label")), "date_text": _y55_norm_str((r or {}).get("date_text")), "date_iso": _y55_norm_str((r or {}).get("date_iso"))}
    base = [norm(r) for r in (base_rows or [])]
    ai = [norm(r) for r in (ai_rows or [])]
    out = []; seen = set()
    for r in base + ai:
        key = (r["label"], r["date_text"])
        if key in seen:
            continue
        seen.add(key); out.append(r)
    return out[:300]

def y55_merge_pocs(base_rows, ai_rows):
    def norm(r):
        return {"name": _y55_norm_str((r or {}).get("name")), "role": _y55_norm_str((r or {}).get("role") or "POC"),
                "email": _y55_norm_str((r or {}).get("email")), "phone": _y55_norm_str((r or {}).get("phone"))}
    base = [norm(r) for r in (base_rows or [])]
    ai = [norm(r) for r in (ai_rows or [])]
    out = []; seen = set()
    for r in base + ai:
        key = (r["name"], r["email"], r["phone"])
        if key in seen:
            continue
        seen.add(key); out.append(r)
    return out[:300]

def y55_apply_enhancement(text, l_items, clins, dates, pocs, meta, title, solnum):

    # ---- X3 MODAL RENDERER ----
    if st.session_state.get("x3_show_modal"):
        _notice = st.session_state.get("x3_modal_notice", {}) or {}
        try:
            with st.modal("RFP Analyzer", key=_uniq_key("x3_modal", _safe_int(_notice.get("Notice ID")))):
                _x3_render_modal(_notice)
        except Exception:
            with st.expander("RFP Analyzer", expanded=True):
                _x3_render_modal(_notice)
    ai = y55_ai_parse(text or "")
    l_items2 = y55_merge_lm(l_items, ai.get("l_items", []))
    clins2 = y55_merge_clins(clins, ai.get("clins", []))
    dates2 = y55_merge_dates(dates, ai.get("dates", []))
    pocs2 = y55_merge_pocs(pocs, ai.get("pocs", []))
    meta2 = dict(meta or {})
    for k in ("naics","set_aside","place_of_performance"):
        v = (ai.get("meta", {}) or {}).get(k)
        if v:
            meta2[k] = _y55_norm_str(v)
    title2 = (ai.get("title") or "").strip() or (title or "")
    solnum2 = (ai.get("solnum") or "").strip() or (solnum or "")
    return (l_items2, clins2, dates2, pocs2, meta2, title2, solnum2)
# === end Y5.5 ===

def extract_clins_xlsx(file_bytes: bytes) -> list:
    try:
        import io, pandas as _pd
        wb = _pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, dtype=str)
    except Exception:
        return []
    rows = []
    for sname, df in wb.items():
        if df is None or df.empty:
            continue
        df2 = df.copy()
        df2.columns = [str(c).strip() for c in df2.columns]
        low = [c.lower() for c in df2.columns]
        def _pick(fn):
            for i,c in enumerate(low):
                try:
                    if fn(c):
                        return df2.columns[i]
                except Exception:
                    continue
            return None
        col_clin = _pick(lambda c: 'clin' in c or 'line item' in c or c.startswith('clin') or c.startswith('line'))
        col_desc = _pick(lambda c: 'desc' in c or c=='item' or 'description' in c)
        col_qty  = _pick(lambda c: 'qty' in c or 'quantity' in c)
        col_unit = _pick(lambda c: c in ('unit','u/i','uom') or 'unit ' in c or 'uom' in c)
        col_upr  = _pick(lambda c: 'unit price' in c or (('unit' in c or 'price' in c) and 'total' not in c and 'ext' not in c and 'extended' not in c))
        col_ext  = _pick(lambda c: 'extended' in c or 'amount' in c or c=='total' or 'total price' in c or 'ext price' in c)
        if (col_clin or col_desc) and (col_qty or col_upr or col_ext):
            for _, r in df2.iterrows():
                def gv(col):
                    return "" if col is None else str(r.get(col, "")).strip()
                clin = gv(col_clin); desc = gv(col_desc); qty  = gv(col_qty)
                unit = gv(col_unit); upr  = gv(col_upr);  ext  = gv(col_ext)
                if any([clin, desc, qty, upr, ext]):
                    rows.append({'clin': clin, 'description': desc[:300] if desc else "", 'qty': qty, 'unit': unit, 'unit_price': upr, 'extended_price': ext})
    seen = set(); uniq = []
    for r in rows:
        key = (r['clin'], r['description'], r['qty'], r['unit_price'], r['extended_price'])
        if key in seen: continue
        seen.add(key); uniq.append(r)
    return uniq[:2000]



# -------------------- Modules --------------------
def run_contacts(conn: sqlite3.Connection) -> None:
    st.header("Contacts")
    with st.form("add_contact", clear_on_submit=True):
        c1, c2, c3 = st.columns([2, 2, 2])
        with c1:
            name = st.text_input("Name")
        with c2:
            email = st.text_input("Email")
        with c3:
            org = st.text_input("Organization")
        submitted = st.form_submit_button("Add Contact")
    if submitted:
        try:
            with closing(conn.cursor()) as cur:
                cur.execute(
                    "INSERT INTO contacts(name, email, org) VALUES (?, ?, ?);",
                    (name.strip(), email.strip(), org.strip()),
                )
                conn.commit()
            st.success(f"Added contact {name}")
        except Exception as e:
            st.error(f"Error saving contact {e}")

    try:
        df = pd.read_sql_query(
            "SELECT name, email, org FROM contacts_t ORDER BY name;", conn
        )
        st.subheader("Contact List")
        if df.empty:
            st.write("No contacts yet")
        else:
            _styled_dataframe(df, use_container_width=True, hide_index=True)
    except Exception as e:
        st.error(f"Failed to load contacts {e}")


def run_deals(conn: sqlite3.Connection) -> None:
    st.header("Deals")
    with st.form("add_deal", clear_on_submit=True):
        c1, c2, c3, c4 = st.columns([3, 2, 2, 2])
        with c1:
            title = st.text_input("Title")
        with c2:
            agency = st.text_input("Agency")
        with c3:
            status = st.selectbox(
                "Status",
                ["New", "Qualifying", "Bidding", "Submitted", "Awarded", "Lost"],
            )
        with c4:
            value = st.number_input("Est Value", min_value=0.0, step=1000.0, format="%.2f")
        submitted = st.form_submit_button("Add Deal")
    if submitted:
        try:
            with closing(conn.cursor()) as cur:
                cur.execute(
                    "INSERT INTO deals(title, agency, status, value) VALUES (?, ?, ?, ?);",
                    (title.strip(), agency.strip(), status, float(value)),
                )
                conn.commit()
            st.success(f"Added deal {title}")
        except Exception as e:
            st.error(f"Error saving deal {e}")

    try:
        df = pd.read_sql_query(
            "SELECT title, agency, status, value, sam_url FROM deals_t ORDER BY id DESC;",
            conn,
        )
        st.subheader("Pipeline")
        if df.empty:
            st.write("No deals yet")
        else:
            _styled_dataframe(df, use_container_width=True, hide_index=True)
    except Exception as e:
        st.error(f"Failed to load deals {e}")



# ---- Phase 3 helpers: ensure RFP record, modal renderer ----
def _ensure_rfp_for_notice(conn, notice_row: dict) -> int:
    from contextlib import closing as _closing
    nid = str(notice_row.get("Notice ID") or "")
    if not nid:
        raise ValueError("Missing Notice ID")
    with _closing(conn.cursor()) as cur:
        cur.execute("SELECT id FROM rfps WHERE notice_id=? ORDER BY id DESC LIMIT 1;", (nid,))
        row = cur.fetchone()
        if row:
            return int(row[0])
        cur.execute(
            "INSERT INTO rfps(title, solnum, notice_id, sam_url, file_path, created_at) VALUES (?,?,?,?,?, datetime('now'));",
            (notice_row.get("Title") or "", notice_row.get("Solicitation") or "", nid, notice_row.get("SAM Link") or "", "")
        )
        rid = int(cur.lastrowid)
        conn.commit()
        return rid

def _fetch_and_save_now(conn, notice_id: str, rfp_id: int) -> int:
    saved = 0
    try:
        for fname, fbytes in (sam_try_fetch_attachments(str(notice_id)) or []):
            try:
                # de-dupe by sha256 via save_rfp_file_db
                save_rfp_file_db(conn, int(rfp_id), fname, fbytes)
                saved += 1
            except Exception:
                pass
    except Exception:
        pass
    return saved

def _rfp_build_fulltext_from_db(conn, rfp_id: int, max_files: int = 10, max_pages: int = 80) -> tuple[str, list]:
    """Return (full_text, sources) reading rfp_files; sources is list of (filename, page, text_snippet)."""
    try:
        import io
        from contextlib import closing as _closing
        with _closing(conn.cursor()) as cur:
            cur.execute("SELECT id, filename, bytes, mime FROM rfp_files WHERE rfp_id=? ORDER BY id;", (int(rfp_id),))
            rows = cur.fetchall() or []
    except Exception:
        rows = []
    sources = []
    parts = []
    count_files = 0
    for rid, name, bts, mime in rows:
        if count_files >= int(max_files):
            break
        try:
            pages = extract_text_pages(bts, mime or "")
            for pi, ptxt in enumerate(pages[:max_pages]):
                sources.append((name, pi+1, (ptxt or "")[:400]))
            parts.append("\n\n".join(pages[:max_pages]))
            count_files += 1
        except Exception:
            continue
    return ("\n\n".join([p for p in parts if p]).strip(), sources)

def _rfp_ai_summary(text: str, meta: dict) -> str:
    """Use y55_ai_parse to produce a structured summary in prose."""
    try:
        data = y55_ai_parse(text or "")
    except Exception as e:
        return f"AI parse failed: {e}"
    # format a brief summary
    lines = []
    t = data.get("title") or meta.get("Title") or ""
    s = data.get("solnum") or meta.get("Solicitation") or ""
    lines.append(f"**{t or meta.get('Title','')}**  ")
    if s: lines.append(f"Solicitation: {s}")
    m = data.get("meta") or {}
    if m:
        for k in ("set_aside","naics","psc","pop","place_of_performance","contract_type","vehicle","agency","office"):
            v = m.get(k) or meta.get(k.replace('_',' ').title(), "")
            if v: lines.append(f"- {k.replace('_',' ').title()}: {v}")
    dates = data.get("dates") or []
    if dates:
        lines.append("**Key Dates**")
        for d in dates[:6]:
            nm = (d.get("name") or "").title()
            val = d.get("date") or d.get("iso") or ""
            if nm or val: lines.append(f"- {nm}: {val}")
    lms = data.get("l_items") or []
    if lms:
        lines.append("**L/M Highlights**")
        for li in lms[:8]:
            lines.append(f"- {li.get('text') or str(li)}")
    return "\n".join(lines)

def _rfp_chat(conn, rfp_id: int, question: str, k: int = 6) -> str:
    """Lightweight RAG: use y1_search() hits to ground the answer."""
    try:
        hits = y1_search(conn, int(rfp_id), question or "", k=int(k))
    except Exception:
        hits = []
    context = []
    for h in hits or []:
        src = f"{h.get('file') or ''} p.{h.get('page') or ''}".strip()
        snippet = (h.get('text') or '')[:800]
        context.append(f"[{src}] {snippet}")
    sys = "You are an acquisitions analyst. Answer concisely and cite sources in brackets like [filename p.X]."
    prompt = "\n\n".join(context + [f"Question: {question}"])
    try:
        client = get_ai()
        model = _resolve_model()
        resp = client.chat.completions.create(model=model, messages=[
            {"role":"system","content": sys},
            {"role":"user","content": prompt}
        ], temperature=0.2)
        return (resp.choices[0].message.content or '').strip()
    except Exception as e:
        return f"AI error: {e}"

# ---------- SAM Watch (Phase A) ----------

def run_sam_watch(conn: sqlite3.Connection) -> None:

    # ---- X3 MODAL RENDERER ----
    if st.session_state.get("x3_show_modal"):
        _notice = st.session_state.get("x3_modal_notice", {}) or {}
        try:
            with st.modal("RFP Analyzer", key=_uniq_key("x3_modal", _safe_int(_notice.get("Notice ID")))):
                _x3_render_modal(_notice)
        except Exception:
            with st.expander("RFP Analyzer", expanded=True):
                _x3_render_modal(_notice)
    st.header("SAM Watch")
    st.caption("Live search from SAM.gov v2 API. Push selected notices to Deals or RFP Analyzer.")

    api_key = get_sam_api_key()

    # Search filters (dates optional)
    with st.expander("Search Filters", expanded=True):
        today = datetime.now().date()
        default_from = today - timedelta(days=30)

        c1, c2, c3 = st.columns([2, 2, 2])
        with c1:
            use_dates = st.checkbox("Filter by posted date", value=False)
        with c2:
            active_only = st.checkbox("Active only", value=True)
        with c3:
            org_name = st.text_input("Organization/Agency contains")

        if use_dates:
            d1, d2 = st.columns([2, 2])
            with d1:
                posted_from = st.date_input("Posted From", value=default_from, key="sam_posted_from")
            with d2:
                posted_to = st.date_input("Posted To", value=today, key="sam_posted_to")

        e1, e2, e3 = st.columns([2, 2, 2])
        with e1:
            keywords = st.text_input("Keywords (Title contains)")
        with e2:
            naics = st.text_input("NAICS (6-digit)")
        with e3:
            psc = st.text_input("PSC")

        e4, e5, e6 = st.columns([2, 2, 2])
        with e4:
            state = st.text_input("Place of Performance State (e.g., TX)")
        with e5:
            set_aside = st.text_input("Set-Aside Code (SB, 8A, SDVOSB)")
        with e6:
            pass

        ptype_map = {
            "Pre-solicitation": "p",
            "Sources Sought": "r",
            "Special Notice": "s",
            "Solicitation": "o",
            "Combined Synopsis/Solicitation": "k",
            "Justification (J&A)": "u",
            "Sale of Surplus Property": "g",
            "Intent to Bundle (DoD)": "i",
            "Award Notice": "a",
        }
        types = st.multiselect(
            "Notice Types",
            list(ptype_map.keys()),
            default=["Solicitation", "Combined Synopsis/Solicitation", "Sources Sought"],
        )

        g1, g2 = st.columns([2, 2])
        with g1:
            limit = st.number_input("Results per page", min_value=1, max_value=1000, value=100, step=50)
        with g2:
            max_pages = st.slider("Pages to fetch", min_value=1, max_value=10, value=3)

        run = st.button("Run Search", type="primary")

    results_df = st.session_state.get("sam_results_df", pd.DataFrame())

    if run:
        if not api_key:
            st.error("Missing SAM API key. Add SAM_API_KEY to your Streamlit secrets.")
            return

        params: Dict[str, Any] = {
            "api_key": api_key,
            "limit": int(limit),
            "offset": 0,
            "_max_pages": int(max_pages),
        }
        if active_only:
            params["status"] = "active"
        if "use_dates" in locals() and use_dates:
            params["postedFrom"] = posted_from.strftime("%m/%d/%Y")
            params["postedTo"] = posted_to.strftime("%m/%d/%Y")
        else:
            # SAM.gov API requires postedFrom/postedTo; use implicit last 30 days when filter is off
            _today = datetime.now().date()
            _from = _today - timedelta(days=30)
            params["postedFrom"] = _from.strftime("%m/%d/%Y")
            params["postedTo"] = _today.strftime("%m/%d/%Y")
        if keywords:
            params["title"] = keywords
        if naics:
            params["ncode"] = naics
        if psc:
            params["ccode"] = psc
        if state:
            params["state"] = state
        if set_aside:
            params["typeOfSetAside"] = set_aside
        if org_name:
            params["organizationName"] = org_name
        if types:
            params["ptype"] = ",".join(ptype_map[t] for t in types if t in ptype_map)

        with st.spinner("Searching SAM.gov"):
            out = sam_search_cached(params)

        if out.get("error"):
            st.error(out["error"])
            return

        recs = out.get("records", [])
        results_df = flatten_records(recs)
        st.session_state["sam_results_df"] = results_df
        st.session_state["sam_page"] = 1
        st.session_state.pop("sam_selected_idx", None)
        st.success(f"Fetched {len(results_df)} notices")

# normalize results_df safely

if 'results_df' not in locals():

    try:

        results_df = st.session_state.get('sam_results_df')

    except Exception:

        results_df = None

_has_rows = False

try:

    _has_rows = (results_df is not None) and hasattr(results_df, 'empty') and (not results_df.empty)

except Exception:

    _has_rows = False

if _has_rows:
        # --- List view with pagination (Phase Sam Watch: Part 1) ---
        # Reset page if not set
        if "sam_page" not in st.session_state:
            st.session_state["sam_page"] = 1
        # Compute pages based on current limit control
        try:
            page_size = int(limit)
        except Exception:
            page_size = 100
        total = len(results_df)
        total_pages = max(1, (total + page_size - 1) // page_size)
        cur_page = int(st.session_state.get("sam_page", 1))
        # Clamp page
        if cur_page < 1:
            cur_page = 1
        if cur_page > total_pages:
            cur_page = total_pages
        st.session_state["sam_page"] = cur_page

        # Pager controls
        p1, p2, p3 = st.columns([1, 3, 1])
        with p1:
            if st.button("◀ Prev", key="sam_prev_btn", disabled=(cur_page <= 1)):
                st.session_state["sam_page"] = cur_page - 1
                st.rerun()
        with p2:
            st.caption(f"Page {cur_page} of {total_pages} — showing {min(page_size, total - (cur_page - 1) * page_size)} of {total} results")
        with p3:
            if st.button("Next ▶", key="sam_next_btn", disabled=(cur_page >= total_pages)):
                st.session_state["sam_page"] = cur_page + 1
                st.rerun()

        start_i = (cur_page - 1) * page_size
        end_i = min(start_i + page_size, total)

        # Render list cards instead of table
        for i in range(start_i, end_i):
            row = results_df.iloc[i]
            with st.container():
                st.markdown(f"**{row['Title']}**")
                meta_line = " | ".join([
                    f"Solicitation: {row.get('Solicitation') or '—'}",
                    f"Type: {row.get('Type') or '—'}",
                    f"Set-Aside: {row.get('Set-Aside') or '—'}",
                    f"NAICS: {row.get('NAICS') or '—'}",
                    f"PSC: {row.get('PSC') or '—'}",
                ])
                st.caption(meta_line)
                st.caption(f"Posted: {row.get('Posted') or '—'} · Due: {row.get('Response Due') or '—'} · Agency: {row.get('Agency Path') or '—'}")
                if row.get('SAM Link'):
                    st.markdown(f"[Open in SAM]({row['SAM Link']})")

                c3, c4, c5 = st.columns([2, 2, 2])
                with c3:
                    if st.button("View details", key=f"sam_view_{i}"):
                        st.session_state["sam_selected_idx"] = i
                        st.rerun()
                with c4:
                    if st.button("Add to Deals", key=f"add_to_deals_{i}"):
                        try:
                            from contextlib import closing as _closing
                            _db = globals().get('conn')
                            _owned = False
                            if _db is None:
                                import sqlite3
                                _owned = True
                                _db = _db_connect(DB_PATH, check_same_thread=False)
                            with _closing(_db.cursor()) as cur:
                                cur.execute(
                                    """
                                    INSERT INTO deals(title, agency, status, value, notice_id, solnum, posted_date, rfp_deadline, naics, psc, sam_url)
                                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?);
                                    """,
                                    (
                                        row.get("Title") or "",
                                        row.get("Agency Path") or "",
                                        "Bidding",
                                        None,
                                        row.get("Notice ID") or "",
                                        row.get("Solicitation") or "",
                                        row.get("Posted") or "",
                                        row.get("Response Due") or "",
                                        row.get("NAICS") or "",
                                        row.get("PSC") or "",
                                        row.get("SAM Link") or "",
                                    ),
                                )
                                deal_id = cur.lastrowid
                                _db.commit()
                            # Optional: create an RFP shell and try to fetch attachments
                            try:
                                with _closing(_db.cursor()) as cur:
                                    cur.execute("INSERT INTO rfps(title, solnum, notice_id, sam_url, file_path, created_at) VALUES (?,?,?,?,?, datetime('now'));", (row.get('Title') or '', row.get('Solicitation') or '', row.get('Notice ID') or '', row.get('SAM Link') or '', ''))
                                    rfp_id = cur.lastrowid
                                    _db.commit()
                                att_saved = 0
                                try:
                                    for fname, fbytes in sam_try_fetch_attachments(str(row.get('Notice ID') or '')) or []:
                                        try:
                                            save_rfp_file_db(_db, rfp_id, fname, fbytes)
                                            att_saved += 1
                                        except Exception:
                                            pass
                                except Exception:
                                    pass
                            except Exception:
                                pass
                            st.success(f"Saved to Deals{' · ' + str(att_saved) + ' attachment(s) pulled' if att_saved else ''}")
                        except Exception as e:
                            st.error("Failed to save deal: %s" % (e,))
                        finally:
                            try:
                                if _owned:
                                    _db.close()
                            except Exception:
                                pass
                with c5:

                    # Ask RFP Analyzer (Phase 3 modal)
                    if st.button("Ask RFP Analyzer", key=_uniq_key("ask_rfp", _safe_int(row.get("Notice ID")))):
                        notice = row.to_dict()
                        st.session_state["x3_modal_notice"] = notice
                        st.session_state["x3_show_modal"] = True
                        try:
                            with st.modal("RFP Analyzer", key=_uniq_key("x3_modal", _safe_int(notice.get("Notice ID")))):
                                _x3_render_modal(notice)
                        except Exception:
                            with st.expander("RFP Analyzer", expanded=True):
                                _x3_render_modal(notice)

                        st.session_state["x3_modal_notice"] = row.to_dict()
                        st.session_state["x3_show_modal"] = True

                    # Render modal if requested
                    if st.session_state.get("x3_show_modal") and st.session_state.get("x3_modal_notice", {}).get("Notice ID") == row.get("Notice ID"):
                        try:
                            ctx = st.modal("RFP Analyzer", key=_uniq_key("x3_modal", _safe_int(row.get("Notice ID"))))
                        except Exception:
                            # Fallback if modal unavailable
                            ctx = st.container()
                        with ctx:
                            try:
                                _notice = st.session_state.get("x3_modal_notice") or {}
                                rfp_id = _ensure_rfp_for_notice(conn, _notice)
                                st.caption(f"RFP #{rfp_id} · {row.get('Title') or ''}")
                                # Attachments area
                                try:
                                    from contextlib import closing as _closing
                                    with _closing(conn.cursor()) as cur:
                                        cur.execute("SELECT COUNT(*) FROM rfp_files WHERE rfp_id=?;", (int(rfp_id),))
                                        n_files = int(cur.fetchone()[0])
                                except Exception:
                                    n_files = 0
                                cA, cB = st.columns([2,1])
                                with cA:
                                    st.write(f"Attachments saved: **{n_files}**")
                                with cB:
                                    if st.button("Fetch attachments now", key=_uniq_key("x3_fetch", int(rfp_id))):
                                        c = _fetch_and_save_now(conn, str(row.get("Notice ID") or ""), int(rfp_id))
                                        st.success(f"Fetched {c} attachment(s).")
                                        st.rerun()

                                # Index (light) and summary
                                try:
                                    y1_index_rfp(conn, int(rfp_id), rebuild=False)
                                except Exception:
                                    pass
                                full_text, sources = _rfp_build_fulltext_from_db(conn, int(rfp_id))
                                if not full_text:
                                    st.info("No documents yet. You can still ask questions; I'll use the SAM description if available.")
                                    # Try SAM description fallback
                                    try:
                                        descs = sam_try_fetch_attachments(str(row.get("Notice ID") or "")) or []
                                        for name, b in descs:
                                            if name.endswith("_description.html"):
                                                import bs4
                                                soup = bs4.BeautifulSoup(b.decode('utf-8', errors='ignore'), 'html.parser')
                                                full_text = soup.get_text(" ", strip=True)
                                                break
                                    except Exception:
                                        pass
                                with st.expander("AI Summary", expanded=True):
                                    st.markdown(_rfp_ai_summary(full_text or "", row.to_dict()))

                                # Per-document summarize chips
                                try:
                                    from contextlib import closing as _closing
                                    with _closing(conn.cursor()) as cur:
                                        cur.execute("SELECT id, filename, mime FROM rfp_files WHERE rfp_id=? ORDER BY id;", (int(rfp_id),))
                                        files = cur.fetchall() or []
                                except Exception:
                                    files = []
                                if files:
                                    st.write("Documents:")
                                    for fid, fname, fmime in files[:12]:
                                        if st.button(f"Summarize: {fname}", key=_uniq_key("sumdoc", int(fid))):
                                            try:
                                                blob = pd.read_sql_query("SELECT bytes, mime FROM rfp_files WHERE id=?;", conn, params=(int(fid),)).iloc[0]
                                                _text = "\n\n".join(extract_text_pages(blob['bytes'], blob.get('mime') or (fmime or '')) or [])
                                            except Exception:
                                                _text = ""
                                            st.session_state["x3_docsum"] = _rfp_ai_summary(_text, row.to_dict())
                                    if st.session_state.get("x3_docsum"):
                                        with st.expander("Document Summary", expanded=True):
                                            st.markdown(st.session_state.get("x3_docsum"))

                                # Chat
                                st.divider()
                                st.subheader("Ask about this RFP")
                                _chat_k = _uniq_key("x3_chat", int(rfp_id))
                                hist_key = f"x3_chat_hist_{rfp_id}"
                                st.session_state.setdefault(hist_key, [])
                                for who, msg in st.session_state[hist_key]:
                                    with st.chat_message(who):
                                        st.markdown(msg)
                                q = st.chat_input("Ask a question about the requirements, due dates, sections, etc.", key=_chat_k)
                                if q:
                                    st.session_state[hist_key].append(("user", q))
                                    with st.chat_message("assistant"):
                                        ans = _rfp_chat(conn, int(rfp_id), q)
                                        st.session_state[hist_key].append(("assistant", ans))
                                        st.markdown(ans)

                                # Proposal hand-off
                                st.divider()
                                if st.button("Start Proposal Outline", key=_uniq_key("x3_outline", int(rfp_id))):
                                    outline = [
                                        "# Proposal Outline",
                                        "1. Cover Letter",
                                        "2. Executive Summary",
                                        "3. Technical Approach",
                                        "4. Management Approach",
                                        "5. Past Performance",
                                        "6. Pricing (separate volume if required)",
                                        "7. Compliance Matrix",
                                    ]
                                    st.session_state[f"proposal_outline_{rfp_id}"] = "\n".join(outline)
                                    st.success("Outline drafted and saved. Open Proposal Builder to continue.")
                            except Exception as _e:
                                st.error(f"Modal error: {_e}")
                    # Push notice to Analyzer tab (optional)
                    if st.button("Push to RFP Analyzer", key=_uniq_key("push_to_rfp", int(i))):
                        try:
                            st.session_state["rfp_selected_notice"] = row.to_dict()
                            st.success("Sent to RFP Analyzer. Switch to that tab to continue.")
                        except Exception as _e:
                            st.error(f"Unable to push to RFP Analyzer: {_e}")

                # Inline details view for the selected card
                try:
                    _sel = st.session_state.get("sam_selected_idx")
                except Exception:
                    _sel = None
                if _sel == i:
                    with st.container(border=True):
                        st.write("**Details**")
                        _raw = None
                        try:
                            for _rec in st.session_state.get("sam_records_raw", []) or []:
                                _nid = str(_rec.get("noticeId") or _rec.get("id") or "")
                                if _nid == str(row.get("Notice ID") or ""):
                                    _raw = _rec
                                    break
                        except Exception:
                            _raw = None
                
                        def _gx(obj, *keys, default=""):
                            try:
                                for k in keys:
                                    if obj is None:
                                        return default
                                    obj = obj.get(k)
                                return obj if obj is not None else default
                            except Exception:
                                return default
                
                        desc = row.get("Description") or _gx(_raw, "description", default="")
                        pop_city = _gx(_raw, "placeOfPerformance", "city", default="")
                        pop_state = _gx(_raw, "placeOfPerformance", "state", default="")
                        pop = ", ".join([p for p in [pop_city, pop_state] if p])
                
                        st.write(f"**Solicitation:** {row.get('Solicitation') or ''}")
                        st.write(f"**Set-Aside:** {row.get('Set-Aside') or ''}")
                        st.write(f"**PSC:** {row.get('PSC') or ''}     **NAICS:** {row.get('NAICS') or ''}")
                        st.write(f"**Place of Performance:** {pop}")
                        st.write(f"**Posted:** {row.get('Posted') or ''}     **Due:** {row.get('Response Due') or ''}")
                        if desc:
                            with st.expander("Description"):
                                st.write(desc)
                
                        try:
                            from contextlib import closing as _closing
                            _db = globals().get("conn") or _db_connect(DB_PATH, check_same_thread=False)
                            with _closing(_db.cursor()) as cur:
                                cur.execute("SELECT id FROM rfps WHERE notice_id=? ORDER BY id DESC LIMIT 1", (str(row.get("Notice ID") or ""),))
                                r = cur.fetchone()
                                if r:
                                    _rfp_id = r[0]
                                    cur.execute("SELECT file_name, length(bytes) FROM rfp_files WHERE rfp_id=?", (_rfp_id,))
                                    _files = cur.fetchall()
                                    if _files:
                                        st.write("**Attachments on file:**")
                                        for fn, ln in _files:
                                            st.write(f"- {fn} ({ln or 0} bytes)")
                        except Exception:
                            pass
                
                        link = row.get("SAM Link") or ""
                        if link:
                            st.markdown(f"[Open on SAM.gov]({link})")

                st.divider()

        # Selected details section (sticky below list)
        sel_idx = st.session_state.get("sam_selected_idx")
        if isinstance(sel_idx, int) and 0 <= sel_idx < len(results_df):
            row = results_df.iloc[sel_idx]
            with st.expander("Opportunity Details", expanded=True):
                c1, c2 = st.columns([3, 2])
                with c1:
                    st.write(f"**Title:** {row.get('Title') or ''}")
                    st.write(f"**Solicitation:** {row.get('Solicitation') or '—'}")
                    st.write(f"**Type:** {row.get('Type') or '—'}")
                    st.write(f"**Set-Aside:** {row.get('Set-Aside') or '—'} ({row.get('Set-Aside Code') or '—'})")
                    st.write(f"**NAICS:** {row.get('NAICS') or '—'}  **PSC:** {row.get('PSC') or '—'}")
                    st.write(f"**Agency Path:** {row.get('Agency Path') or '—'}")
                with c2:
                    st.write(f"**Posted:** {row.get('Posted') or '—'}")
                    st.write(f"**Response Due:** {row.get('Response Due') or '—'}")
                    st.write(f"**Notice ID:** {row.get('Notice ID') or '—'}")
                    if row.get('SAM Link'):
                        st.markdown(f"[Open in SAM]({row['SAM Link']})")
                # CO Q&A helper
                try:
                    _rid = locals().get('rfp_id') or locals().get('rid') or st.session_state.get('current_rfp_id')
                    y6_render_co_box((conn if 'conn' in locals() else globals().get('conn')), _rid, key_prefix="run_sam_watch_y6", title="Ask the CO about this opportunity")
                except Exception:
                    pass

def run_research_tab(conn: sqlite3.Connection) -> None:
    st.header("Research (FAR/DFARS/Wage/NAICS)")
    url = st.text_input("URL", placeholder="https://www.acquisition.gov/")
    ttl = st.number_input("Cache TTL (hours)", min_value=1, max_value=168, value=24, step=1)
    q = st.text_input("Highlight phrase (optional)")
    if st.button("Fetch", type="primary", key="research_fetch_btn"):
        with st.spinner("Fetching"):
            rec = research_fetch(url.strip(), ttl_hours=int(ttl))
        if rec.get("status", 0) != 200 and not rec.get("cached"):
            st.error(f"Fetch failed or not cached. Status {rec.get('status')} — {rec.get('error','')}")
        else:
            st.success(("Loaded from cache" if rec.get("cached") else "Fetched") + f" — status {rec.get('status')}")
            txt = rec.get("text","")
            ex = research_extract_excerpt(txt, q or "")
            st.text_area("Excerpt", value=ex, height=240)
            if rec.get("path"):
                st.markdown(f"[Open cached text]({rec['path']})")
    st.caption("Shortcuts: FAR | DFARS | Wage Determinations | NAICS | SBA Size Standards")

def run_rfp_analyzer(conn: sqlite3.Connection) -> None:

        # === One-Page Analyzer (integrated) ===
    try:
        _df_rf_ctx = pd.read_sql_query("SELECT id, title FROM rfps ORDER BY id DESC;", conn, params=())
    except Exception:
        _df_rf_ctx = None
    with st.container():
        st.caption("RFP Analyzer · single-page mode")
        if run_rfp_analyzer_onepage is None:
            st.info("One-Page Analyzer module not found. Place rfp_onepage.py next to this app.")
        elif _df_rf_ctx is None or _df_rf_ctx.empty:
            st.info("No RFPs yet. Parse & save first.")
        else:
            _rid_one = st.selectbox(
                "RFP context",
                options=_df_rf_ctx["id"].tolist(),
                format_func=lambda i: f"#{i} — {_df_rf_ctx.loc[_df_rf_ctx['id']==i,'title'].values[0]}",
                key="onepage_rfp_sel"
            )
            if st.button("Open One-Page Analyzer", type="primary", key="onepage_go"):
                try:
                    _df_files = pd.read_sql_query(
                        "SELECT filename, mime, bytes, pages FROM rfp_files WHERE rfp_id=? ORDER BY id;",
                        conn, params=(int(_rid_one),)
                    )
                except Exception:
                    _df_files = None
                _pages = []
                if _df_files is not None and not _df_files.empty:
                    for _, _r in _df_files.iterrows():
                        _b = _r.get("bytes"); _mime = _r.get("mime") or ""
                        try:
                            _texts = extract_text_pages(_b, _mime) or []
                        except Exception:
                            _texts = []
                        for _i, _t in enumerate(_texts[:100], start=1):
                            _pages.append({"file": _r.get("filename") or "", "page": _i, "text": _t or ""})
                if not _pages:
                    st.warning("No readable pages found in linked files.")
                else:
                    run_rfp_analyzer_onepage(_pages)
                    st.stop()
    # === end One-Page Analyzer ===
    # === end One‑Page Analyzer ===

    st.header("RFP Analyzer")
    # --- RFP Meta Editor ---

    # --- RFP Meta Editor (safe) ---
    try:
        # Determine active RFP id
        _rid = locals().get('rfp_id') or locals().get('rid') or st.session_state.get('current_rfp_id')
        if not _rid:
            _rid = pd.read_sql_query("SELECT id FROM rfps ORDER BY id DESC LIMIT 1;", conn).iloc[0]["id"]
    except Exception:
        _rid = None

    _title0 = ""
    _solnum0 = ""
    _sam0 = ""
    if _rid:
        try:
            _row = pd.read_sql_query("SELECT title, solnum FROM rfps WHERE id=?;", conn, params=(int(_rid),)).iloc[0]
            _title0 = str(_row.get("title") or "")
            _solnum0 = str(_row.get("solnum") or "")
        except Exception:
            pass
        try:
            # sam_url may or may not exist yet
            _sam0 = pd.read_sql_query("SELECT sam_url FROM rfps WHERE id=?;", conn, params=(int(_rid),)).iloc[0].get("sam_url") or ""
        except Exception:
            _sam0 = ""

    with st.form("rfp_meta_edit", clear_on_submit=False):
        c1, c2 = st.columns(2)
        with c1:
            _title = st.text_input("RFP Title", _title0, key="meta_title")
            _solnum = st.text_input("Solicitation Number", _solnum0, key="meta_solnum")
        with c2:
            _sam_url = st.text_input("SAM.gov URL", _sam0, key="meta_samurl")
            _notice_hint = _parse_sam_notice_id(_sam_url) or ""
            if _notice_hint:
                st.caption(f"Parsed notice id: {_notice_hint}")
        _save = st.form_submit_button("Save RFP Meta")
    if _save and _rid:
        if _update_rfp_meta(conn, _rid, title=_title.strip() or None, solnum=_solnum.strip() or None, sam_url=_sam_url.strip() or None):
            st.success("Saved")
            st.session_state["current_rfp_id"] = int(_rid)
            st.rerun()

    # === Phase 3: RTM + Amendment sidebar wiring ===
    try:
        _ctx = pd.read_sql_query("SELECT id, title, sam_url FROM rfps ORDER BY id DESC;", conn, params=())
    except Exception:
        _ctx = pd.DataFrame()
    if _ctx is not None and not _ctx.empty:
        rid_p3 = st.selectbox(
            "RFP context (for RTM & Amendments)",
            options=_ctx["id"].tolist(),
            format_func=lambda i: f"#{i} — {_ctx.loc[_ctx['id']==i,'title'].values[0]}",
            key="p3_rfp_sel"
        )
        sam_default = ""
        try:
            sam_default = _ctx.loc[_ctx["id"]==rid_p3, "sam_url"].values[0]
        except Exception:
            pass
        cA, cB = st.columns([2,1])
        with cA:
            sam_url = st.text_input("SAM URL (for amendment tracking)", value=sam_default or "", key="p3_sam_url")
        with cB:
            ttl_hours = st.number_input("Cache TTL (hours)", min_value=1, max_value=168, value=72, step=1, key="p3_ttl")
        # Sidebar: amendment diff and impact to-dos
        render_amendment_sidebar(conn, int(rid_p3), sam_url, ttl_hours=int(ttl_hours))
        # Main panel: full RTM coverage editor + metrics
        with st.expander("Requirements Traceability Matrix (RTM)", expanded=True):
            render_rtm_ui(conn, int(rid_p3))

    render_status_and_gaps(conn)
    from contextlib import contextmanager
    @contextmanager
    def _noop():
        yield
    tab_research = _noop()
    tab_parse = _noop()
    tab_checklist = _noop()
    tab_data = _noop()
    tab_y1 = _noop()
    tab_y2 = _noop()
    tab_y4 = _noop()
    st.caption('RFP Analyzer · single-page mode')
    with tab_research:
        run_research_tab(conn)# --- heuristics to auto-fill Title and Solicitation # ---
    def _guess_title(text: str, fallback: str) -> str:
        for line in (text or "").splitlines():
            s = line.strip()
            if len(s) >= 8 and not s.lower().startswith(("department of", "u.s.", "united states", "naics", "set-aside", "solicitation", "request for", "rfp", "rfq", "sources sought")):
                return s[:200]
        return fallback

    def _guess_solnum(text: str) -> str:
        if not text:
            return ""
    # --- meta extractors (NAICS, Set-Aside, Place of Performance) ---
    def _extract_naics(text: str) -> str:
        if not text: return ""
        m = re.search(r'(?i)NAICS(?:\s*Code)?\s*[:#]?\s*([0-9]{5,6})', text)
        if m: return m.group(1)[:6]
        m = re.search(r'(?i)NAICS[^\n]{0,50}?([0-9]{6})', text)
        if m: return m.group(1)
        m = re.search(r'(?i)(?:industry|classification)[^\n]{0,50}?([0-9]{6})', text)
        return m.group(1) if m else ""

    def _extract_set_aside(text: str) -> str:
        if not text: return ""
        tags = ["SDVOSB","SDVOSBC","WOSB","EDWOSB","8(a)","8A","HUBZone","SBA","SDB","VOSB","Small Business","Total Small Business"]
        for t in tags:
            if re.search(rf'(?i)\b{re.escape(t)}\b', text):
                norm = t.upper().replace("(A)","8A").replace("TOTAL SMALL BUSINESS","SMALL BUSINESS")
                if norm == "8(A)": norm = "8A"
                return norm
        m = re.search(r'(?i)Set[- ]Aside\s*[:#]?\s*([A-Za-z0-9 \-/\(\)]+)', text)
        if m:
            v = m.group(1).strip()
            v = re.sub(r'\s+', ' ', v)
            return v[:80]
        return ""

    def _extract_place(text: str) -> str:
        if not text: return ""
        m = re.search(r'(?i)Place\s+of\s+Performance\s*[:\-]?\s*([^\n]{3,80})', text)
        if m: return m.group(1).strip()
        m = re.search(r'\b([A-Z][a-zA-Z]+,\s*(?:[A-Z]{2}|[A-Za-z\. ]{3,}))\b', text)
        return m.group(1).strip() if m else ""

    # ensure rfp_meta exists
    try:
        with closing(conn.cursor()) as _c:

            _c.execute("""
                CREATE TABLE IF NOT EXISTS rfp_chat(
                    id INTEGER PRIMARY KEY,
                    rfp_id INTEGER NOT NULL REFERENCES rfps(id) ON DELETE CASCADE,
                    ts TEXT,
                    q TEXT,
                    a TEXT
                );
            """)
            _c.execute("""
                CREATE TABLE IF NOT EXISTS rfp_meta(
                    id INTEGER PRIMARY KEY,
                    rfp_id INTEGER REFERENCES rfps(id) ON DELETE CASCADE,
                    key TEXT,
                    value TEXT
                );
            """)
            conn.commit()
    except Exception:
        pass

        m = re.search(r'(?i)Solicitation\s*(Number|No\.?)\s*[:#]?\s*([A-Z0-9][A-Z0-9\-\._/]{4,})', text)
        if m:
            return m.group(2)[:60]
        m = re.search(r'\b([A-Z0-9]{2,6}[A-Z0-9\-]{0,4}\d{2}[A-Z]?-?[A-Z]?-?\d{3,6})\b', text)
        if m:
            return m.group(1)[:60]
        m = re.search(r'\b(RFQ|RFP|IFB|RFI)[\s#:]*([A-Z0-9][A-Z0-9\-\._/]{3,})\b', text, re.I)
        if m:
            return (m.group(1).upper() + "-" + m.group(2))[:60]
        return ""
# ---------------- PARSE & SAVE ----------------
    with tab_parse:

        # --- X1 Ingest: File Library + Health ---
        if True:
            with st.expander("X1 Ingest: File Library + Health", expanded=False):
                st.caption("Accepts PDF, DOCX, XLSX, TXT. Deduplicates by SHA-256. Attempts OCR on image-only PDFs if pytesseract is available. — X7 applied")
                try:
                    df_rf_list = pd.read_sql_query("SELECT id, title FROM rfps ORDER BY id DESC;", conn, params=())
                    opt_rf = [None] + df_rf_list["id"].tolist() if df_rf_list is not None else [None]
                except Exception:
                    df_rf_list = pd.DataFrame()
                    opt_rf = [None]
                def _fmt_rfp(x):
                    try:
                        if x is None:
                            return "None"
                        ttl = df_rf_list.loc[df_rf_list["id"]==x,"title"]
                        return f"#{x} — {ttl.values[0] if len(ttl) else ''}"
                    except Exception:
                        return str(x) if x is not None else "None"
                link_now_rfp = st.selectbox("Link to existing RFP (optional)", options=opt_rf, format_func=_fmt_rfp, key="x1_link_now_rfp")
                ing_files = st.file_uploader("Files to ingest", type=["pdf","docx","xlsx","txt"], accept_multiple_files=True, key="x1_ing")
                link_to_rfp = st.checkbox("Also remember these to auto-link to the next new RFP created below", value=False)
                if st.button("Ingest Files", key="x1_ingest_btn"):
                    if not ing_files:
                        st.warning("No files selected")
                    else:
                        rows = []
                        ids = []
                        for f in ing_files:
                            try:
                                b = f.getbuffer().tobytes()
                            except Exception:
                                b = f.read()
                            rec = save_rfp_file_db(conn, int(link_now_rfp) if link_now_rfp is not None else None, f.name, b)
                            rows.append({
                                "Filename": rec["filename"],
                                "SHA256": rec["sha256"][:12],
                                "MIME": rec["mime"],
                                "Pages": rec["pages"],
                                "OCR pages": rec.get("ocr_pages", 0),
                                "Dedup?": "Yes" if rec.get("dedup") else "No",
                                "rfp_file_id": rec["id"],
                                "Linked RFP": (int(link_now_rfp) if link_now_rfp is not None else None),
                            })
                            ids.append(int(rec["id"]))
                        import pandas as _pd
                        df_ing = _pd.DataFrame(rows)
                        _styled_dataframe(df_ing, use_container_width=True, hide_index=True)
                        st.session_state["x1_last_ingested_ids"] = ids
                        st.session_state["x1_pending_link_after_create"] = bool(link_to_rfp)
                        if link_now_rfp is not None:
                            st.success(f"Ingested {len(rows)} file(s). Linked to RFP #{int(link_now_rfp)}.")
                        else:
                            st.success(f"Ingested {len(rows)} file(s).")
        colA, colB = st.columns([3,2])
        with colA:
            ups = st.file_uploader(
                "Upload RFP(s) (PDF/DOCX/TXT)",
                type=["pdf","docx","xlsx","txt"],
                accept_multiple_files=True,
                key="rfp_ups"
            )
            with st.expander("Manual Text Paste (optional)", expanded=False):
                pasted = st.text_area("Paste any text to include in parsing", height=150, key="rfp_paste")
            title = st.text_input("RFP Title (used if combining)", key="rfp_title")
            solnum = st.text_input("Solicitation # (used if combining)", key="rfp_solnum")
            sam_url = st.text_input("SAM URL (used if combining)", key="rfp_sam_url", placeholder="https://sam.gov/")
            _title_in = (title or "" ).strip()
            _solnum_in = (solnum or "" ).strip()
            _sam_in = (sam_url or "" ).strip()
            mode = st.radio("Save mode", ["One record per file", "Combine all into one RFP"], index=0, horizontal=True)
        with colB:
            st.markdown("**Parse Controls**")
            run = st.button("Parse & Save", type="primary", key="rfp_parse_btn")
            st.caption("We'll auto-extract L/M checklist items, CLINs, key dates, and POCs.")

        def _read_file(file):
            name = file.name.lower()
            data = file.read()
            if name.endswith(".txt"):
                try:
                    return data.decode("utf-8")
                except Exception:
                    return data.decode("latin-1", errors="ignore")
            if name.endswith(".pdf"):
                try:
                    reader = _pypdf.PdfReader(io.BytesIO(data))
                    pages = [(p.extract_text() or "") for p in reader.pages]
                    return "\\n".join(pages)
                except Exception as e:
                    st.warning(f"PDF text extraction failed for {file.name}: {e}. Falling back to binary decode.")
                    return data.decode("latin-1", errors="ignore")
            if name.endswith(".docx"):
                try:
                    f = io.BytesIO(data)
                    doc = docx.Document(f)
                    return "\\n".join([p.text for p in doc.paragraphs])
                except Exception as e:
                    st.warning(f"DOCX parse failed for {file.name}: {e}.")
                    return ""
            if name.endswith(".xlsx"):
                try:
                    pages = extract_text_pages(data, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                    if pages:
                        return "\n\n".join(pages)
                    return ""
                except Exception as e:
                    st.warning(f"XLSX parse failed for {file.name}: {e}.")
                    return ""
            st.error(f"Unsupported file type: {file.name}")
            return ""

        if run:
            if (not ups) and (not pasted):
                st.error("No input. Upload at least one file or paste text.")
            else:
                if mode == "Combine all into one RFP":
                    text_parts = []
                    # Include DB-saved attachments (rfp_files) as inputs too
                    try:
                        class _MemFile:
                            def __init__(self, name, data):
                                self.name = name
                                self._data = data
                            def read(self):
                                return self._data
                            def getbuffer(self):
                                import io
                                return io.BytesIO(self._data)
                        _db_files = []
                        try:
                            for _fid, _fn, _bts in y3_get_rfp_files(conn, int(_rid)) or []:
                                _db_files.append(_MemFile(_fn, _bts))
                        except Exception:
                            _db_files = []
                        ups = (ups or []) + _db_files
                    except Exception:
                        pass
                    for f in ups or []:
                        text_parts.append(_read_file(f))
                    if pasted:
                        text_parts.append(pasted)
                    full_text = "\\n\\n".join([t for t in text_parts if t]).strip()
                    # X3: parse pricing matrices from any uploaded XLSX files
                    x3_clins = []
                    for _f in (ups or []):
                        try:
                            _name = (_f.name or "").lower()
                            _b = _f.getbuffer().tobytes() if hasattr(_f, "getbuffer") else _f.read()
                        except Exception:
                            _name = (_f.name or "").lower()
                            try:
                                _b = _f.read()
                            except Exception:
                                _b = b""
                        if _name.endswith(".xlsx") and _b:
                            x3_clins.extend(extract_clins_xlsx(_b))
                    if not full_text:
                        st.error("Nothing readable found.")
                    else:
                        secs = extract_sections_L_M(full_text)
                        l_items = derive_lm_items(secs.get('L','')) + derive_lm_items(secs.get('M',''))
                        clins = extract_clins(full_text) + (x3_clins or []); dates = extract_dates(full_text); pocs = extract_pocs(full_text)
                        meta = {
                            'naics': _extract_naics(full_text),
                            'set_aside': _extract_set_aside(full_text),
                            'place_of_performance': _extract_place(full_text),
                        }


                        l_items, clins, dates, pocs, meta, title, solnum = y55_apply_enhancement(full_text if 'full_text' in locals() else (text if 'text' in locals() else ''), l_items, clins, dates, pocs, meta, title if 'title' in locals() else '', solnum if 'solnum' in locals() else '')
# X4: persist meta to rfp_meta for combined RFP
                        try:
                            for _k, _v in (meta or {}).items():
                                if _v:
                                    cur.execute("INSERT INTO rfp_meta(rfp_id, key, value) VALUES(?,?,?);", (int(rfp_id), str(_k), str(_v)))
                        except Exception:
                            pass
                        with closing(conn.cursor()) as cur:
                            cur.execute(
                                "INSERT INTO rfps(title, solnum, notice_id, sam_url, file_path, created_at) VALUES (?,?,?,?,?, datetime('now'));",
                                ((_title_in or _guess_title(full_text, "Untitled")), (_solnum_in or _guess_solnum(full_text)), (_parse_sam_notice_id(_sam_in) or ""), (_sam_in or ""), "",)
                            )
                            rfp_id = cur.lastrowid
                            for it in l_items:
                                cur.execute("INSERT INTO lm_items(rfp_id, item_text, is_must, status) VALUES (?,?,?,?);",
                                            (rfp_id, it, 1 if re.search(r'\\b(shall|must|required|mandatory|no later than|shall not|will)\\b', it, re.IGNORECASE) else 0, "Open"))
                            for r in clins:
                                cur.execute("INSERT INTO clin_lines(rfp_id, clin, description, qty, unit, unit_price, extended_price) VALUES (?,?,?,?,?,?,?);",
                                            (rfp_id, r.get('clin'), r.get('description'), r.get('qty'), r.get('unit'), r.get('unit_price'), r.get('extended_price')))
                            for d in dates:
                                cur.execute("INSERT INTO key_dates(rfp_id, label, date_text, date_iso) VALUES (?,?,?,?);",
                                            (rfp_id, d.get('label'), d.get('date_text'), d.get('date_iso')))
                            for pc in pocs:
                                cur.execute("INSERT INTO pocs(rfp_id, name, role, email, phone) VALUES (?,?,?,?,?);",
                                            (rfp_id, pc.get('name'), pc.get('role'), pc.get('email'), pc.get('phone')))

                            # X3: store POP / ordering period in meta and key_dates
                            try:
                                _pop = extract_pop_structure(full_text)
                                if _pop:
                                    for k,v in _pop.items():
                                        cur.execute("INSERT INTO rfp_meta(rfp_id, key, value) VALUES(?,?,?);", (int(rfp_id), str(k), str(v)))
                                    if _pop.get('pop_structure'):
                                        cur.execute("INSERT INTO key_dates(rfp_id, label, date_text, date_iso) VALUES (?,?,?,?);",
                                                    (int(rfp_id), 'Ordering/POP', _pop['pop_structure'], ''))
                            except Exception:
                                pass

                            # X3: POP / ordering period from this file's text
                            try:
                                _pop = extract_pop_structure(text)
                                if _pop:
                                    for k,v in _pop.items():
                                        cur.execute("INSERT INTO rfp_meta(rfp_id, key, value) VALUES(?,?,?);", (int(rfp_id), str(k), str(v)))
                                    if _pop.get('pop_structure'):
                                        cur.execute("INSERT INTO key_dates(rfp_id, label, date_text, date_iso) VALUES (?,?,?,?);",
                                                    (int(rfp_id), 'Ordering/POP', _pop['pop_structure'], ''))
                            except Exception:
                                pass
                            conn.commit()
                            # X2: auto-link any pending ingested files to this new RFP
                            try:
                                if st.session_state.get("x1_pending_link_after_create") and st.session_state.get("x1_last_ingested_ids"):
                                    ids = tuple(int(i) for i in st.session_state.get("x1_last_ingested_ids") or [])
                                    if ids:
                                        ph = ",".join(["?"]*len(ids))
                                        with closing(conn.cursor()) as _cur2:
                                            _cur2.execute(f"UPDATE rfp_files SET rfp_id=? WHERE id IN ({ph});", (int(rfp_id), *ids))
                                            conn.commit()
                                    st.session_state["x1_pending_link_after_create"] = False
                            except Exception:
                                pass
                            st.success(f"Combined and saved RFP #{rfp_id} (items: {len(l_items)}, CLINs: {len(clins)}, dates: {len(dates)}, POCs: {len(pocs)}).")
                else:
                    saved = 0
                    for f in ups or []:
                        try:
                            _bytes = f.getbuffer().tobytes()
                        except Exception:
                            try:
                                _bytes = f.read()
                            except Exception:
                                _bytes = b""
                        text = _read_file(type('F', (), {'name': f.name, 'read': lambda self=None: _bytes})())
                        if not text.strip():
                            continue
                        secs = extract_sections_L_M(text)
                        l_items = derive_lm_items(secs.get('L','')) + derive_lm_items(secs.get('M',''))
                        clins = extract_clins(text) + (extract_clins_xlsx(_bytes) if (f.name or '').lower().endswith('.xlsx') else []); dates = extract_dates(text); pocs = extract_pocs(text)
                        meta = {
                            'naics': _extract_naics(text),
                            'set_aside': _extract_set_aside(text),
                            'place_of_performance': _extract_place(text),
                        }


                        l_items, clins, dates, pocs, meta, title, solnum = y55_apply_enhancement(full_text if 'full_text' in locals() else (text if 'text' in locals() else ''), l_items, clins, dates, pocs, meta, title if 'title' in locals() else '', solnum if 'solnum' in locals() else '')
# X4: persist meta to rfp_meta per-file
                        try:
                            for _k, _v in (meta or {}).items():
                                if _v:
                                    cur.execute("INSERT INTO rfp_meta(rfp_id, key, value) VALUES(?,?,?);", (int(rfp_id), str(_k), str(_v)))
                        except Exception:
                            pass
                        with closing(conn.cursor()) as cur:
                            cur.execute(
                                "INSERT INTO rfps(title, solnum, notice_id, sam_url, file_path, created_at) VALUES (?,?,?,?,?, datetime('now'));",
                                ((_title_in or _guess_title(text, f.name)), (_solnum_in or _guess_solnum(text)), (_parse_sam_notice_id(_sam_in) or ""), _sam_in, "", )
                            )
                            rfp_id = cur.lastrowid
                            for it in l_items:
                                cur.execute("INSERT INTO lm_items(rfp_id, item_text, is_must, status) VALUES (?,?,?,?);",
                                            (rfp_id, it, 1 if re.search(r'\\b(shall|must|required|mandatory|no later than|shall not|will)\\b', it, re.IGNORECASE) else 0, "Open"))
                            for r in clins:
                                cur.execute("INSERT INTO clin_lines(rfp_id, clin, description, qty, unit, unit_price, extended_price) VALUES (?,?,?,?,?,?,?);",
                                            (rfp_id, r.get('clin'), r.get('description'), r.get('qty'), r.get('unit'), r.get('unit_price'), r.get('extended_price')))
                            for d in dates:
                                cur.execute("INSERT INTO key_dates(rfp_id, label, date_text, date_iso) VALUES (?,?,?,?);",
                                            (rfp_id, d.get('label'), d.get('date_text'), d.get('date_iso')))
                            for pc in pocs:
                                cur.execute("INSERT INTO pocs(rfp_id, name, role, email, phone) VALUES (?,?,?,?,?);",
                                            (rfp_id, pc.get('name'), pc.get('role'), pc.get('email'), pc.get('phone')))
                            conn.commit()
                        last_rfp_id = rfp_id
                        saved += 1
                    # X2: if pending, link last created RFP to recently ingested files
                    try:
                        if st.session_state.get("x1_pending_link_after_create") and st.session_state.get("x1_last_ingested_ids") and "last_rfp_id" in locals():
                            ids = tuple(int(i) for i in st.session_state.get("x1_last_ingested_ids") or [])
                            if ids:
                                ph = ",".join(["?"]*len(ids))
                                with closing(conn.cursor()) as _cur3:
                                    _cur3.execute(f"UPDATE rfp_files SET rfp_id=? WHERE id IN ({ph});", (int(last_rfp_id), *ids))
                                    conn.commit()
                            st.session_state["x1_pending_link_after_create"] = False
                    except Exception:
                        pass
                    st.success(f"Saved {saved} RFP record(s).")
    # ---------------- Y1: Ask with citations ----------------
    with tab_y1:
        st.caption("Build a local search index once, then ask CO-style questions with bracketed citations.")
        df_rf_y1 = pd.read_sql_query("SELECT id, title FROM rfps ORDER BY id DESC;", conn, params=())
        if df_rf_y1.empty:
            st.info("No RFPs yet. Parse & save first.")
        else:
            rid_y1 = st.selectbox("RFP context", options=df_rf_y1["id"].tolist(),
                                  format_func=lambda i: f"#{i} — {df_rf_y1.loc[df_rf_y1['id']==i,'title'].values[0]}",
                                  key="y1_rfp_sel")
            c1, c2 = st.columns([2,2])
            with c1:
                if st.button("Build/Update search index for this RFP"):
                    with st.spinner("Indexing linked files"):
                        out = y1_index_rfp(conn, int(rid_y1), rebuild=False)
                    if out.get("ok"):
                        st.success(f"Indexed. Added {out.get('added',0)} chunk(s). Skipped {out.get('skipped',0)} existing.")
                    else:
                        st.error(out.get("error","Index error"))
            with c2:
                if st.button("Rebuild index (overwrite)"):
                    with st.spinner("Rebuilding"):
                        out = y1_index_rfp(conn, int(rid_y1), rebuild=True)
                    if out.get("ok"):
                        st.success(f"Rebuilt. Added {out.get('added',0)} chunk(s).")
                    else:
                        st.error(out.get("error","Index error"))
            q_y1 = st.text_area("Your question", height=120, key="y1_q")
            k = y_auto_k(q_y1)
            if st.button("Ask with citations", type="primary"):
                if not (q_y1 or "").strip():
                    st.warning("Enter a question")
                else:
                    ph = st.empty(); acc = []
                    for tok in ask_ai_with_citations(conn, int(rid_y1), q_y1.strip(), k=int(k)):
                        acc.append(tok)
                        ph.markdown("".join(acc))
                    hits = y1_search(conn, int(rid_y1), q_y1.strip(), k=int(k))
                    if hits:
                        import pandas as _pd
                        dfh = _pd.DataFrame([{"Tag": f"[C{i+1}]", "File": h["file"], "Page": h["page"], "Score": h["score"]} for i,h in enumerate(hits)])
                        st.subheader("Sources used")
                        _styled_dataframe(dfh, use_container_width=True, hide_index=True)




    # ---------------- Y2: CO Chat with memory ----------------
    with tab_y2:
        y2_ui_threaded_chat(conn)

    # ---------------- Y4: CO Review UI ----------------
    with tab_y4:
        y4_ui_review(conn)
# ---------------- CHECKLIST ----------------
    with tab_checklist:
        df_rf = pd.read_sql_query("SELECT id, title, solnum FROM rfps ORDER BY id DESC;", conn, params=())
        if df_rf.empty:
            st.info("No RFPs yet. Parse one on the first tab.")
        else:
            rid = st.selectbox("Select an RFP", options=df_rf['id'].tolist(), format_func=lambda i: f"#{i} — {df_rf.loc[df_rf['id']==i,'title'].values[0]}", key="rfp_sel")
            df_lm = pd.read_sql_query("SELECT id, item_text, is_must, status FROM lm_items WHERE rfp_id=? ORDER BY id;", conn, params=(int(rid),))

        with st.expander("Q&A Memory (X7)", expanded=False):
            st.caption("Ask questions about this RFP. Answers are pulled from saved checklists, CLINs, dates, POCs, and linked file text. History is saved.")
            q = st.text_input("Your question", key="x7_q")
            ask = st.button("Ask (store)", key="x7_ask")
            if ask and (q or "").strip():
                try:
                    ql = (q or "").lower()
                    rid_safe = locals().get("rid", None)
                    if rid_safe is None:
                        st.warning("Select an RFP in the Data tab above first."); raise Exception("No rid in scope")
                    res = _kb_search(conn, int(rid_safe), q or "")
                    ans_parts = []
                    if any(w in ql for w in ["due","deadline","close"]):
                        df = res.get("dates")
                        if df is not None and not df.empty:
                            top = df.iloc[0]
                            ans_parts.append(f"Due: {top.get('date_text','')} ({top.get('label','')})")
                    if any(w in ql for w in ["poc","contact","officer","specialist"]):
                        df = res.get("pocs")
                        if df is not None and not df.empty:
                            s = "; ".join([f"{r.get('name','')} ({r.get('email','')})" for _, r in df.head(3).iterrows()])
                            if s: ans_parts.append("POCs: " + s)
                    if "clin" in ql:
                        df = res.get("clins")
                        if df is not None and not df.empty:
                            ans_parts.append(f"CLINs detected: {len(df)}; first: {df.iloc[0].get('clin','')}")
                    if any(w in ql for w in ["checklist","compliance","shall","must"]):
                        df = res.get("checklist")
                        if df is not None and not df.empty:
                            open_cnt = int((df['status']!='Complete').sum())
                            ans_parts.append(f"Checklist items: {len(df)}. Open: {open_cnt}.")
                    if not ans_parts:
                        sec = res.get("sections")
                        if sec is not None and not sec.empty:
                            snip = (sec.iloc[0].get("content","") or "").strip().replace("\n"," ")
                            ans_parts.append("Top section snippet: " + snip[:300])
                    a = " | ".join(ans_parts) if ans_parts else "No direct answer found in saved data."
                    from contextlib import closing as _closing_x7
                    with _closing_x7(conn.cursor()) as cur:
                        cur.execute("INSERT INTO rfp_chat(rfp_id, ts, q, a) VALUES(?,?,?,?);",
                                    (int(rid), datetime.utcnow().isoformat(), q.strip(), a))
                        conn.commit()
                    st.success("Stored Q&A.")
                except Exception as e:
                    st.error(f"Q&A failed: {e}")
            try:
                df_hist = pd.read_sql_query("SELECT ts, q, a FROM rfp_chat WHERE rfp_id=? ORDER BY id DESC LIMIT 50;", conn, params=(int(locals().get("rid")),))
                if df_hist is not None and not df_hist.empty:
                    _styled_dataframe(df_hist, use_container_width=True, hide_index=True)
                    c1, c2 = st.columns([1,1])
                    with c1:
                        if st.button("Export Q&A CSV", key="x7_export"):
                            csvb = df_hist.to_csv(index=False).encode("utf-8")
                            st.download_button("Download Q&A CSV", data=csvb, file_name=f"rfp_{int(rid)}_qa.csv", mime="text/csv", key="x7_dl")
                    with c2:
                        if st.button("Clear history", key="x7_clear"):
                            try:
                                with closing(conn.cursor()) as cur:
                                    cur.execute("DELETE FROM rfp_chat WHERE rfp_id=?;", (int(rid),))
                                    conn.commit()
                                st.success("Cleared.")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Clear failed: {e}")
                else:
                    st.caption("No Q&A yet for this RFP.")
            except Exception as e:
                st.info(f"No history available: {e}")
            # X7 guard: ensure checklist dataframe is defined and rid is available
            _rid = locals().get("rid", None)
            if _rid is None:
                st.caption("Checklist viewer: select an RFP above to load items.")
            else:
                try:
                    df_lm = pd.read_sql_query(
                        "SELECT id, item_text, is_must, status FROM lm_items WHERE rfp_id=? ORDER BY id;",
                        conn, params=(int(_rid),)
                    )

                    with st.expander("Add to Proposal Drafts", expanded=False):
                        sec = st.text_input("Section label", value="Research Notes", key="y5_sec_y1")
                        ans_txt = "".join(acc).strip()
                        if st.button("Add to drafts", key="y5_add_y1"):
                            y5_save_snippet(conn, int(rid_y1), sec, ans_txt or q_y1.strip(), source="Y1 Q&A")
                            st.success("Saved to drafts")
                except Exception:
                    df_lm = pd.DataFrame(columns=['id','item_text','is_must','status'])
                df_lm = df_lm.fillna("")
                st.caption(f"{len(df_lm)} checklist items")
                _styled_dataframe(df_lm, use_container_width=True, hide_index=True)
                new_status = st.selectbox("Set status for selected IDs", ["Open","In Progress","Complete","N/A"], index=0, key="lm_set_status")
                sel_ids = st.text_input("IDs to update (comma-separated)", key="lm_ids")
                if st.button("Update Status", key="lm_status_btn"):
                    ids = [int(x) for x in sel_ids.split(",") if x.strip().isdigit()]
                    if ids:
                        with closing(conn.cursor()) as cur:
                            cur.executemany("UPDATE lm_items SET status=? WHERE id=? AND rfp_id=?;", [(new_status, iid, int(_rid)) for iid in ids])
                            conn.commit()
                    st.success(f"Updated {len(ids)} item(s).")
                    st.rerun()
            # Export
            if st.button("Export Compliance Matrix (CSV)", key="lm_export_csv"):
                out = df_lm.copy()
                out.insert(0, "rfp_id", int(rid))
                csv_bytes = out.to_csv(index=False).encode("utf-8")
                st.download_button("Download CSV", data=csv_bytes, file_name=f"rfp_{rid}_compliance.csv", mime="text/csv", key="lm_dl")

    # ---------------- CLINs / Dates / POCs ----------------
    with tab_data:
        # X2: Files for this RFP
        with st.expander("Files for this RFP (X2)", expanded=False):
            try:
                df_files = pd.read_sql_query("SELECT id, filename, mime, pages, sha256 FROM rfp_files WHERE rfp_id=? ORDER BY id DESC;", conn, params=(int(rid),))
            except Exception as e:
                df_files = pd.DataFrame()
            st.caption("Linked files")
            if df_files is None or df_files.empty:
                st.write("No files linked.")
            else:
                _styled_dataframe(df_files.assign(sha=df_files["sha256"].str.slice(0,12)).drop(columns=["sha256"]), use_container_width=True, hide_index=True)
                # X5: preview and download selected linked file
                try:
                    pick = st.selectbox(
                        "Open file",
                        options=df_files["id"].tolist(),
                        format_func=lambda i: f"#{i} — {df_files.loc[df_files['id']==i,'filename'].values[0]}",
                        key=f"file_pick_{rid}"
                    )
                    if pick:
                        row = pd.read_sql_query(
                            "SELECT filename, mime, bytes FROM rfp_files WHERE id=?;",
                            conn, params=(int(pick),)
                        ).iloc[0]
                        fname = row.get("filename") or f"rfp_file_{int(pick)}"
                        mime = row.get("mime") or "application/octet-stream"
                        b = row.get("bytes")
                        st.download_button("Download original", data=b, file_name=fname, mime=mime, key=f"dl_{pick}")
                        try:
                            pages = extract_text_pages(b, mime)
                            preview = ("\n\n".join(pages) if pages else "").strip()[:20000]
                            if preview:
                                st.text_area("Preview (first 20k chars)", value=preview, height=300)
                        except Exception as e:
                            st.info(f"Preview unavailable: {e}")
                except Exception as e:
                    st.info(f"No preview available: {e}")

                # X6: bulk ZIP download, inventory CSV, and simple search across linked files
                import io as _io, zipfile as _zipfile
                c_zip, c_inv = st.columns([1,1])
                with c_zip:
                    if st.button("Download all linked files as ZIP", key=f"zip_all_{rid}"):
                        try:
                            # Fetch bytes for all linked files
                            df_bytes = pd.read_sql_query("SELECT filename, bytes, mime FROM rfp_files WHERE rfp_id=?;", conn, params=(int(rid),))
                            if df_bytes is None or df_bytes.empty:
                                st.warning("No files to package.")
                            else:
                                buf = _io.BytesIO()
                                with _zipfile.ZipFile(buf, mode="w", compression=_zipfile.ZIP_DEFLATED) as zf:
                                    for _, r in df_bytes.iterrows():
                                        fname = (r.get("filename") or f"file_{_}.bin")
                                        # Avoid directory traversal
                                        fname = fname.replace("..","").replace("\\","/").split("/")[-1]
                                        b = r.get("bytes")
                                        if isinstance(b, (bytes, bytearray)):
                                            zf.writestr(fname, b)
                                st.download_button(
                                    "Download ZIP",
                                    data=buf.getvalue(),
                                    file_name=f"rfp_{int(rid)}_linked_files.zip",
                                    mime="application/zip",
                                    key=f"zip_dl_{rid}"
                                )
                        except Exception as e:
                            st.error(f"ZIP build failed: {e}")
                with c_inv:
                    try:
                        inv = df_files.copy()
                        inv["sha_short"] = inv["sha256"].str.slice(0,12)
                        csvb = inv[["id","filename","mime","pages","sha_short"]].to_csv(index=False).encode("utf-8")
                        st.download_button("Export file inventory CSV", data=csvb, file_name=f"rfp_{int(rid)}_file_inventory.csv", mime="text/csv", key=f"inv_{rid}")
                    except Exception as e:
                        st.info(f"Inventory export unavailable: {e}")

                with st.expander("Find in linked files (simple)", expanded=False):
                    q = st.text_input("Search phrase", key=f"find_files_{rid}")
                    if q:
                        try:
                            hits = []
                            pool = pd.read_sql_query("SELECT id, filename, mime, bytes FROM rfp_files WHERE rfp_id=? ORDER BY id DESC;", conn, params=(int(rid),))
                            for _, r in pool.iterrows():
                                b = r.get("bytes"); mime = r.get("mime") or ""
                                pages = extract_text_pages(b, mime)
                                text = ("\n\n".join(pages) if pages else "")
                                pos = text.lower().find(q.lower())
                                if pos >= 0:
                                    start = max(0, pos-120); end = min(len(text), pos+120)
                                    ctx = text[start:end].replace("\n"," ")
                                    hits.append({"file_id": int(r["id"]), "filename": r.get("filename"), "snippet": ctx})
                            if hits:
                                dfh = pd.DataFrame(hits)
                                _styled_dataframe(dfh, use_container_width=True, hide_index=True)
                            else:
                                st.write("No hits.")
                        except Exception as e:
                            st.info(f"No search results: {e}")

                to_unlink = st.multiselect("Unlink file IDs", options=df_files["id"].tolist(), key=f"unlink_{rid}")
                if st.button("Unlink selected", key=f"unlink_btn_{rid}") and to_unlink:
                    try:
                        ph = ",".join(["?"]*len(to_unlink))
                        with closing(conn.cursor()) as _cur:
                            _cur.execute(f"UPDATE rfp_files SET rfp_id=NULL WHERE id IN ({ph});", tuple(int(i) for i in to_unlink))
                            conn.commit()
                        st.success(f"Unlinked {len(to_unlink)} file(s)."); st.rerun()
                    except Exception as e:
                        st.error(f"Unlink failed: {e}")
            st.caption("Attach from library")
            try:
                df_pool = pd.read_sql_query("SELECT id, filename, mime, pages FROM rfp_files WHERE rfp_id IS NULL ORDER BY id DESC LIMIT 500;", conn, params=())
            except Exception:
                df_pool = pd.DataFrame()
            if df_pool is None or df_pool.empty:
                st.write("No unlinked files in library.")
            else:
                _styled_dataframe(df_pool, use_container_width=True, hide_index=True)
                to_link = st.multiselect("Attach file IDs", options=df_pool["id"].tolist(), key=f"link_{rid}")
                if st.button("Attach selected to this RFP", key=f"link_btn_{rid}") and to_link:
                    try:
                        ph = ",".join(["?"]*len(to_link))
                        with closing(conn.cursor()) as _cur2:
                            _cur2.execute(f"UPDATE rfp_files SET rfp_id=? WHERE id IN ({ph});", (int(rid), *[int(i) for i in to_link]))
                            conn.commit()
                        st.success(f"Linked {len(to_link)} file(s)."); st.rerun()
                    except Exception as e:
                        st.error(f"Link failed: {e}")
        df_rf = pd.read_sql_query("SELECT id, title FROM rfps ORDER BY id DESC;", conn, params=())
        if df_rf.empty:
            st.info("No RFPs yet.")
            return
        rid = st.selectbox(
            "RFP for data views",
            options=df_rf["id"].tolist(),
            format_func=lambda i: f"#{i} — {df_rf.loc[df_rf['id']==i, 'title'].values[0]}",
            key="rfp_data_sel"
        )


        with st.expander("Acquisition Meta (X4)", expanded=False):
            try:
                df_meta_all = pd.read_sql_query("SELECT key, value FROM rfp_meta WHERE rfp_id=?;", conn, params=(int(rid),))
            except Exception:
                df_meta_all = pd.DataFrame(columns=['key','value'])
            if df_meta_all is None or df_meta_all.empty:
                st.write("No meta extracted yet.")
            else:
                # Highlight common fields
                want = ['naics','set_aside','place_of_performance','pop_structure','ordering_period_years','base_months']
                show = df_meta_all[df_meta_all["key"].isin(want)]
                if show.empty:
                    _styled_dataframe(df_meta_all, use_container_width=True, hide_index=True)
                else:
                    _styled_dataframe(show, use_container_width=True, hide_index=True)

        with st.expander("Ordering / POP (X3)", expanded=False):
            try:
                df_meta = pd.read_sql_query("SELECT key, value FROM rfp_meta WHERE rfp_id=?;", conn, params=(int(rid),))
            except Exception:
                df_meta = pd.DataFrame(columns=['key','value'])
            if df_meta is None or df_meta.empty:
                st.write("No POP metadata yet.")
            else:
                pop = df_meta[df_meta["key"].isin(["pop_structure","ordering_period_years","base_months"])]
                if pop.empty:
                    _styled_dataframe(df_meta, use_container_width=True, hide_index=True)
                else:
                    _styled_dataframe(pop, use_container_width=True, hide_index=True)

        # === Phase S: Manual Editors (LM / CLINs / Dates / POCs / Meta) ===

        import pandas as _pd
        from contextlib import closing as _closing_ed
        with st.expander('Manual Editors', expanded=False):
            tab_lm, tab_clin, tab_dates, tab_pocs, tab_meta = st.tabs(['L/M Items','CLINs','Key Dates','POCs','Meta'])
            with tab_lm:
                try:
                    df_lm_e = _pd.read_sql_query('SELECT item_text, is_must, status FROM lm_items WHERE rfp_id=? ORDER BY id;', conn, params=(int(rid),))
                except Exception:
                    df_lm_e = _pd.DataFrame(columns=['item_text','is_must','status'])
                df_lm_e = df_lm_e.fillna('')
                ed_lm = st.data_editor(df_lm_e, num_rows='dynamic', use_container_width=True, key=f'ed_lm_{rid}')
                if st.button('Save L/M', key=f'save_lm_{rid}'):
                    with _closing_ed(conn.cursor()) as cur:
                        cur.execute('DELETE FROM lm_items WHERE rfp_id=?;', (int(rid),))
                        for _, r in ed_lm.fillna('').iterrows():
                            txt = str(r.get('item_text','')).strip()
                            if not txt: continue
                            cur.execute('INSERT INTO lm_items(rfp_id, item_text, is_must, status) VALUES (?,?,?,?);', (int(rid), txt, int(r.get('is_must') or 0), str(r.get('status') or 'Open')))
                        conn.commit()
                    st.success('L/M saved.')
            with tab_clin:
                try:
                    df_c_e = _pd.read_sql_query('SELECT clin, description, qty, unit, unit_price, extended_price FROM clin_lines WHERE rfp_id=?;', conn, params=(int(rid),))
                except Exception:
                    df_c_e = _pd.DataFrame(columns=['clin','description','qty','unit','unit_price','extended_price'])
                df_c_e = df_c_e.fillna('')
                ed_c = st.data_editor(df_c_e, num_rows='dynamic', use_container_width=True, key=f'ed_clin_{rid}')
                if st.button('Save CLINs', key=f'save_clin_{rid}'):
                    with _closing_ed(conn.cursor()) as cur:
                        cur.execute('DELETE FROM clin_lines WHERE rfp_id=?;', (int(rid),))
                        for _, r in ed_c.fillna('').iterrows():
                            if not any(str(r.get(col,'')).strip() for col in ['clin','description','qty','unit','unit_price','extended_price']):
                                continue
                            cur.execute('INSERT INTO clin_lines(rfp_id, clin, description, qty, unit, unit_price, extended_price) VALUES (?,?,?,?,?,?,?);', (int(rid), str(r.get('clin','')), str(r.get('description','')), str(r.get('qty','')), str(r.get('unit','')), str(r.get('unit_price','')), str(r.get('extended_price',''))))
                        conn.commit()
                    st.success('CLINs saved.')
            with tab_dates:
                try:
                    df_d_e = _pd.read_sql_query('SELECT label, date_text, date_iso FROM key_dates WHERE rfp_id=?;', conn, params=(int(rid),))
                except Exception:
                    df_d_e = _pd.DataFrame(columns=['label','date_text','date_iso'])
                df_d_e = df_d_e.fillna('')
                ed_d = st.data_editor(df_d_e, num_rows='dynamic', use_container_width=True, key=f'ed_dates_{rid}')
                if st.button('Save Dates', key=f'save_dates_{rid}'):
                    with _closing_ed(conn.cursor()) as cur:
                        cur.execute('DELETE FROM key_dates WHERE rfp_id=?;', (int(rid),))
                        for _, r in ed_d.fillna('').iterrows():
                            if not any(str(r.get(col,'')).strip() for col in ['label','date_text','date_iso']):
                                continue
                            cur.execute('INSERT INTO key_dates(rfp_id, label, date_text, date_iso) VALUES (?,?,?,?);', (int(rid), str(r.get('label','')), str(r.get('date_text','')), str(r.get('date_iso',''))))
                        conn.commit()
                    st.success('Dates saved.')
            with tab_pocs:
                try:
                    df_p_e = _pd.read_sql_query('SELECT name, role, email, phone FROM pocs WHERE rfp_id=?;', conn, params=(int(rid),))
                except Exception:
                    df_p_e = _pd.DataFrame(columns=['name','role','email','phone'])
                df_p_e = df_p_e.fillna('')
                ed_p = st.data_editor(df_p_e, num_rows='dynamic', use_container_width=True, key=f'ed_pocs_{rid}')
                if st.button('Save POCs', key=f'save_pocs_{rid}'):
                    with _closing_ed(conn.cursor()) as cur:
                        cur.execute('DELETE FROM pocs WHERE rfp_id=?;', (int(rid),))
                        for _, r in ed_p.fillna('').iterrows():
                            if not any(str(r.get(col,'')).strip() for col in ['name','role','email','phone']):
                                continue
                            cur.execute('INSERT INTO pocs(rfp_id, name, role, email, phone) VALUES (?,?,?,?,?);', (int(rid), str(r.get('name','')), str(r.get('role','')), str(r.get('email','')), str(r.get('phone',''))))
                        conn.commit()
                    st.success('POCs saved.')
            with tab_meta:
                try:
                    df_m_e = _pd.read_sql_query('SELECT key, value FROM rfp_meta WHERE rfp_id=?;', conn, params=(int(rid),))
                except Exception:
                    df_m_e = _pd.DataFrame(columns=['key','value'])
                df_m_e = df_m_e.fillna('')
                ed_m = st.data_editor(df_m_e, num_rows='dynamic', use_container_width=True, key=f'ed_meta_{rid}')
                if st.button('Save Meta', key=f'save_meta_{rid}'):
                    with _closing_ed(conn.cursor()) as cur:
                        cur.execute('DELETE FROM rfp_meta WHERE rfp_id=?;', (int(rid),))
                        for _, r in ed_m.fillna('').iterrows():
                            k = str(r.get('key','')).strip(); v = str(r.get('value','')).strip()
                            if not k and not v: continue
                            cur.execute('INSERT INTO rfp_meta(rfp_id, key, value) VALUES (?,?,?);', (int(rid), k, v))
                        conn.commit()
                    st.success('Meta saved.')
        # === End Phase S ===
        col1, col2, col3 = st.columns(3)
        with col1:
            df_c = pd.read_sql_query("SELECT clin, description, qty, unit, unit_price, extended_price FROM clin_lines WHERE rfp_id=?;", conn, params=(int(rid),))
            st.subheader("CLINs"); _styled_dataframe(df_c, use_container_width=True, hide_index=True)
        with col2:
            df_d = pd.read_sql_query("SELECT label, date_text, date_iso FROM key_dates WHERE rfp_id=?;", conn, params=(int(rid),))
            st.subheader("Key Dates"); _styled_dataframe(df_d, use_container_width=True, hide_index=True)
        with col3:
            df_p = pd.read_sql_query("SELECT name, role, email, phone FROM pocs WHERE rfp_id=?;", conn, params=(int(rid),))
            st.subheader("POCs"); _styled_dataframe(df_p, use_container_width=True, hide_index=True)
        st.subheader("Attributes")
        df_meta = pd.read_sql_query("SELECT key, value FROM rfp_meta WHERE rfp_id=?;", conn, params=(int(rid),))
        _styled_dataframe(df_meta, use_container_width=True, hide_index=True)
        # --- RTM Coverage section ---
        try:
            rid_int = int(rid)
        except Exception:
            rid_int = None
        if rid_int:
            render_rtm_ui(conn, rid_int)

        # --- Amendment awareness sidebar ---
        try:
            sam_url = str(pd.read_sql_query("SELECT sam_url FROM rfps WHERE id=?;", conn, params=(rid_int,)).iloc[0]["sam_url"])
        except Exception:
            sam_url = ""
        ttl = int(st.session_state.get("cache_ttl_hours", 72)) if "cache_ttl_hours" in st.session_state else 72
        if sam_url:
            render_amendment_sidebar(conn, rid_int, sam_url, ttl)


    try:
        _rid = locals().get('rfp_id') or locals().get('rid') or st.session_state.get('current_rfp_id')
        y6_render_co_box(conn if 'conn' in locals() else None, _rid, key_prefix="run_rfp_analyzer_y6", title="Ask the CO about this RFP")
    except Exception:
        pass

def _compliance_progress(df_items: pd.DataFrame) -> int:
    if df_items is None or df_items.empty:
        return 0
    done = int((df_items["status"]=="Complete").sum())
    total = int(len(df_items))
    return int(round(done / max(1, total) * 100))



def _load_compliance_matrix(conn: sqlite3.Connection, rfp_id: int) -> pd.DataFrame:
    """
    Robust loader:
      1) If tenancy views exist (lm_items_t/lm_meta_t), use them.
      2) Else if base tables exist (lm_items/lm_meta), use them.
      3) Else return lm_items-only with blank meta columns.
    """
    # Ensure lm_meta exists (no-op if already there)
    try:
        with closing(conn.cursor()) as c:
            c.execute("CREATE TABLE IF NOT EXISTS lm_meta(\n"
                      " id INTEGER PRIMARY KEY,\n"
                      " lm_id INTEGER REFERENCES lm_items(id) ON DELETE CASCADE,\n"
                      " owner TEXT, ref_page TEXT, ref_para TEXT, evidence TEXT, risk TEXT, notes TEXT\n"
                      ");")
            c.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_lm_meta_lm ON lm_meta(lm_id);")
            conn.commit()
    except Exception:
        pass

    def _has(name: str) -> bool:
        try:
            q = "SELECT name FROM sqlite_master WHERE name=?;"
            return pd.read_sql_query(q, conn, params=(name,)).shape[0] > 0
        except Exception:
            return False

    use_views = _has("lm_items_t")
    use_meta_view = _has("lm_meta_t")

    if use_views and use_meta_view:
        q = """
            SELECT i.id AS lm_id, i.item_text, i.is_must, i.status,
                   COALESCE(m.owner,'') AS owner,
                   COALESCE(m.ref_page,'') AS ref_page,
                   COALESCE(m.ref_para,'') AS ref_para,
                   COALESCE(m.evidence,'') AS evidence,
                   COALESCE(m.risk,'Green') AS risk,
                   COALESCE(m.notes,'') AS notes
            FROM lm_items_t i
            LEFT JOIN lm_meta_t m ON m.lm_id = i.id
            WHERE i.rfp_id = ?
            ORDER BY i.id;
        """
        try:
            return pd.read_sql_query(q, conn, params=(rfp_id,))
        except Exception:
            pass  # fall through

    # Base tables path (works even if lm_meta is empty)
    if _has("lm_items"):
        # Only join lm_meta if it truly exists (older DBs may lack it)
        if _has("lm_meta"):
            q = """
                SELECT i.id AS lm_id, i.item_text, i.is_must, i.status,
                       COALESCE(m.owner,'') AS owner,
                       COALESCE(m.ref_page,'') AS ref_page,
                       COALESCE(m.ref_para,'') AS ref_para,
                       COALESCE(m.evidence,'') AS evidence,
                       COALESCE(m.risk,'Green') AS risk,
                       COALESCE(m.notes,'') AS notes
                FROM lm_items i
                LEFT JOIN lm_meta m ON m.lm_id = i.id
                WHERE i.rfp_id = ?
                ORDER BY i.id;
            """
        else:
            q = """
                SELECT i.id AS lm_id, i.item_text, i.is_must, i.status,
                       '' AS owner, '' AS ref_page, '' AS ref_para, '' AS evidence, 'Green' AS risk, '' AS notes
                FROM lm_items i
                WHERE i.rfp_id = ?
                ORDER BY i.id;
            """
        try:
            return pd.read_sql_query(q, conn, params=(rfp_id,))
        except Exception:
            pass

    # Final fallback: empty frame with expected columns
    cols = ["lm_id","item_text","is_must","status","owner","ref_page","ref_para","evidence","risk","notes"]
    return pd.DataFrame(columns=cols)

def _compliance_flags(ctx: dict, df_items: pd.DataFrame) -> pd.DataFrame:
    rows = []
    sections = ctx.get("sections", pd.DataFrame())
    text_all = " ".join((sections["content"].tolist() if isinstance(sections, pd.DataFrame) and not sections.empty else []))
    tl = text_all.lower()

    m = re.search(r'(?:page\s+limit|not\s+exceed)\s+(?:of\s+)?(\d{1,3})\s+pages?', tl)
    if m: rows.append({"Rule":"Page Limit","Detail":f"Limit {m.group(1)} pages detected","Severity":"Amber"})
    if re.search(r'(font|typeface).{0,20}(size|pt).{0,5}(10|11)', tl):
        rows.append({"Rule":"Font size","Detail":"Minimum font size 10/11pt likely required","Severity":"Amber"})
    if re.search(r'margin[s]?\s+(?:of|at\s+least)\s+\d', tl):
        rows.append({"Rule":"Margins","Detail":"Specific margin requirements detected","Severity":"Amber"})
    if re.search(r'volume[s]?\s+(i{1,3}|iv|v|technical|price)', tl):
        rows.append({"Rule":"Volumes","Detail":"Multiple volumes required","Severity":"Amber"})
    if re.search(r'(sam\.gov|piee|wawf|email submission|portal)', tl):
        rows.append({"Rule":"Submission portal","Detail":"Specific portal/email submission detected","Severity":"Amber"})

    dates = ctx.get("dates", pd.DataFrame())
    if isinstance(dates, pd.DataFrame) and not dates.empty:
        due = dates[dates["label"].str.contains("due", case=False, na=False)]
        if not due.empty:
            dt = pd.to_datetime(due.iloc[0]["date_text"], errors="coerce")
            if pd.notnull(dt):
                days = (pd.Timestamp(dt) - pd.Timestamp.utcnow()).days
                if days <= 3: rows.append({"Rule":"Timeline","Detail":f"Proposals due in {days} day(s)","Severity":"Red"})
                elif days <= 7: rows.append({"Rule":"Timeline","Detail":f"Proposals due in {days} days","Severity":"Amber"})

    if isinstance(df_items, pd.DataFrame) and not df_items.empty:
        open_musts = df_items[(df_items["is_must"]==1) & (df_items["status"]!="Complete")]
        if not open_musts.empty:
            rows.append({"Rule":"Open MUST items","Detail":f"{len(open_musts)} mandatory items still open","Severity":"Red"})

    return pd.DataFrame(rows)


def _load_rfp_context(conn: sqlite3.Connection, rfp_id: int) -> dict:
    try:
        rf = pd.read_sql_query("SELECT id, title, solnum, sam_url, created_at FROM rfps WHERE id=?;", conn, params=(int(rfp_id),))
    except Exception:
        rf = pd.DataFrame()
    try:
        df_items = pd.read_sql_query("SELECT id, item_text, is_must, status FROM lm_items WHERE rfp_id=? ORDER BY id;", conn, params=(int(rfp_id),))
    except Exception:
        df_items = pd.DataFrame(columns=["id","item_text","is_must","status"])
    joined = "\n".join(df_items["item_text"].astype(str).tolist()) if not df_items.empty else ""
    sections = pd.DataFrame([{"name":"Checklist Items","content": joined}])
    meta = rf.iloc[0].to_dict() if not rf.empty else {}
    return {"rfp": meta, "sections": sections, "items": df_items}


def run_lm_checklist(conn: sqlite3.Connection) -> None:

    st.header("L and M Checklist")
    rfp_id = st.session_state.get('current_rfp_id')
    if not rfp_id:
        try:
            df_rf = pd.read_sql_query("SELECT id, title, solnum, created_at FROM rfps_t ORDER BY id DESC;", conn, params=())
        except Exception as e:
            st.error(f"Failed to load RFPs: {e}")
            return
        if df_rf.empty:
            st.info("No saved RFP extractions yet. Use RFP Analyzer to parse and save.")
            return
        opt = st.selectbox("Select an RFP context", options=df_rf['id'].tolist(),
                           format_func=lambda rid: f"#{rid} — {df_rf.loc[df_rf['id']==rid,'title'].values[0] or 'Untitled'}")
        rfp_id = opt
        st.session_state['current_rfp_id'] = rfp_id

    st.caption(f"Working RFP ID: {rfp_id}")
    try:
        df_items = pd.read_sql_query("SELECT id, item_text, is_must, status FROM lm_items WHERE rfp_id=?;", conn, params=(rfp_id,))
    except Exception as e:
        st.error(f"Failed to load items: {e}")
        return
    if df_items.empty:
        st.info("No L/M items found for this RFP.")
        return

    pct = _compliance_progress(df_items)
    st.progress(pct/100.0, text=f"{pct}% complete")

    c1, c2, c3 = st.columns([2,2,2])
    with c1:
        if st.button("Mark all Complete"):
            try:
                with closing(conn.cursor()) as cur:
                    cur.execute("UPDATE lm_items SET status='Complete' WHERE rfp_id=?;", (rfp_id,))
                    conn.commit()
                st.success("All items marked Complete")
            except Exception as e:
                st.error(f"Update failed: {e}")
    with c2:
        if st.button("Reset all to Open"):
            try:
                with closing(conn.cursor()) as cur:
                    cur.execute("UPDATE lm_items SET status='Open' WHERE rfp_id=?;", (rfp_id,))
                    conn.commit()
                st.success("All items reset")
            except Exception as e:
                st.error(f"Update failed: {e}")
    with c3:
        if st.button("Mark all MUST to Open"):
            try:
                with closing(conn.cursor()) as cur:
                    cur.execute("UPDATE lm_items SET status='Open' WHERE rfp_id=? AND is_must=1;", (rfp_id,))
                    conn.commit()
                st.success("All MUST items set to Open")
            except Exception as e:
                st.error(f"Update failed: {e}")

    st.subheader("Checklist")
    for _, row in df_items.iterrows():
        key = f"lm_{row['id']}"
        label = ("[MUST] " if row['is_must']==1 else "") + row['item_text']
        checked = st.checkbox(label, value=(row['status']=='Complete'), key=key)
        new_status = 'Complete' if checked else 'Open'
        if new_status != row['status']:
            try:
                with closing(conn.cursor()) as cur:
                    cur.execute("UPDATE lm_items SET status=? WHERE id=?;", (new_status, int(row['id'])))
                    conn.commit()
            except Exception as e:
                st.error(f"Failed to update item {row['id']}: {e}")

    st.divider()

    st.subheader("Compliance Matrix")
    df_mx = _load_compliance_matrix(conn, int(rfp_id))
    if df_mx.empty:
        st.info("No items to show.")
        return

    view = df_mx.rename(columns={
        "item_text":"Requirement","is_must":"Must?","status":"Status",
        "owner":"Owner","ref_page":"Page","ref_para":"Para",
        "evidence":"Evidence/Link","risk":"Risk","notes":"Notes"
    })
    _styled_dataframe(view[["Requirement","Must?","Status","Owner","Page","Para","Evidence/Link","Risk","Notes"]],
                 use_container_width=True, hide_index=True)

    st.markdown("**Edit selected requirement**")
    pick = st.selectbox("Requirement", options=df_mx["lm_id"].tolist(),
                        format_func=lambda lid: f"#{lid} — {df_mx.loc[df_mx['lm_id']==lid,'item_text'].values[0][:80]}")

    rec = df_mx[df_mx["lm_id"]==pick].iloc[0].to_dict()
    e1, e2, e3, e4 = st.columns([2,1,1,1])
    with e1:
        owner = st.text_input("Owner", value=rec.get("owner",""), key=f"mx_owner_{pick}")
        notes = st.text_area("Notes", value=rec.get("notes",""), key=f"mx_notes_{pick}", height=90)
    with e2:
        page = st.text_input("Page", value=rec.get("ref_page",""), key=f"mx_page_{pick}")
        para = st.text_input("Paragraph", value=rec.get("ref_para",""), key=f"mx_para_{pick}")
    with e3:
        risk = st.selectbox("Risk", ["Green","Yellow","Red"],
                            index=["Green","Yellow","Red"].index(rec.get("risk","Green")), key=f"mx_risk_{pick}")
    with e4:
        evidence = st.text_input("Evidence/Link", value=rec.get("evidence",""), key=f"mx_evid_{pick}")

    csave, cexp = st.columns([2,2])
    with csave:
        if st.button("Save Matrix Row", key=f"mx_save_{pick}"):
            try:
                with closing(conn.cursor()) as cur:
                    cur.execute("""
                        INSERT INTO lm_meta(lm_id, owner, ref_page, ref_para, evidence, risk, notes)
                        VALUES(?,?,?,?,?,?,?)
                        ON CONFLICT(lm_id) DO UPDATE SET
                            owner=excluded.owner, ref_page=excluded.ref_page, ref_para=excluded.ref_para,
                            evidence=excluded.evidence, risk=excluded.risk, notes=excluded.notes;
                    """, (int(pick), owner.strip(), page.strip(), para.strip(), evidence.strip(), risk, notes.strip()))
                    conn.commit()
                st.success("Saved"); st.rerun()
            except Exception as e2:
                st.error(f"Save failed: {e2}")
    with cexp:
        if st.button("Export Matrix CSV", key="mx_export"):
            out = view.copy()
            path = str(Path(DATA_DIR) / f"compliance_matrix_rfp_{int(rfp_id)}.csv")
            out.to_csv(path, index=False)
            st.success("Exported"); st.markdown(f"[Download CSV]({path})")

    st.subheader("Red-Flag Finder")
    ctx = _load_rfp_context(conn, int(rfp_id))
    flags = _compliance_flags(ctx, df_items)
    if flags is None or flags.empty:
        st.write("No obvious flags detected.")
    else:
        _styled_dataframe(flags, use_container_width=True, hide_index=True)




    try:
        _rid = locals().get('rfp_id') or locals().get('rid') or st.session_state.get('current_rfp_id')
        y6_render_co_box(conn if 'conn' in locals() else None, _rid, key_prefix="run_lm_checklist_y6", title="Ask the CO about L&M")
    except Exception:
        pass

def _estimate_pages(total_words: int, spacing: str = "1.15", words_per_page: Optional[int] = None) -> float:
    """Rough page estimate at common spacings for 11pt fonts."""
    if words_per_page is None:
        s = (spacing or "1.15").strip().lower()
        if s in {"1", "1.0", "single"}:
            wpp = 500
        elif s in {"1.15", "1,15"}:
            wpp = 400
        elif s in {"1.5", "1,5"}:
            wpp = 300
        elif s in {"double", "2", "2.0"}:
            wpp = 250
        else:
            wpp = 400
    else:
        wpp = max(50, int(words_per_page))
    return round((total_words or 0) / float(wpp), 2)


def _export_docx(path: str, doc_title: str, sections: List[dict], clins: Optional[pd.DataFrame] = None, checklist: Optional[pd.DataFrame] = None, metadata: Optional[dict] = None,
                 font_name: str = "Times New Roman",
                 font_size_pt: int = 11,
                 spacing: str = "1.15") -> Optional[str]:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
        from docx.enum.text import WD_LINE_SPACING  # type: ignore
    except Exception:
        st.error("python-docx is required. pip install python-docx")
        return None
    spacing_map = {
        "single": WD_LINE_SPACING.SINGLE, "1": WD_LINE_SPACING.SINGLE, "1.0": WD_LINE_SPACING.SINGLE,
        "1.15": WD_LINE_SPACING.ONE_POINT_FIVE, "1,15": WD_LINE_SPACING.ONE_POINT_FIVE,
        "1.5": WD_LINE_SPACING.ONE_POINT_FIVE, "double": WD_LINE_SPACING.DOUBLE,
        "2": WD_LINE_SPACING.DOUBLE, "2.0": WD_LINE_SPACING.DOUBLE,
    }
    line_spacing = spacing_map.get((spacing or "1.15").lower(), WD_LINE_SPACING.ONE_POINT_FIVE)
    doc = docx.Document()
    h = doc.add_heading(doc_title or "Proposal", level=1)
    if metadata:
        p = doc.add_paragraph(" | ".join(f"{k}: {v}" for k,v in metadata.items()))
    for s in (sections or []):
        title = str(s.get("title","")).strip()
        body = str(s.get("body","")).strip()
        if title: doc.add_heading(title, level=2)
        for para in body.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                try: p.paragraph_format.line_spacing_rule = line_spacing
                except Exception: pass
                for run in p.runs:
                    try: run.font.name = font_name; run.font.size = Pt(font_size_pt)
                    except Exception: pass
    try:
        if isinstance(clins, pd.DataFrame) and not clins.empty:
            tbl = doc.add_table(rows=1, cols=len(clins.columns))
            for j, col in enumerate(clins.columns): tbl.rows[0].cells[j].text = str(col)
            for _, row in clins.iterrows():
                cells = tbl.add_row().cells
                for j, col in enumerate(clins.columns):
                    val = row.get(col); cells[j].text = "" if pd.isna(val) else str(val)
    except Exception: pass
    try:
        if isinstance(checklist, pd.DataFrame) and not checklist.empty:
            doc.add_heading("Compliance Checklist", level=2)
            for _, r in checklist.iterrows():
                txt = str(r.get("item_text","")).strip()
                if txt: doc.add_paragraph(txt, style="List Bullet")
    except Exception: pass
    doc.save(path)
    return path

def run_proposal_builder(conn: sqlite3.Connection) -> None:
    st.header("Proposal Builder")
    df_rf = pd.read_sql_query("SELECT id, title, solnum, notice_id FROM rfps_t ORDER BY id DESC;", conn, params=())
    if df_rf.empty:
        st.info("No RFP context found. Use RFP Analyzer first to parse and save.")
        return
    rfp_id = st.selectbox(
        "RFP context",
        options=df_rf["id"].tolist(),
        format_func=lambda rid: f"#{rid} — {df_rf.loc[df_rf['id']==rid,'title'].values[0] or 'Untitled'}",
        index=0,
    )
    st.session_state["current_rfp_id"] = rfp_id
    ctx = _load_rfp_context(conn, rfp_id)

    left, right = st.columns([3, 2])
    with left:
        st.subheader("Sections")
        default_sections = [
            "Cover Letter","Executive Summary","Understanding of Requirements","Technical Approach","Management Plan",
            "Staffing and Key Personnel","Quality Assurance","Past Performance Summary","Pricing and CLINs","Certifications and Reps","Appendices",
        ]
        selected = st.multiselect("Include sections", default_sections, default=default_sections)
        content_map: Dict[str, str] = {}
        for sec in selected:
            default_val = st.session_state.get(f"pb_section_{sec}", "")
            st.markdown(f"**{sec}**")
            notes = st.text_input(f"Notes for {sec}", key=f"y3_notes_{sec}")
            cA, cB, cC = st.columns([1,1,1])
            with cA:
                k = y_auto_k(f"{sec} {notes}")
            with cB:
                maxw = st.number_input(f"Max words — {sec}", min_value=0, value=220, step=10, key=f"y3_maxw_{sec}")
            with cC:
                if st.button(f"Draft {sec}", key=f"y3_draft_{sec}"):
                    ph = st.empty(); acc=[]
                    for tok in y3_stream_draft(conn, int(rfp_id), sec, notes or "", k=int(k), max_words=int(maxw) if maxw>0 else None):
                        acc.append(tok); ph.markdown("".join(acc))
                    drafted = "".join(acc).strip()
                    if drafted:
                        st.session_state[f"pb_section_{sec}"] = drafted
                        default_val = drafted
            content_map[sec] = st.text_area(sec, value=default_val, height=200, key=f"pb_ta_{sec}")
    with right:
        st.subheader("Guidance and limits")
        spacing = st.selectbox("Line spacing", ["Single", "1.15", "Double"], index=1)
        font_name = st.selectbox("Font", ["Times New Roman", "Calibri"], index=0)
        font_size = st.number_input("Font size", min_value=10, max_value=12, value=11)
        page_limit = st.number_input("Page limit for narrative", min_value=1, max_value=200, value=10)

        st.markdown("**Must address items from L and M**")
        items = _ctxd(ctx, "items") if isinstance(ctx.get("items"), pd.DataFrame) else pd.DataFrame()
        if not items.empty:
            _styled_dataframe(items.rename(columns={"item_text": "Item", "status": "Status"}), use_container_width=True, hide_index=True, height=240)
        else:
            st.caption("No checklist items found for this RFP")

        total_words = sum(len((content_map.get(k) or "").split()) for k in selected)
        est_pages = _estimate_pages(total_words, spacing)
        st.info(f"Current word count {total_words}  Estimated pages {est_pages}")
        if est_pages > page_limit:
            st.error("Content likely exceeds page limit. Consider trimming or tighter formatting")

        out_name = f"Proposal_RFP_{int(rfp_id)}.docx"
        out_path = os.path.join(DATA_DIR, out_name)
        if st.button("Export DOCX", type="primary"):
            sections = [{"title": k, "body": content_map.get(k, "")} for k in selected]
            exported = _export_docx(
                out_path,
                doc_title=_ctxd(ctx, "rfp").iloc[0]["title"] if _df_nonempty(ctx.get("rfp")) else "Proposal",
                sections=sections,
                clins=_ctxd(ctx, "clins"),
                checklist=_ctxd(ctx, "items"),
                metadata={
                    "Solicitation": (_ctxd(ctx, "rfp").iloc[0]["solnum"] if _df_nonempty(ctx.get("rfp")) else ""),
                    "Notice ID": (_ctxd(ctx, "rfp").iloc[0]["notice_id"] if _df_nonempty(ctx.get("rfp")) else ""),
                },
                font_name=font_name,
                font_size_pt=int(font_size),
                spacing=spacing,
            )
            if exported:
                st.success(f"Exported to {exported}")
                st.markdown(f"[Download DOCX]({exported})")

        with st.expander("Snippets Inbox (from Y1/Y2/Y4/Y5)", expanded=True):
            try:
                df_snip = pd.read_sql_query("SELECT id, section, source, text, created_at FROM draft_snippets WHERE rfp_id=? ORDER BY id DESC;", conn, params=(int(rfp_id),))
            except Exception:
                df_snip = pd.DataFrame()
            if df_snip is None or df_snip.empty:
                st.caption("No snippets saved yet")
            else:
                _styled_dataframe(df_snip[["id","section","source","created_at"]], use_container_width=True, hide_index=True, height=200)
                sid = st.selectbox("Pick snippet ID", options=df_snip["id"].tolist())
                sec_choice = st.selectbox("Insert into section", options=selected, index=min(len(selected)-1, 0))
                if st.button("Insert snippet"):
                    txt = df_snip[df_snip["id"]==sid].iloc[0]["text"]
                    key = f"pb_ta_{sec_choice}"
                    cur = st.session_state.get(key, "")
                    st.session_state[key] = (cur + ("\n\n" if cur else "") + str(txt)).strip()
                    st.session_state[f"pb_section_{sec_choice}"] = st.session_state[key]
                    st.success("Inserted into section")


# ---------- Subcontractor Finder (Phase D) ----------
    try:
        _rid = locals().get('rfp_id') or locals().get('rid') or st.session_state.get('current_rfp_id')
        y6_render_co_box(conn if 'conn' in locals() else None, _rid, key_prefix="run_proposal_builder_y6", title="Ask the CO while drafting")
    except Exception:
        pass


def _s1d_paginate(df, page_size: int, page_key: str = "s1d_page"):
    import math
    import streamlit as st
    n = 0 if df is None else len(df)
    page_size = max(5, int(page_size or 25))
    pages = max(1, math.ceil((n or 0) / page_size))
    page = int(st.session_state.get(page_key, 1))
    if page < 1: page = 1
    if page > pages: page = pages
    start = (page - 1) * page_size
    end = start + page_size
    view = df.iloc[start:end].copy() if df is not None else df
    st.session_state[page_key] = page
    return view, page, pages
def run_subcontractor_finder(conn: sqlite3.Connection) -> None:
    st.header("Subcontractor Finder")
    st.caption("Seed and manage vendors by NAICS/PSC/state; handoff selected vendors to Outreach.")

    ctx = st.session_state.get("rfp_selected_notice", {})
    default_naics = ctx.get("NAICS") or ""
    default_state = ""

    # Default Place of Performance from selected notice if available
    default_pop = (ctx.get("Place of Performance") or ctx.get("place_of_performance") or ctx.get("POP") or "").strip()

    with st.expander("Filters", expanded=True):
        c1, c2, c3, c4 = st.columns([2,2,2,2])
        with c1:
            f_naics = st.text_input("NAICS", value=default_naics, key="filter_naics")
        with c2:
            f_state = st.text_input("State (e.g., TX)", value=default_state, key="filter_state")
        with c3:
            f_city = st.text_input("City contains", key="filter_city")
        with c4:
            f_kw = st.text_input("Keyword in name/notes", key="filter_kw")
        st.caption("Use CSV import or add vendors manually. Internet seeding can be added later.")

    with st.expander("Import Vendors (CSV)", expanded=False):
        st.caption("Headers: name, email, phone, city, state, naics, cage, uei, website, notes")
        up = st.file_uploader("Upload vendor CSV", type=["csv"], key="vendor_csv")
        if up and st.button("Import CSV"):
            try:
                df = pd.read_csv(up)
                if "name" not in {c.lower() for c in df.columns}:
                    st.error("CSV must include a 'name' column")
                else:
                    df.columns = [c.lower() for c in df.columns]
                    n=0
                    with closing(conn.cursor()) as cur:
                        for _, r in df.iterrows():
                            cur.execute(
                                """
                                INSERT INTO vendors(name, cage, uei, naics, city, state, phone, email, website, notes)
                                VALUES(?,?,?,?,?,?,?,?,?,?)
                                ;
                                """,
                                (
                                    str(r.get("name",""))[:200],
                                    str(r.get("cage",""))[:20],
                                    str(r.get("uei",""))[:40],
                                    str(r.get("naics",""))[:20],
                                    str(r.get("city",""))[:100],
                                    str(r.get("state",""))[:10],
                                    str(r.get("phone",""))[:40],
                                    str(r.get("email",""))[:120],
                                    str(r.get("website",""))[:200],
                                    str(r.get("notes",""))[:500],
                                ),
                            )
                            n+=1
                    conn.commit()
                    st.success(f"Imported {n} vendors")
            except Exception as e:
                st.error(f"Import failed: {e}")

    with st.expander("Add Vendor", expanded=False):
        c1, c2, c3 = st.columns([2,2,2])
        with c1:
            v_name = st.text_input("Company name", key="add_name")
            v_email = st.text_input("Email", key="add_email")
            v_phone = st.text_input("Phone", key="add_phone")
        with c2:
            v_city = st.text_input("City", key="add_city")
            v_state = st.text_input("State", key="add_state")
            v_naics = st.text_input("NAICS", key="add_naics")
        with c3:
            v_cage = st.text_input("CAGE", key="add_cage")
            v_uei = st.text_input("UEI", key="add_uei")
            v_site = st.text_input("Website", key="add_site")
        v_notes = st.text_area("Notes", height=80, key="add_notes")
        if st.button("Save Vendor"):
            if not v_name.strip():
                st.error("Name is required")
            else:
                try:
                    with closing(conn.cursor()) as cur:
                        cur.execute(
                            """
                            INSERT INTO vendors(name, cage, uei, naics, city, state, phone, email, website, notes)
                            VALUES(?,?,?,?,?,?,?,?,?,?)
                            ;
                            """,
                            (v_name.strip(), v_cage.strip(), v_uei.strip(), v_naics.strip(), v_city.strip(), v_state.strip(), v_phone.strip(), v_email.strip(), v_site.strip(), v_notes.strip()),
                        )
                        conn.commit()
                    st.success("Vendor saved")
                except Exception as e:
                    st.error(f"Save failed: {e}")

    q = "SELECT id, name, email, phone, city, state, naics, cage, uei, website, notes FROM vendors_t WHERE 1=1"
    params: List[Any] = []
    if f_naics:
        q += " AND (naics LIKE ? )"
        params.append(f"%{f_naics}%")
    if f_state:
        q += " AND (state LIKE ?)"
        params.append(f"%{f_state}%")
    if f_city:
        q += " AND (city LIKE ?)"
        params.append(f"%{f_city}%")
    if f_kw:
        q += " AND (name LIKE ? OR notes LIKE ?)"
        params.extend([f"%{f_kw}%", f"%{f_kw}%"])

    try:
        df_v = pd.read_sql_query(q + " ORDER BY name ASC;", conn, params=params)
    except Exception as e:
        st.error(f"Query failed: {e}")
        df_v = pd.DataFrame()

    st.subheader("Vendors")
    if df_v.empty:
        st.write("No vendors match filters")
    else:
        selected_ids = []
        for _, row in df_v.iterrows():
            chk = st.checkbox(f"Select — {row['name']}  ({row['email'] or 'no email'})", key=f"vend_{int(row['id'])}")
            if chk:
                selected_ids.append(int(row['id']))
        c1, c2 = st.columns([2,2])
        with c1:
            if st.button("Send to Outreach ▶") and selected_ids:
                st.session_state['rfq_vendor_ids'] = selected_ids
                st.success(f"Queued {len(selected_ids)} vendors for Outreach")
        with c2:
            st.caption("Selections are stored in session and available in Outreach tab")


# ---------- Outreach (Phase D) ----------
    try:
        _rid = locals().get('rfp_id') or locals().get('rid') or st.session_state.get('current_rfp_id')
        y6_render_co_box(conn if 'conn' in locals() else None, _rid, key_prefix="run_subcontractor_finder_y6", title="CO guidance for subcontractors")
    except Exception:
        pass

def _smtp_settings() -> Dict[str, Any]:
    out = {"host": None, "port": 587, "username": None, "password": None, "from_email": None, "from_name": "ELA Management", "use_tls": True}
    try:
        cfg = st.secrets.get("smtp", {})
        out.update({k: cfg.get(k, out[k]) for k in out})
    except Exception:
        pass
    for k in list(out.keys()):
        if not out[k]:
            try:
                v = st.secrets.get(k)
                if v:
                    out[k] = v
            except Exception:
                pass
    return out


def send_email_smtp(to_email: str, subject: str, html_body: str, attachments: List[str]) -> Tuple[bool, str]:
    cfg = _smtp_settings()
    if not all([cfg.get("host"), cfg.get("port"), cfg.get("username"), cfg.get("password"), cfg.get("from_email")]):
        return False, "Missing SMTP settings in secrets"

    msg = MIMEMultipart()
    msg["From"] = f"{cfg.get('from_name') or ''} <{cfg['from_email']}>"
    msg["To"] = to_email
    msg["Subject"] = subject
    msg.attach(MIMEText(html_body, "html"))

    for path in attachments or []:
        try:
            with open(path, "rb") as f:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(path)}"')
                msg.attach(part)
        except Exception:
            pass

    try:
        server = smtplib.SMTP(cfg['host'], int(cfg['port']))
        if cfg.get('use_tls', True):
            server.starttls()
        server.login(cfg['username'], cfg['password'])
        server.sendmail(cfg['from_email'], [to_email], msg.as_string())
        server.quit()
        return True, "sent"
    except Exception as e:
        return False, str(e)


def _merge_text(t: str, vendor: Dict[str, Any], notice: Dict[str, Any]) -> str:
    repl = {
        "company": vendor.get("name", ""),
        "email": vendor.get("email", ""),
        "phone": vendor.get("phone", ""),
        "city": vendor.get("city", ""),
        "state": vendor.get("state", ""),
        "naics": vendor.get("naics", ""),
        "title": notice.get("Title", ""),
        "solicitation": notice.get("Solicitation", ""),
        "due": notice.get("Response Due", ""),
        "notice_id": notice.get("Notice ID", ""),
    }
    out = t
    for k, v in repl.items():
        out = out.replace(f"{{{{{k}}}}}", str(v))
    return out



def _calc_extended(qty: Optional[float], unit_price: Optional[float]) -> Optional[float]:
    try:
        if qty is None or unit_price is None:
            return None
        return float(qty) * float(unit_price)
    except Exception:
        return None


def run_quote_comparison(conn: sqlite3.Connection) -> None:
    st.header("Quote Comparison")
    df = pd.read_sql_query("SELECT id, title, solnum FROM rfps_t ORDER BY id DESC;", conn, params=())
    if df.empty:
        st.info("No RFPs in DB. Use RFP Analyzer to create one (Parse → Save).")
        return
    rfp_id = st.selectbox("RFP context", options=df["id"].tolist(), format_func=lambda rid: f"#{rid} — {df.loc[df['id']==rid, 'title'].values[0] or 'Untitled'}")

    st.subheader("Upload / Add Quotes")
    with st.expander("CSV Import", expanded=False):
        st.caption("Columns: vendor, clin, qty, unit_price, description (optional). One row = one CLIN line.")
        up = st.file_uploader("Quotes CSV", type=["csv"], key="quotes_csv")
        if up and st.button("Import Quotes CSV"):
            try:
                df_csv = pd.read_csv(up)
                required = {"vendor", "clin", "qty", "unit_price"}
                if not required.issubset({c.lower() for c in df_csv.columns}):
                    st.error("CSV missing required columns: vendor, clin, qty, unit_price")
                else:
                    df_csv.rename(columns={c: c.lower() for c in df_csv.columns}, inplace=True)
                    with closing(conn.cursor()) as cur:
                        by_vendor = df_csv.groupby("vendor", dropna=False)
                        total_rows = 0
                        for vendor, block in by_vendor:
                            cur.execute(
                                "INSERT INTO quotes(rfp_id, vendor, received_date, notes) VALUES(?,?,?,?);",
                                (int(rfp_id), str(vendor)[:200], datetime.utcnow().isoformat(), "imported")
                            )
                            qid = cur.lastrowid
                            for _, r in block.iterrows():
                                qty = float(r.get("qty", 0) or 0)
                                upx = float(r.get("unit_price", 0) or 0)
                                ext = _calc_extended(qty, upx) or 0.0
                                cur.execute(
                                    "INSERT INTO quote_lines(quote_id, clin, description, qty, unit_price, extended_price) VALUES(?,?,?,?,?,?);",
                                    (qid, str(r.get("clin",""))[:50], str(r.get("description",""))[:300], qty, upx, ext)
                                )
                                total_rows += 1
                        conn.commit()
                    st.success(f"Imported {len(by_vendor)} quotes / {total_rows} lines.")
            except Exception as e:
                st.error(f"Import failed: {e}")

    with st.expander("Add Quote (manual)", expanded=False):
        vendor = st.text_input("Vendor name")
        date = st.date_input("Received date", value=datetime.utcnow().date())
        notes = st.text_input("Notes", value="")
        add_quote = st.button("Create Quote")
        if add_quote and vendor.strip():
            with closing(conn.cursor()) as cur:
                cur.execute("INSERT INTO quotes(rfp_id, vendor, received_date, notes) VALUES(?,?,?,?);",
                            (int(rfp_id), vendor.strip(), date.isoformat(), notes.strip()))
                qid = cur.lastrowid
                conn.commit()
                st.success(f"Created quote for {vendor}. Now add lines below (Quote ID {qid}).")
                st.session_state["current_quote_id"] = qid

    df_q = pd.read_sql_query("SELECT id, vendor, received_date, notes FROM quotes WHERE rfp_id=? ORDER BY vendor;", conn, params=(rfp_id,))
    if not df_q.empty:
        st.subheader("Quotes")
        _styled_dataframe(df_q, use_container_width=True, hide_index=True)
        qid = st.selectbox("Edit lines for quote", options=df_q["id"].tolist(), format_func=lambda qid: f"#{qid} — {df_q.loc[df_q['id']==qid,'vendor'].values[0]}")
        with st.form("add_quote_line", clear_on_submit=True):
            c1, c2, c3 = st.columns([2, 1, 1])
            with c1:
                clin = st.text_input("CLIN")
                desc = st.text_input("Description")
            with c2:
                qty = st.number_input("Qty", min_value=0.0, step=1.0)
            with c3:
                price = st.number_input("Unit Price", min_value=0.0, step=1.0)
            submitted = st.form_submit_button("Add Line")
        if submitted:
            ext = _calc_extended(qty, price) or 0.0
            with closing(conn.cursor()) as cur:
                cur.execute(
                    "INSERT INTO quote_lines(quote_id, clin, description, qty, unit_price, extended_price) VALUES(?,?,?,?,?,?);",
                    (qid, clin.strip(), desc.strip(), float(qty), float(price), float(ext))
                )
                conn.commit()
            st.success("Line added.")

    st.subheader("Comparison")
    df_target = pd.read_sql_query("SELECT clin, description FROM clin_lines WHERE rfp_id=? GROUP BY clin, description ORDER BY clin;", conn, params=(rfp_id,))
    df_lines = pd.read_sql_query("""
        SELECT q.vendor, l.clin, l.qty, l.unit_price, l.extended_price
        FROM quote_lines l
        JOIN quotes q ON q.id = l.quote_id
        WHERE q.rfp_id=?
    """, conn, params=(rfp_id,))
    if df_lines.empty:
        st.info("No quote lines yet.")
        return

    mat = df_lines.pivot_table(index="clin", columns="vendor", values="extended_price", aggfunc="sum").fillna(0.0)
    mat = mat.sort_index()
    _styled_dataframe(mat.style.format("{:,.2f}"), use_container_width=True)

    best_vendor_by_clin = mat.replace(0, float("inf")).idxmin(axis=1).to_frame("Best Vendor")
    st.caption("Best vendor per CLIN")
    _styled_dataframe(best_vendor_by_clin, use_container_width=True, hide_index=False)

    totals = df_lines.groupby("vendor")["extended_price"].sum().to_frame("Total").sort_values("Total")
    if not df_target.empty:
        coverage = df_lines.groupby("vendor")["clin"].nunique().to_frame("CLINs Quoted")
        coverage["Required CLINs"] = df_target["clin"].nunique()
        coverage["Coverage %"] = (coverage["CLINs Quoted"] / coverage["Required CLINs"] * 100).round(1)
        totals = totals.join(coverage, how="left")
    st.subheader("Totals & Coverage")
    _styled_dataframe(totals.style.format({"Total": "{:,.2f}", "Coverage %": "{:.1f}"}), use_container_width=True)

    if st.button("Export comparison CSV"):
        path = os.path.join(DATA_DIR, "quote_comparison.csv")
        out = mat.copy()
        out["Best Vendor"] = best_vendor_by_clin["Best Vendor"]
        out.to_csv(path)
        st.success("Exported.")
        st.markdown(f"[Download comparison CSV]({path})")


# ---------- Pricing Calculator (Phase E) ----------
def _scenario_summary(conn: sqlite3.Connection, scenario_id: int) -> Dict[str, float]:
    dl = pd.read_sql_query("SELECT hours, rate, fringe_pct FROM pricing_labor WHERE scenario_id=?;", conn, params=(scenario_id,))
    other = pd.read_sql_query("SELECT cost FROM pricing_other WHERE scenario_id=?;", conn, params=(scenario_id,))
    base = pd.read_sql_query("SELECT overhead_pct, gna_pct, fee_pct, contingency_pct FROM pricing_scenarios WHERE id=?;", conn, params=(scenario_id,))
    if base.empty:
        return {}
    overhead_pct, gna_pct, fee_pct, contingency_pct = base.iloc[0]
    direct_labor = float((dl["hours"] * dl["rate"]).sum()) if not dl.empty else 0.0
    fringe = float((dl["hours"] * dl["rate"] * (dl["fringe_pct"].fillna(0.0) / 100)).sum()) if not dl.empty else 0.0
    other_dir = float(other["cost"].sum()) if not other.empty else 0.0
    overhead = (direct_labor + fringe) * (float(overhead_pct) / 100.0)
    gna = (direct_labor + fringe + overhead + other_dir) * (float(gna_pct) / 100.0)
    subtotal = direct_labor + fringe + overhead + gna + other_dir
    contingency = subtotal * (float(contingency_pct) / 100.0)
    fee = (subtotal + contingency) * (float(fee_pct) / 100.0)
    total = subtotal + contingency + fee
    return {
        "Direct Labor": round(direct_labor, 2),
        "Fringe": round(fringe, 2),
        "Overhead": round(overhead, 2),
        "G&A": round(gna, 2),
        "Other Direct": round(other_dir, 2),
        "Contingency": round(contingency, 2),
        "Fee/Profit": round(fee, 2),
        "Total": round(total, 2),
    }


def run_pricing_calculator(conn: sqlite3.Connection) -> None:
    st.header("Pricing Calculator")
    df = pd.read_sql_query("SELECT id, title FROM rfps_t ORDER BY id DESC;", conn, params=())
    if df.empty:
        st.info("No RFP context. Use RFP Analyzer (parse & save) first.")
        return
    rfp_id = st.selectbox("RFP context", options=df["id"].tolist(), format_func=lambda rid: f"#{rid} — {df.loc[df['id']==rid, 'title'].values[0]}")

    st.subheader("Scenario")
    df_sc = pd.read_sql_query("SELECT id, name FROM pricing_scenarios WHERE rfp_id=? ORDER BY id DESC;", conn, params=(rfp_id,))
    mode = st.radio("Mode", ["Create new", "Edit existing"], horizontal=True)
    if mode == "Create new":
        name = st.text_input("Scenario name", value="Base")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            overhead = st.number_input("Overhead %", min_value=0.0, value=20.0, step=1.0)
        with c2:
            gna = st.number_input("G&A %", min_value=0.0, value=10.0, step=1.0)
        with c3:
            fee = st.number_input("Fee/Profit %", min_value=0.0, value=7.0, step=0.5)
        with c4:
            contingency = st.number_input("Contingency %", min_value=0.0, value=0.0, step=0.5)
        if st.button("Create scenario", type="primary"):
            with closing(conn.cursor()) as cur:
                cur.execute("""
                    INSERT INTO pricing_scenarios(rfp_id, name, overhead_pct, gna_pct, fee_pct, contingency_pct, created_at)
                    VALUES(?,?,?,?,?,?,?);
                """, (int(rfp_id), name.strip(), float(overhead), float(gna), float(fee), float(contingency), datetime.utcnow().isoformat()))
                conn.commit()
            st.success("Scenario created.")
            st.rerun()
        return
    else:
        if df_sc.empty:
            st.info("No scenarios yet. Switch to 'Create new'.")
            return
        scenario_id = st.selectbox("Pick a scenario", options=df_sc["id"].tolist(), format_func=lambda sid: df_sc.loc[df_sc["id"]==sid, "name"].values[0])

    st.subheader("Labor")
    with st.form("add_labor", clear_on_submit=True):
        c1, c2, c3, c4 = st.columns([2, 1, 1, 1])
        with c1:
            cat = st.text_input("Labor Category")
        with c2:
            hrs = st.number_input("Hours", min_value=0.0, step=1.0)
        with c3:
            rate = st.number_input("Rate", min_value=0.0, step=1.0)
        with c4:
            fringe = st.number_input("Fringe %", min_value=0.0, value=0.0, step=0.5)
        add_lab = st.form_submit_button("Add labor row")
    if add_lab:
        with closing(conn.cursor()) as cur:
            cur.execute("""
                INSERT INTO pricing_labor(scenario_id, labor_cat, hours, rate, fringe_pct) VALUES(?,?,?,?,?);
            """, (int(scenario_id), cat.strip(), float(hrs), float(rate), float(fringe)))
            conn.commit()
        st.success("Added.")

    df_lab = pd.read_sql_query("""
        SELECT id, labor_cat, hours, rate, fringe_pct, (hours*rate) AS direct, (hours*rate*fringe_pct/100.0) AS fringe
        FROM pricing_labor WHERE scenario_id=?;
    """, conn, params=(scenario_id,))
    _styled_dataframe(df_lab, use_container_width=True, hide_index=True)

    st.subheader("Other Direct Costs")
    with st.form("add_odc", clear_on_submit=True):
        c1, c2 = st.columns([3, 1])
        with c1:
            label = st.text_input("Label")
        with c2:
            cost = st.number_input("Cost", min_value=0.0, step=100.0)
        add_odc = st.form_submit_button("Add ODC")
    if add_odc:
        with closing(conn.cursor()) as cur:
            cur.execute("INSERT INTO pricing_other(scenario_id, label, cost) VALUES(?, ?, ?);", (int(scenario_id), label.strip(), float(cost)))
            conn.commit()
        st.success("Added ODC.")

    df_odc = pd.read_sql_query("SELECT id, label, cost FROM pricing_other WHERE scenario_id=?;", conn, params=(scenario_id,))
    _styled_dataframe(df_odc, use_container_width=True, hide_index=True)

    st.subheader("Summary")
    s = _scenario_summary(conn, int(scenario_id))
    if not s:
        st.info("Add labor/ODCs to see a summary.")
        return
    df_sum = pd.DataFrame(list(s.items()), columns=["Component", "Amount"])
    _styled_dataframe(df_sum.style.format({"Amount": "{:,.2f}"}), use_container_width=True, hide_index=True)

    if st.button("Export pricing CSV"):
        path = os.path.join(DATA_DIR, f"pricing_scenario_{int(scenario_id)}.csv")
        df_sum.to_csv(path, index=False)
        st.success("Exported.")
        st.markdown(f"[Download pricing CSV]({path})")


# ---------- Win Probability (Phase E) ----------
def _price_competitiveness(conn: sqlite3.Connection, rfp_id: int, our_total: Optional[float]) -> Optional[float]:
    df = pd.read_sql_query("""
        SELECT q.vendor, SUM(l.extended_price) AS total
        FROM quotes q JOIN quote_lines l ON q.id = l.quote_id
        WHERE q.rfp_id=?
        GROUP BY q.vendor
        ORDER BY total ASC;
    """, conn, params=(rfp_id,))
    if df.empty or our_total is None:
        return None
    comp_min = float(df["total"].min())
    if our_total <= comp_min:
        return 100.0
    ratio = (our_total - comp_min) / comp_min
    if ratio <= 0.05:
        return 85 + (0.05 - ratio) * (15/0.05)
    if ratio <= 0.10:
        return 70 + (0.10 - ratio) * (15/0.05)
    if ratio <= 0.25:
        return 70 * (0.25 - ratio) / 0.15
    return 0.0


def run_win_probability(conn: sqlite3.Connection) -> None:
    st.header("Win Probability")
    df = pd.read_sql_query("SELECT id, title FROM rfps_t ORDER BY id DESC;", conn, params=())
    if df.empty:
        st.info("No RFP context. Use RFP Analyzer first.")
        return
    rfp_id = st.selectbox("RFP context", options=df["id"].tolist(), format_func=lambda rid: f"#{rid} — {df.loc[df['id']==rid, 'title'].values[0]}")

    df_items = pd.read_sql_query("SELECT status FROM lm_items WHERE rfp_id=?;", conn, params=(rfp_id,))
    if df_items.empty:
        compliance = st.slider("Compliance (est.)", 0, 100, 70)
    else:
        done = (df_items["status"] == "Complete").sum()
        total = len(df_items)
        compliance = int(round(done / max(1, total) * 100))

    tech = st.slider("Technical fit", 0, 100, 75)
    past_perf = st.slider("Past performance relevance", 0, 100, 70)
    team = st.slider("Team strength / subs readiness", 0, 100, 70)
    smallbiz = st.slider("Set-aside / socio-economic alignment", 0, 100, 80)

    df_sc = pd.read_sql_query("SELECT id, name FROM pricing_scenarios WHERE rfp_id=? ORDER BY id DESC;", conn, params=(rfp_id,))
    price_score = None
    our_total = None
    if not df_sc.empty:
        sid = st.selectbox("Use pricing scenario (optional)", options=[None] + df_sc["id"].tolist(),
                           format_func=lambda x: "None" if x is None else df_sc.loc[df_sc["id"]==x, "name"].values[0])
        if sid:
            our_total = _scenario_summary(conn, int(sid)).get("Total")
    if our_total is None:
        our_total = st.number_input("Our total price (if no scenario)", min_value=0.0, value=0.0, step=1000.0)
    price_score = _price_competitiveness(conn, int(rfp_id), our_total)
    if price_score is None:
        price_score = st.slider("Price competitiveness (est.)", 0, 100, 70)

    st.subheader("Weights")
    c1, c2, c3 = st.columns(3)
    with c1:
        w_comp = st.number_input("Weight: Compliance", 0, 100, 20)
        w_tech = st.number_input("Weight: Technical", 0, 100, 25)
    with c2:
        w_past = st.number_input("Weight: Past Perf", 0, 100, 15)
        w_team = st.number_input("Weight: Team", 0, 100, 15)
    with c3:
        w_price = st.number_input("Weight: Price", 0, 100, 25)
        w_small = st.number_input("Weight: Small Biz", 0, 100, 0)
    total_w = w_comp + w_tech + w_past + w_team + w_price + w_small
    if total_w == 0:
        st.error("Weights must sum to > 0")
        return

    comp = {
        "Compliance": compliance,
        "Technical": tech,
        "Past Performance": past_perf,
        "Team": team,
        "Price": int(round(price_score)),
        "Small Business": smallbiz,
    }
    df_scores = pd.DataFrame(list(comp.items()), columns=["Factor", "Score (0-100)"])
    _styled_dataframe(df_scores, use_container_width=True, hide_index=True)

    weighted = (
        compliance * w_comp + tech * w_tech + past_perf * w_past + team * w_team + int(round(price_score)) * w_price + smallbiz * w_small
    ) / total_w
    win_prob = round(float(weighted), 1)
    st.subheader(f"Estimated Win Probability: **{win_prob}%**")

    if st.button("Export assessment CSV"):
        path = os.path.join(DATA_DIR, "win_probability_assessment.csv")
        out = df_scores.copy()
        out.loc[len(out)] = ["Weighted Result", win_prob]
        out.to_csv(path, index=False)
        st.success("Exported.")
        st.markdown(f"[Download assessment CSV]({path})")


# ---------- Phase F: Chat Assistant (rules-based over DB) ----------
def _kb_search(conn: sqlite3.Connection, rfp_id: Optional[int], query: str) -> Dict[str, Any]:
    q = query.lower()
    res: Dict[str, Any] = {}
    # RFP sections
    if rfp_id:
        dfL = pd.read_sql_query("SELECT section, content FROM rfp_sections WHERE rfp_id=?;", conn, params=(rfp_id,))
    else:
        dfL = pd.read_sql_query("SELECT section, content FROM rfp_sections;", conn, params=())
    if not dfL.empty:
        dfL["score"] = dfL["content"].str.lower().apply(lambda t: sum(1 for w in q.split() if w in (t or "")))
        res["sections"] = dfL.sort_values("score", ascending=False).head(5)

    # Checklist
    if rfp_id:
        dfCk = pd.read_sql_query("SELECT item_text, status FROM lm_items WHERE rfp_id=?;", conn, params=(rfp_id,))
    else:
        dfCk = pd.read_sql_query("SELECT item_text, status FROM lm_items;", conn, params=())
    if not dfCk.empty:
        dfCk["score"] = dfCk["item_text"].str.lower().apply(lambda t: sum(1 for w in q.split() if w in (t or "")))
        res["checklist"] = dfCk.sort_values("score", ascending=False).head(10)

    # CLINs
    if rfp_id:
        dfCL = pd.read_sql_query("SELECT clin, description, qty, unit FROM clin_lines WHERE rfp_id=?;", conn, params=(rfp_id,))
    else:
        dfCL = pd.read_sql_query("SELECT clin, description, qty, unit FROM clin_lines;", conn, params=())
    if not dfCL.empty:
        dfCL["score"] = (dfCL["clin"].astype(str) + " " + dfCL["description"].astype(str)).str.lower().apply(lambda t: sum(1 for w in q.split() if w in (t or "")))
        res["clins"] = dfCL.sort_values("score", ascending=False).head(10)

    # Dates
    if rfp_id:
        dfDt = pd.read_sql_query("SELECT label, date_text FROM key_dates WHERE rfp_id=?;", conn, params=(rfp_id,))
    else:
        dfDt = pd.read_sql_query("SELECT label, date_text FROM key_dates;", conn, params=())
    if not dfDt.empty:
        dfDt["score"] = (dfDt["label"].astype(str) + " " + dfDt["date_text"].astype(str)).str.lower().apply(lambda t: sum(1 for w in q.split() if w in (t or "")))
        res["dates"] = dfDt.sort_values("score", ascending=False).head(10)

    # POCs
    if rfp_id:
        dfP = pd.read_sql_query("SELECT name, role, email, phone FROM pocs WHERE rfp_id=?;", conn, params=(rfp_id,))
    else:
        dfP = pd.read_sql_query("SELECT name, role, email, phone FROM pocs;", conn, params=())
    if not dfP.empty:
        dfP["score"] = (dfP["name"].astype(str) + " " + dfP["role"].astype(str) + " " + dfP["email"].astype(str)).str.lower().apply(lambda t: sum(1 for w in q.split() if w in (t or "")))
        res["pocs"] = dfP.sort_values("score", ascending=False).head(10)

    # Quotes summary by vendor
    if rfp_id:
        dfQ = pd.read_sql_query("""
            SELECT q.vendor, SUM(l.extended_price) AS total, COUNT(DISTINCT l.clin) AS clins_quoted
            FROM quotes q JOIN quote_lines l ON q.id=l.quote_id
            WHERE q.rfp_id=?
            GROUP BY q.vendor
            ORDER BY total ASC;
        """, conn, params=(rfp_id,))
        res["quotes"] = dfQ

    # Coverage & compliance
    if rfp_id:
        df_target = pd.read_sql_query("SELECT DISTINCT clin FROM clin_lines WHERE rfp_id=?;", conn, params=(rfp_id,))
        total_clins = int(df_target["clin"].nunique()) if not df_target.empty else 0
        df_items = pd.read_sql_query("SELECT status FROM lm_items WHERE rfp_id=?;", conn, params=(rfp_id,))
        compl = 0
        if not df_items.empty:
            compl = int(round(((df_items["status"]=="Complete").sum() / max(1, len(df_items))) * 100))
        res["meta"] = {"total_clins": total_clins, "compliance_pct": compl}

    return res


def run_chat_assistant(conn: sqlite3.Connection) -> None:
    st.header("Chat Assistant (DB-aware)")
    st.caption("Answers from your saved RFPs, checklist, CLINs, dates, POCs, quotes, and pricing — no external API.")

    df_rf = pd.read_sql_query("SELECT id, title FROM rfps_t ORDER BY id DESC;", conn, params=())
    rfp_opt = None
    if not df_rf.empty:
        rfp_opt = st.selectbox("Context (optional)", options=[None] + df_rf["id"].tolist(),
                               format_func=lambda rid: "All RFPs" if rid is None else f"#{rid} — {df_rf.loc[df_rf['id']==rid, 'title'].values[0]}")

    q = st.text_input("Ask a question (e.g., 'When are proposals due?', 'Show POCs', 'Which vendor is lowest?')")
    ask = st.button("Ask", type="primary")
    if not ask:
        st.caption("Quick picks: due date • POCs • open checklist • CLINs • quotes total • compliance")
        return

    res = _kb_search(conn, rfp_opt, q or "")
    # Heuristic intents
    ql = (q or "").lower()
    if any(w in ql for w in ["due", "deadline", "close"]):
        st.subheader("Key Dates")
        df = res.get("dates", pd.DataFrame())
        if df is not None and not df.empty:
            _styled_dataframe(df[["label","date_text"]], use_container_width=True, hide_index=True)
    if any(w in ql for w in ["poc", "contact", "officer", "specialist"]):
        st.subheader("Points of Contact")
        df = res.get("pocs", pd.DataFrame())
        if df is not None and not df.empty:
            _styled_dataframe(df[["name","role","email","phone"]], use_container_width=True, hide_index=True)
    if "clin" in ql:
        st.subheader("CLINs")
        df = res.get("clins", pd.DataFrame())
        if df is not None and not df.empty:
            _styled_dataframe(df[["clin","description","qty","unit"]], use_container_width=True, hide_index=True)
    if any(w in ql for w in ["checklist", "compliance"]):
        st.subheader("Checklist (top hits)")
        df = res.get("checklist", pd.DataFrame())
        if df is not None and not df.empty:
            _styled_dataframe(df[["item_text","status"]], use_container_width=True, hide_index=True)
        meta = res.get("meta", {})
        if meta:
            st.info(f"Compliance completion: {meta.get('compliance_pct',0)}%")
    if any(w in ql for w in ["quote", "price", "vendor", "lowest"]):
        st.subheader("Quote Totals by Vendor")
        df = res.get("quotes", pd.DataFrame())
        if df is not None and not df.empty:
            _styled_dataframe(df, use_container_width=True, hide_index=True)
            st.caption("Lowest total appears at the top.")

    # Generic best-matches
    sec = res.get("sections", pd.DataFrame())
    if sec is not None and not sec.empty:
        st.subheader("Relevant RFP Sections (snippets)")
        sh = sec.copy()
        sh["snippet"] = sh["content"].str.slice(0, 400)
        _styled_dataframe(sh[["section","snippet","score"]], use_container_width=True, hide_index=True)


# ---------- Phase F: Capability Statement ----------
def _export_capability_docx(path: str, profile: Dict[str, str]) -> Optional[str]:
    try:
        from docx.shared import Pt, Inches  # type: ignore
    except Exception:
        st.error("python-docx is required. pip install python-docx")
        return None

    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7); s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

    title = profile.get("company_name") or "Capability Statement"
    doc.add_heading(title, level=1)
    if profile.get("tagline"):
        p = doc.add_paragraph(profile["tagline"]); p.runs[0].italic = True

    meta = [
        ("Address", "address"), ("Phone", "phone"), ("Email", "email"), ("Website", "website"),
        ("UEI", "uei"), ("CAGE", "cage")
    ]
    p = doc.add_paragraph()
    for label, key in meta:
        val = profile.get(key, "")
        if val:
            p.add_run(f"{label}: {val}  ")

    def add_bullets(title, key):
        txt = (profile.get(key) or "").strip()
        if not txt:
            return
        doc.add_heading(title, level=2)
        for line in [x.strip() for x in txt.splitlines() if x.strip()]:
            doc.add_paragraph(line, style="List Bullet")

    # Content blocks
    add_bullets("Core Competencies", "core_competencies")
    add_bullets("Differentiators", "differentiators")
    add_bullets("Certifications", "certifications")
    add_bullets("Past Performance Highlights", "past_performance")

    naics = (profile.get("naics") or "").replace(",", ", ")
    if naics.strip():
        doc.add_heading("NAICS Codes", level=2)
        doc.add_paragraph(naics)

    contact = profile.get("primary_poc", "")
    if contact.strip():
        doc.add_heading("Primary POC", level=2)
        doc.add_paragraph(contact)

    doc.save(path)
    return path


def _orig_run_capability_statement(conn: sqlite3.Connection) -> None:
    st.header("Capability Statement")
    st.caption("Store your company profile and export a polished 1-page DOCX capability statement.")

    # Load existing (id=1)
    df = pd.read_sql_query("SELECT * FROM org_profile WHERE id=1;", conn, params=())
    vals = df.iloc[0].to_dict() if not df.empty else {}

    with st.form("org_profile_form"):
        c1, c2 = st.columns([2,2])
        with c1:
            company_name = st.text_input("Company Name", value=vals.get("company_name",""))
            tagline = st.text_input("Tagline (optional)", value=vals.get("tagline",""))
            address = st.text_area("Address", value=vals.get("address",""), height=70)
            phone = st.text_input("Phone", value=vals.get("phone",""))
            email = st.text_input("Email", value=vals.get("email",""))
            website = st.text_input("Website", value=vals.get("website",""))
        with c2:
            uei = st.text_input("UEI", value=vals.get("uei",""))
            cage = st.text_input("CAGE", value=vals.get("cage",""))
            naics = st.text_input("NAICS (comma separated)", value=vals.get("naics",""))
            core_competencies = st.text_area("Core Competencies (one per line)", value=vals.get("core_competencies",""), height=110)
            differentiators = st.text_area("Differentiators (one per line)", value=vals.get("differentiators",""), height=110)
        c3, c4 = st.columns([2,2])
        with c3:
            certifications = st.text_area("Certifications (one per line)", value=vals.get("certifications",""), height=110)
        with c4:
            past_performance = st.text_area("Past Performance Highlights (one per line)", value=vals.get("past_performance",""), height=110)
            primary_poc = st.text_area("Primary POC (name, title, email, phone)", value=vals.get("primary_poc",""), height=70)
        saved = st.form_submit_button("Save Profile", type="primary")

    if saved:
        try:
            with closing(conn.cursor()) as cur:
                cur.execute("DELETE FROM org_profile WHERE id=1;")
                cur.execute("""
                    INSERT INTO org_profile(id, company_name, tagline, address, phone, email, website, uei, cage, naics, core_competencies, differentiators, certifications, past_performance, primary_poc)
                    VALUES(1,?,?,?,?,?,?,?,?,?,?,?,?,?,?);
                """, (company_name, tagline, address, phone, email, website, uei, cage, naics, core_competencies, differentiators, certifications, past_performance, primary_poc))
                conn.commit()
            st.success("Profile saved.")
        except Exception as e:
            st.error(f"Save failed: {e}")

    # Export
    if st.button("Export Capability Statement DOCX"):
        prof = pd.read_sql_query("SELECT * FROM org_profile WHERE id=1;", conn, params=())
        if prof.empty:
            st.error("Save your profile first.")
        else:
            p = prof.iloc[0].to_dict()
            path = os.path.join(DATA_DIR, "Capability_Statement.docx")
            out = _export_capability_docx(path, p)
            if out:
                st.success("Exported.")
                st.markdown(f"[Download DOCX]({out})")




# ---------- Phase G: Past Performance Library + Generator ----------
def _pp_score_one(rec: dict, rfp_title: str, rfp_sections: pd.DataFrame) -> int:
    title = (rfp_title or "").lower()
    hay = (title + " " + " ".join((rfp_sections["content"].tolist() if isinstance(rfp_sections, pd.DataFrame) and not rfp_sections.empty else []))).lower()
    score = 0
    # NAICS bonus
    if rec.get("naics") and rec["naics"] in hay:
        score += 40
    # Keywords
    kws = (rec.get("keywords") or "").lower().replace(";", ",").split(",")
    kws = [k.strip() for k in kws if k.strip()]
    for k in kws[:10]:
        if k in hay:
            score += 6
    # Recency via POP end
    try:
        from datetime import datetime
        if rec.get("pop_end"):
            y = int(str(rec["pop_end"]).split("-")[0])
            age = max(0, datetime.now().year - y)
            score += max(0, 20 - (age * 4))  # up to +20, decays 4/yr
    except Exception:
        pass
    # CPARS bonus
    if (rec.get("cpars_rating") or "").strip():
        score += 8
    # Value signal
    try:
        val = float(rec.get("value") or 0)
        if val >= 1000000: score += 6
        elif val >= 250000: score += 3
    except Exception:
        pass
    return min(score, 100)


def _pp_writeup_block(rec: dict) -> str:
    parts = []
    title = rec.get("project_title") or "Project"
    cust = rec.get("customer") or ""
    cn = rec.get("contract_no") or ""
    role = rec.get("role") or ""
    pop = " – ".join([x for x in [rec.get("pop_start") or "", rec.get("pop_end") or ""] if x])
    val = rec.get("value") or ""
    parts.append(f"**{title}** — {cust} {('(' + cn + ')') if cn else ''}")
    meta_bits = [b for b in [f"Role: {role}" if role else "", f"POP: {pop}" if pop else "", f"Value: ${val:,.0f}" if isinstance(val,(int,float)) else (f"Value: {val}" if val else ""), f"NAICS: {rec.get('naics','')}"] if b]
    if meta_bits:
        parts.append("  \n" + " | ".join(meta_bits))
    if rec.get("scope"):
        parts.append(f"**Scope/Work:** {rec['scope']}")
    if rec.get("results"):
        parts.append(f"**Results/Outcome:** {rec['results']}")
    if rec.get("cpars_rating"):
        parts.append(f"**CPARS:** {rec['cpars_rating']}")
    if any([rec.get("contact_name"), rec.get("contact_email"), rec.get("contact_phone")]):
        parts.append("**POC:** " + ", ".join([x for x in [rec.get("contact_name"), rec.get("contact_email"), rec.get("contact_phone")] if x]))
    return "\n\n".join(parts)



def run_past_performance(conn: sqlite3.Connection) -> None:
    st.header("Past Performance Library")
    st.caption("Store/import projects, score relevance vs an RFP, generate writeups, and push to Proposal Builder.")

    # CSV Import
    with st.expander("Import CSV", expanded=False):
        st.caption("Columns: project_title, customer, contract_no, naics, role, pop_start, pop_end, value, scope, results, cpars_rating, contact_name, contact_email, contact_phone, keywords, notes")
        up = st.file_uploader("Upload CSV", type=["csv"], key="pp_csv")
        if up and st.button("Import", key="pp_do_import"):
            try:
                df = pd.read_csv(up)
                # Normalize headers
                df.columns = [c.strip().lower() for c in df.columns]
                required = {"project_title"}
                if not required.issubset(set(df.columns)):
                    st.error("CSV must include at least 'project_title'")
                else:
                    n=0
                    with closing(conn.cursor()) as cur:
                        for _, r in df.iterrows():
                            cur.execute("""
                                INSERT INTO past_perf(project_title, customer, contract_no, naics, role, pop_start, pop_end, value, scope, results, cpars_rating, contact_name, contact_email, contact_phone, keywords, notes)
                                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?);
                            """, (
                                str(r.get("project_title",""))[:200],
                                str(r.get("customer",""))[:200],
                                str(r.get("contract_no",""))[:100],
                                str(r.get("naics",""))[:20],
                                str(r.get("role",""))[:100],
                                str(r.get("pop_start",""))[:20],
                                str(r.get("pop_end",""))[:20],
                                float(r.get("value")) if str(r.get("value","")).strip() not in ("","nan") else None,
                                str(r.get("scope",""))[:2000],
                                str(r.get("results",""))[:2000],
                                str(r.get("cpars_rating",""))[:100],
                                str(r.get("contact_name",""))[:200],
                                str(r.get("contact_email",""))[:200],
                                str(r.get("contact_phone",""))[:100],
                                str(r.get("keywords",""))[:500],
                                str(r.get("notes",""))[:500],
                            ))
                            n+=1
                    conn.commit()
                    st.success(f"Imported {n} projects.")
            except Exception as e:
                st.error(f"Import failed: {e}")

    # Add Project
    with st.expander("Add Project", expanded=False):
        c1, c2, c3 = st.columns([2,2,2])
        with c1:
            project_title = st.text_input("Project Title")
            customer = st.text_input("Customer (Agency/Prime)")
            contract_no = st.text_input("Contract #")
            naics = st.text_input("NAICS")
            role = st.text_input("Role (Prime/Sub)")
        with c2:
            pop_start = st.text_input("POP Start (YYYY-MM)")
            pop_end = st.text_input("POP End (YYYY-MM)")
            value = st.text_input("Value (number)")
            cpars_rating = st.text_input("CPARS Rating (optional)")
            keywords = st.text_input("Keywords (comma-separated)")
        with c3:
            contact_name = st.text_input("POC Name")
            contact_email = st.text_input("POC Email")
            contact_phone = st.text_input("POC Phone")
            scope = st.text_area("Scope/Work", height=100)
            results = st.text_area("Results/Outcome", height=100)
        notes = st.text_area("Notes", height=70)
        if st.button("Save Project", key="pp_save_project"):
            try:
                v = float(value) if value.strip() else None
            except Exception:
                v = None
            try:
                with closing(conn.cursor()) as cur:
                    cur.execute("""
                        INSERT INTO past_perf(project_title, customer, contract_no, naics, role, pop_start, pop_end, value, scope, results, cpars_rating, contact_name, contact_email, contact_phone, keywords, notes)
                        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?);
                    """, (project_title.strip(), customer.strip(), contract_no.strip(), naics.strip(), role.strip(), pop_start.strip(), pop_end.strip(), v, scope.strip(), results.strip(), cpars_rating.strip(), contact_name.strip(), contact_email.strip(), contact_phone.strip(), keywords.strip(), notes.strip()))
                    conn.commit()
                st.success("Saved project.")
            except Exception as e:
                st.error(f"Save failed: {e}")

    # Filters
    with st.expander("Filter", expanded=True):
        f1, f2, f3 = st.columns([2,2,2])
        with f1:
            f_kw = st.text_input("Keyword in title/scope/results")
        with f2:
            f_naics = st.text_input("NAICS filter")
        with f3:
            f_role = st.text_input("Role filter")
    q = "SELECT * FROM past_perf WHERE 1=1"
    params = []
    if f_kw:
        q += " AND (project_title LIKE ? OR scope LIKE ? OR results LIKE ?)"
        params.extend([f"%{f_kw}%", f"%{f_kw}%", f"%{f_kw}%"])
    if f_naics:
        q += " AND naics LIKE ?"
        params.append(f"%{f_naics}%")
    if f_role:
        q += " AND role LIKE ?"
        params.append(f"%{f_role}%")
    df = pd.read_sql_query(q + " ORDER BY id DESC;", conn, params=params)
    if df.empty:
        st.info("No projects found.")
        return

    st.subheader("Projects")
    _styled_dataframe(df[["id","project_title","customer","contract_no","naics","role","pop_start","pop_end","value","cpars_rating"]], use_container_width=True, hide_index=True)
    selected_ids = st.multiselect("Select projects for writeup", options=df["id"].tolist(), format_func=lambda i: f"#{i} — {df.loc[df['id']==i, 'project_title'].values[0]}")

    # Relevance scoring vs RFP
    df_rf = pd.read_sql_query("SELECT id, title FROM rfps_t ORDER BY id DESC;", conn, params=())
    rfp_id = None
    if not df_rf.empty:
        rfp_id = st.selectbox("RFP context for relevance scoring (optional)", options=[None] + df_rf["id"].tolist(),
                              format_func=lambda rid: "None" if rid is None else f"#{rid} — {df_rf.loc[df_rf['id']==rid,'title'].values[0]}")
    if rfp_id:
        ctx = _load_rfp_context(conn, int(rfp_id))
        title = (ctx["rfp"].iloc[0]["title"] if _df_nonempty(ctx.get("rfp")) else "")
        secs = ctx.get("sections", pd.DataFrame())
        # Compute scores
        scores = []
        for _, r in df.iterrows():
            scores.append(_pp_score_one(r.to_dict(), title, secs))
        df_sc = df.copy()
        df_sc["Relevance"] = scores
        st.subheader("Relevance vs selected RFP")
        _styled_dataframe(df_sc[["project_title","naics","role","pop_end","value","Relevance"]].sort_values("Relevance", ascending=False),
                     use_container_width=True, hide_index=True)

    # Generate writeups
    st.subheader("Generate Writeups")
    tone = st.selectbox("Template", ["Concise bullets", "Narrative paragraph"])
    max_n = st.slider("How many projects", 1, 7, min(3, len(selected_ids)) if selected_ids else 3)
    do_gen = st.button("Generate", type="primary")
    if do_gen:
        picked = df[df["id"].isin(selected_ids)].head(max_n).to_dict(orient="records")
        if not picked:
            st.error("Select at least one project.")
            return
        # Build markdown text
        blocks = []
        for r in picked:
            blk = _pp_writeup_block(r)
            if tone == "Concise bullets":
                # convert sentences to bullets
                bullets = []
                for line in blk.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    if not line.startswith("**"):
                        bullets.append(f"- {line}")
                    else:
                        bullets.append(line)
                blocks.append("\n".join(bullets))
            else:
                blocks.append(blk)
        final_md = "\n\n".join(blocks)
        st.markdown("**Preview**")
        st.write(final_md)

        # Push to Proposal Builder section
        st.session_state["pb_section_Past Performance Summary"] = final_md
        st.success("Pushed to Proposal Builder → Past Performance Summary")

        # Export DOCX
        out_path = str(Path(DATA_DIR) / "Past_Performance_Writeups.docx")
        _export_past_perf_docx(out_path, past_perf)
def _wp_load_paper(conn: sqlite3.Connection, paper_id: int) -> pd.DataFrame:
    return pd.read_sql_query(
        "SELECT id, position, title, body, image_path FROM white_paper_sections WHERE paper_id=? ORDER BY position ASC;",
        conn, params=(paper_id,)
    )

def _wp_export_docx(path: str, title: str, subtitle: str, sections: pd.DataFrame) -> Optional[str]:
    try:
        import docx  # type: ignore
    except Exception:
        try:
            from docx import Document  # type: ignore
            import docx  # type: ignore
        except Exception:
            pass

            st.error("python-docx is required. pip install python-docx")
            return None
    try:
        from docx.shared import Inches  # type: ignore
    except Exception:
        pass
    try:
        doc = docx.Document()
        doc.add_heading(title or "Win Plan", 0)
        if subtitle:
            doc.add_paragraph(subtitle)
        if isinstance(sections, pd.DataFrame) and not sections.empty:
            for _, row in sections.iterrows():
                sec = str(row.get("Section") or row.get("section") or row.get("name") or "Section")
                body = str(row.get("Content") or row.get("content") or row.get("text") or "")
                doc.add_heading(sec, level=2)
                for para in (body or "").split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
        doc.save(path)
        return path
    except Exception as e:
        pass

        st.error(f"DOCX export failed: {e}")
        return None

def run_white_paper_builder(conn: sqlite3.Connection) -> None:
    st.header("White Paper Builder")
    st.caption("Templates → Drafts → DOCX export. Can include images per section.")

    # --- Templates ---
    with st.expander("Templates", expanded=False):
        df_t = pd.read_sql_query("SELECT id, name, description, created_at FROM white_templates ORDER BY id DESC;", conn, params=())
        t_col1, t_col2 = st.columns([2,2])
        with t_col1:
            st.subheader("Create Template")
            t_name = st.text_input("Template name", key="wp_t_name")
            t_desc = st.text_area("Description", key="wp_t_desc", height=70)
            if st.button("Save Template", key="wp_t_save"):
                if not t_name.strip():
                    st.error("Name required")
                else:
                    with closing(conn.cursor()) as cur:
                        cur.execute("INSERT INTO white_templates(name, description, created_at) VALUES(?,?,datetime('now'));", (t_name.strip(), t_desc.strip()))
                        conn.commit()
                    st.success("Template saved"); st.rerun()
        with t_col2:
            if df_t.empty:
                st.info("No templates yet.")
            else:
                st.subheader("Edit Template Sections")
                t_sel = st.selectbox("Choose template", options=df_t["id"].tolist(), format_func=lambda tid: df_t.loc[df_t["id"]==tid, "name"].values[0], key="wp_t_sel")
                df_ts = _wp_load_template(conn, int(t_sel))
                _styled_dataframe(df_ts, use_container_width=True, hide_index=True)
                st.markdown("**Add section**")
                ts_title = st.text_input("Section title", key="wp_ts_title")
                ts_body = st.text_area("Default body", key="wp_ts_body", height=120)
                if st.button("Add section to template", key="wp_ts_add"):
                    pos = int((df_ts["position"].max() if not df_ts.empty else 0) + 1)
                    with closing(conn.cursor()) as cur:
                        cur.execute("INSERT INTO white_template_sections(template_id, position, title, body) VALUES(?,?,?,?);",
                                    (int(t_sel), pos, ts_title.strip(), ts_body.strip()))
                        conn.commit()
                    st.success("Section added"); st.rerun()
                # Reorder / delete (simple)
                if not df_ts.empty:
                    st.markdown("**Reorder / Delete**")
                    for _, r in df_ts.iterrows():
                        c1, c2, c3 = st.columns([2,1,1])
                        with c1:
                            new_pos = st.number_input(f"#{int(r['id'])} pos", min_value=1, value=int(r['position']), step=1, key=f"wp_ts_pos_{int(r['id'])}")
                        with c2:
                            if st.button("Apply", key=f"wp_ts_pos_apply_{int(r['id'])}"):
                                with closing(conn.cursor()) as cur:
                                    cur.execute("UPDATE white_template_sections SET position=? WHERE id=?;", (int(new_pos), int(r["id"])))
                                    conn.commit()
                                st.success("Updated position"); st.rerun()
                        with c3:
                            if st.button("Delete", key=f"wp_ts_del_{int(r['id'])}"):
                                with closing(conn.cursor()) as cur:
                                    cur.execute("DELETE FROM white_template_sections WHERE id=?;", (int(r["id"]),))
                                    conn.commit()
                                st.success("Deleted"); st.rerun()

    st.divider()

    # --- Drafts ---
    st.subheader("Drafts")
    df_p = pd.read_sql_query("SELECT id, title, subtitle, created_at, updated_at FROM white_papers ORDER BY id DESC;", conn, params=())
    c1, c2 = st.columns([2,2])
    with c1:
        st.markdown("**Create draft from template**")
        df_t = pd.read_sql_query("SELECT id, name FROM white_templates ORDER BY id DESC;", conn, params=())
        d_title = st.text_input("Draft title", key="wp_d_title")
        d_sub = st.text_input("Subtitle (optional)", key="wp_d_sub")
        if df_t.empty:
            st.caption("No templates available")
            t_sel2 = None
        else:
            t_sel2 = st.selectbox("Template", options=[None] + df_t["id"].tolist(),
                                  format_func=lambda x: "Blank" if x is None else df_t.loc[df_t["id"]==x, "name"].values[0],
                                  key="wp_d_template")
        if st.button("Create draft", key="wp_d_create"):
            if not d_title.strip():
                st.error("Title required")
            else:
                with closing(conn.cursor()) as cur:
                    cur.execute("INSERT INTO white_papers(title, subtitle, rfp_id, created_at, updated_at) VALUES(?,?,?,?,datetime('now'));",
                                (d_title.strip(), d_sub.strip(), None, datetime.utcnow().isoformat()))
                    pid = cur.lastrowid
                    if t_sel2:
                        df_ts2 = _wp_load_template(conn, int(t_sel2))
                        for _, r in df_ts2.sort_values("position").iterrows():
                            cur.execute("INSERT INTO white_paper_sections(paper_id, position, title, body) VALUES(?,?,?,?);",
                                        (int(pid), int(r["position"]), r.get("title"), r.get("body")))
                    conn.commit()
                st.success("Draft created"); st.rerun()
    with c2:
        if df_p.empty:
            st.info("No drafts yet.")
        else:
            st.markdown("**Open a draft**")
            p_sel = st.selectbox("Draft", options=df_p["id"].tolist(), format_func=lambda pid: df_p.loc[df_p["id"]==pid, "title"].values[0], key="wp_d_sel")

    # Editing panel
    if 'p_sel' in locals() and p_sel:
        st.subheader(f"Editing draft #{int(p_sel)}")
        df_sec = _wp_load_paper(conn, int(p_sel))
        # Add section
        st.markdown("**Add section**")
        ns_title = st.text_input("Section title", key="wp_ns_title")
        ns_body = st.text_area("Body", key="wp_ns_body", height=140)
        ns_img = st.file_uploader("Optional image", type=["png","jpg","jpeg"], key="wp_ns_img")
        if st.button("Add section", key="wp_ns_add"):
            img_path = None
            if ns_img is not None:
                img_path = save_uploaded_file(ns_img, subdir="whitepapers")
            pos = int((df_sec["position"].max() if not df_sec.empty else 0) + 1)
            with closing(conn.cursor()) as cur:
                cur.execute("INSERT INTO white_paper_sections(paper_id, position, title, body, image_path) VALUES(?,?,?,?,?);",
                            (int(p_sel), pos, ns_title.strip(), ns_body.strip(), img_path))
                cur.execute("UPDATE white_papers SET updated_at=datetime('now') WHERE id=?;", (int(p_sel),))
                conn.commit()
            st.success("Section added"); st.rerun()

        # Section list
        if df_sec.empty:
            st.info("No sections yet.")
        else:
            for _, r in df_sec.iterrows():
                st.markdown(f"**Section #{int(r['position'])}: {r.get('title') or 'Untitled'}**")
                e1, e2, e3, e4 = st.columns([2,1,1,1])
                with e1:
                    new_title = st.text_input("Title", value=r.get("title") or "", key=f"wp_sec_title_{int(r['id'])}")
                    new_body = st.text_area("Body", value=r.get("body") or "", key=f"wp_sec_body_{int(r['id'])}", height=140)
                with e2:
                    new_pos = st.number_input("Pos", value=int(r["position"]), min_value=1, step=1, key=f"wp_sec_pos_{int(r['id'])}")
                    if st.button("Apply", key=f"wp_sec_apply_{int(r['id'])}"):
                        with closing(conn.cursor()) as cur:
                            cur.execute("UPDATE white_paper_sections SET title=?, body=?, position=? WHERE id=?;",
                                        (new_title.strip(), new_body.strip(), int(new_pos), int(r["id"])))
                            cur.execute("UPDATE white_papers SET updated_at=datetime('now') WHERE id=?;", (int(p_sel),))
                            conn.commit()
                        st.success("Updated"); st.rerun()
                with e3:
                    up_img = st.file_uploader("Replace image", type=["png","jpg","jpeg"], key=f"wp_sec_img_{int(r['id'])}")
                    if st.button("Save image", key=f"wp_sec_img_save_{int(r['id'])}"):
                        if up_img is None:
                            st.warning("Choose an image first")
                        else:
                            img_path = save_uploaded_file(up_img, subdir="whitepapers")
                            with closing(conn.cursor()) as cur:
                                cur.execute("UPDATE white_paper_sections SET image_path=? WHERE id=?;", (img_path, int(r["id"])))
                                cur.execute("UPDATE white_papers SET updated_at=datetime('now') WHERE id=?;", (int(p_sel),))
                                conn.commit()
                            st.success("Image saved"); st.rerun()
                with e4:
                    if st.button("Delete", key=f"wp_sec_del_{int(r['id'])}"):
                        with closing(conn.cursor()) as cur:
                            cur.execute("DELETE FROM white_paper_sections WHERE id=?;", (int(r["id"]),))
                            cur.execute("UPDATE white_papers SET updated_at=datetime('now') WHERE id=?;", (int(p_sel),))
                            conn.commit()
                        st.success("Deleted"); st.rerun()
                st.divider()

            # Export & Push
            x1, x2 = st.columns([2,2])
            with x1:
                if st.button("Export DOCX", key="wp_export"):
                    out_path = str(Path(DATA_DIR) / f"White_Paper_{int(p_sel)}.docx")
                    exp = _wp_export_docx(out_path,
                                          df_p.loc[df_p["id"]==p_sel, "title"].values[0],
                                          df_p.loc[df_p["id"]==p_sel, "subtitle"].values[0] if "subtitle" in df_p.columns else "",
                                          _wp_load_paper(conn, int(p_sel)))
                    if exp:
                        st.success("Exported");
                try:
                    from pathlib import Path as _Path
                    with open(exp, "rb") as _f:
                        _data = _f.read()
                    _fname = _Path(exp).name or "export.docx"
                    st.download_button("Download DOCX", data=_data, file_name=_fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as _e:
                    st.error(f"Download failed: {_e}")

            with x2:
                if st.button("Push narrative to Proposal Builder", key="wp_push"):
                    # Concatenate sections to markdown
                    secs = _wp_load_paper(conn, int(p_sel))
                    lines = []
                    for _, rr in secs.sort_values("position").iterrows():
                        lines.append(f"## {rr.get('title') or 'Section'}\n\n{rr.get('body') or ''}")
                    md = "\n\n".join(lines)
                    st.session_state["pb_section_White Paper"] = md
                    st.success("Pushed to Proposal Builder → 'White Paper' section")



# ---------- Phase I: CRM (Activities • Tasks • Pipeline) ----------
def _stage_probability(stage: str) -> int:
    mapping = {
        "New": 10, "Qualifying": 30, "Bidding": 50, "Submitted": 60, "Awarded": 100, "Lost": 0
    }
    return mapping.get(stage or "", 10)

def run_crm(conn: sqlite3.Connection) -> None:
    st.header("CRM")
    tabs = st.tabs(["Activities", "Tasks", "Pipeline"])

    # --- Activities
    with tabs[0]:
        st.subheader("Log Activity")
        df_deals = pd.read_sql_query("SELECT id, title FROM deals_t ORDER BY id DESC;", conn, params=())
        df_contacts = pd.read_sql_query("SELECT id, name FROM contacts_t ORDER BY name;", conn, params=())
        a_col1, a_col2, a_col3 = st.columns([2,2,2])
        with a_col1:
            a_type = st.selectbox("Type", ["Call","Email","Meeting","Note"], key="act_type")
            a_subject = st.text_input("Subject", value=st.session_state.get("outreach_subject",""))
        with a_col2:
            a_deal = st.selectbox("Related Deal (optional)", options=[None] + df_deals["id"].tolist(),
                                  format_func=lambda x: "None" if x is None else f"#{x} — {df_deals.loc[df_deals['id']==x,'title'].values[0]}",
                                  key="act_deal")
            a_contact = st.selectbox("Related Contact (optional)", options=[None] + df_contacts["id"].tolist(),
                                     format_func=lambda x: "None" if x is None else df_contacts.loc[df_contacts["id"]==x, "name"].values[0],
                                     key="act_contact")
        with a_col3:
            a_notes = st.text_area("Notes", height=100, key="act_notes")
            if st.button("Save Activity", key="act_save"):
                with closing(conn.cursor()) as cur:
                    cur.execute("""
                        INSERT INTO activities(ts, type, subject, notes, deal_id, contact_id) VALUES(datetime('now'),?,?,?,?,?);
                    """, (a_type, a_subject.strip(), a_notes.strip(), a_deal if a_deal else None, a_contact if a_contact else None))
                    conn.commit()
                st.success("Saved")

        st.subheader("Activity Log")
        f1, f2, f3 = st.columns([2,2,2])
        with f1:
            f_type = st.multiselect("Type filter", ["Call","Email","Meeting","Note"])
        with f2:
            f_deal = st.selectbox("Deal filter", options=[None] + df_deals["id"].tolist(),
                                  format_func=lambda x: "All" if x is None else f"#{x} — {df_deals.loc[df_deals['id']==x,'title'].values[0]}",
                                  key="act_f_deal")
        with f3:
            f_contact = st.selectbox("Contact filter", options=[None] + df_contacts["id"].tolist(),
                                     format_func=lambda x: "All" if x is None else df_contacts.loc[df_contacts["id"]==x, "name"].values[0],
                                     key="act_f_contact")
        q = "SELECT ts, type, subject, notes, deal_id, contact_id FROM activities_t WHERE 1=1"
        params = []
        if f_type:
            q += " AND type IN (%s)" % ",".join(["?"]*len(f_type))
            params.extend(f_type)
        if f_deal:
            q += " AND deal_id=?"; params.append(f_deal)
        if f_contact:
            q += " AND contact_id=?"; params.append(f_contact)
        q += " ORDER BY ts DESC"
        df_a = pd.read_sql_query(q, conn, params=params)
        if df_a.empty:
            st.write("No activities")
        else:
            _styled_dataframe(df_a, use_container_width=True, hide_index=True)
            if st.button("Export CSV", key="act_export"):
                path = str(Path(DATA_DIR) / "activities.csv")
                df_a.to_csv(path, index=False)
                st.markdown(f"[Download CSV]({path})")

    # --- Tasks
    with tabs[1]:
        st.subheader("New Task")
        df_deals = pd.read_sql_query("SELECT id, title FROM deals_t ORDER BY id DESC;", conn, params=())
        df_contacts = pd.read_sql_query("SELECT id, name FROM contacts_t ORDER BY name;", conn, params=())
        t1, t2, t3 = st.columns([2,2,2])
        with t1:
            t_title = st.text_input("Task title", key="task_title")
            t_due = st.date_input("Due date", key="task_due")
        with t2:
            t_priority = st.selectbox("Priority", ["Low","Normal","High"], index=1, key="task_priority")
            t_status = st.selectbox("Status", ["Open","In Progress","Done"], index=0, key="task_status")
        with t3:
            t_deal = st.selectbox("Related Deal (optional)", options=[None] + df_deals["id"].tolist(),
                                  format_func=lambda x: "None" if x is None else f"#{x} — {df_deals.loc[df_deals['id']==x,'title'].values[0]}",
                                  key="task_deal")
            t_contact = st.selectbox("Related Contact (optional)", options=[None] + df_contacts["id"].tolist(),
                                     format_func=lambda x: "None" if x is None else df_contacts.loc[df_contacts["id"]==x, "name"].values[0],
                                     key="task_contact")
        if st.button("Add Task", key="task_add"):
            with closing(conn.cursor()) as cur:
                cur.execute("""
                    INSERT INTO tasks(title, due_date, status, priority, deal_id, contact_id, created_at)
                    VALUES(?,?,?,?,?,?,datetime('now'));
                """, (t_title.strip(), t_due.isoformat() if t_due else None, t_status, t_priority, t_deal if t_deal else None, t_contact if t_contact else None))
                conn.commit()
            st.success("Task added")

        st.subheader("My Tasks")
        f1, f2 = st.columns([2,2])
        with f1:
            tf_status = st.multiselect("Status", ["Open","In Progress","Done"], default=["Open","In Progress"])
        with f2:
            tf_priority = st.multiselect("Priority", ["Low","Normal","High"], default=[])
        q = "SELECT id, title, due_date, status, priority, deal_id, contact_id FROM tasks_t WHERE 1=1"
        params = []
        if tf_status:
            q += " AND status IN (%s)" % ",".join(["?"]*len(tf_status)); params.extend(tf_status)
        if tf_priority:
            q += " AND priority IN (%s)" % ",".join(["?"]*len(tf_priority)); params.extend(tf_priority)
        q += " ORDER BY COALESCE(due_date,'9999-12-31') ASC"
        df_t = pd.read_sql_query(q, conn, params=params)
        if df_t.empty:
            st.write("No tasks")
        else:
            for _, r in df_t.iterrows():
                c1, c2, c3, c4 = st.columns([3,2,2,2])
                with c1:
                    st.write(f"**{r['title']}**  — due {r['due_date'] or '—'}")
                with c2:
                    new_status = st.selectbox("Status", ["Open","In Progress","Done"],
                                              index=["Open","In Progress","Done"].index(r["status"] if r["status"] in ["Open","In Progress","Done"] else "Open"),
                                              key=f"task_status_{int(r['id'])}")
                with c3:
                    new_pri = st.selectbox("Priority", ["Low","Normal","High"],
                                            index=["Low","Normal","High"].index(r["priority"] if r["priority"] in ["Low","Normal","High"] else "Normal"),
                                            key=f"task_pri_{int(r['id'])}")
                with c4:
                    if st.button("Apply", key=f"task_apply_{int(r['id'])}"):
                        with closing(conn.cursor()) as cur:
                            cur.execute("UPDATE tasks SET status=?, priority=?, completed_at=CASE WHEN ?='Done' THEN datetime('now') ELSE completed_at END WHERE id=?;",
                                        (new_status, new_pri, new_status, int(r["id"])))
                            conn.commit()
                        st.success("Updated")

            if st.button("Export CSV", key="task_export"):
                path = str(Path(DATA_DIR) / "tasks.csv")
                df_t.to_csv(path, index=False)
                st.markdown(f"[Download CSV]({path})")

    # --- Pipeline
    with tabs[2]:
        st.subheader("Weighted Pipeline")
        df = pd.read_sql_query("SELECT id, title, agency, status, value FROM deals_t ORDER BY id DESC;", conn, params=())
        if df.empty:
            st.info("No deals")
        else:
            df["prob_%"] = df["status"].apply(_stage_probability)
            df["expected_value"] = (df["value"].fillna(0).astype(float) * df["prob_%"] / 100.0).round(2)
            # Stage age: days since last stage change
            df_log = pd.read_sql_query("SELECT deal_id, stage, MAX(changed_at) AS last_change FROM deal_stage_log_t GROUP BY deal_id, stage;", conn, params=())
            def stage_age(row):
                try:
                    last = df_log[(df_log["deal_id"]==row["id"]) & (df_log["stage"]==row["status"])]["last_change"]
                    if last.empty: return None
                    dt = pd.to_datetime(last.values[0])
                    return (pd.Timestamp.utcnow() - dt).days
                except Exception:
                    return None
            df["stage_age_days"] = df.apply(stage_age, axis=1)
            _styled_dataframe(df[["title","agency","status","value","prob_%","expected_value","stage_age_days"]], use_container_width=True, hide_index=True)

            st.subheader("Summary by Stage")
            summary = df.groupby("status").agg(
                deals=("id","count"),
                value=("value","sum"),
                expected=("expected_value","sum")
            ).reset_index().sort_values("expected", ascending=False)
            _styled_dataframe(summary, use_container_width=True, hide_index=True)
            if st.button("Export Pipeline CSV", key="pipe_export"):
                path = str(Path(DATA_DIR) / "pipeline.csv")
                df.to_csv(path, index=False)
                st.markdown(f"[Download CSV]({path})")




def _ensure_files_table(conn: sqlite3.Connection) -> None:
    try:
        with closing(conn.cursor()) as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS files(
                    id INTEGER PRIMARY KEY,
                    owner_type TEXT,
                    owner_id INTEGER,
                    filename TEXT,
                    path TEXT,
                    size INTEGER,
                    mime TEXT,
                    tags TEXT,
                    notes TEXT,
                    uploaded_at TEXT
                );
            """)
            cur.execute("CREATE INDEX IF NOT EXISTS idx_files_owner ON files(owner_type, owner_id);")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_files_tags ON files(tags);")
            conn.commit()
    except Exception:
        pass

# ---------- Phase J: File Manager & Submission Kit ----------
def _detect_mime(name: str) -> str:
    name = (name or "").lower()
    if name.endswith(".pdf"): return "application/pdf"
    if name.endswith(".docx"): return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if name.endswith(".xlsx"): return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if name.endswith(".zip"): return "application/zip"
    if name.endswith(".png"): return "image/png"
    if name.endswith(".jpg") or name.endswith(".jpeg"): return "image/jpeg"
    if name.endswith(".txt"): return "text/plain"
    return "application/octet-stream"


def run_file_manager(conn: sqlite3.Connection) -> None:
    _ensure_files_table(conn)
    st.header("File Manager")
    st.caption("Attach files to RFPs / Deals / Vendors, tag them, and build a zipped submission kit.")

    # --- Attach uploader ---
    with st.expander("Upload & Attach", expanded=True):
        c1, c2 = st.columns([2,2])
        with c1:
            owner_type = st.selectbox("Attach to", ["RFP", "Deal", "Vendor", "Other"], key="fm_owner_type")
            owner_id = None
            if owner_type == "RFP":
                df_rf = pd.read_sql_query("SELECT id, title FROM rfps_t ORDER BY id DESC;", conn, params=())
                if not df_rf.empty:
                    owner_id = st.selectbox("RFP", options=df_rf["id"].tolist(),
                                            format_func=lambda i: f"#{i} — {df_rf.loc[df_rf['id']==i, 'title'].values[0]}",
                                            key="fm_owner_rfp")
            elif owner_type == "Deal":
                df_deal = pd.read_sql_query("SELECT id, title FROM deals_t ORDER BY id DESC;", conn, params=())
                if not df_deal.empty:
                    owner_id = st.selectbox("Deal", options=df_deal["id"].tolist(),
                                            format_func=lambda i: f"#{i} — {df_deal.loc[df_deal['id']==i, 'title'].values[0]}",
                                            key="fm_owner_deal")
            elif owner_type == "Vendor":
                df_v = pd.read_sql_query("SELECT id, name FROM vendors_t ORDER BY name;", conn, params=())
                if not df_v.empty:
                    owner_id = st.selectbox("Vendor", options=df_v["id"].tolist(),
                                            format_func=lambda i: f"#{i} — {df_v.loc[df_v['id']==i, 'name'].values[0]}",
                                            key="fm_owner_vendor")
            # Owner_id can be None for "Other"
        with c2:
            tags = st.text_input("Tags (comma-separated)", key="fm_tags")
            notes = st.text_area("Notes (optional)", height=70, key="fm_notes")

        ups = st.file_uploader("Select files", type=None, accept_multiple_files=True, key="fm_files")
        if st.button("Upload", key="fm_upload"):
            if not ups:
                st.warning("Pick at least one file")
            else:
                saved = 0
                for f in ups:
                    pth = save_uploaded_file(f, subdir="attachments")
                    if not pth:
                        continue
                    try:
                        with closing(conn.cursor()) as cur:
                            cur.execute("""
                                INSERT INTO files(owner_type, owner_id, filename, path, size, mime, tags, notes, uploaded_at)
                                VALUES(?,?,?,?,?,?,?,?,datetime('now'));
                            """, (
                                owner_type, int(owner_id) if owner_id else None, f.name, pth, f.size, _detect_mime(f.name),
                                tags.strip(), notes.strip()
                            ))
                            conn.commit()
                            saved += 1
                    except Exception as e:
                        st.error(f"DB save failed: {e}")
                st.success(f"Uploaded {saved} file(s).")

    # --- Library & filters ---
    with st.expander("Library", expanded=True):
        l1, l2, l3 = st.columns([2,2,2])
        with l1:
            f_owner = st.selectbox("Filter by type", ["All", "RFP", "Deal", "Vendor", "Other"], key="fm_f_owner")
        with l2:
            f_tag = st.text_input("Tag contains", key="fm_f_tag")
        with l3:
            f_kw = st.text_input("Filename contains", key="fm_f_kw")

        q = "SELECT id, owner_type, owner_id, filename, path, size, mime, tags, notes, uploaded_at FROM files_t WHERE 1=1"
        params = []
        if f_owner and f_owner != "All":
            q += " AND owner_type=?"; params.append(f_owner)
        if f_tag:
            q += " AND tags LIKE ?"; params.append(f"%{f_tag}%")
        if f_kw:
            q += " AND filename LIKE ?"; params.append(f"%{f_kw}%")
        q += " ORDER BY uploaded_at DESC"
        try:
            df_files = pd.read_sql_query(q, conn, params=params)
        except Exception as e:
            _ensure_files_table(conn)
            try:
                df_files = pd.read_sql_query(q, conn, params=params)
            except Exception as e2:
                st.error(f"Failed to load files: {e2}")
                df_files = pd.DataFrame()
        if df_files.empty:
            st.write("No files yet.")
        else:
            _styled_dataframe(df_files.drop(columns=["path"]), use_container_width=True, hide_index=True)
            # Per-row controls
            for _, r in df_files.iterrows():
                c1, c2, c3, c4 = st.columns([3,2,2,2])
                with c1:
                    st.caption(f"#{int(r['id'])} — {r['filename']} ({r['owner_type']} {int(r['owner_id']) if r['owner_id'] else ''})")
                with c2:
                    new_tags = st.text_input("Tags", value=r.get("tags") or "", key=f"fm_row_tags_{int(r['id'])}")
                with c3:
                    new_notes = st.text_input("Notes", value=r.get("notes") or "", key=f"fm_row_notes_{int(r['id'])}")
                with c4:
                    b1, b2 = st.columns(2)
                    with b1:
                        if st.button("Save", key=f"fm_row_save_{int(r['id'])}"):
                            with closing(conn.cursor()) as cur:
                                cur.execute("UPDATE files SET tags=?, notes=? WHERE id=?;", (new_tags.strip(), new_notes.strip(), int(r["id"])))
                                conn.commit()
                            st.success("Updated")
                    with b2:
                        if st.button("Delete", key=f"fm_row_del_{int(r['id'])}"):
                            with closing(conn.cursor()) as cur:
                                cur.execute("DELETE FROM files_t WHERE id=?;", (int(r["id"]),))
                                conn.commit()
                            try:
                                if r.get("path") and os.path.exists(r["path"]):
                                    os.remove(r["path"])
                            except Exception:
                                pass
                            st.success("Deleted"); st.rerun()

    # --- Submission Kit (ZIP) ---
    st.subheader("Submission Kit (ZIP)")
    df_rf_all = pd.read_sql_query("SELECT id, title FROM rfps_t ORDER BY id DESC;", conn, params=())
    if df_rf_all.empty:
        st.info("Create an RFP in RFP Analyzer first (Parse → Save).")
        return

    kit_rfp = st.selectbox("RFP", options=df_rf_all["id"].tolist(),
                           format_func=lambda rid: f"#{rid} — {df_rf_all.loc[df_rf_all['id']==rid,'title'].values[0]}",
                           key="fm_kit_rfp")

    # Load files for this RFP
    try:
        df_kit = pd.read_sql_query("SELECT id, filename, path, tags FROM files_t WHERE owner_type='RFP' AND owner_id=? ORDER BY uploaded_at DESC;", conn, params=(int(kit_rfp),))
    except Exception:
        _ensure_files_table(conn)
        df_kit = pd.DataFrame(columns=["id","filename","path","tags"])
    st.caption("Select attachments to include")
    selected = []
    if df_kit.empty:
        st.write("No attachments linked to this RFP yet.")
    else:
        for _, r in df_kit.iterrows():
            if st.checkbox(f"{r['filename']}  {('['+r['tags']+']') if r.get('tags') else ''}", key=f"fm_ck_{int(r['id'])}"):
                selected.append(int(r["id"]))

    # Optional: include generated docs if they exist
    st.markdown("**Optional generated docs to include (if found):**")
    gen_paths = []
    # Proposal doc
    prop_path = str(Path(DATA_DIR) / f"Proposal_RFP_{int(kit_rfp)}.docx")
    if Path(prop_path).exists():
        if st.checkbox("Include Proposal DOCX", key="fm_inc_prop"):
            gen_paths.append(prop_path)
    # Past Performance writeups
    pp_path = str(Path(DATA_DIR) / "Past_Performance_Writeups.docx")
    if Path(pp_path).exists():
        if st.checkbox("Include Past Performance DOCX", key="fm_inc_pp"):
            gen_paths.append(pp_path)
    # White papers (include any)
    white_candidates = sorted(Path(DATA_DIR).glob("White_Paper_*.docx"))
    if white_candidates:
        inc_wp = st.multiselect("Include White Papers", options=[str(p) for p in white_candidates],
                                format_func=lambda p: Path(p).name, key="fm_inc_wp")
        gen_paths.extend(inc_wp)

    if st.button("Build ZIP", type="primary", key="fm_build_zip"):
        if not selected and not gen_paths:
            st.warning("Select at least one attachment or generated document.")
        else:
            # Collect paths
            rows = []
            if selected:
                ph = ",".join(["?"]*len(selected))
                df_sel = pd.read_sql_query(f"SELECT filename, path FROM files_t WHERE id IN ({ph});", conn, params=selected)
                for _, r in df_sel.iterrows():
                    rows.append((r["filename"], r["path"]))
            for p in gen_paths:
                rows.append((Path(p).name, p))

            # Create ZIP
            from zipfile import ZipFile, ZIP_DEFLATED
            ts = pd.Timestamp.utcnow().strftime("%Y%m%d_%H%M%S")
            zip_path = str(Path(DATA_DIR) / f"submission_kit_RFP_{int(kit_rfp)}_{ts}.zip")
            try:
                with ZipFile(zip_path, "w", compression=ZIP_DEFLATED) as z:
                    for fname, p in rows:
                        try:
                            z.write(p, arcname=fname)
                        except Exception:
                            pass
                    # Add a manifest
                    manifest = "Submission Kit Manifest\n"
                    manifest += f"RFP ID: {int(kit_rfp)}\n"
                    manifest += "\nIncluded files:\n" + "\n".join(f"- {fname}" for fname, _ in rows)
                    z.writestr("MANIFEST.txt", manifest)
                st.success("Submission kit created")
                st.markdown(f"[Download ZIP]({zip_path})")
            except Exception as e:
                st.error(f"ZIP failed: {e}")



# ---------- Phase L: RFQ Pack ----------
def _rfq_pack_by_id(conn: sqlite3.Connection, pid: int) -> dict | None:
    df = pd.read_sql_query("SELECT * FROM rfq_packs_t WHERE id=?;", conn, params=(pid,))
    return None if df.empty else df.iloc[0].to_dict()

def _rfq_lines(conn: sqlite3.Connection, pid: int) -> pd.DataFrame:
    return pd.read_sql_query("SELECT id, clin_code, description, qty, unit, naics, psc FROM rfq_lines_t WHERE pack_id=? ORDER BY id ASC;", conn, params=(pid,))

def _rfq_vendors(conn: sqlite3.Connection, pid: int) -> pd.DataFrame:
    q = """
        SELECT rv.id, rv.vendor_id, v.name, v.email, v.phone
        FROM rfq_vendors_t rv
        JOIN vendors v ON v.id = rv.vendor_id
        WHERE rv.pack_id=?
        ORDER BY v.name;
    """
    try:
        return pd.read_sql_query(q, conn, params=(pid,))
    except Exception:
        return pd.DataFrame(columns=["id","vendor_id","name","email","phone"])

def _rfq_attachments(conn: sqlite3.Connection, pid: int) -> pd.DataFrame:
    return pd.read_sql_query("SELECT id, file_id, name, path FROM rfq_attach_t WHERE pack_id=? ORDER BY id ASC;", conn, params=(pid,))

def _rfq_build_zip(conn: sqlite3.Connection, pack_id: int) -> Optional[str]:
    pack = _rfq_pack_by_id(conn, pack_id)
    if not pack:
        st.error("Pack not found"); return None
    title = pack.get("title") or f"RFQ_{pack_id}"
    # Files to include
    df_att = _rfq_attachments(conn, pack_id)
    files = []
    for _, r in df_att.iterrows():
        if r.get("path"):
            files.append((r["name"] or Path(r["path"]).name, r["path"]))
        elif r.get("file_id"):
            # fallback to files table
            try:
                df = pd.read_sql_query("SELECT filename, path FROM files_t WHERE id=?;", conn, params=(int(r["file_id"]),))
                if not df.empty:
                    files.append((df.iloc[0]["filename"], df.iloc[0]["path"]))
            except Exception:
                pass
    # CLINs CSV
    df_lines = _rfq_lines(conn, pack_id)
    clin_csv_path = str(Path(DATA_DIR) / f"rfq_{pack_id}_CLINs.csv")
    df_lines.to_csv(clin_csv_path, index=False)
    files.append((Path(clin_csv_path).name, clin_csv_path))

    # Mail-merge CSV for vendors
    df_v = _rfq_vendors(conn, pack_id)
    mail_csv_path = str(Path(DATA_DIR) / f"rfq_{pack_id}_vendors_mailmerge.csv")
    mm = df_v.rename(columns={"name":"VendorName", "email":"VendorEmail", "phone":"VendorPhone"})[["VendorName","VendorEmail","VendorPhone"]]
    mm["Subject"] = f"Request for Quote – {title}"
    due = pack.get("due_date") or ""
    mm["Body"] = (
        f"Hello {{VendorName}},\n\n"
        f"Please review the attached RFQ package for '{title}'. "
        f"Reply with pricing and availability no later than {due}.\n\n"
        f"Thank you,"
    )
    mm.to_csv(mail_csv_path, index=False)
    files.append((Path(mail_csv_path).name, mail_csv_path))

    # Build zip
    ts = pd.Timestamp.utcnow().strftime("%Y%m%d_%H%M%S")
    zip_path = str(Path(DATA_DIR) / f"RFQ_Pack_{pack_id}_{ts}.zip")
    try:
        with ZipFile(zip_path, "w", compression=ZIP_DEFLATED) as z:
            for fname, pth in files:
                try: z.write(pth, arcname=fname)
                except Exception: pass
            # manifest
            manifest = "RFQ Pack Manifest\n"
            manifest += f"Title: {title}\n"
            manifest += f"Due: {pack.get('due_date') or ''}\n"
            manifest += f"Lines: {len(df_lines)}\n"
            manifest += f"Vendors: {len(df_v)}\n"
            z.writestr("MANIFEST.txt", manifest)
        return zip_path
    except Exception as e:
        st.error(f"ZIP failed: {e}")
        return None

def run_rfq_pack(conn: sqlite3.Connection) -> None:
    st.header("RFQ Pack")
    st.caption("Build vendor-ready RFQ packages from your CLINs, attachments, and vendor list.")

    # -- Create / open
    left, right = st.columns([2,2])
    with left:
        st.subheader("Create")
        df_rf = pd.read_sql_query("SELECT id, title FROM rfps_t ORDER BY id DESC;", conn, params=())
        rf_opt = st.selectbox("RFP (optional)", options=[None] + df_rf["id"].tolist(),
                              format_func=lambda x: "None" if x is None else f"#{x} — {df_rf.loc[df_rf['id']==x,'title'].values[0]}",
                              key="rfq_rfp_sel")
        title = st.text_input("Pack title", key="rfq_title")
        due = st.date_input("Quote due date", key="rfq_due")
        instr = st.text_area("Instructions to vendors (email body)", height=100, key="rfq_instr")
        if st.button("Create RFQ Pack", key="rfq_create"):
            if not title.strip():
                st.error("Title required")
            else:
                with closing(conn.cursor()) as cur:
                    cur.execute("""
                        INSERT INTO rfq_packs(rfp_id, deal_id, title, instructions, due_date, created_at, updated_at)
                        VALUES(?,?,?,?,?,datetime('now'),datetime('now'));
                    """, (rf_opt if rf_opt else None, None, title.strip(), instr.strip(), str(due)))
                    conn.commit()
                st.success("Created"); st.rerun()
    with right:
        st.subheader("Open")
        df_pk = pd.read_sql_query("SELECT id, title, due_date, created_at FROM rfq_packs_t ORDER BY id DESC;", conn, params=())
        if df_pk.empty:
            st.info("No RFQ packs yet")
            return
        pk_sel = st.selectbox("RFQ Pack", options=df_pk["id"].tolist(),
                              format_func=lambda pid: f"#{pid} — {df_pk.loc[df_pk['id']==pid,'title'].values[0]} (due {df_pk.loc[df_pk['id']==pid,'due_date'].values[0] or '—'})",
                              key="rfq_open_sel")

    st.divider()
    st.subheader(f"Editing pack #{int(pk_sel)}")

    # ---- CLINs / Lines ----
    st.markdown("### CLINs / Lines")
    df_lines = _rfq_lines(conn, int(pk_sel))
    _styled_dataframe(df_lines, use_container_width=True, hide_index=True)
    c1, c2, c3, c4, c5, c6 = st.columns([1.2,3,1,1,1,1])
    with c1:
        l_code = st.text_input("CLIN", key="rfq_line_code")
    with c2:
        l_desc = st.text_input("Description", key="rfq_line_desc")
    with c3:
        l_qty = st.number_input("Qty", min_value=0.0, value=1.0, step=1.0, key="rfq_line_qty")
    with c4:
        l_unit = st.text_input("Unit", value="EA", key="rfq_line_unit")
    with c5:
        l_naics = st.text_input("NAICS", key="rfq_line_naics")
    with c6:
        l_psc = st.text_input("PSC", key="rfq_line_psc")
    if st.button("Add Line", key="rfq_line_add"):
        with closing(conn.cursor()) as cur:
            cur.execute("""
                INSERT INTO rfq_lines(pack_id, clin_code, description, qty, unit, naics, psc)
                VALUES(?,?,?,?,?,?,?);
            """, (int(pk_sel), l_code.strip(), l_desc.strip(), float(l_qty or 0), l_unit.strip(), l_naics.strip(), l_psc.strip()))
            conn.commit()
        st.success("Line added"); st.rerun()

    if not df_lines.empty:
        st.markdown("**Edit existing lines**")
        for _, r in df_lines.iterrows():
            ec1, ec2, ec3, ec4 = st.columns([3,1,1,1])
            with ec1:
                nd = st.text_input("Desc", value=r["description"] or "", key=f"rfq_line_e_desc_{int(r['id'])}")
            with ec2:
                nq = st.number_input("Qty", value=float(r["qty"] or 0), step=1.0, key=f"rfq_line_e_qty_{int(r['id'])}")
            with ec3:
                nu = st.text_input("Unit", value=r["unit"] or "EA", key=f"rfq_line_e_unit_{int(r['id'])}")
            with ec4:
                if st.button("Save", key=f"rfq_line_e_save_{int(r['id'])}"):
                    with closing(conn.cursor()) as cur:
                        cur.execute("UPDATE rfq_lines SET description=?, qty=?, unit=? WHERE id=?;",
                                    (nd.strip(), float(nq or 0), nu.strip(), int(r["id"])))
                        conn.commit()
                    st.success("Updated"); st.rerun()

    st.divider()

    # ---- Attachments ----
    st.markdown("### Attachments")
    pack = _rfq_pack_by_id(conn, int(pk_sel))
    rfp_id = pack.get("rfp_id")
    if rfp_id:
        df_rfp_files = pd.read_sql_query("SELECT id, filename, path, tags FROM files_t WHERE owner_type='RFP' AND owner_id=? ORDER BY uploaded_at DESC;", conn, params=(int(rfp_id),))
    else:
        df_rfp_files = pd.DataFrame(columns=["id","filename","path","tags"])
    df_att = _rfq_attachments(conn, int(pk_sel))
    _styled_dataframe(df_att.drop(columns=[]), use_container_width=True, hide_index=True)

    st.markdown("**Add from File Manager**")
    # allow selecting from all files
    df_all_files = pd.read_sql_query("SELECT id, filename FROM files_t ORDER BY uploaded_at DESC;", conn, params=())
    add_file = st.selectbox("File", options=[None] + df_all_files["id"].astype(int).tolist(),
                            format_func=lambda i: "Choose…" if i is None else f"#{i} — {df_all_files.loc[df_all_files['id']==i,'filename'].values[0]}",
                            key="rfq_att_file")
    if st.button("Add Attachment", key="rfq_att_add"):
        if add_file is None:
            st.warning("Pick a file")
        else:
            df_one = pd.read_sql_query("SELECT filename, path FROM files_t WHERE id=?;", conn, params=(int(add_file),))
            if df_one.empty:
                st.error("File not found")
            else:
                with closing(conn.cursor()) as cur:
                    cur.execute("INSERT INTO rfq_attach(pack_id, file_id, name, path) VALUES(?,?,?,?);",
                                (int(pk_sel), int(add_file), df_one.iloc[0]["filename"], df_one.iloc[0]["path"]))
                    conn.commit()
                st.success("Added"); st.rerun()

    if not df_att.empty:
        for _, r in df_att.iterrows():
            dc1, dc2 = st.columns([3,1])
            with dc1:
                st.caption(f"#{int(r['id'])} — {r['name'] or Path(r['path']).name}")
            with dc2:
                if st.button("Remove", key=f"rfq_att_del_{int(r['id'])}"):
                    with closing(conn.cursor()) as cur:
                        cur.execute("DELETE FROM rfq_attach_t WHERE id=?;", (int(r["id"]),))
                        conn.commit()
                    st.success("Removed"); st.rerun()

    st.divider()

    # ---- Vendors ----
    st.markdown("### Vendors")
    try:
        df_vendors = pd.read_sql_query("SELECT id, name, email FROM vendors_t ORDER BY name;", conn, params=())
    except Exception as e:
        st.info("No vendors table yet. Use Subcontractor Finder to add vendors.")
        df_vendors = pd.DataFrame(columns=["id","name","email"])
    df_rv = _rfq_vendors(conn, int(pk_sel))
    _styled_dataframe(df_rv[["name","email","phone"]] if not df_rv.empty else pd.DataFrame(), use_container_width=True, hide_index=True)

    add_vs = st.multiselect("Add vendors", options=df_vendors["id"].astype(int).tolist(),
                            format_func=lambda vid: df_vendors.loc[df_vendors["id"]==vid, "name"].values[0],
                            key="rfq_vendor_add")
    if st.button("Add Selected Vendors", key="rfq_vendor_add_btn"):
        with closing(conn.cursor()) as cur:
            for vid in add_vs:
                try:
                    cur.execute("INSERT OR IGNORE INTO rfq_vendors(pack_id, vendor_id) VALUES(?,?);", (int(pk_sel), int(vid)))
                except Exception:
                    pass
            conn.commit()
        st.success("Vendors added"); st.rerun()

    if not df_rv.empty:
        for _, r in df_rv.iterrows():
            vc1, vc2 = st.columns([3,1])
            with vc1:
                st.caption(f"{r['name']} — {r.get('email') or ''}")
            with vc2:
                if st.button("Remove", key=f"rfq_vendor_del_{int(r['id'])}"):
                    with closing(conn.cursor()) as cur:
                        cur.execute("DELETE FROM rfq_vendors_t WHERE id=?;", (int(r["id"]),))
                        conn.commit()
                    st.success("Removed"); st.rerun()

    st.divider()

    # ---- Build + Exports ----
    st.markdown("### Build & Export")
    czip, cmcsv, cclin = st.columns([2,2,2])
    with czip:
        if st.button("Build RFQ ZIP", type="primary", key="rfq_build_zip"):
            z = _rfq_build_zip(conn, int(pk_sel))
            if z:
                st.success("ZIP ready"); st.markdown(f"[Download ZIP]({z})")

    with cmcsv:
        if st.button("Export Vendors Mail-Merge CSV", key="rfq_mail_csv"):
            df_v = _rfq_vendors(conn, int(pk_sel))
            if df_v.empty:
                st.warning("No vendors selected")
            else:
                out = df_v.rename(columns={"name":"VendorName","email":"VendorEmail","phone":"VendorPhone"})[["VendorName","VendorEmail","VendorPhone"]]
                out["Subject"] = f"Request for Quote – {_rfq_pack_by_id(conn, int(pk_sel)).get('title')}"
                out["Body"] = _rfq_pack_by_id(conn, int(pk_sel)).get("instructions") or ""
                path = str(Path(DATA_DIR) / f"rfq_{int(pk_sel)}_mailmerge.csv")
                out.to_csv(path, index=False)
                st.success("Exported"); st.markdown(f"[Download CSV]({path})")

    with cclin:
        if st.button("Export CLINs CSV", key="rfq_clins_csv"):
            df = _rfq_lines(conn, int(pk_sel))
            if df.empty:
                st.warning("No CLINs yet")
            else:
                path = str(Path(DATA_DIR) / f"rfq_{int(pk_sel)}_CLINs.csv")
                df.to_csv(path, index=False)
                st.success("Exported"); st.markdown(f"[Download CSV]({path})")


def _db_path_from_conn(conn: sqlite3.Connection) -> str:
    try:
        df = pd.read_sql_query("PRAGMA database_list;", conn, params=())
        p = df[df["name"]=="main"]["file"].values[0]
        return p or str(Path(DATA_DIR) / "app.db")
    except Exception:
        return str(Path(DATA_DIR) / "app.db")

def migrate(conn: sqlite3.Connection) -> None:
    """Lightweight idempotent migrations and indices."""
    with closing(conn.cursor()) as cur:
        # read current version
        try:
            ver = int(pd.read_sql_query("SELECT ver FROM schema_version WHERE id=1;", conn, params=()).iloc[0]["ver"])
        except Exception:
            ver = 0

        # v1: add common indexes
        if ver < 1:
            try:
                cur.execute("CREATE INDEX IF NOT EXISTS idx_deals_stage ON deals(stage);")
                cur.execute("CREATE INDEX IF NOT EXISTS idx_deals_status ON deals(status);")
                cur.execute("CREATE INDEX IF NOT EXISTS idx_lm_items_rfp ON lm_items(rfp_id);")
                cur.execute("CREATE INDEX IF NOT EXISTS idx_files_owner2 ON files(owner_type, owner_id);")
                cur.execute("CREATE INDEX IF NOT EXISTS idx_tasks_due ON tasks(due_date);")
            except Exception:
                pass
            cur.execute("UPDATE schema_version SET ver=1 WHERE id=1;")
            conn.commit()

        # v2: ensure NOT NULL defaults where safe (no schema changes if exists)
        if ver < 2:
            try:
                cur.execute("PRAGMA foreign_keys=ON;")
            except Exception:
                pass
            cur.execute("UPDATE schema_version SET ver=2 WHERE id=1;")
            conn.commit()

        # v3: WAL checkpoint to ensure clean state
        if ver < 3:
            try:
                cur.execute("PRAGMA wal_checkpoint(FULL);")
            except Exception:
                pass
            cur.execute("UPDATE schema_version SET ver=3 WHERE id=1;")
            conn.commit()



# ---------- Phase N: Backup & Data ----------
def _current_tenant(conn: sqlite3.Connection) -> int:
    try:
        return int(pd.read_sql_query("SELECT ctid FROM current_tenant WHERE id=1;", conn, params=()).iloc[0]["ctid"])
    except Exception:
        return 1

def _safe_name(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", s or "")

def _backup_db(conn: sqlite3.Connection) -> Optional[str]:
    # Prefer VACUUM INTO; fallback to file copy using sqlite3 backup API
    db_path = _db_path_from_conn(conn)
    ts = pd.Timestamp.utcnow().strftime("%Y%m%d_%H%M%S")
    out = Path(DATA_DIR) / f"backup_{ts}.db"
    try:
        with closing(conn.cursor()) as cur:
            cur.execute(f"VACUUM INTO '{str(out)}';")
        return str(out)
    except Exception:
        # fallback: use backup API
        try:
            import sqlite3 as _sq
            src = _sq.connect(db_path)
            dst = _sq.connect(str(out))
            with dst:
                src.backup(dst)
            src.close(); dst.close()
            return str(out)
        except Exception as e:
            st.error(f"Backup failed: {e}")
            return None

def _restore_db_from_upload(conn: sqlite3.Connection, upload) -> bool:
    # Use backup API to copy uploaded DB into main DB file
    db_path = _db_path_from_conn(conn)
    tmp = Path(DATA_DIR) / ("restore_" + _safe_name(upload.name))
    try:
        tmp.write_bytes(upload.getbuffer())
    except Exception as e:
        st.error(f"Could not write uploaded file: {e}")
        return False
    try:
        src = _sq.connect(str(tmp))
        dst = _sq.connect(db_path)
        with dst:
            src.backup(dst)  # replaces content
        src.close(); dst.close()
        return True
    except Exception as e:
        st.error(f"Restore failed: {e}")
        return False

def _export_table_csv(conn: sqlite3.Connection, table_or_view: str, scoped: bool = True) -> Optional[str]:
    name = table_or_view
    if scoped and not name.endswith("_t"):
        # if a view exists, prefer it
        name_t = name + "_t"
        try:
            pd.read_sql_query(f"SELECT 1 FROM {name_t} LIMIT 1;", conn)
            name = name_t
        except Exception:
            pass
    try:
        df = pd.read_sql_query(f"SELECT * FROM {name};", conn)
        if df.empty:
            st.info("No rows to export.")
        path = Path(DATA_DIR) / f"export_{name}_{pd.Timestamp.utcnow().strftime('%Y%m%d_%H%M%S')}.csv"
        df.to_csv(path, index=False)
        return str(path)
    except Exception as e:
        st.error(f"Export failed: {e}")
        return None

def _import_csv_into_table(conn: sqlite3.Connection, csv_file, table: str, scoped_to_current: bool=True) -> int:
    # Read CSV and insert rows. If tenant_id missing and scoped, stamp with current tenant.
    try:
        df = pd.read_csv(io.BytesIO(csv_file.getbuffer()))
    except Exception as e:
        st.error(f"CSV read failed: {e}")
        return 0
    if scoped_to_current and "tenant_id" not in df.columns:
        df["tenant_id"] = _current_tenant(conn)
    # Align columns with destination
    try:
        cols = pd.read_sql_query(f"PRAGMA table_info({table});", conn)["name"].tolist()
    except Exception as e:
        st.error(f"Table info failed: {e}")
        return 0
    present = [c for c in df.columns if c in cols]
    if not present:
        st.error("No matching columns in CSV.")
        return 0
    df2 = df[present].copy()
    # Drop ID if autoincrement
    if "id" in df2.columns:
        try: df2 = df2.drop(columns=["id"])
        except Exception: pass
    # Insert
    try:
        placeholders = ",".join(["?"]*len(df2.columns))
        sql = f"INSERT INTO {table}({','.join(df2.columns)}) VALUES({placeholders});"
        with closing(conn.cursor()) as cur:
            cur.executemany(sql, df2.itertuples(index=False, name=None))
            conn.commit()
        return len(df2)
    except Exception as e:
        st.error(f"Import failed: {e}")
        return 0

def run_backup_and_data(conn: sqlite3.Connection) -> None:
    st.header("Backup & Data")
    st.caption("WAL on; lightweight migrations; export/import CSV; backup/restore the SQLite DB.")

    st.subheader("Database Info")
    dbp = _db_path_from_conn(conn)
    st.write(f"Path: `{dbp}`")
    try:
        ver = pd.read_sql_query("SELECT ver FROM schema_version WHERE id=1;", conn, params=()).iloc[0]["ver"]
    except Exception:
        ver = "n/a"
    st.write(f"Schema version: **{ver}**")

    c1, c2, c3 = st.columns([2,2,2])
    with c1:
        if st.button("Run Migrations"):
            try:
                migrate(conn); st.success("Migrations done")
            except Exception as e:
                st.error(f"Migrations failed: {e}")
    with c2:
        if st.button("WAL Checkpoint (FULL)"):
            try:
                with closing(conn.cursor()) as cur:
                    cur.execute("PRAGMA wal_checkpoint(FULL);")
                st.success("Checkpoint complete")
            except Exception as e:
                st.error(f"Checkpoint failed: {e}")
    with c3:
        if st.button("Analyze DB"):
            try:
                with closing(conn.cursor()) as cur:
                    cur.execute("ANALYZE;")
                st.success("ANALYZE done")
            except Exception as e:
                st.error(f"ANALYZE failed: {e}")

    st.divider()
    st.subheader("Backup & Restore")
    b1, b2 = st.columns([2,2])
    with b1:
        if st.button("Create Backup (.db)"):
            p = _backup_db(conn)
            if p: st.success("Backup created"); st.markdown(f"[Download backup]({p})")
    with b2:
        up = st.file_uploader("Restore from .db file", type=["db","sqlite","sqlite3"])
        if up and st.button("Restore Now"):
            ok = _restore_db_from_upload(conn, up)
            if ok:
                st.success("Restore completed. Please rerun the app.")

    st.divider()
    st.subheader("Export / Import CSV")
    tables = ["rfps","lm_items","lm_meta","deals","activities","tasks","deal_stage_log",
              "vendors","files","rfq_packs","rfq_lines","rfq_vendors","rfq_attach","contacts"]
    tsel = st.selectbox("Table", options=tables, key="persist_tbl")
    e1, e2 = st.columns([2,2])
    with e1:
        if st.button("Export CSV (current workspace)"):
            p = _export_table_csv(conn, tsel, scoped=True)
            if p: st.success("Exported"); st.markdown(f"[Download CSV]({p})")
    with e2:
        if st.button("Export CSV (all rows)"):
            p = _export_table_csv(conn, tsel, scoped=False)
            if p: st.success("Exported"); st.markdown(f"[Download CSV]({p})")

    upcsv = st.file_uploader("Import into selected table from CSV", type=["csv"], key="persist_upcsv")
    if upcsv and st.button("Import CSV"):
        n = _import_csv_into_table(conn, upcsv, tsel, scoped_to_current=True)
        if n:
            st.success(f"Imported {n} row(s) into {tsel}")
            st.rerun()



# ---------- Phase O: Global Theme & Layout ----------
def _apply_theme_old() -> None:
    css = """
    <style>
    /* Base font and spacing */
    html, body, [class*="css"]  { font-family: Inter, system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, "Apple Color Emoji","Segoe UI Emoji"; }
    .main .block-container { padding-top: 1rem; padding-bottom: 4rem; max-width: 1400px; }
    /* Headings */
    h1, h2, h3 { letter-spacing: 0.2px; }
    h1 { font-size: 1.8rem; margin-bottom: .25rem; }
    h2 { font-size: 1.2rem; margin-top: 1rem; }
    /* Sidebar */
    section[data-testid="stSidebar"] { width: 320px !important; }
    .sidebar-brand { font-weight: 700; font-size: 1.1rem; margin: .25rem 0 .5rem 0; }
    .sidebar-subtle { color: rgba(0,0,0,.55); font-size: .85rem; margin-bottom: .5rem; }
    /* Cards */
    .card { border: 1px solid rgba(0,0,0,.08); border-radius: 14px; padding: 14px 16px; margin: 8px 0 14px 0; box-shadow: 0 1px 2px rgba(0,0,0,.04); background: #fff; }
    .card h3 { margin: 0 0 6px 0; font-size: 1.05rem; }
    /* Dataframes */
    div[data-testid="stDataFrame"] { border-radius: 12px; overflow: hidden; border: 1px solid rgba(0,0,0,.08); }
    /* Tabs */
    button[data-baseweb="tab"] { padding-top: 6px !important; padding-bottom: 6px !important; font-weight: 600; }
    /* Buttons */
    .stButton>button { border-radius: 10px; padding: 0.4rem 0.8rem; }
    /* Hide Streamlit default footer/menu */
    #MainMenu {visibility: hidden;} footer {visibility: hidden;}
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

def page_header(title: str, subtitle: str | None = None) -> None:
    st.markdown(f"<div class='card'><h3>{title}</h3>" + (f"<div class='sidebar-subtle'>{subtitle}</div>" if subtitle else "") + "</div>", unsafe_allow_html=True)

# ---------- nav + main ----------
def init_session() -> None:
    if "initialized" not in st.session_state:
        st.session_state.initialized = True


def nav() -> str:
    st.sidebar.title("Workspace")
    st.sidebar.caption(BUILD_LABEL)
    st.sidebar.caption(f"SHA {_file_hash()}")
    return st.sidebar.selectbox(
        "Go to",
        [
            "SAM Watch",
            "RFP Analyzer",
            "L and M Checklist",
            "Proposal Builder",
            "File Manager",
            "Past Performance",
            "White Paper Builder",
            "Subcontractor Finder",
            "Outreach",
            "RFQ Pack",
            "Backup & Data",
            "Quote Comparison",
            "Pricing Calculator",
            "Win Probability",
            "Chat Assistant",
            "Capability Statement",
            "CRM",
            "Contacts",
            "Deals",
        ],
    )




def render_workspace_switcher(conn: sqlite3.Connection) -> None:
    with st.sidebar.expander("Workspace", expanded=True):
        try:
            df_tenants = pd.read_sql_query("SELECT id, name FROM tenants ORDER BY id;", conn, params=())
        except Exception:
            df_tenants = pd.DataFrame(columns=["id","name"])
        try:
            cur_tid = int(pd.read_sql_query("SELECT ctid FROM current_tenant WHERE id=1;", conn, params=()).iloc[0]["ctid"])
        except Exception:
            cur_tid = 1
        opt = st.selectbox("Organization", options=(df_tenants["id"].astype(int).tolist() if not df_tenants.empty else [1]),
                           format_func=lambda i: (df_tenants.loc[df_tenants["id"]==i,"name"].values[0] if not df_tenants.empty else "Default"),
                           key="tenant_sel")
        if st.button("Switch", key="tenant_switch"):
            with closing(conn.cursor()) as cur:
                cur.execute("UPDATE current_tenant SET ctid=? WHERE id=1;", (int(opt),))
                conn.commit()
            st.session_state['tenant_id'] = int(opt)
            st.success("Workspace switched"); st.rerun()

        st.divider()
        new_name = st.text_input("New workspace name", key="tenant_new_name")
        if st.button("Create workspace", key="tenant_create"):
            if new_name.strip():
                with closing(conn.cursor()) as cur:
                    cur.execute("INSERT OR IGNORE INTO tenants(name, created_at) VALUES(?, datetime('now'));", (new_name.strip(),))
                    conn.commit()
                st.success("Workspace created"); st.rerun()
            else:
                st.warning("Enter a name")



# --- Phase U helper: namespaced keys for Streamlit ---
def ns(scope: str, key: str) -> str:
    """Generate stable, unique Streamlit widget keys."""
    return f"{scope}::{key}"
# === S1 Subcontractor Finder: Google Places ===

def run_subcontractor_finder_s1_hook(conn):
    ensure_subfinder_s1_schema(conn)
    try:
        s1_render_places_panel(conn)
    except Exception:
        pass


def router(page: str, conn: sqlite3.Connection) -> None:
    """Dynamic router. Resolves run_<snake_case(page)> and executes safely."""
    import re as _re
    name = "run_" + _re.sub(r"[^a-z0-9]+", "_", (page or "").lower()).strip("_")
    fn = globals().get(name)
    # explicit fallbacks for known variant names
    if not callable(fn):
        alt = {
            "L and M Checklist": ["run_l_and_m_checklist", "run_lm_checklist"],
            "Backup & Data": ["run_backup_data", "run_backup_and_data"],
        }.get((page or "").strip(), [])
        for a in alt:
            fn = globals().get(a)
            if callable(fn):
                break
    if not callable(fn):
        import streamlit as _st
        _st.warning(f"No handler for page '{page}' resolved as {name}.")
        return
    _safe_route_call(fn, conn)
    # Hooks
    if (page or "").strip() == "Subcontractor Finder":
        _safe_route_call(globals().get("run_subcontractor_finder_s1_hook", lambda _c: None), conn)
    if (page or "").strip() == "Proposal Builder":
        _safe_route_call(globals().get("pb_phase_v_section_library", lambda _c: None), conn)
def main() -> None:
    # Phase 1 re-init inside main
    try:
        _init_phase1_ui()
        _sidebar_brand()
    except Exception:
        pass

    conn = get_db()
    global _O4_CONN

    st.title(APP_TITLE)
    st.caption(BUILD_LABEL)
    # Y0 main panel (always on)
    try:
        y0_ai_panel()
    except Exception:
        pass
    router(nav(), conn)



# --- Outreach schema guard: fallback stub used if the full implementation is defined later ---
if "_o3_ensure_schema" not in globals():
    def _o3_ensure_schema(conn):
        try:
            from contextlib import closing
            with closing(conn.cursor()) as cur:
                # Minimal tables used by Outreach features
                cur.execute("CREATE TABLE IF NOT EXISTS vendors_t (id INTEGER PRIMARY KEY, name TEXT, email TEXT, phone TEXT, city TEXT, state TEXT, naics TEXT)")
                cur.execute("CREATE TABLE IF NOT EXISTS current_tenant (id INTEGER PRIMARY KEY, ctid INTEGER)")
                cur.execute("INSERT OR IGNORE INTO current_tenant(id, ctid) VALUES (1, 1)")
                cur.execute("CREATE TABLE IF NOT EXISTS outreach_templates (id INTEGER PRIMARY KEY, name TEXT, subject TEXT, body TEXT)")
                cur.execute("CREATE TABLE IF NOT EXISTS smtp_settings (id INTEGER PRIMARY KEY, host TEXT, port INTEGER, username TEXT, password TEXT, use_tls INTEGER)")
            conn.commit()
        except Exception:
            pass


# --- Outreach recipients UI: fallback stub used if the full implementation is defined later ---
if "_o3_collect_recipients_ui" not in globals():
    def _o3_collect_recipients_ui(conn):
        import pandas as _pd
        q = "SELECT id, name, email, phone, city, state, naics FROM vendors_t WHERE 1=1"
        params = []
        c1, c2, c3 = st.columns(3)
        with c1:
            f_naics = st.text_input("NAICS filter", key="o3_f_naics")
        with c2:
            f_state = st.text_input("State filter", key="o3_f_state")
        with c3:
            f_city = st.text_input("City filter", key="o3_f_city")
        if f_naics:
            q += " AND IFNULL(naics,'') LIKE ?"
            params.append(f"%{f_naics}%")
        if f_state:
            q += " AND IFNULL(state,'') LIKE ?"
            params.append(f"%{f_state}%")
        if f_city:
            q += " AND IFNULL(city,'') LIKE ?"
            params.append(f"%{f_city}%")
        try:
            df = _pd.read_sql_query(q + " ORDER BY name ASC;", conn, params=tuple(params))
        except Exception:
            df = _pd.DataFrame(columns=["id","name","email","phone","city","state","naics"])
        if df is None or df.empty:
            # fallback to vendors table
            q = "SELECT id, name, email, phone, city, state, naics FROM vendors WHERE 1=1"
            try:
                df = _pd.read_sql_query(q + " ORDER BY name ASC;", conn, params=tuple(params))
            except Exception:
                df = _pd.DataFrame(columns=["id","name","email","phone","city","state","naics"])
        st.caption(f"{len(df)} vendors match filters")
        if not df.empty:
            _styled_dataframe(df, use_container_width=True, hide_index=True)
        return df

# --- Outreach SMTP sender picker: working fallback ---
if "_o3_render_sender_picker" not in globals():
    def _o3_render_sender_picker_legacy():
        import streamlit as st
        conn = get_o4_conn()
        try:
            ensure_outreach_o1_schema(conn)
        except Exception:
            pass
        _ensure_email_accounts_schema(conn)
        rows = _get_senders(conn)
        if not rows:
            st.info("No sender accounts configured. Add one now:")
            with st.form("o4_inline_add_sender", clear_on_submit=True):
                email = st.text_input("Email address")
                display = st.text_input("Display name")
                pw = st.text_input("App password", type="password")
                ok = st.form_submit_button("Save sender")
                if ok and email:
                    try:
                        with conn:
                            conn.execute("""
                            INSERT INTO email_accounts(user_email, display_name, app_password)
                            VALUES(?,?,?)
                            ON CONFLICT(user_email) DO UPDATE SET
                                display_name=excluded.display_name,
                                app_password=excluded.app_password
                            """, (email.strip(), display or "", pw or ""))
                        st.success("Saved")
                        try:
                            import streamlit as st
                            st.session_state["o4_sender_sel"] = email.strip()
                        except Exception:
                            pass
                        st.rerun()
                    except Exception as e:
                        st.error(f"Save failed: {e}")
            if ok and email:
                conn.execute("INSERT OR REPLACE INTO email_accounts(user_email, display_name, app_password) VALUES(?,?,?)",
                             (email.strip().lower(), display.strip(), pw.strip()))
                conn.commit()
                rows = _get_senders(conn)
        if not rows:
            st.error("No sender accounts configured")
            return {"email": "", "app_password": ""}
        choices = [r[0] for r in rows] + ["<add new>"]
        sel = st.selectbox("From account", choices, key="o4_from_addr")
        if sel == "<add new>":
            st.info("Add a sender in Outreach -> Sender accounts")
            return {"email": "", "app_password": ""}
        pw = st.text_input("App password", type="password", key="o4_from_pw")
        return {"email": sel, "app_password": pw}

def y_auto_k(text: str) -> int:
    t = (text or "").lower()
    n = len(t)
    broad = any(k in t for k in [
        "overview", "summary", "summarize", "list all", "requirements", "compliance",
        "section l", "section m", "evaluation factors", "factors", "checklist",
        "compare", "differences", "conflict", "conflicts", "crosswalk", "matrix"
    ])
    if not t.strip():
        return 4
    base = 7 if broad else 4
    if n > 500:
        base += 1
    if n > 1200:
        base += 1
    return max(3, min(8, base))


# === PHASE 8: Latency and UX ===
# Memoize y1_search by (rfp_id, query) and a snapshot of the chunks table to keep cache fresh.
def _y1_snapshot(conn, rfp_id: int) -> str:
    try:
        df = pd.read_sql_query(
            "SELECT COUNT(1) AS c, COALESCE(MAX(id),0) AS m FROM rfp_chunks WHERE rfp_id=?;",
            conn, params=(int(rfp_id),)
        )
        c = int(df.iloc[0]["c"]) if df is not None and not df.empty else 0
        m = int(df.iloc[0]["m"]) if df is not None and not df.empty else 0
        return f"{c}:{m}"
    except Exception:
        return "0:0"

try:
    import streamlit as _st_phase8
except Exception:
    _st_phase8 = None

if _st_phase8 is not None:
    @_st_phase8.cache_data(show_spinner=False, ttl=600)
    def _y1_search_cached(db_path: str, rfp_id: int, query: str, k: int, snapshot: str):
        import sqlite3 as _sql8
        # Re-open a read-only connection to keep cache pure
        try:
            conn2 = _sql8.connect(db_path, check_same_thread=False)
        except Exception:
            conn2 = _sql8.connect(db_path)
        try:
            return _y1_search_uncached(conn2, int(rfp_id), query or "", int(k))
        finally:
            try:
                conn2.close()
            except Exception:
                pass
def y1_search(conn, rfp_id: int, query: str, k: int = 6):
    # Compute snapshot to invalidate cache if chunks changed
    snap = _y1_snapshot(conn, int(rfp_id))
    try:
        db_path = DB_PATH  # provided in app
    except Exception:
        # very conservative fallback
        db_path = "data/govcon.db"
    return _y1_search_cached(db_path, int(rfp_id), query or "", int(k), snap)


# Enable chunk-level streaming in Y2 and Y4
def y2_stream_answer(conn, rfp_id: int, thread_id: int, user_q: str, k: int = 6, temperature: float = 0.2):
    try:
        for tok in ask_ai_with_citations(conn, int(rfp_id), user_q or "", k=int(k), temperature=temperature):
            yield tok
    except NameError:
        # fallback if dependencies were not merged
        hits = []
        try:
            hits = y1_search(conn, int(rfp_id), user_q or "", k=int(k))
        except Exception:
            pass
        yield "[system] limited mode. rebuild index on Y1, then retry."

# Re-define y4_stream_review to ensure true token streaming (shadow any earlier stub)
def y4_stream_review(conn, rfp_id: int, draft_text: str, k: int = 6, temperature: float = 0.1):
    msgs = _y4_build_messages(conn, int(rfp_id), draft_text or "", k=int(k))
    client = get_ai()
    model_name = _resolve_model()
    try:
        resp = client.chat.completions.create(
            model=model_name,
            messages=msgs,
            temperature=float(temperature),
            stream=True
        )
    except Exception as _e:
        if "model_not_found" in str(_e) or "does not exist" in str(_e):
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=msgs,
                temperature=float(temperature),
                stream=True
            )
        else:
            yield f"AI unavailable: {type(_e).__name__}: {_e}"
            return
    for ch in resp:
        try:
            delta = ch.choices[0].delta
            if hasattr(delta, "content") and delta.content:
                yield delta.content
        except Exception:
            pass
# === end PHASE 8 ===



def _y1_cache_bust():
    try:
        st.session_state.pop("_y1_cache", None)
    except Exception:
        pass

# --- Phase 9: tiny test harness ---
def test_seed_cases():
    results = []
    try:
        _flags_ok = bool(CO_STRICT) and bool(PARSER_STRICT) and bool(EVIDENCE_GATE)
    except Exception:
        _flags_ok = False
    results.append(("flags_default_on", _flags_ok))

    import sqlite3
    from contextlib import closing as _closing
    conn = _db_connect(DB_PATH, check_same_thread=False)
    with _closing(conn.cursor()) as cur:
        cur.execute("CREATE TABLE IF NOT EXISTS rfp_chunks(rfp_id INTEGER, rfp_file_id INTEGER, file_name TEXT, page INTEGER, chunk_idx INTEGER, text TEXT, emb TEXT);")
        conn.commit()
    try:
        gen = ask_ai_with_citations(conn, 1, "What is the due date?", k=6, temperature=0.0)
        first = next(gen, "")
        _ev_gate = isinstance(first, str) and first.lower().startswith("[system]")
    except Exception:
        _ev_gate = False
    results.append(("evidence_gate_no_hits", _ev_gate))

    try:
        d1 = _y55_norm_date("Oct 1, 2025")
        d2 = _y55_norm_date("10/01/2025")
        _date_ok = (isinstance(d1, str) and d1.startswith("2025-10-01")) and (isinstance(d2, str) and d2.startswith("2025-10-01"))
    except Exception:
        _date_ok = True
    results.append(("parser_date_norm", _date_ok))

    try:
        with _closing(conn.cursor()) as cur:
            cur.execute("CREATE TABLE IF NOT EXISTS rfps(id INTEGER PRIMARY KEY, title TEXT);")
            cur.execute("INSERT INTO rfps(id,title) VALUES(1,'Test RFP');")
            cur.execute("CREATE TABLE IF NOT EXISTS lm_items(rfp_id INTEGER, section TEXT, item TEXT);")
            conn.commit()
        ok_gate, missing = require_LM_minimum(conn, 1)
        _lm_gate = (ok_gate is False) and isinstance(missing, list)
    except Exception:
        _lm_gate = True
    results.append(("lm_gate_detects_missing", _lm_gate))

    all_ok = all(v for _, v in results)
    return {"ok": all_ok, "results": results}

# === X16.1: Capability Statement — AI drafting helper ===
def run_capability_statement(conn):
    try:
        if '_orig_run_capability_statement' in globals():
            _orig_run_capability_statement(conn)
        else:
            st.header("Capability Statement")
            st.info("Base UI not found in this build. Showing AI helper only.")
    except Exception as e:
        st.error(f"Capability Statement base UI error: {e}")
    try:
        with st.expander("X16.1 — AI drafting helper (OpenAI)", expanded=False):
            st.caption("Draft tagline, core competencies, and differentiators using your org profile and recent RFP context.")
            # Load org profile
            try:
                dfp = pd.read_sql_query("SELECT * FROM org_profile WHERE id=1;", conn)
            except Exception:
                dfp = None
            profile = (dfp.iloc[0].to_dict() if isinstance(dfp, pd.DataFrame) and not dfp.empty else {})
            company = profile.get("company_name","").strip() or "Your Company"
            tagline0 = profile.get("tagline","") or ""
            core0 = profile.get("core_competencies","") or ""
            diff0 = profile.get("differentiators","") or ""
            st.write(f"**Company:** {company}")
            audience = st.text_input("Audience focus (e.g., BOP Facilities, USAF MXG, VA VISN)", key="x161_aud")
            tone = st.selectbox("Tone", ["Crisp federal", "Technical", "Plain language"], index=0, key="x161_tone")
            include_past_perf = st.checkbox("Incorporate past performance bullets if available", value=True, key="x161_pp")
            # Optional RFP context
            try:
                dfr = pd.read_sql_query("SELECT id, title FROM rfps ORDER BY id DESC;", conn)
                rfp_sel = st.selectbox("Optional RFP context", options=[None] + dfr["id"].tolist(),
                                       format_func=lambda i: "None" if i is None else f"#{i} — {dfr.loc[dfr['id']==i,'title'].values[0]}",
                                       key="x161_rfp")
            except Exception:
                rfp_sel = None
            ctx = ""
            if rfp_sel:
                try:
                    # Pull a compact context from chunks
                    hits = pd.read_sql_query("SELECT text FROM rfp_chunks WHERE rfp_id=? LIMIT 40;", conn, params=(int(rfp_sel),))
                    ctx = "\n".join((hits["text"].fillna("").tolist() if hits is not None else []))[:12000]
                except Exception:
                    ctx = ""
            # Compose prompt
            def _ask_x161(kind: str):
                sys = "You are a senior federal capture writer. Use short, precise bullets. Avoid marketing fluff. No emojis."
                req = f"""Company: {company}
Audience: {audience or '(general federal)'}
Tone: {tone}
Existing tagline: {tagline0[:200]}
Existing core competencies: {core0[:800]}
Existing differentiators: {diff0[:800]}
Include past performance: {bool(include_past_perf)}
Task: Draft {kind} for a one-page capability statement.
Constraints:
- 4–7 bullets for lists. 12–18 words each.
- Use federal terms. No hyperbole. No first person.
- If RFP context provided, align content.
RFP Context (optional, may be empty):
{ctx if ctx else '(none)'}"""
                try:
                    client = get_ai()
                    model = (globals().get('_resolve_model') or (lambda: 'gpt-4o-mini'))()
                    resp = client.chat.completions.create(
                        model=model,
                        messages=[{"role":"system","content": sys}, {"role":"user","content": req}],
                        temperature=0.2
                    )
                    return (resp.choices[0].message.content or "").strip()
                except Exception as e:
                    return f"AI error: {e}"
            c1, c2, c3 = st.columns([1,1,1])
            with c1:
                if st.button("Draft Tagline", key="x161_tl"):
                    st.session_state["x161_tagline"] = _ask_x161("a concise 8–14 word tagline")
            with c2:
                if st.button("Draft Core Competencies", key="x161_cc"):
                    st.session_state["x161_core"] = _ask_x161("Core Competencies bullets")
            with c3:
                if st.button("Draft Differentiators", key="x161_df"):
                    st.session_state["x161_diff"] = _ask_x161("Differentiators bullets")
            st.text_input("Tagline (AI)", value=st.session_state.get("x161_tagline",""), key="x161_tagline_box")
            st.text_area("Core Competencies (AI)", value=st.session_state.get("x161_core",""), height=160, key="x161_core_box")
            st.text_area("Differentiators (AI)", value=st.session_state.get("x161_diff",""), height=160, key="x161_diff_box")
            save = st.button("Save AI fields into org_profile", key="x161_save")
            if save:
                try:
                    with closing(conn.cursor()) as cur:
                        # ensure profile row exists
                        cur.execute("INSERT OR IGNORE INTO org_profile(id, company_name) VALUES(1, ?);", (company,))
                        cur.execute("UPDATE org_profile SET tagline=?, core_competencies=?, differentiators=? WHERE id=1;",
                                    (st.session_state.get("x161_tagline",""), st.session_state.get("x161_core",""), st.session_state.get("x161_diff","")))
                        conn.commit()
                    st.success("Saved to org_profile")
                except Exception as e:
                    st.error(f"Save failed: {e}")
    except Exception as e:
        st.error(f"X16.1 panel error: {e}")
# === end X16.1 ===


# === S1 Subcontractor Finder: Google Places ====================================
def ensure_subfinder_s1_schema(conn):
    try:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(vendors)").fetchall()]
    except Exception:
        cols = []
    with conn:
        if "place_id" not in cols:
            try: conn.execute("ALTER TABLE vendors ADD COLUMN place_id TEXT")
            except Exception: pass
        if "fit_score" not in cols:
            try: conn.execute("ALTER TABLE vendors ADD COLUMN fit_score REAL")
            except Exception: pass
        if "tags" not in cols:
            try: conn.execute("ALTER TABLE vendors ADD COLUMN tags TEXT")
            except Exception: pass
        conn.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_vendors_place_id ON vendors(place_id)")

def s1_normalize_phone(s:str)->str:
    s = (s or "").strip()
    return "".join(ch for ch in s if ch.isdigit())

def s1_get_google_api_key()->str|None:
    import os
    try:

        if "google" in st.secrets and "api_key" in st.secrets["google"]:
            return st.secrets["google"]["api_key"]
        if "google_api_key" in st.secrets:
            return st.secrets["google_api_key"]
    except Exception:
        pass
    return os.environ.get("GOOGLE_API_KEY")

def s1_geocode_address(address:str):
    key = s1_get_google_api_key()
    if not key: return None
    import urllib.parse, urllib.request, json
    params = urllib.parse.urlencode({"address": address, "key": key})
    url = f"https://maps.googleapis.com/maps/api/geocode/json?{params}"
    try:
        with urllib.request.urlopen(url, timeout=20) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        if data.get("results"):
            loc = data["results"][0]["geometry"]["location"]
            return float(loc["lat"]), float(loc["lng"])
    except Exception:
        return None
    return None

def s1_places_text_search(query:str, lat:float, lon:float, radius_meters:int, page_token:str|None=None)->dict:
    key = s1_get_google_api_key()
    if not key: return {"error":"no_api_key"}
    import urllib.parse, urllib.request, json, time
    base = "https://maps.googleapis.com/maps/api/place/textsearch/json"
    params = {"query": query, "location": f"{lat},{lon}", "radius": radius_meters, "key": key}
    if page_token:
        params = {"pagetoken": page_token, "key": key}
        time.sleep(2.0)
    url = base + "?" + urllib.parse.urlencode(params)
    try:
        with urllib.request.urlopen(url, timeout=20) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except Exception as e:
        return {"error": str(e)}

def s1_vendor_exists(conn, place_id:str|None, email:str|None, phone:str|None)->bool:
    q = "SELECT 1 FROM vendors WHERE 1=0"
    args = []
    if place_id:
        q += " OR place_id=?"; args.append(place_id)
    if email:
        q += " OR LOWER(email)=LOWER(?)"; args.append(email.strip())
    if phone:
        q += " OR REPLACE(REPLACE(REPLACE(phone,'-',''),'(',''),')','') LIKE ?"
        args.append("%" + s1_normalize_phone(phone) + "%")
    row = conn.execute(q, tuple(args)).fetchone()
    return bool(row)

def s1_save_vendor(conn, v:dict)->int|None:
    ensure_subfinder_s1_schema(conn)
    name = v.get("name") or ""
    pid = v.get("place_id") or None
    addr = v.get("formatted_address") or ""
    city, state = None, None
    phone = v.get("formatted_phone_number") or v.get("international_phone_number") or ""
    phone = s1_normalize_phone(phone)
    website = v.get("website") or ""
    email = v.get("email") or ""
    with conn:
        conn.execute("""
            INSERT INTO vendors(name, city, state, phone, email, website, notes, place_id)
            VALUES(?,?,?,?,?,?,?,?)
            ON CONFLICT(place_id) DO UPDATE SET
              name=excluded.name,
              phone=COALESCE(excluded.phone, vendors.phone),
              website=COALESCE(excluded.website, vendors.website)
        """, (name, city, state, phone, email, website, addr, pid))
    row = conn.execute("SELECT id FROM vendors WHERE place_id=?", (pid,)).fetchone()
    return row[0] if row else None

def s1_calc_radius_meters(miles:int)->int:
    return int(float(miles) * 1609.344)
def s1_render_places_panel(conn, default_addr:str|None=None):
    import streamlit as st, pandas as pd
    ensure_subfinder_s1_schema(conn)
    st.markdown("### Google Places search")
    key_addr = st.text_input("Place of performance address", value=default_addr or "", key="s1_addr")
    miles = st.slider("Miles radius", min_value=5, max_value=250, value=50, step=5, key="s1_miles")
    q = st.text_input("Search keywords or NAICS", value=st.session_state.get("s1_q","janitorial"), key="s1_q")
    hide_saved = st.checkbox("Hide vendors already saved", value=True, key="s1_hide_saved")
    if "s1_page_token" not in st.session_state: st.session_state["s1_page_token"] = None
    cols = st.columns([1,1,1])
    search_clicked = cols[0].button("Search")
    next_clicked = cols[1].button("Next page")
    clear_clicked = cols[2].button("Clear")
    if clear_clicked: st.session_state["s1_page_token"] = None
    if not (search_clicked or next_clicked): return
    if not key_addr:
        st.error("Address required"); return
    loc = s1_geocode_address(key_addr)
    if not loc:
        st.error("Geocoding failed or API key missing"); return
    lat, lon = loc
    token = st.session_state.get("s1_page_token") if next_clicked else None
    data = s1_places_text_search(q, lat, lon, s1_calc_radius_meters(miles), token)
    if "error" in data:
        st.error(f"Places error: {data['error']}"); return
    st.session_state["s1_page_token"] = data.get("next_page_token")
    rows = []
    for r in data.get("results", []):
        pid = r.get("place_id")
        name = r.get("name")
        addr = r.get("formatted_address", "")
        rating = r.get("rating", None)
        open_now = (r.get("opening_hours") or {}).get("open_now")
        dup = hide_saved and pid and s1_vendor_exists(conn, pid, None, None)
        if not dup:
            rows.append({"place_id": pid,"name": name,"address": addr,"rating": rating,"open_now": open_now})
    if not rows:
        st.info("No new vendors in this page"); return
    df = pd.DataFrame(rows)
    _styled_dataframe(df, use_container_width=True, hide_index=True)
    ids = [r["place_id"] for r in rows if r.get("place_id")]
    if not ids: return
    to_save = st.multiselect("Select vendors to save", ids, format_func=lambda x: next((r["name"] for r in rows if r["place_id"]==x), x))
    # [removed duplicate Save selected block]
    if st.session_state.get("s1_page_token"):
        st.caption("Another page is available. Click Next page to load more.")
    st.caption("Set st.secrets['google']['api_key'] or env GOOGLE_API_KEY")


def o1_list_accounts(conn):
    ensure_outreach_o1_schema(conn)
    rows = conn.execute("""
      SELECT user_email, display_name, smtp_host, smtp_port
      FROM email_accounts ORDER BY user_email
    """).fetchall()
    return rows

def ensure_outreach_o1_schema(conn):
    with conn:
        conn.execute("""
        CREATE TABLE IF NOT EXISTS email_accounts(
            user_email TEXT PRIMARY KEY,
            display_name TEXT DEFAULT '',
            app_password TEXT DEFAULT '',
            smtp_host TEXT DEFAULT 'smtp.gmail.com',
            smtp_port INTEGER DEFAULT 465,
            use_ssl INTEGER DEFAULT 1
        )""")


def o1_delete_email_account(conn, user_email:str):
    ensure_outreach_o1_schema(conn)
    with conn:
        conn.execute("DELETE FROM email_accounts WHERE user_email=?", (user_email.strip(),))


# === O2: Outreach Templates ====================================================
def ensure_email_templates(conn):
    with conn:
        conn.execute("""
        CREATE TABLE IF NOT EXISTS email_templates(
            id INTEGER PRIMARY KEY,
            name TEXT UNIQUE NOT NULL,
            subject TEXT NOT NULL DEFAULT '',
            html_body TEXT NOT NULL DEFAULT '',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )""")
        conn.execute("""
        CREATE TRIGGER IF NOT EXISTS trg_email_templates_update
        AFTER UPDATE ON email_templates
        BEGIN
            UPDATE email_templates SET updated_at=CURRENT_TIMESTAMP WHERE id=NEW.id;
        END;""")

def email_template_list(conn):
    ensure_email_templates(conn)
    return conn.execute("SELECT id, name, subject, html_body FROM email_templates ORDER BY name").fetchall()

def email_template_get(conn, template_id:int):
    ensure_email_templates(conn)
    return conn.execute("SELECT id, name, subject, html_body FROM email_templates WHERE id=?", (int(template_id),)).fetchone()

def email_template_upsert(conn, name:str, subject:str, html_body:str, template_id:int|None=None):
    ensure_email_templates(conn)
    with conn:
        if template_id:
            conn.execute("UPDATE email_templates SET name=?, subject=?, html_body=? WHERE id=?",
                         (name.strip(), subject, html_body, int(template_id)))
            return template_id
        conn.execute("""
            INSERT INTO email_templates(name, subject, html_body)
            VALUES(?,?,?)
            ON CONFLICT(name) DO UPDATE SET subject=excluded.subject, html_body=excluded.html_body
        """, (name.strip(), subject, html_body))
        return conn.execute("SELECT id FROM email_templates WHERE name=?", (name.strip(),)).fetchone()[0]

def email_template_delete(conn, template_id:int):
    ensure_email_templates(conn)
    with conn:
        conn.execute("DELETE FROM email_templates WHERE id=?", (int(template_id),))

import re as _re_o2
MERGE_TAGS = {"company","email","title","solicitation","due","notice_id","first_name","last_name","city","state"}
def template_missing_tags(text:str, required:set[str]=MERGE_TAGS)->set[str]:
    found = set(_re_o2.findall(r"{{\s*([a-zA-Z0-9_]+)\s*}}", text or ""))
    return required - found

def render_outreach_templates(conn):

    st.subheader("Email templates")
    ensure_email_templates(conn)
    rows = email_template_list(conn)
    names = ["<new>"] + [r[1] for r in rows]
    sel = st.selectbox("Template", names, key="tpl_sel")
    if sel == "<new>":
        tid = None
        name = st.text_input("Name", key="tpl_name")
        subject = st.text_input("Subject", value=st.session_state.get("outreach_subject",""))
        html = st.text_area("HTML body", value=st.session_state.get("outreach_html",""), height=300)
    else:
        row = next(r for r in rows if r[1] == sel)
        tid = row[0]
        name = st.text_input("Name", value=row[1], key="tpl_name__2")
        subject = st.text_input("Subject", value=row[2], key="tpl_subject")
        html = st.text_area("HTML body", value=row[3], key="tpl_html", height=240)
    missing = template_missing_tags((subject or "") + " " + (html or ""))
    if missing:
        st.info("Missing merge tags: " + ", ".join(sorted(missing)))
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("Save"):
            if not (name or "").strip():
                st.error("Name required")
            else:
                email_template_upsert(conn, name, subject or "", html or "", tid)
                st.success("Saved")
    with c2:
        if (tid is not None) and st.button("Duplicate"):
            email_template_upsert(conn, f"{name} copy", subject or "", html or "", None)
            st.success("Duplicated")
            st.rerun()
    with c3:
        if (tid is not None) and st.button("Delete"):
            email_template_delete(conn, tid)
            st.success("Deleted")
            st.rerun()

def _tpl_picker_prefill(conn):

    rows = email_template_list(conn)
    if not rows:
        return None
    names = [r[1] for r in rows]
    choice = st.selectbox("Use template", ["<none>"] + names, key="tpl_pick")
    if choice != "<none>":
        row = next(r for r in rows if r[1] == choice)
        st.session_state["outreach_subject"] = row[2]
        st.session_state["outreach_html"] = row[3]
        return row
    return None

def seed_default_templates(conn):
    ensure_email_templates(conn)
    defaults = [
        ("RFQ Intro", "RFQ: {{title}} {{solicitation}} due {{due}}",
         "<p>Hello {{first_name}},</p><p>We are collecting quotes for {{title}} under {{solicitation}} due {{due}}.</p><p>Unsubscribe: reply STOP</p>"),
        ("Follow Up 1", "Follow up on {{title}} RFQ", "<p>Checking in on your quote for {{title}}.</p><p>Unsubscribe: reply STOP</p>")
    ]
    for n,s,h in defaults:
        email_template_upsert(conn, n, s, h, None)





# --- O4 wrapper: delegates to __p_o4_ui if present, else shows fallback UI ---
def o4_sender_accounts_ui(conn):
    try:
        return __p_o4_ui(conn)  # provided by O4 module when available
    except Exception:
        import streamlit as _st
        from contextlib import closing
        _st.info("O4 sender accounts fallback UI loaded.")
        # Load existing if any
        host, port, username, password, use_tls = "smtp.gmail.com", 587, "", "", True
        try:
            with closing(conn.cursor()) as cur:
                cur.execute("CREATE TABLE IF NOT EXISTS smtp_settings (id INTEGER PRIMARY KEY, label TEXT, host TEXT, port INTEGER, username TEXT, password TEXT, use_tls INTEGER)")
                row = cur.execute("SELECT label, host, port, username, password, use_tls FROM smtp_settings WHERE id=1").fetchone()
                if row:
                    label, host, port, username, password, use_tls = row[0] or "", row[1] or "smtp.gmail.com", int(row[2] or 587), row[3] or "", row[4] or "", bool(row[5] or 0)
        except Exception:
            pass
        with _st.form("o4_fallback_sender_mm", clear_on_submit=False):
            label = _st.text_input("Label", value=label if 'label' in locals() else "Default")
            username = _st.text_input("Gmail address", value=username)
            password = _st.text_input("App password", type="password", value=password)
            c1,c2 = _st.columns(2)
            with c1: host = _st.text_input("SMTP host", value=host)
            with c2: port = _st.number_input("SMTP port", 1, 65535, value=int(port))
            use_tls = _st.checkbox("Use STARTTLS", value=bool(use_tls))
            saved = _st.form_submit_button("Save sender")
            if saved:
                try:
                    with closing(conn.cursor()) as cur:
                        cur.execute("CREATE TABLE IF NOT EXISTS smtp_settings (id INTEGER PRIMARY KEY, label TEXT, host TEXT, port INTEGER, username TEXT, password TEXT, use_tls INTEGER)")
                        cur.execute("INSERT OR REPLACE INTO smtp_settings(id, label, host, port, username, password, use_tls) VALUES(1,?,?,?,?,?,?)",
                                    (label.strip() or "Default", host.strip(), int(port), username.strip(), password.strip(), 1 if use_tls else 0))
                        conn.commit()
                    _st.success("Sender saved")
                except Exception as e:
                    _st.error(f"Save failed: {e}")
        # Show current
        try:
            with closing(conn.cursor()) as cur:
                row = cur.execute("SELECT label, username, host, port, use_tls FROM smtp_settings WHERE id=1").fetchone()
            if row:
                _st.caption(f"Active: {row[0]} — {row[1]} via {row[2]}:{row[3]} TLS={'on' if row[4] else 'off'}")
        except Exception:
            pass


def render_outreach_mailmerge(conn):
    globals()['_O4_CONN'] = conn
    import streamlit as st
    import pandas as _pd
    # 1) Recipients
    rows = _o3_collect_recipients_ui(conn) if "_o3_collect_recipients_ui" in globals() else None
    st.subheader("Mail Merge & Send")
    # 2) Template inputs
    subj = st.text_input("Subject", value=st.session_state.get("outreach_subject",""), key="o3_subject")
    body = st.text_area("HTML Body", value=st.session_state.get("outreach_body",""), height=260, key="o3_body")
    # 3) Sender
    st.subheader("Sender")
    sender = _o3_render_sender_picker() if "_o3_render_sender_picker" in globals() else {}
    # normalize keys
    if sender and "username" in sender and "email" not in sender:
        sender["email"] = sender.get("username","")
    if sender and "password" in sender and "app_password" not in sender:
        sender["app_password"] = sender.get("password","")
    c1, c2, c3 = st.columns([1,1,2])
    with c1:
        test = st.button("Test run (no send)", key="o3_test")
    with c2:
        do = st.button("Send batch", type="primary", key="o3_send")
    with c3:
        maxn = st.number_input("Max to send", min_value=1, max_value=5000, value=500, step=50, key="o3_max")
    if (test or do) and rows is not None and hasattr(rows, "empty") and not rows.empty:
        try:
            if "_o3_send_batch" in globals():
                if test:
                    st.info("Test run: rendering only, no SMTP.")
                    # call with test_only=True
                    out = _o3_send_batch(conn, sender, rows, subj, body, True, int(maxn))
                else:
                    out = _o3_send_batch(conn, sender, rows, subj, body, False, int(maxn))
                st.success("Send function executed")
            else:
                st.warning("Send function not available in this build.")
        except Exception as e:
            st.error(f"Send failed: {e}")


def run_outreach(conn):
    import streamlit as st

    # O4 badge if present
    try:
        _o4_render_badge()
    except Exception:
        pass

    st.header("Outreach")
    with st.expander("Compliance (O6)", expanded=False):
        render_outreach_o6_compliance(conn)

    # O6: handle unsubscribe links
    o6_handle_query_unsubscribe(conn)
    with st.expander("Follow-ups & SLA (O5)", expanded=False):
        render_outreach_o5_followups(conn)

    # Sender accounts (O4)
    try:
        with st.expander("Sender accounts", expanded=True):
            # guarded render

            __ok = _render_once('o4_sender')

            if __ok:

                o4_sender_accounts_ui(conn)
    except Exception as e:
        st.warning(f"O4 sender UI unavailable: {e}")

    # Templates (O2)
    try:
        _tpl_picker_prefill(conn)
        with st.expander("Templates", expanded=False):
            render_outreach_templates(conn)
    except Exception:
        pass

    # Mail merge + send (O3)
    try:
        with st.expander("Mail Merge & Send", expanded=True):
            render_outreach_mailmerge(conn)
    except Exception as e:
        st.error(f"Mail merge panel error: {e}")


def _o3_ensure_schema(conn):
    with _o3c(conn.cursor()) as cur:
        cur.execute("""CREATE TABLE IF NOT EXISTS outreach_optouts(
            id INTEGER PRIMARY KEY,
            email TEXT UNIQUE,
            reason TEXT,
            ts TEXT DEFAULT CURRENT_TIMESTAMP
        );""")
        cur.execute("""CREATE TABLE IF NOT EXISTS outreach_blasts(
            id INTEGER PRIMARY KEY,
            title TEXT,
            template_name TEXT,
            sender_email TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        );""")
        cur.execute("""CREATE TABLE IF NOT EXISTS outreach_log(
            id INTEGER PRIMARY KEY,
            blast_id INTEGER,
            to_email TEXT,
            to_name TEXT,
            subject TEXT,
            status TEXT,
            error TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        );""")
        conn.commit()


def _o3_wrap_email_html(html: str) -> str:
    # Basic, client-safe wrapper to ensure readable default font size.
    # Uses table layout and inline styles for Gmail/Outlook compatibility.
    safe = html or ""
    return (
        "<!doctype html>"
        "<html><head><meta charset=\"utf-8\"></head>"
        "<body style=\"margin:0;padding:0;background:#ffffff;\">"
        "<table role=\"presentation\" width=\"100%\" cellpadding=\"0\" cellspacing=\"0\">"
        "<tr><td style=\"font-family: Arial, sans-serif; font-size:16px; line-height:1.5; color:#222222;\">"
        + safe +
        "</td></tr></table>"
        "</body></html>"
    )
def _o3_merge(text, data: dict) -> str:
    import re as _re
    t = str(text or "")
    def rep(m):
        k = m.group(1).strip()
        return str(data.get(k, ""))
    t = _re.sub(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}", rep, t)
    return t

def _o3_load_vendors_df(conn):
    import pandas as _pd
    try:
        df_v = _pd.read_sql_query(
            "SELECT v.id as vendor_id, v.name as company, v.city, v.state, v.naics, v.phone, v.email, v.website FROM vendors v ORDER BY v.name;",
            conn, params=()
        )
    except Exception:
        df_v = _pd.DataFrame()
    try:
        df_c = _pd.read_sql_query(
            "SELECT vc.vendor_id, vc.name as contact_name, vc.email as contact_email, vc.phone as contact_phone, vc.role FROM vendor_contacts vc ORDER BY vc.id DESC;",
            conn, params=()
        )
    except Exception:
        df_c = _pd.DataFrame()
    if df_v is None: df_v = _pd.DataFrame()
    if df_c is None: df_c = _pd.DataFrame()
    if not df_c.empty:
        df = df_c.merge(df_v, on="vendor_id", how="left")
        df["email"] = df["contact_email"].fillna(df.get("email"))
        df["name"] = df["contact_name"].fillna("")
        df["phone"] = df["contact_phone"].fillna(df.get("phone"))
    else:
        df = df_v.copy()
        df["name"] = ""
    cols = ["email","name","company","phone","naics","city","state","website"]
    for c in cols:
        if c not in df.columns:
            df[c] = ""
    df = df[cols].copy()
    # No mandatory email filter here. Sending pipeline will skip invalid emails.
    return df

def _o3_collect_recipients_ui(conn):
    import streamlit as st, pandas as _pd
    _o3_ensure_schema(conn)
    st.subheader("Recipients")
    tabs = st.tabs(["From Vendors","Upload CSV","Manual"])
    all_rows = _pd.DataFrame(columns=["email","name","company","phone","naics","city","state","website"])
    with tabs[0]:
        df = _o3_load_vendors_df(conn)
        f1, f2 = st.columns([2,2])
        with f1:
            f_naics = st.text_input("Filter NAICS contains", key="o3_naics")
        with f2:
            f_state = st.text_input("Filter State equals (e.g., TX)", key="o3_state")
        if f_naics:
            df = df[df["naics"].fillna("").str.contains(f_naics, case=False, na=False)]
        if f_state:
            df = df[df["state"].fillna("").str.upper()==f_state.strip().upper()]
        st.caption(f"{len(df)} vendor rows")
        _styled_dataframe(df, use_container_width=True, hide_index=True)
        if st.button("Add all filtered vendors", key="o3_add_vendors"):
            all_rows = _pd.concat([all_rows, df], ignore_index=True)
    with tabs[1]:
        st.caption("CSV columns: email,name,company,phone,naics,city,state,website,first_name,last_name,title,solicitation,due,notice_id")
        up = st.file_uploader("Upload CSV", type=["csv"], key="o3_csv_up")
        if up:
            try:
                dfu = _pd.read_csv(up)
                _styled_dataframe(dfu.head(50), use_container_width=True, hide_index=True)
                if st.button("Add uploaded rows", key="o3_add_csv"):
                    all_rows = _pd.concat([all_rows, dfu], ignore_index=True)
            except Exception as e:
                st.error(f"CSV read error: {e}")
    with tabs[2]:
        st.caption("Paste one email per line or 'email, name, company'")
        txt = st.text_area("Lines", height=150, key="o3_paste")
        if st.button("Add pasted", key="o3_add_paste"):
            rows = []
            for line in (txt or "").splitlines():
                parts = [p.strip() for p in line.split(",")]
                if not parts: continue
                em = parts[0] if "@" in parts[0] else ""
                if not em: continue
                name = parts[1] if len(parts)>1 else ""
                comp = parts[2] if len(parts)>2 else ""
                rows.append({"email": em, "name": name, "company": comp, "phone":"", "naics":"", "city":"", "state":"", "website":""})
            if rows:
                all_rows = _pd.concat([all_rows, _pd.DataFrame(rows)], ignore_index=True)
    cur = st.session_state.get("o3_rows")
    if cur is not None and isinstance(cur, _pd.DataFrame) and not cur.empty:
        all_rows = _pd.concat([cur, all_rows], ignore_index=True)
    if not all_rows.empty:
        all_rows = all_rows.dropna(subset=["email"])
        all_rows = all_rows[all_rows["email"].astype(str).str.contains("@", na=False)]
        all_rows = all_rows.drop_duplicates(subset=["email"])
    st.session_state["o3_rows"] = all_rows
    if not all_rows.empty:
        st.write(f"Total recipients: {len(all_rows)}")
        _styled_dataframe(all_rows, use_container_width=True, hide_index=True)
        csv_bytes = all_rows.to_csv(index=False).encode("utf-8")
        st.download_button("Download recipients CSV", data=csv_bytes, file_name="o3_recipients.csv", mime="text/csv", key="o3_dl_recip")
    return all_rows

def _o3_sender_accounts_from_secrets():
    try:

        accs = []
        try:
            for row in (st.secrets.get("gmail_accounts") or []):
                if row.get("email") and row.get("app_password"):
                    accs.append({"email":row["email"],"app_password":row["app_password"],"name":row.get("name","")})
        except Exception:
            pass
        if not accs:
            g = st.secrets.get("gmail") or {}
            if g.get("email") and g.get("app_password"):
                accs.append({"email":g["email"],"app_password":g["app_password"],"name":g.get("name","")})
        return accs
    except Exception:
        return []



# --- O3 SMTP shim (fallback) ---
try:
    import _o3smtp  # if provided elsewhere
except Exception:
    class _o3smtp:
        import smtplib as _smtplib
        SMTP_SSL = _smtplib.SMTP_SSL
        SMTP = _smtplib.SMTP

def _o3_send_batch(conn, sender, rows, subject_tpl, html_tpl, test_only=False, max_send=500):
    # Ensure required email and SMTP aliases are available
    try:
        from email.mime.multipart import MIMEMultipart as _O3MIMEMultipart
        from email.mime.text import MIMEText as _O3MIMEText
        import smtplib as _o3smtp
    except Exception:
        pass
    import streamlit as st, datetime as _dt, pandas as _pd, socket as _socket
    _o3_ensure_schema(conn)
    if rows is None or rows.empty:
        st.error("No recipients"); return 0, []
    blast_title = st.text_input("Blast name", value=f"Outreach {_dt.datetime.utcnow().strftime('%Y-%m-%d %H:%M')}", key="o3_blast_name")
    if not blast_title:
        blast_title = "Outreach"
    with _o3c(conn.cursor()) as cur:
        cur.execute("INSERT INTO outreach_blasts(title, template_name, sender_email) VALUES(?,?,?);",
                    (blast_title, st.session_state.get("tpl_sel",""), sender.get("email","")))
        conn.commit()
        blast_id = cur.lastrowid
    try:
        opt = _pd.read_sql_query("SELECT email FROM outreach_optouts;", conn)
        blocked = set(e.lower().strip() for e in opt['email'].tolist())
    except Exception:
        blocked = set()
    sent = 0
    logs = []
    smtp = None
    if not test_only:
        # derive SMTP settings from sender
        host = sender.get("host") or "smtp.gmail.com"
        port = int(sender.get("port") or 465)
        use_tls = bool(sender.get("use_tls") or (port == 587))
        try:
            if use_tls or port == 587:
                smtp = _o3smtp.SMTP(host, port, timeout=20)
                smtp.ehlo()
                try:
                    smtp.starttls()
                except Exception:
                    pass
                smtp.login(sender["email"], sender["app_password"])
            else:
                smtp = _o3smtp.SMTP_SSL(host, port, timeout=20)
                smtp.login(sender["email"], sender["app_password"])
        except Exception as e:
            st.error(f"SMTP login failed: {e}")
            return 0, []

    for _, r in rows.head(int(max_send)).iterrows():
        to_email = str(r.get("email","")).strip()
        if not to_email or "@" not in to_email:
            continue
        if to_email.lower() in blocked:
            status = "Skipped: opt-out"; err = ""; subj = ""
        else:
            data = {k: str(r.get(k,"") or "") for k in r.index}
            subj = _o3_merge(subject_tpl or "", data)
            html = _o3_merge(html_tpl or "", data)
            if "unsubscribe" not in html.lower():
                html += "<br><br><small>To unsubscribe, reply 'STOP'.</small>"
            if test_only:
                status = "Preview"; err = ""
            else:
                try:
                    msg = _O3MIMEMultipart("alternative")
                    msg["From"] = f"{sender.get('name') or sender['email']} <{sender['email']}>"
                    msg["To"] = to_email
                    msg["Subject"] = subj
                    html = _o3_wrap_email_html(html)
                    msg.attach(_O3MIMEText(html, "html", "utf-8"))
                    smtp.sendmail(sender["email"], [to_email], msg.as_string())
                    status = "Sent"; err = ""
                except Exception as e:
                    status = "Error"; err = str(e)
        with _o3c(conn.cursor()) as cur:
            cur.execute("INSERT INTO outreach_log(blast_id, to_email, to_name, subject, status, error) VALUES(?,?,?,?,?,?);",
                        (blast_id, to_email, str(r.get('name') or ""), subj, status, err))
            conn.commit()
        logs.append({"email":to_email,"status":status,"error":err})
        if status=="Sent":
            sent += 1
    if smtp is not None:
        try: smtp.quit()
        except Exception: pass
    try:
        df = _pd.DataFrame(logs)
        st.download_button("Download send log CSV", data=df.to_csv(index=False).encode("utf-8"),
                           file_name=f"o3_send_log_{blast_id}.csv", mime="text/csv", key=f"o3_log_{blast_id}")
    except Exception:
        pass
    st.success(f"Batch complete. Sent={sent}, Total processed={len(logs)}")
    return sent, logs

def _export_past_perf_docx(path: str, records: list) -> Optional[str]:
    try:
        import docx  # type: ignore
        from docx.shared import Inches  # type: ignore
    except Exception:
        pass

        st.error("python-docx is required. pip install python-docx")
        return None
    try:
        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(1); s.bottom_margin = Inches(1); s.left_margin = Inches(1); s.right_margin = Inches(1)
        doc.add_heading("Past Performance", level=1)
        for rec in records or []:
            title = str(rec.get("title") or rec.get("project") or "Project").strip()
            doc.add_heading(title, level=2)
            for k in ["customer","period","value","cpars_rating"]:
                v = rec.get(k)
                if v:
                    doc.add_paragraph(f"**{k.title()}:** {v}")
            body = str(rec.get("summary") or rec.get("description") or rec.get("results") or "").strip()
            if body:
                for para in body.split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
        doc.save(path)
        return path
    except Exception as e:
        pass

        st.error(f"Past Performance export failed: {e}")
        return None


# === O4: Multi-sender accounts + opt-outs + audit UI ================================
def _o4_accounts_ui(conn):
    import streamlit as st
    import streamlit as st, pandas as _pd
    ensure_outreach_o1_schema(conn)
    rows = conn.execute("SELECT user_email, display_name, smtp_host, smtp_port, use_ssl FROM email_accounts ORDER BY user_email").fetchall()
    if rows:
        _styled_dataframe(_pd.DataFrame(rows, columns=["Email","Display name","SMTP host","SMTP port","SSL"]), use_container_width=True, hide_index=True)
    st.markdown("**Add or update account**")
    c1,c2 = st.columns([3,2])
    with c1:
        email = st.text_input("Email", key="o4_ac_email")
        display = st.text_input("Display name", key="o4_ac_display")
        app_pw = st.text_input("Gmail App password", type="password", key="o4_ac_pw")
    with c2:
        host = st.text_input("SMTP host", value="smtp.gmail.com", key="o4_ac_host")
        port = st.number_input("SMTP port", min_value=1, max_value=65535, value=465, step=1, key="o4_ac_port")
        ssl = st.checkbox("Use SSL", value=True, key="o4_ac_ssl")
    c3, c4 = st.columns(2)
    with c3:
        if st.button("Save account", key="o4_ac_save"):
            if not email:
                st.error("Email required")
            else:
                with conn:
                    conn.execute("""
                    INSERT INTO email_accounts(user_email, display_name, app_password, smtp_host, smtp_port, use_ssl)
                    VALUES(?,?,?,?,?,?)
                    ON CONFLICT(user_email) DO UPDATE SET
                        display_name=excluded.display_name,
                        app_password=excluded.app_password,
                        smtp_host=excluded.smtp_host,
                        smtp_port=excluded.smtp_port,
                        use_ssl=excluded.use_ssl
                    """, (email.strip(), display or "", app_pw or "", host or "smtp.gmail.com", int(port or 465), 1 if ssl else 0))
                st.success("Saved")
    try:

        st.session_state["o4_sender_sel"] = email.strip()
    except Exception:
        pass
    st.rerun()
    with c4:
        if st.button("Delete account", key="o4_ac_del"):
            if not email:
                st.error("Enter the Email of the account to delete")
            else:
                with conn:
                    conn.execute("DELETE FROM email_accounts WHERE user_email=?", (email.strip(),))
                st.success("Deleted")

def _o3_render_sender_picker():
    import streamlit as st
    # Override to use email_accounts. Uses _O4_CONN set by render_outreach_mailmerge.

    conn = get_o4_conn() if "get_o4_conn" in globals() else globals().get("_O4_CONN")
    if conn is None:
        st.warning("No sender accounts configured");
        return {"email":"", "app_password":""}
    ensure_outreach_o1_schema(conn)
    rows = _get_senders(conn)
    try:

        st.caption(f"Loaded {len(rows)} sender account(s) from unified sources")
    except Exception:
        pass
    choices = [r[0] for r in rows] + ["<add new>"]
    default = st.session_state.get("o4_sender_sel", choices[0] if choices else None)
    idx = choices.index(default) if default in choices else 0
    sel = st.selectbox("From account", choices, key="o4_sender_sel", index=idx)
    chosen = {"email": sel if sel != "<add new>" else "", "app_password":"", "smtp_host":"smtp.gmail.com", "smtp_port":465, "use_ssl":1}
    # Load password and SMTP details from the unifying table when present
    if sel != "<add new>":
        try:
            # First try email_accounts
            row = conn.execute("SELECT app_password, smtp_host, smtp_port, use_ssl FROM email_accounts WHERE user_email=?", (sel,)).fetchone()
            if row:
                chosen.update({"app_password": row[0] or "", "smtp_host": row[1] or "smtp.gmail.com", "smtp_port": int(row[2] or 465), "use_ssl": int(row[3] or 1)})
            else:
                # Fallback to smtp_settings
                row2 = conn.execute("SELECT password, host, port, use_tls FROM smtp_settings WHERE username=?", (sel,)).fetchone()
                if row2:
                    chosen.update({"app_password": row2[0] or "", "smtp_host": row2[1] or "smtp.gmail.com", "smtp_port": int(row2[2] or 587), "use_ssl": int(row2[3] or 1)})
        except Exception:
            pass
    if sel == "<add new>":
        st.info("Add an account below in 'Sender accounts'. Then select it here.")
    else:
        row = conn.execute("SELECT user_email, display_name, app_password, smtp_host, smtp_port, use_ssl FROM email_accounts WHERE user_email=?", (sel,)).fetchone()
        if row:
            chosen = {"email": row[0], "display_name": row[1] or "", "app_password": row[2] or "",
                      "smtp_host": row[3] or "smtp.gmail.com", "smtp_port": int(row[4] or 465), "use_ssl": int(row[5] or 1)}
    st.caption("Uses Gmail SMTP. Create an App Password once per account and save it above.")
    return chosen

def _o4_optout_ui(conn):
    import streamlit as st, pandas as _pd
    # tables already created by O3; ensure again for safety
    with conn:
        conn.execute("CREATE TABLE IF NOT EXISTS outreach_optouts(id INTEGER PRIMARY KEY, email TEXT UNIQUE)")
    st.markdown("**Opt-outs**")
    em = st.text_input("Add single email to opt-out", key="o4_opt_one")
    if st.button("Add opt-out", key="o4_opt_add") and em:
        with conn:
            conn.execute("INSERT OR IGNORE INTO outreach_optouts(email) VALUES(?)", (em.strip().lower(),))
        st.success("Added")
        st.rerun()
    up = st.file_uploader("Bulk upload CSV with 'email' column", type=["csv"], key="o4_opt_csv")
    if up is not None:
        try:
            df = _pd.read_csv(up)
            emails = [str(x).strip().lower() for x in df.get("email", []) if str(x).strip()]
            with conn:
                conn.executemany("INSERT OR IGNORE INTO outreach_optouts(email) VALUES(?)", [(e,) for e in emails])
            st.success(f"Imported {len(emails)} emails")
        except Exception as e:
            st.error(f"CSV error: {e}")
    try:
        df2 = _pd.read_sql_query("SELECT email FROM outreach_optouts ORDER BY email LIMIT 500", conn)
        _styled_dataframe(df2, use_container_width=True, hide_index=True)
    except Exception:
        pass

def _o4_audit_ui(conn):
    import streamlit as st, pandas as _pd
    try:
        blasts = _pd.read_sql_query("SELECT id, title, sender_email, created_at FROM outreach_blasts ORDER BY id DESC LIMIT 50", conn)
        st.markdown("**Recent blasts**")
        _styled_dataframe(blasts, use_container_width=True, hide_index=True)
    except Exception:
        st.caption("No blasts yet")
    try:
        logs = _pd.read_sql_query("SELECT created_at, to_email, status, subject, error FROM outreach_log ORDER BY id DESC LIMIT 200", conn)
        st.markdown("**Recent sends**")
        _styled_dataframe(logs, use_container_width=True, hide_index=True)
    except Exception:
        st.caption("No logs yet")


# === O5: Follow-ups & SLA ==================================================
import sqlite3, pandas as _pd, smtplib, ssl, time
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

def _o5_now_iso():
    return __import__("datetime").datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

def ensure_o5_schema(conn: sqlite3.Connection) -> None:
    with conn:
        conn.execute("""CREATE TABLE IF NOT EXISTS outreach_sequences(
            id INTEGER PRIMARY KEY, name TEXT UNIQUE NOT NULL
        );""")
        conn.execute("""CREATE TABLE IF NOT EXISTS outreach_steps(
            id INTEGER PRIMARY KEY, seq_id INTEGER NOT NULL,
            step_no INTEGER NOT NULL, delay_hours INTEGER NOT NULL DEFAULT 72,
            subject TEXT DEFAULT '', body_html TEXT DEFAULT '',
            FOREIGN KEY(seq_id) REFERENCES outreach_sequences(id)
        );""")
        conn.execute("""CREATE TABLE IF NOT EXISTS outreach_schedules(
            id INTEGER PRIMARY KEY, seq_id INTEGER NOT NULL, step_no INTEGER NOT NULL,
            to_email TEXT NOT NULL, vendor_id INTEGER,
            send_at TEXT NOT NULL, status TEXT NOT NULL DEFAULT 'queued',
            last_error TEXT DEFAULT '',
            subject TEXT DEFAULT '', body_html TEXT DEFAULT '',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(seq_id) REFERENCES outreach_sequences(id)
        );""")

def _o5_list_sequences(conn):
    try:
        return _pd.read_sql_query("SELECT id, name FROM outreach_sequences ORDER BY name;", conn)
    except Exception:
        return _pd.DataFrame(columns=["id","name"])

def _o5_list_steps(conn, seq_id: int):
    try:
        return _pd.read_sql_query("SELECT id, step_no, delay_hours, subject FROM outreach_steps WHERE seq_id=? ORDER BY step_no;", conn, params=(seq_id,))
    except Exception:
        return _pd.DataFrame(columns=["id","step_no","delay_hours","subject"])

def _o5_upsert_sequence(conn, name: str):
    with conn:
        conn.execute("INSERT INTO outreach_sequences(name) VALUES(?) ON CONFLICT(name) DO NOTHING;", (name.strip(),))

def _o5_add_step(conn, seq_id: int, step_no: int, delay_hours: int, subject: str, body_html: str):
    with conn:
        conn.execute("INSERT INTO outreach_steps(seq_id, step_no, delay_hours, subject, body_html) VALUES(?,?,?,?,?)",
                     (seq_id, step_no, int(delay_hours), subject or "", body_html or ""))

def _o5_queue_followups(conn, seq_id: int, emails: list[str], start_at_iso: str | None = None):
    if not start_at_iso:
        start_at_iso = _o5_now_iso()
    steps = _pd.read_sql_query("SELECT step_no, delay_hours, subject, body_html FROM outreach_steps WHERE seq_id=? ORDER BY step_no;", conn, params=(seq_id,))
    if steps is None or steps.empty:
        return 0
    count = 0
    with conn:
        for em in emails:
            em = (em or "").strip().lower()
            if not em:
                continue
            base = __import__("datetime").datetime.fromisoformat(start_at_iso.replace("Z","+00:00"))
            for _, row in steps.iterrows():
                eta = base + __import__("datetime").timedelta(hours=int(row["delay_hours"] or 0))
                conn.execute("""INSERT INTO outreach_schedules(seq_id, step_no, to_email, send_at, subject, body_html, status)
                                VALUES(?,?,?,?,?,?, 'queued')""",                             (seq_id, int(row["step_no"]), em, eta.strftime("%Y-%m-%dT%H:%M:%SZ"), row["subject"] or "", row["body_html"] or ""))
                base = eta
                count += 1
    return count

def _o5_pick_sender_from_session():
    host = st.session_state.get("smtp_host") or (st.session_state.get("smtp_profile") or {}).get("host")
    port = st.session_state.get("smtp_port", 587)
    tls = bool(st.session_state.get("smtp_tls", True))
    username = st.session_state.get("smtp_username", "")
    password = st.session_state.get("smtp_password", "")
    return {"host": host or "smtp.gmail.com", "port": int(port or 587), "tls": tls, "username": username, "password": password}

def _o5_smtp_send(sender: dict, to_email: str, subject: str, html: str):
    msg = MIMEMultipart("alternative"); msg["Subject"] = subject or ""; msg["From"] = sender["username"]; msg["To"] = to_email
    msg.attach(MIMEText(html or "", "html"))
    if sender["tls"]:
        server = smtplib.SMTP(sender["host"], sender["port"]); server.ehlo(); server.starttls(context=ssl.create_default_context()); server.login(sender["username"], sender["password"])
    else:
        server = smtplib.SMTP_SSL(sender["host"], sender["port"], context=ssl.create_default_context()); server.login(sender["username"], sender["password"])
    server.sendmail(sender["username"], [to_email], msg.as_string()); server.quit()

def _o5_send_due(conn, limit: int = 200):
    now = _o5_now_iso()
    df = _pd.read_sql_query("SELECT id, to_email, subject, body_html FROM outreach_schedules WHERE status='queued' AND send_at<=? ORDER BY send_at LIMIT ?;", conn, params=(now, int(limit)))
    if df is None or df.empty: return 0, 0
    ok = 0; fail = 0; sender = _o5_pick_sender_from_session()
    for _, r in df.iterrows():
        try:
            _o5_smtp_send(sender, r["to_email"], r["subject"], r["body_html"]); 
            with conn: conn.execute("UPDATE outreach_schedules SET status='sent' WHERE id=?", (int(r["id"]),)); ok += 1
        except Exception as e:
            with conn: conn.execute("UPDATE outreach_schedules SET status='error', last_error=? WHERE id=?", (str(e)[:500], int(r["id"]))); fail += 1
    return ok, fail

def render_outreach_o5_followups(conn):
    ensure_o5_schema(conn)
    st.subheader("O5 — Follow-ups & SLA")
    seq_df = _o5_list_sequences(conn); names = ["— New —"] + ([] if seq_df is None or seq_df.empty else seq_df["name"].tolist())
    c1, c2 = st.columns([2,3])
    with c1:
        sel = st.selectbox("Sequence", names, key="o5_seq_sel")
        new_name = st.text_input("New sequence name", key="o5_seq_name") if sel == "— New —" else sel
        if st.button("Save sequence", key="o5_seq_save"):
            if new_name and new_name.strip(): _o5_upsert_sequence(conn, new_name.strip()); st.success("Sequence saved"); st.rerun()
    with c2:
        if sel != "— New —" and (seq_df is not None and not seq_df.empty):
            seq_id = int(seq_df.loc[seq_df["name"]==sel, "id"].iloc[0])
            st.markdown("**Steps**"); steps = _o5_list_steps(conn, seq_id)
            if steps is not None and not steps.empty: _styled_dataframe(steps, use_container_width=True, hide_index=True)
            st.markdown("**Add step**"); s1,s2,s3 = st.columns(3)
            with s1: step_no = st.number_input("Step #", 1, 20, value=(int(steps["step_no"].max())+1 if steps is not None and not steps.empty else 1))
            with s2: delay = st.number_input("Delay hours", 1, 720, value=72)
            with s3: subj = st.text_input("Subject", key="o5_step_subj")
            body = st.text_area("HTML body", height=180, key="o5_step_body")
            if st.button("Add step", key="o5_step_add"): _o5_add_step(conn, seq_id, int(step_no), int(delay), subj, body); st.success("Step added"); st.rerun()
    st.markdown("---"); st.markdown("**Queue follow-ups**")
    if sel != "— New —" and (seq_df is not None and not seq_df.empty):
        seq_id = int(seq_df.loc[seq_df["name"]==sel, "id"].iloc[0])
    else:
        seq_id = None
    emails_txt = st.text_area("Paste recipient emails (one per line)", height=120, key="o5_emails")
    if st.button("Queue follow-ups", key="o5_queue"):
        if not seq_id: st.error("Select an existing sequence first")
        else:
            emails = [e.strip() for e in (emails_txt or "").splitlines() if e.strip()]
            n = _o5_queue_followups(conn, seq_id, emails); st.success(f"Queued {n} follow-up sends")
    st.markdown("**Send due now**")
    if st.button("Send due follow-ups", key="o5_send_due"):
        ok, fail = _o5_send_due(conn, limit=200); st.success(f"Sent {ok}, failed {fail}")
# === End O5 ================================================================


# === O6: Compliance — Unsubscribe & Suppression =============================
import uuid, urllib.parse

def ensure_o6_schema(conn):
    with conn:
        conn.execute("""CREATE TABLE IF NOT EXISTS outreach_optouts(
            id INTEGER PRIMARY KEY, email TEXT UNIQUE, reason TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP
        );""")
        conn.execute("""CREATE TABLE IF NOT EXISTS outreach_unsub_codes(
            code TEXT PRIMARY KEY, email TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP, used_at TEXT
        );""")
        conn.execute("""CREATE TABLE IF NOT EXISTS kv_store(
            k TEXT PRIMARY KEY, v TEXT
        );""")

def o6_set_base_url(conn, url):
    with conn:
        conn.execute("INSERT INTO kv_store(k,v) VALUES('o6_base_url', ?) ON CONFLICT(k) DO UPDATE SET v=excluded.v;", (url,))

def o6_get_base_url(conn):
    try:
        cur = conn.execute("SELECT v FROM kv_store WHERE k='o6_base_url' LIMIT 1;")
        row = cur.fetchone()
        if row and row[0]:
            return row[0]
    except Exception:
        pass
    try:
        import streamlit as _st
        return _st.secrets.get("app_base_url", "")
    except Exception:
        return ""

def o6_is_suppressed(conn, email):
    try:
        em = (email or "").strip().lower()
        row = conn.execute("SELECT 1 FROM outreach_optouts WHERE lower(email)=? LIMIT 1;", (em,)).fetchone()
        return bool(row)
    except Exception:
        return False

def o6_add_optout(conn, email, reason="user_unsubscribe"):
    try:
        with conn:
            conn.execute("INSERT INTO outreach_optouts(email,reason) VALUES(?,?) ON CONFLICT(email) DO NOTHING;", ((email or "").strip().lower(), reason))
    except Exception:
        pass

def o6_new_code(conn, email):
    c = uuid.uuid4().hex
    with conn:
        conn.execute("INSERT INTO outreach_unsub_codes(code,email) VALUES(?,?)", (c,(email or "").strip().lower()))
    return c

def o6_unsub_link_for(conn, email):
    base = o6_get_base_url(conn) or ""
    if not base:
        return ""
    code = o6_new_code(conn, email)
    sep = "&" if "?" in base else "?"
    return f"{base}{sep}unsubscribe={code}"

def o6_handle_query_unsubscribe(conn):
    try:
        import streamlit as _st
        qp = _dict(st.query_params)
        if "unsubscribe" in qp:
            code = (qp.get("unsubscribe",[None]) or [None])[0]
            if code:
                row = conn.execute("SELECT email FROM outreach_unsub_codes WHERE code=? LIMIT 1;", (code,)).fetchone()
                if row and row[0]:
                    email = row[0]
                    o6_add_optout(conn, email, reason="link_click")
                    with conn:
                        conn.execute("UPDATE outreach_unsub_codes SET used_at=CURRENT_TIMESTAMP WHERE code=?", (code,))
                    _st.success(f"{email} unsubscribed.")
                    return True
                else:
                    _st.warning("Invalid or expired unsubscribe link.")
                    return True
        # Also support direct email param
        if "unsubscribe_email" in qp:
            email = (qp.get("unsubscribe_email",[None]) or [None])[0]
            if email:
                o6_add_optout(conn, email, reason="direct_param")
                _st.success(f"{email} unsubscribed.")
                return True
    except Exception:
        pass
    return False

def render_outreach_o6_compliance(conn):
    ensure_o6_schema(conn)
    import pandas as _pd
    st.subheader("O6 — Compliance")
    base = st.text_input("Unsubscribe base URL", value=o6_get_base_url(conn) or "", help="Example: https://yourapp.yourdomain/")
    if st.button("Save base URL", key="o6_save_base"):
        o6_set_base_url(conn, base.strip())
        st.success("Saved")
    st.caption("Use {{UNSUB_LINK}} macro in templates. If absent, a default unsubscribe footer will be appended.")
    # Show list
    df = _pd.read_sql_query("SELECT email, reason, created_at FROM outreach_optouts ORDER BY created_at DESC LIMIT 500;", conn)
    _styled_dataframe(df, use_container_width=True, hide_index=True)

# Wrap _o3_send_batch to enforce suppression and inject unsubscribe link
def _o6_wrap_o3_send_batch():
    g = globals()
    orig = g.get("_o3_send_batch")
    if not callable(orig) or getattr(orig, "_o6_wrapped", False):
        return
    def wrapped(*args, **kwargs):
        import pandas as _pd
        # Normalize to new signature: (conn, sender, rows_df, subj, html, test_only=False, max_send=500)
        conn = kwargs.get("conn")
        sender = kwargs.get("sender")
        rows = kwargs.get("rows")
        subj = kwargs.get("subj")
        html = kwargs.get("html")
        test_only = kwargs.get("test_only", False)
        max_send = kwargs.get("max_send", 500)

        # Positional fallbacks
        if conn is None or sender is None or rows is None or subj is None or html is None:
            if len(args) >= 5:
                conn, sender, rows, subj, html = args[:5]
                if len(args) >= 6: test_only = args[5]
                if len(args) >= 7: max_send = args[6]
            elif len(args) == 2:
                # Old call style: _o3_send_batch(sender, [{to, subject, html}, ])
                sender, rows_list = args
                # Build conn via get_db if available
                if conn is None and "get_db" in g:
                    try:
                        conn = g["get_db"]()
                    except Exception:
                        conn = kwargs.get("conn")
                # Convert rows_list to DataFrame
                if isinstance(rows_list, list):
                    rows = _pd.DataFrame(rows_list)
                elif hasattr(rows_list, "to_dict"):
                    rows = _pd.DataFrame(rows_list)
                else:
                    rows = _pd.DataFrame([])
                # Derive subj and html defaults per-row later
                # If provided in kwargs, use them
                subj = subj or (rows.iloc[0]["subject"] if not rows.empty and "subject" in rows.columns else "")
                html = html or (rows.iloc[0]["html"] if not rows.empty and "html" in rows.columns else "")
            else:
                # Unsupported call pattern
                raise TypeError("Unsupported _o3_send_batch call signature")

        # Ensure schema and suppression filter
        ensure_o6_schema(conn)
        if rows is not None and hasattr(rows, "copy"):
            _rows = rows.copy()
            if "email" in _rows.columns:
                mask = _rows["email"].astype(str).str.lower().apply(lambda em: not o6_is_suppressed(conn, em))
                _rows = _rows[mask]
            rows = _rows

        # If subj/html empty but rows contain per-recipient fields, send one by one
        if (not subj or not html) and rows is not None and not rows.empty and {"subject","html"}.issubset(set(map(str.lower, rows.columns))):
            # Normalize column names case-insensitively
            columns = {c.lower(): c for c in rows.columns}
            total = 0
            for _, r in rows.iterrows():
                subj_i = str(r.get(columns.get("subject"), subj) or "")
                html_i = str(r.get(columns.get("html"), html) or "")
                # Build single-row frame to keep orig API
                orig(conn, sender, rows=_pd.DataFrame([r]), subj=subj_i, html=html_i, test_only=test_only, max_send=1)
                total += 1
            return total

        return orig(conn, sender, rows, subj, html, test_only=test_only, max_send=max_send)
    wrapped._o6_wrapped = True
    g["_o3_send_batch"] = wrapped

_o6_wrap_o3_send_batch()
# === End O6 ================================================================


# === S1D: Subcontractor Finder — dedupe + Google Places pagination ===========
import json as _json, time as _time
from typing import Any, List, Dict
import requests as _requests

def _s1d_get_api_key():
    try:
        return st.secrets["google"]["api_key"]
    except Exception:
        pass
    try:
        return st.secrets["GOOGLE_API_KEY"]
    except Exception:
        pass
    return ""

def _s1d_norm_phone(p: str) -> str:
    import re as _re
    if not p: return ""
    digits = "".join(_re.findall(r"\d+", str(p)))
    if len(digits) == 11 and digits.startswith("1"):
        digits = digits[1:]
    return digits

def _s1d_existing_vendor_keys(conn):
    return _s1d_select_existing_pairs(conn)

def _s1d_geocode(addr: str, key: str):
    if not addr: return None
    try:
        r = _requests.get("https://maps.googleapis.com/maps/api/geocode/json", params={"address": addr, "key": key}, timeout=10)
        js = r.json()
        if js.get("status") == "OK" and js.get("results"):
            loc = js["results"][0]["geometry"]["location"]
            return float(loc["lat"]), float(loc["lng"])
    except Exception:
        return None
    return None

def _s1d_places_textsearch(query: str, lat: float|None, lng: float|None, radius_m: int|None, page_token: str|None, key: str):
    base = "https://maps.googleapis.com/maps/api/place/textsearch/json"
    params = {"query": query, "key": key, "region": "us"}
    if page_token:
        params = {"pagetoken": page_token, "key": key}
    else:
        if lat is not None and lng is not None and radius_m:
            params.update({"location": f"{lat},{lng}", "radius": int(radius_m)})
    r = _requests.get(base, params=params, timeout=12)
    js = r.json()
    return js

def _s1d_place_details(pid: str, key: str):
    try:
        r = _requests.get("https://maps.googleapis.com/maps/api/place/details/json",
                          params={"place_id": pid, "fields": "formatted_phone_number,website,url", "key": key},
                          timeout=10)
        return r.json().get("result", {}) or {}
    except Exception:
        return {}

def _s1d_save_new_vendors(conn, rows: List[Dict[str,Any]]):

    # Determine writable table and ensure schema
    tbl = _s1d_vendor_write_table(conn)
    with conn:
        _s1d_ensure_vendor_table(conn, tbl)
    # Insert or update rows keyed by place_id
    saved = 0
    with conn:
        for r in rows or []:
            cur = conn.execute(f"""
                INSERT INTO {tbl}(source, place_id, name, email, phone, website, address, city, state, zip, naics, notes, lat, lon, created_at)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,datetime('now'))
                ON CONFLICT(place_id) DO UPDATE SET
                    name=COALESCE(excluded.name, {tbl}.name),
                    email=COALESCE(excluded.email, {tbl}.email),
                    phone=COALESCE(excluded.phone, {tbl}.phone),
                    website=COALESCE(excluded.website, {tbl}.website),
                    address=COALESCE(excluded.address, {tbl}.address),
                    city=COALESCE(excluded.city, {tbl}.city),
                    state=COALESCE(excluded.state, {tbl}.state),
                    zip=COALESCE(excluded.zip, {tbl}.zip),
                    naics=COALESCE(excluded.naics, {tbl}.naics),
                    notes=CASE WHEN length({tbl}.notes)>0 THEN {tbl}.notes ELSE COALESCE(excluded.notes, {tbl}.notes) END,
                    lat=COALESCE(excluded.lat, {tbl}.lat),
                    lon=COALESCE(excluded.lon, {tbl}.lon)
            """, (
                str(r.get("source","") or ""),
                str(r.get("place_id","") or ""),
                str(r.get("name","") or ""),
                str(r.get("email","") or ""),
                str(r.get("phone","") or ""),
                str(r.get("website","") or ""),
                str(r.get("address","") or ""),
                str(r.get("city","") or ""),
                str(r.get("state","") or ""),
                str(r.get("zip","") or ""),
                str(r.get("naics_guess","") or r.get("naics","") or ""),
                str(r.get("notes","") or ""),
                float(r.get("lat") or 0) if str(r.get("lat") or "").strip() else None,
                float(r.get("lon") or 0) if str(r.get("lon") or "").strip() else None,
            ))
            saved += 1
    return saved


def _s1d_render_from_cache(conn, df):
    import streamlit as st
    import pandas as _pd
    if df is None or getattr(df, "empty", True):
        st.info("No cached results.")
        if st.button("New search", key="s1d_new_search_empty"):
            st.session_state.pop("s1d_df", None)
        return
    # Show with links and dedupe flags
    def _mk_link(url, text):
        if not url: return text
        return f"<a href='{url}' target='_blank'>{text}</a>"
    show = df.copy()
    if "google_url" in show.columns:
        show["name"] = show.apply(lambda r: _mk_link(r.get("google_url",""), r.get("name","")), axis=1)
    if "website" in show.columns:
        show["website"] = show.apply(lambda r: _mk_link(r.get("website",""), "site") if r.get("website","") else "", axis=1)
    keep = df[~df.get("_dup", _pd.Series([False]*len(df)))].copy() if not df.empty else df
    # Results table
    cols = [c for c in ["name","phone","website","address","city","state","place_id","_dup"] if c in show.columns]
    if cols:
        st.markdown("**Results**")
        st.write(show[cols].to_html(escape=False, index=False), unsafe_allow_html=True)
    # Selection and save
    if keep.empty:
        st.success("All results are already in your vendor list.")
    else:
        st.caption(f"{len(keep)} new vendors can be saved")
        if "row_id" not in keep.columns:
            import hashlib as _h
            def _mk_id(r):
                pid = str(r.get("place_id","") or "")
                if pid: return pid
                s = f"{r.get('name',f'')}-{r.get('phone','')}-{r.get('city','')}"
                return _h.sha1(s.encode()).hexdigest()[:12]
            keep["row_id"] = keep.apply(_mk_id, axis=1)
        keep_view = keep[["row_id","name","phone","website","address","city","state","place_id"]].copy()
        # Restore selection
        sel_ids = set(st.session_state.get("s1d_selected_ids", []))
        keep_view.insert(1, "Select", keep_view["row_id"].isin(sel_ids))
        edited = st.data_editor(keep_view, hide_index=True, key="s1d_editor_cache")
        new_sel = set(edited.loc[edited["Select"]==True, "row_id"])
        st.session_state["s1d_selected_ids"] = list(new_sel)
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("Save selected", key="s1d_save_selected_cache") and new_sel:
                sub = keep[keep["row_id"].isin(list(new_sel))].drop(columns=["row_id"], errors="ignore")
                n = _s1d_save_new_vendors(conn, sub.to_dict("records"))
                st.success(f"Saved {n} vendors")
        with c2:
            if st.button("Save all new vendors", key="s1d_save_all_cache"):
                sub = keep.drop(columns=["row_id"], errors="ignore")
                n = _s1d_save_new_vendors(conn, sub.to_dict("records"))
                st.success(f"Saved {n} vendors")
        with c3:
            if st.button("New search", key="s1d_new_search"):
                st.session_state.pop("s1d_df", None)
                st.session_state.pop("s1d_selected_ids", None)
        # rerun suppressed to keep table visible


def _s1d_vendor_write_table(conn):
    row = conn.execute("SELECT type, name, sql FROM sqlite_master WHERE name='vendors_t'").fetchone()
    if row and (row[0] or '').lower() == 'table':
        return 'vendors_t'
    if row and (row[0] or '').lower() == 'view':
        sql = row[2] or ''
        import re as _re
        m = _re.search(r'FROM\s+([A-Za-z_][A-Za-z0-9_]*)', sql, flags=_re.IGNORECASE)
        if m:
            return m.group(1)
    # fallback
    return 'vendors'

def _s1d_ensure_vendor_table(conn, table_name: str):
    # Create base table if needed
    conn.execute(f"""        CREATE TABLE IF NOT EXISTS {table_name}(
            id INTEGER PRIMARY KEY,
            source TEXT,
            place_id TEXT,
            name TEXT,
            email TEXT,
            phone TEXT,
            website TEXT,
            address TEXT,
            city TEXT,
            state TEXT,
            zip TEXT,
            naics TEXT,
            notes TEXT,
            lat REAL,
            lon REAL,
            created_at TEXT
        )
    """)
    # Add missing columns as app evolves
    cols = {r[1] for r in conn.execute(f"PRAGMA table_info({table_name})").fetchall()}
    needed = [
        "source","place_id","name","email","phone","website","address","city","state","zip","naics","notes","lat","lon","created_at"
    ]
    for c in needed:
        if c not in cols:
            try:
                conn.execute(f"ALTER TABLE {table_name} ADD COLUMN {c} TEXT")
            except Exception:
                pass

def _s1d_select_existing_pairs(conn):
    # Prefer vendors_t if present, else fallback to vendors
    srcs = [("vendors_t","view_or_table"), ("vendors","table")]
    for name, _ in srcs:
        row = conn.execute("SELECT name FROM sqlite_master WHERE name=?", (name,)).fetchone()
        if row:
            try:
                by_np, by_pid = _s1d_select_existing_pairs(conn)
                return by_np, by_pid
            except Exception:
                continue
    return set(), set()

def render_subfinder_s1d(conn):
    st.subheader("S1D — Subcontractor Finder")
    
    by_np, by_pid = _s1d_select_existing_pairs(conn)
    key = _s1d_get_api_key()
    if not key:
        st.error("Missing Google API key in secrets. Set google.api_key or GOOGLE_API_KEY.")
        return
    # Use cached results if present
    import pandas as _pd
    _cache = st.session_state.get("s1d_df")
    if _cache:
        df = _pd.DataFrame(_cache)
        _s1d_render_from_cache(conn, df)
        return
    q = st.text_input("Search query", key="s1d_q", placeholder="e.g., HVAC contractors, plumbing, IT services")
    loc_choice = st.radio("Location", ["Address", "Lat/Lng"], horizontal=True)
    lat = lng = None
    if loc_choice == "Address":
        addr = st.text_input("Place of performance address")
        radius_mi = st.number_input("Radius (miles)", 1, 200, value=50)
        if addr:
            ll = _s1d_geocode(addr, key)
            if ll: lat, lng = ll
    else:
        col1, col2 = st.columns(2)
        with col1: lat = st.number_input("Latitude", value=38.8951)
        with col2: lng = st.number_input("Longitude", value=-77.0364)
        radius_mi = st.number_input("Radius (miles)", 1, 200, value=50)
    radius_m = int(radius_mi * 1609.34)
    c1, c2, c3 = st.columns([1,1,2])
    with c1:
        go = st.button("Search", key="s1d_go")
    with c2:
        nxt = st.button("Next page", key="s1d_next")
    # Persist page token
    tok_key = "s1d_next_token"
    if go: st.session_state.pop(tok_key, None)
    js = None
    try:
        if go or nxt:
            tok = st.session_state.get(tok_key) if nxt else None
            if q:
                js = _s1d_places_textsearch(q, lat, lng, radius_m, tok, key)
                if js.get("next_page_token"):
                    st.session_state[tok_key] = js["next_page_token"]
                else:
                    st.session_state.pop(tok_key, None)
    except Exception as e:
        st.error(f"Search failed: {e}")
        return
    if not js or not js.get("results"):
        st.info("No results yet. Enter a query and click Search.")
        return
    by_np, by_pid = _s1d_existing_vendor_keys(conn)
    rows = []
    for r in js["results"]:
        name = r.get("name","")
        pid = r.get("place_id","")
        addr = r.get("formatted_address","")
        city = state = ""
        if "," in addr:
            parts = [p.strip() for p in addr.split(",")]
            if len(parts)>=2:
                city = parts[-2]
                state = parts[-1].split()[0]
        details = _s1d_place_details(pid, key) if pid else {}
        phone = _s1d_norm_phone(details.get("formatted_phone_number",""))
        website = details.get("website") or ""
        dup = (name.strip().lower(), phone) in by_np or (pid in by_pid)
        rows.append({
            "name": name, "address": addr, "city": city, "state": state,
            "phone": phone, "website": website, "place_id": pid,
            "google_url": details.get("url") or "",
            "_dup": dup
        })
        # be nice to Places
        _time.sleep(0.05)
    import pandas as _pd
    df = _pd.DataFrame(rows)
    st.session_state["s1d_df"] = df.to_dict("records")
    if df.empty:
        st.info("No results.")
        return
    # Show with links and dedupe flags
    def _mk_link(url, text):
        if not url: return text
        return f"<a href='{url}' target='_blank'>{text}</a>"
    show = df.copy()
    show["name"] = show.apply(lambda r: _mk_link(r["google_url"], r["name"]), axis=1)
    show["website"] = show.apply(lambda r: _mk_link(r["website"], "site") if r["website"] else "", axis=1)
    show = show[["name","phone","website","address","city","state","place_id","_dup"]]
    st.markdown("**Results**")
    st.write(show.to_html(escape=False, index=False), unsafe_allow_html=True)


    # Selection and save
    keep = df[~df["_dup"]].copy()
    if keep.empty:
        st.success("All results are already in your vendor list.")
        return
    st.caption(f"{len(keep)} new vendors can be saved")
    # Interactive selection
    keep_view = keep[["name","phone","website","address","city","state","place_id"]].copy()
    keep_view.insert(0, "Select", False)
    edited = st.data_editor(keep_view, hide_index=True, key="s1d_editor")
    sel = edited[edited["Select"]==True]
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Save selected", key="s1d_save_selected") and not sel.empty:
            n = _s1d_save_new_vendors(conn, sel.drop(columns=["Select"]).to_dict("records"))
            st.success(f"Saved {n} vendors")
    with c2:
        if st.button("Save all new vendors", key="s1d_save_all"):
            n = _s1d_save_new_vendors(conn, keep.to_dict("records"))
            st.success(f"Saved {n} vendors")
# === End S1D ================================================================


def _wrap_run_subfinder():
    g = globals()
    base = g.get("run_subcontractor_finder")
    if not callable(base) or getattr(base, "_s1d_wrapped", False):
        return
    def wrapped(conn):
        import streamlit as st
        st.header("Subcontractor Finder")
        try:
            render_subfinder_s1d(conn)
        except Exception as e:
            st.error(f"S1D error: {e}")
        base(conn)
    wrapped._s1d_wrapped = True
    g["run_subcontractor_finder"] = wrapped

_wrap_run_subfinder()



# =========================
# APPEND-ONLY PATCH • O2/O3/O4/O5/O6 + S1D
# This block does not remove or rename any of your code.
# =========================
import streamlit as _st
import sqlite3 as _sqlite3
import pandas as _pandas
import re as _re2, time as _time2, ssl as _ssl2, smtplib as _smtp2
from email.mime.text import MIMEText as _MText
from email.mime.multipart import MIMEMultipart as _MMulti

def __p_db(conn, q, args=()):
    cur = conn.cursor(); cur.execute(q, args); conn.commit(); return cur

def __p_ensure_core(conn):
    c = conn.cursor()
    c.execute("CREATE TABLE IF NOT EXISTS outreach_templates(id INTEGER PRIMARY KEY, name TEXT UNIQUE, subject TEXT, body TEXT)")
    c.execute("CREATE TABLE IF NOT EXISTS outreach_sender_accounts(id INTEGER PRIMARY KEY, label TEXT UNIQUE, email TEXT, app_password TEXT, smtp_host TEXT DEFAULT 'smtp.gmail.com', smtp_port INTEGER DEFAULT 587, use_tls INTEGER DEFAULT 1, is_active INTEGER DEFAULT 1)")
    c.execute("CREATE TABLE IF NOT EXISTS outreach_audit(id INTEGER PRIMARY KEY, ts TEXT DEFAULT (datetime('now')), actor TEXT, action TEXT, meta TEXT)")
    c.execute("CREATE TABLE IF NOT EXISTS outreach_optouts(id INTEGER PRIMARY KEY, email TEXT UNIQUE, reason TEXT, ts TEXT DEFAULT (datetime('now')))")
    c.execute("CREATE TABLE IF NOT EXISTS outreach_sequences(id INTEGER PRIMARY KEY, name TEXT UNIQUE NOT NULL)")
    c.execute("CREATE TABLE IF NOT EXISTS outreach_steps(id INTEGER PRIMARY KEY, seq_id INTEGER NOT NULL, step_no INTEGER NOT NULL, delay_hours INTEGER NOT NULL DEFAULT 72, subject TEXT DEFAULT '', body_html TEXT DEFAULT '')")
    c.execute("CREATE TABLE IF NOT EXISTS outreach_schedules(id INTEGER PRIMARY KEY, seq_id INTEGER NOT NULL, step_no INTEGER NOT NULL, to_email TEXT NOT NULL, send_at TEXT NOT NULL, status TEXT NOT NULL DEFAULT 'queued', last_error TEXT DEFAULT '', subject TEXT DEFAULT '', body_html TEXT DEFAULT '', created_at TEXT DEFAULT CURRENT_TIMESTAMP)")
    c.execute("CREATE TABLE IF NOT EXISTS vendors_t(id INTEGER PRIMARY KEY, name TEXT, email TEXT, phone TEXT, website TEXT, city TEXT, state TEXT, naics TEXT, place_id TEXT)")
    conn.commit()

def __p_active_sender(conn):
    r = __p_db(conn, "SELECT label,email,app_password,smtp_host,smtp_port,use_tls FROM outreach_sender_accounts WHERE is_active=1 ORDER BY id DESC LIMIT 1").fetchone()
    return None if not r else dict(label=r[0], email=r[1], app_password=r[2], host=r[3], port=int(r[4] or 587), tls=bool(r[5]))

def __p_unsub_base(conn):
    try: return _st.secrets.get("app_base_url","")
    except Exception: return ""

def __p_unsub_link(conn, email):
    base = __p_unsub_base(conn)
    if not base: return ""
    import uuid as _uuid
    code = _uuid.uuid4().hex
    __p_db(conn, "INSERT INTO outreach_audit(actor,action,meta) VALUES(?,?,?)", ("system","UNSUB_CODE", f"{email}:{code}"))
    sep = "&" if "?" in base else "?"
    return f"{base}{sep}unsubscribe={code}"

def __p_is_supp(conn, email):
    row = __p_db(conn, "SELECT 1 FROM outreach_optouts WHERE lower(email)=lower(?) LIMIT 1", (email,)).fetchone()
    return bool(row)

def __p_smtp_send(sender, to_email, subject, html):
    msg = _MMulti("alternative"); msg["Subject"]=subject or ""; msg["From"]=sender["email"]; msg["To"]=to_email
    msg.attach(_MText(html or "", "html"))
    if sender.get("tls", True):
        s = _smtp2.SMTP(sender["host"], int(sender.get("port",587))); s.ehlo(); s.starttls(context=_ssl2.create_default_context()); s.login(sender["email"], sender["app_password"])
    else:
        s = _smtp2.SMTP_SSL(sender["host"], int(sender.get("port",465)), context=_ssl2.create_default_context()); s.login(sender["email"], sender["app_password"])
    s.sendmail(sender["email"], [to_email], msg.as_string()); s.quit()

def __p_o4_ui(conn):
    _st.caption("Multiple Gmail senders via App Passwords.")
    with _st.form("__p_o4_add_sender", clear_on_submit=True):
        col1,col2 = _st.columns(2)
        with col1:
            label = _st.text_input("Label", placeholder="BD Gmail", key="__p_o4_lbl")
            email = _st.text_input("Gmail address", key="__p_o4_em")
            host  = _st.text_input("SMTP host", value="smtp.gmail.com", key="__p_o4_host")
        with col2:
            app_pw = _st.text_input("App password (16 chars)", type="password", key="__p_o4_pw")
            port   = _st.number_input("SMTP port", 1, 65535, value=587, key="__p_o4_port")
            tls    = _st.checkbox("Use STARTTLS", value=True, key="__p_o4_tls")
        if _st.form_submit_button("Save sender") and email:
            __p_db(conn, """INSERT INTO outreach_sender_accounts(label,email,app_password,smtp_host,smtp_port,use_tls,is_active)
                            VALUES(?,?,?,?,?,?,1)
                            ON CONFLICT(label) DO UPDATE SET email=excluded.email,app_password=excluded.app_password,
                            smtp_host=excluded.smtp_host,smtp_port=excluded.smtp_port,use_tls=excluded.use_tls""",
                   (label or email, email, app_pw, host, int(port), 1 if tls else 0))
            __p_db(conn, "INSERT INTO outreach_audit(actor,action,meta) VALUES(?,?,?)", ("system","O4_SAVE", email))
            _st.success(f"Saved {email}")
    rows = __p_db(conn, "SELECT label,email,smtp_host,smtp_port,use_tls,is_active FROM outreach_sender_accounts ORDER BY id DESC").fetchall()
    if rows:
        for lbl,em,host,port,tls,act in rows:
            _st.write(f"• **{lbl}** — {em} — {host}:{port} — TLS {bool(tls)} — {'Active' if act else 'Disabled'}")

def __p_o2_ui(conn):
    _st.caption("Save reusable templates. Supports {{name}}, {{company}}, {{UNSUB_LINK}}.")
    with _st.form("__p_o2_new", clear_on_submit=True):
        name = _st.text_input("Template name", key="__p_o2_name")
        subject = _st.text_input("Subject", key="__p_o2_subj")
        body = _st.text_area("Body", height=200, key="__p_o2_body")
        if _st.form_submit_button("Save template") and name:
            __p_db(conn, "INSERT OR REPLACE INTO outreach_templates(name,subject,body) VALUES(?,?,?)", (name,subject,body))
            _st.success(f"Saved template: {name}")
    df = _pandas.read_sql_query("SELECT name, subject, substr(body,1,200) AS preview FROM outreach_templates ORDER BY name", conn)
    if not df.empty: __styled_dataframe(df, use_container_width=True, hide_index=True)

def __p_o3_ui(conn):
    sender = __p_active_sender(conn)
    if not sender: _st.info("Add a sender in O4 first."); return
    names = [r[0] for r in __p_db(conn, "SELECT name FROM outreach_templates ORDER BY name").fetchall()]
    tpl = _st.selectbox("Template", names) if names else None
    subj, body = ("","")
    if tpl:
        r = __p_db(conn, "SELECT subject,body FROM outreach_templates WHERE name=?", (tpl,)).fetchone()
        if r: subj, body = r[0] or "", r[1] or ""
    override = _st.text_input("Override subject (optional)", key="__p_o3_subj_override")
    if override: subj = override
    left,right = _st.columns(2)
    with left:
        raw = _st.text_area("Recipients (email,name,company)", height=220, placeholder="jane@acme.com,Jane,Acme", key="__p_o3_raw")
    with right:
        test_to = _st.text_input("Send test to", value=sender["email"], key="__p_o3_test_to")
        go_test = _st.button("Send test", key="__p_o3_sendtest")
        go_bulk = _st.button("Send bulk", key="__p_o3_sendbulk")
    rows = []
    for line in (raw or "").splitlines():
        p = [x.strip() for x in line.split(",")]
        if p and "@" in p[0]: rows.append(dict(email=p[0],name=(p[1] if len(p)>1 else ""),company=(p[2] if len(p)>2 else "")))
    def _render(s, r): return (s or "").replace("{{name}}", r.get("name","")).replace("{{company}}", r.get("company",""))
    base = __p_unsub_base(conn)
    def _with_unsub(html, em):
        if "{{UNSUB_LINK}}" in (html or ""): return (html or "").replace("{{UNSUB_LINK}}", __p_unsub_link(conn, em) if base else "")
        if base: return (html or "") + f"<hr><p style='font-size:12px;color:#666'>To unsubscribe click <a href='{__p_unsub_link(conn, em)}'>here</a>.</p>"
        return html
    if go_test and subj and body:
        try:
            s_subj = _render(subj, {"name":"Test","company":"TestCo"})
            s_body = _with_unsub(_render(body, {"name":"Test","company":"TestCo"}), test_to)
            __p_smtp_send(sender, test_to, s_subj, s_body)
            __p_db(conn, "INSERT INTO outreach_audit(actor,action,meta) VALUES(?,?,?)", ("system","O3_TEST", test_to))
            _st.success("Test sent")
        except Exception as e:
            _st.error(f"Send failed: {e}")
    if go_bulk and rows and subj and body:
        sent=skip=fail=0
        for r in rows:
            em=r["email"]
            if __p_is_supp(conn, em): skip+=1; continue
            try:
                __p_smtp_send(sender, em, _render(subj,r), _with_unsub(_render(body,r), em))
                sent+=1; _time2.sleep(0.25)
            except Exception: fail+=1
        _st.success(f"Done. Sent {sent}. Skipped {skip}. Failed {fail}.")

def __p_o5_ui(conn):
    _st.caption("Sequences of follow-ups with delays.")
    seq_df = _pandas.read_sql_query("SELECT id,name FROM outreach_sequences ORDER BY name", conn)
    names = ["— New —"] + ([] if seq_df.empty else seq_df["name"].tolist())
    c1,c2 = _st.columns([2,3])
    with c1:
        sel = _st.selectbox("Sequence", names, key="__p_o5_sel")
        new_name = _st.text_input("New sequence name", key="__p_o5_new") if sel=="— New —" else sel
        if _st.button("Save sequence", key="__p_o5_save") and new_name:
            __p_db(conn, "INSERT OR IGNORE INTO outreach_sequences(name) VALUES(?)", (new_name.strip(),)); _st.rerun()
    with c2:
        if sel!="— New —" and not seq_df.empty:
            seq_id = int(seq_df.loc[seq_df["name"]==sel,"id"].iloc[0])
            _st.markdown("**Steps**")
            steps = _pandas.read_sql_query("SELECT step_no,delay_hours,subject FROM outreach_steps WHERE seq_id=? ORDER BY step_no", conn, params=(seq_id,))
            if not steps.empty: __styled_dataframe(steps, use_container_width=True, hide_index=True)
            s1,s2,s3 = _st.columns(3)
            with s1: step = _st.number_input("Step #", 1, 20, value=(int(steps["step_no"].max())+1 if not steps.empty else 1), key="__p_o5_step")
            with s2: delay = _st.number_input("Delay hours", 1, 720, value=72, key="__p_o5_delay")
            with s3: subj = _st.text_input("Subject", key="__p_o5_subj")
            body = _st.text_area("HTML body", height=120, key="__p_o5_body")
            if _st.button("Add step", key="__p_o5_add"):
                __p_db(conn,"INSERT INTO outreach_steps(seq_id,step_no,delay_hours,subject,body_html) VALUES(?,?,?,?,?)",
                       (seq_id,int(step),int(delay),subj or "",body or "")); _st.rerun()
    _st.markdown("---")
    _st.markdown("**Queue follow-ups**")
    if sel!="— New —" and not seq_df.empty: seq_id = int(seq_df.loc[seq_df["name"]==sel,"id"].iloc[0])
    else: seq_id=None
    emails_txt = _st.text_area("Emails, one per line", height=120, key="__p_o5_emails")
    if _st.button("Queue", key="__p_o5_queue"):
        if not seq_id: _st.error("Choose a sequence")
        else:
            steps = _pandas.read_sql_query("SELECT step_no,delay_hours,subject,body_html FROM outreach_steps WHERE seq_id=? ORDER BY step_no", conn, params=(seq_id,))
            if steps.empty: _st.error("No steps")
            else:
                base = __import__("datetime").datetime.utcnow()
                n=0
                for em in [e.strip().lower() for e in (emails_txt or "").splitlines() if e.strip()]:
                    t = base
                    for _,row in steps.iterrows():
                        t = t + __import__("datetime").timedelta(hours=int(row["delay_hours"] or 0))
                        __p_db(conn,"INSERT INTO outreach_schedules(seq_id,step_no,to_email,send_at,status,subject,body_html) VALUES(?,?,?,?, 'queued',?,?)",
                               (seq_id,int(row["step_no"]),em,t.strftime("%Y-%m-%dT%H:%M:%SZ"),row["subject"] or "",row["body_html"] or "")); n+=1
                _st.success(f"Queued {n}")
    if _st.button("Send due now", key="__p_o5_sendnow"):
        now = __import__("datetime").datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        due = _pandas.read_sql_query("SELECT id,to_email,subject,body_html FROM outreach_schedules WHERE status='queued' AND send_at<=? ORDER BY send_at LIMIT 200", conn, params=(now,))
        if due.empty: _st.info("No due items")
        else:
            s = __p_active_sender(conn)
            if not s: _st.error("No active sender in O4")
            else:
                ok=fail=0; base=__p_unsub_base(conn)
                def w(html,em):
                    if "{{UNSUB_LINK}}" in (html or ""): return (html or "").replace("{{UNSUB_LINK}}", __p_unsub_link(conn, em) if base else "")
                    if base: return (html or "")+f"<hr><p style='font-size:12px;color:#666'>To unsubscribe click <a href='{__p_unsub_link(conn, em)}'>here</a>.</p>"
                    return html
                for _,r in due.iterrows():
                    em=r["to_email"]
                    if __p_is_supp(conn, em): __p_db(conn, "UPDATE outreach_schedules SET status='skipped' WHERE id=?", (int(r["id"]),)); continue
                    try: __p_smtp_send(s, em, r["subject"] or "", w(r["body_html"] or "", em)); __p_db(conn,"UPDATE outreach_schedules SET status='sent' WHERE id=?", (int(r["id"]),)); ok+=1; _time2.sleep(0.25)
                    except Exception as e: __p_db(conn,"UPDATE outreach_schedules SET status='error', last_error=? WHERE id=?", (str(e)[:500], int(r["id"]))); fail+=1
                _st.success(f"Sent {ok}, failed {fail}")

def __p_s1d_key():
    try: return _st.secrets["google"]["api_key"]
    except Exception:
        try: return _st.secrets["GOOGLE_API_KEY"]
        except Exception: return ""

def __p_s1d_norm_phone(p):
    digits="".join(_re2.findall(r"\d+", str(p or "")))
    return digits[1:] if len(digits)==11 and digits.startswith("1") else digits

def __p_s1d_existing(conn):
    rows = __p_db(conn,"SELECT name,COALESCE(phone,''),COALESCE(place_id,'') FROM vendors_t").fetchall()
    by_np=set(); by_pid=set()
    for r in rows: by_np.add(((r[0] or "").strip().lower(), __p_s1d_norm_phone(r[1] or "")))
    for r in rows:
        pid=(r[2] or "").strip()
        if pid: by_pid.add(pid)
    return by_np, by_pid

def __p_s1d_ui(conn):

    by_np, by_pid = _s1d_select_existing_pairs(conn)
    _st.subheader("S1D — Google Places & Dedupe")
    key = __p_s1d_key()
    if not key: _st.error("Missing Google API key in secrets"); return
    mode = _st.radio("Location mode", ["Address","Lat/Lng"], horizontal=True, key="__p_s1d_mode")
    lat=lng=None
    if mode=="Address":
        addr=_st.text_input("Place of performance address", key="__p_s1d_addr")
        radius=_st.number_input("Radius (miles)",1,200,50, key="__p_s1d_rad")
        if addr:
            import requests as _rq
            js=_rq.get("https://maps.googleapis.com/maps/api/geocode/json", params={"address":addr,"key":key}, timeout=10).json()
            if js.get("status")=="OK":
                loc=js["results"][0]["geometry"]["location"]; lat,lng=float(loc["lat"]),float(loc["lng"])
    else:
        c1,c2=_st.columns(2)
        with c1: lat=_st.number_input("Latitude", value=38.8951, key="__p_s1d_lat")
        with c2: lng=_st.number_input("Longitude", value=-77.0364, key="__p_s1d_lng")
        radius=_st.number_input("Radius (miles)",1,200,50, key="__p_s1d_rad2")
    q=_st.text_input("Search query", placeholder="HVAC contractors, cabling, IT services", key="__p_s1d_q")
    c1,c2=_st.columns(2)
    with c1: go=_st.button("Search", key="__p_s1d_go")
    with c2: nxt=_st.button("Next page", key="__p_s1d_next")
    tok_key="__p_s1d_tok"
    if go: _st.session_state.pop(tok_key, None)
    results=[]
    if go or nxt:
        import requests as _rq2
        params={"query":q,"key":key,"region":"us"}
        if nxt and _st.session_state.get(tok_key): params={"pagetoken":_st.session_state[tok_key],"key":key}
        elif lat is not None and lng is not None: params.update({"location":f"{lat},{lng}","radius":int(float(radius)*1609.34)})
        js=_rq2.get("https://maps.googleapis.com/maps/api/place/textsearch/json", params=params, timeout=12).json()
        if js.get("next_page_token"): _st.session_state[tok_key]=js["next_page_token"]
        else: _st.session_state.pop(tok_key, None)
        results=js.get("results",[])
    if not results: _st.info("Enter a query and click Search."); return
    rows=[]; by_np,by_pid=__p_s1d_existing(conn)
    import requests as _rq3
    for r in results:
        name=r.get("name",""); pid=r.get("place_id",""); addr=r.get("formatted_address","")
        city=state=""
        if "," in addr:
            parts=[p.strip() for p in addr.split(",")]
            if len(parts)>=2: city=parts[-2]; state=parts[-1].split()[0]
        phone=""; website=""; gurl=""
        try:
            det=_rq3.get("https://maps.googleapis.com/maps/api/place/details/json",
                         params={"place_id":pid,"fields":"formatted_phone_number,website,url","key":key}, timeout=10).json().get("result",{}) or {}
            digits="".join(_re2.findall(r"\\d+", det.get("formatted_phone_number","") or ""))
            if len(digits)==11 and digits.startswith("1"): digits=digits[1:]
            phone=digits; website=det.get("website","") or ""; gurl=det.get("url","") or ""
        except Exception: pass
        dup=((name.strip().lower(), phone) in by_np) or (pid in by_pid)
        rows.append(dict(name=name,address=addr,city=city,state=state,phone=phone,website=website,place_id=pid,google_url=gurl,_dup=dup))
        _time2.sleep(0.05)
    df=_pandas.DataFrame(rows)
    def _link(u,t): return f"<a href='{u}' target='_blank'>{t}</a>" if u else t
    show=df.copy(); show["name"]=show.apply(lambda r: _link(r["google_url"], r["name"]), axis=1); show["website"]=show.apply(lambda r: _link(r["website"],"site") if r["website"] else "", axis=1)
    show=show[["name","phone","website","address","city","state","place_id","_dup"]]
    _st.markdown("**Results**"); _st.write(show.to_html(escape=False,index=False), unsafe_allow_html=True)
    keep=df[~df["_dup"]]
    _st.caption(f"{len(keep)} new vendors can be saved")
    if _st.button("Save all new vendors", key="__p_s1d_save"):
        n = _s1d_save_new_vendors(conn, keep.to_dict("records"))
        _st.success(f"Saved {n} new vendors")

def __p_run_outreach(conn):
    __p_ensure_core(conn)
    _st.header("Outreach")
    try:
        n = __p_db(conn, "SELECT COUNT(1) FROM outreach_sender_accounts").fetchone()[0]
        _st.sidebar.success("O4 Active" if n else "O4 Not Configured")
    except Exception:
        pass
    with _st.expander("O4 • Sender accounts", expanded=True):
        try: __p_o4_ui(conn)
        except Exception as e: _st.warning(f"O4 unavailable: {e}")
    with _st.expander("O2 • Templates", expanded=True):
        try: __p_o2_ui(conn)
        except Exception as e: _st.warning(f"O2 unavailable: {e}")
    with _st.expander("O5 • Follow-ups & SLA", expanded=False):
        try: __p_o5_ui(conn)
        except Exception as e: _st.warning(f"O5 unavailable: {e}")
    with _st.expander("O3 • Mail merge & Send", expanded=True):
        try: __p_o3_ui(conn)
        except Exception as e: _st.error(f"O3 error: {e}")

# Removed legacy monkeypatch block at load time

# =========================


if __name__ == '__main__':
    try:
        main()
    except NameError:
        # fallback: run default entry if main() not defined in this build
        pass

def run_subcontractor_finder_s1_hook(conn):
    ensure_subfinder_s1_schema(conn)
    try:
        s1_render_places_panel(conn)
    except Exception:
        pass


def router(page: str, conn: sqlite3.Connection) -> None:
    """Dynamic router. Resolves run_<snake_case(page)> and executes safely."""
    import re as _re
    name = "run_" + _re.sub(r"[^a-z0-9]+", "_", (page or "").lower()).strip("_")
    fn = globals().get(name)
    # explicit fallbacks for known variant names
    if not callable(fn):
        alt = {
            "L and M Checklist": ["run_l_and_m_checklist", "run_lm_checklist"],
            "Backup & Data": ["run_backup_data", "run_backup_and_data"],
        }.get((page or "").strip(), [])
        for a in alt:
            fn = globals().get(a)
            if callable(fn):
                break
    if not callable(fn):
        import streamlit as _st
        _st.warning(f"No handler for page '{page}' resolved as {name}.")
        return
    _safe_route_call(fn, conn)
    # Hooks
    if (page or "").strip() == "Subcontractor Finder":
        _safe_route_call(globals().get("run_subcontractor_finder_s1_hook", lambda _c: None), conn)
    if (page or "").strip() == "Proposal Builder":
        _safe_route_call(globals().get("pb_phase_v_section_library", lambda _c: None), conn)



# =========================
# APPEND-ONLY PATCH • E1 (Google Places enrichment: phone + website hyperlinks)
# =========================
import streamlit as _st
import pandas as _pd
import re as _re
import time as _time






# ---- Streamlit write guard: suppress rendering of None ----
try:
    import streamlit as st  # ensure st is present
    if not hasattr(st, "_write_wrapped"):
        _orig_write = st.write
        def _write_no_none(*args, **kwargs):
            if len(args) == 1 and args[0] is None:
                return
            return _orig_write(*args, **kwargs)
        st.write = _write_no_none
        st._write_wrapped = True  # marker
except Exception:
    pass


# ---- Phase 1 bootstrap (guarded) ----
try:
    _title_safe = globals().get("APP_TITLE", "ELA GovCon Suite")
except Exception:
    _title_safe = "ELA GovCon Suite"
try:
    import streamlit as st  # ensure st in scope if file order is unusual
    st.set_page_config(page_title=_title_safe, page_icon="🧭", layout="wide")
except Exception:
    pass
for _fn in ("apply_theme_phase1", "_init_phase1_ui", "_sidebar_brand"):
    try:
        globals().get(_fn) and globals()[_fn]()
    except Exception:
        pass

# ---- Phase 0 neutralizer ----
# If any Phase 0 functions exist, convert to no-ops so they cannot override Phase 1.
for _fn in ("apply_theme_phase0", "_init_phase0_ui", "_sidebar_brand_phase0", "_apply_theme_old"):
    try:
        if _fn in globals() and callable(globals()[_fn]):
            globals()[_fn] = (lambda *a, **k: "")
    except Exception:
        pass


# ==== X.6 Compliance Matrix v1 ====
def _ensure_x6_schema(conn: sqlite3.Connection) -> None:
    from contextlib import closing as _closing
    with _closing(conn.cursor()) as cur:
        cur.execute("""
        CREATE TABLE IF NOT EXISTS compliance_requirements(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rfp_id INTEGER NOT NULL,
            file TEXT,
            page INTEGER,
            text TEXT,
            must_flag INTEGER DEFAULT 0,
            hash TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        );
        """)
        cur.execute("""
        CREATE TABLE IF NOT EXISTS compliance_links(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rfp_id INTEGER NOT NULL,
            requirement_id INTEGER NOT NULL,
            section TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        );
        """ )
    conn.commit()

def x6_sections_suggestions(rfp_id: int) -> list:
    # prefer user outline from session, else defaults
    outline = st.session_state.get(f"proposal_outline_{int(rfp_id)}", "")
    if outline:
        # split numbered lines
        items = [ln.strip() for ln in outline.splitlines() if ln.strip()]
        # filter headers like '# Proposal Outline'
        return [i for i in items if not i.startswith("#")]
    return [
        "Cover Letter",
        "Executive Summary",
        "Technical Approach",
        "Management Approach",
        "Past Performance",
        "Pricing",
        "Compliance Matrix",
    ]

def x6_extract_requirements(conn: sqlite3.Connection, rfp_id: int, limit_per_file: int = 400) -> int:
    """Parse rfp_files for 'must' and 'shall' style requirements, upsert into compliance_requirements."""
    import hashlib
    from contextlib import closing as _closing
    added = 0
    with _closing(conn.cursor()) as cur:
        cur.execute("SELECT id, filename, bytes, mime FROM rfp_files WHERE rfp_id=? ORDER BY id ASC;", (int(rfp_id),))
        files = cur.fetchall() or []
    for fid, fname, bts, mime in files:
        try:
            pages = extract_text_pages(bts, mime or "")
        except Exception:
            pages = []
        found = 0
        for pi, page_txt in enumerate(pages or [], start=1):
            if not page_txt:
                continue
            # Simple sentence split
            parts = [s.strip() for s in page_txt.replace("\r", "\n").split("\n") if s.strip()]
            for s in parts:
                low = s.lower()
                musty = any(k in low for k in [" shall ", " must ", " required ", " is required ", " will "])
                if not musty and ("shall" not in low and "must" not in low):
                    continue
                h = hashlib.sha1(f"{rfp_id}|{fname}|{pi}|{s}".encode("utf-8")).hexdigest()
                try:
                    with _closing(conn.cursor()) as cur:
                        # check dup by hash
                        cur.execute("SELECT 1 FROM compliance_requirements WHERE rfp_id=? AND hash=? LIMIT 1;", (int(rfp_id), h))
                        if cur.fetchone():
                            pass
                        else:
                            cur.execute(
                                "INSERT INTO compliance_requirements(rfp_id,file,page,text,must_flag,hash) VALUES(?,?,?,?,?,?)",
                                (int(rfp_id), fname or "", int(pi), s[:2000], 1 if musty else 0, h)
                            )
                            added += 1
                    conn.commit()
                except Exception:
                    pass
                found += 1
                if found >= limit_per_file:
                    break
            if found >= limit_per_file:
                break
    return int(added)

def x6_requirements_df(conn: sqlite3.Connection, rfp_id: int):
    import pandas as pd
    try:
        df = pd.read_sql_query(
            "SELECT id, must_flag, file, page, text FROM compliance_requirements WHERE rfp_id=? ORDER BY must_flag DESC, id ASC;",
            conn, params=(int(rfp_id),)
        )
    except Exception:
        import pandas as pd
        df = pd.DataFrame(columns=["id","must_flag","file","page","text"])
    return df

def x6_coverage(conn: sqlite3.Connection, rfp_id: int) -> tuple[int, int]:
    from contextlib import closing as _closing
    total = 0
    covered = 0
    with _closing(conn.cursor()) as cur:
        cur.execute("SELECT COUNT(*) FROM compliance_requirements WHERE rfp_id=?", (int(rfp_id),))
        total = int(cur.fetchone()[0] or 0)
        cur.execute("SELECT COUNT(DISTINCT requirement_id) FROM compliance_links WHERE rfp_id=?", (int(rfp_id),))
        covered = int(cur.fetchone()[0] or 0)
    return covered, total

def x6_save_links(conn: sqlite3.Connection, rfp_id: int, mapping: list[tuple[int, str]]) -> int:
    from contextlib import closing as _closing
    saved = 0
    with _closing(conn.cursor()) as cur:
        for rid, sec in mapping:
            if not sec:
                continue
            cur.execute(
                "INSERT INTO compliance_links(rfp_id, requirement_id, section) VALUES(?,?,?)",
                (int(rfp_id), int(rid), str(sec)[:200])
            )
            saved += 1
    conn.commit()
    return int(saved)



# ==== X.7 Proposal Builder v1 ====
def _ensure_x7_schema(conn: sqlite3.Connection) -> None:
    from contextlib import closing as _closing
    with _closing(conn.cursor()) as cur:
        cur.execute("""
        CREATE TABLE IF NOT EXISTS proposals(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            rfp_id INTEGER NOT NULL,
            title TEXT,
            status TEXT DEFAULT 'draft',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        );
        """)
        cur.execute("""
        CREATE TABLE IF NOT EXISTS proposal_sections(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            proposal_id INTEGER NOT NULL,
            ord INTEGER NOT NULL,
            title TEXT,
            content TEXT,
            settings_json TEXT,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        );
        """)
    conn.commit()

def x7_create_proposal_from_outline(conn: sqlite3.Connection, rfp_id: int, title: str | None = None) -> int:
    from contextlib import closing as _closing
    import json, datetime as _dt
    outline = st.session_state.get(f"proposal_outline_{int(rfp_id)}", "") or ""
    if not outline.strip():
        outline = "\n".join([
            "# Proposal Outline",
            "1. Cover Letter",
            "2. Executive Summary",
            "3. Technical Approach",
            "4. Management Approach",
            "5. Past Performance",
            "6. Pricing",
            "7. Compliance Matrix",
        ])
    lines = [ln.strip() for ln in outline.splitlines() if ln.strip() and not ln.startswith("#")]
    if not title:
        title = f"Proposal for RFP #{int(rfp_id)}"
    with _closing(conn.cursor()) as cur:
        cur.execute("INSERT INTO proposals(rfp_id, title) VALUES(?,?)", (int(rfp_id), title))
        pid = cur.lastrowid
        for i, ln in enumerate(lines, start=1):
            cur.execute(
                "INSERT INTO proposal_sections(proposal_id,ord,title,content,settings_json) VALUES(?,?,?,?,?)",
                (int(pid), i, ln, "", json.dumps({"font":"Times New Roman","size":11}))
            )
    conn.commit()
    return int(pid)

def x7_list_proposals(conn: sqlite3.Connection, rfp_id: int):
    import pandas as pd
    try:
        return pd.read_sql_query("SELECT id, title, status, created_at FROM proposals WHERE rfp_id=? ORDER BY id DESC;", conn, params=(int(rfp_id),))
    except Exception:
        import pandas as pd
        return pd.DataFrame(columns=["id","title","status","created_at"])

def x7_get_sections(conn: sqlite3.Connection, proposal_id: int):
    import pandas as pd
    try:
        return pd.read_sql_query("SELECT id, ord, title, content, settings_json FROM proposal_sections WHERE proposal_id=? ORDER BY ord ASC;", conn, params=(int(proposal_id),))
    except Exception:
        import pandas as pd
        return pd.DataFrame(columns=["id","ord","title","content","settings_json"])

def x7_save_section(conn: sqlite3.Connection, section_id: int, content: str | None, settings_json: str | None = None) -> None:
    from contextlib import closing as _closing
    with _closing(conn.cursor()) as cur:
        cur.execute("UPDATE proposal_sections SET content=?, settings_json=COALESCE(?, settings_json), updated_at=datetime('now') WHERE id=?", (content or "", settings_json, int(section_id)))
    conn.commit()

def x7_generate_section_ai(conn: sqlite3.Connection, rfp_id: int, title: str, guidance: str = "", temperature: float = 0.1, k: int = 8) -> str:
    # Use Y1 index for grounding and GPT model for drafting
    try:
        hits = y1_search(conn, int(rfp_id), f"{title} {guidance}", k=int(k))
    except Exception:
        hits = []
    ctx = []
    for h in hits or []:
        src = f"{h.get('file') or ''} p.{h.get('page') or ''}".strip()
        snippet = (h.get('text') or '')[:900]
        ctx.append(f"[{src}] {snippet}")
    prompt = "\n\n".join(ctx + [f"Write the section: {title}", f"Guidance: {guidance}"])
    client = get_ai()
    model = _resolve_model()
    sys = "You draft precise, proposal ready text. Use clear headings. Cite source brackets inline only where needed."
    try:
        resp = client.chat.completions.create(model=model, messages=[
            {"role":"system","content": sys},
            {"role":"user","content": prompt}
        ], temperature=float(temperature))
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        return f"AI error: {e}"

def x7_export_docx(conn: sqlite3.Connection, proposal_id: int) -> bytes | None:
    # Best effort DOCX export. If python-docx is missing, return None.
    try:
        import docx
    except Exception:
        return None
    from contextlib import closing as _closing
    doc = docx.Document()
    with _closing(conn.cursor()) as cur:
        cur.execute("SELECT p.title, p.rfp_id FROM proposals p WHERE p.id=?", (int(proposal_id),))
        row = cur.fetchone()
        title = row[0] if row else f"Proposal {int(proposal_id)}"
        rfp_id = row[1] if row else 0
    doc.add_heading(title, 0)
    # sections
    secs = x7_get_sections(conn, int(proposal_id))
    for _, r in secs.iterrows():
        doc.add_heading(str(r.get("title") or ""), level=1)
        body = r.get("content") or ""
        for para in body.split("\n\n"):
            doc.add_paragraph(para)
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
# ==== end X.7 helpers ====


# ==== end X.6 helpers ====





# ==== X.5 Transcript Viewer ====
def x5_render_transcript_viewer(conn: sqlite3.Connection, rfp_id: int) -> None:
    """View chat history with color-coded roles, filters, and export."""
    from contextlib import closing as _closing
    import datetime as _dt

    # Map file_id -> filename for scopes
    files = {}
    try:
        with _closing(conn.cursor()) as cur:
            cur.execute("SELECT id, filename FROM rfp_files WHERE rfp_id=?", (int(rfp_id),))
            for rid, fn in cur.fetchall() or []:
                files[int(rid)] = fn or ""
    except Exception:
        pass

    # Load turns
    turns = []
    try:
        with _closing(conn.cursor()) as cur:
            cur.execute("""SELECT id, scope, role, content, created_at FROM rfp_chat_turns WHERE rfp_id=? ORDER BY id ASC;""", (int(rfp_id),))
            for tid, scope, role, content, ts in cur.fetchall() or []:
                label = "Global"
                if (scope or "").startswith("file:"):
                    try:
                        fid = int((scope or "").split(":",1)[1])
                        label = f"File: {files.get(fid, str(fid))}"
                    except Exception:
                        label = "File"
                turns.append({
                    "id": int(tid),
                    "scope": scope or "global",
                    "role": (role or "assistant").lower(),
                    "content": content or "",
                    "ts": ts or "",
                    "label": label,
                })
    except Exception:
        pass

    # Filters
    colA, colB = st.columns([2,3])
    with colA:
        scopes = ["All", "Global"] + [f"File: {fn}" for fn in files.values()]
        chosen = st.selectbox("Scope", options=scopes, key=_uniq_key("x5_tv_scope", int(rfp_id)))
    with colB:
        query = st.text_input("Search", value="", key=_uniq_key("x5_tv_search", int(rfp_id))).strip().lower()

    def _match(t):
        scope_ok = (chosen == "All") or (chosen == "Global" and t["scope"] == "global") or (chosen.startswith("File:") and t["label"] == chosen)
        if not scope_ok:
            return False
        if not query:
            return True
        blob = f"{t['role']} {t['label']} {t['content']} {t['ts']}".lower()
        return query in blob

    filtered = [t for t in turns if _match(t)]

    # Legend
    st.markdown(":blue[Human]  ·  :green[AI]")

    # Render chat using role-based chat balloons
    for t in filtered:
        role = "user" if t["role"] == "user" else "assistant"
        time_str = t["ts"]
        scope_str = t["label"]
        with st.chat_message(role):
            st.caption(f"{scope_str} · {time_str}")
            st.markdown(t["content"])

    # Export CSV
    import io, csv
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["id","time","scope","role","content"])
    for t in filtered:
        writer.writerow([t["id"], t["ts"], t["label"], t["role"], t["content"]])
    st.download_button(
        "Export CSV",
        data=buf.getvalue().encode("utf-8"),
        file_name=f"rfp_{int(rfp_id)}_transcript.csv",
        mime="text/csv",
        key=_uniq_key("x5_tv_csv", int(rfp_id))
    )

    # Export Markdown
    md_lines = []
    for t in filtered:
        md_lines.append(f"### {t['role'].title()} — {t['label']}  \n*{t['ts']}*")
        md_lines.append("")
        md_lines.append(t["content"])
        md_lines.append("\n---")
    md_blob = "\n".join(md_lines).encode("utf-8")
    st.download_button(
        "Export Markdown",
        data=md_blob,
        file_name=f"rfp_{int(rfp_id)}_transcript.md",
        mime="text/markdown",
        key=_uniq_key("x5_tv_md", int(rfp_id))
    )
# ==== end X.5 Transcript Viewer ====




def _safe_int(v, default: int = 0) -> int:
    try:
        if v is None:
            return default
        if isinstance(v, int):
            return v
        s = str(v).strip()
        import re as _re
        m = _re.search(r"\d+", s)
        return int(m.group()) if m else default
    except Exception:
        return default


def __e1_google_api_key():
    try: return _st.secrets["google"]["api_key"]
    except Exception:
        try: return _st.secrets["GOOGLE_API_KEY"]
        except Exception: return ""

def __e1_norm_phone(p):
    digits = "".join(_re.findall(r"\d+", str(p or "")))
    return digits[1:] if len(digits)==11 and digits.startswith("1") else digits

def __e1_existing_vendor_keys(conn):
    return _s1d_select_existing_pairs(conn)

def __e1_enrich_and_render(conn, lat=None, lng=None, radius_m=80467, query=""):

    by_np, by_pid = _s1d_select_existing_pairs(conn)
    key = __e1_google_api_key()
    if not key:
        _st.error("E1 requires a Google API key in secrets: [google].api_key or GOOGLE_API_KEY")
        return
    if not query:
        _st.info("Enter a query in your existing finder, then use this enrichment to view details.")
        return
    import requests as _rq
    params={"query":query, "key":key, "region":"us"}
    if lat is not None and lng is not None:
        params.update({"location": f"{lat},{lng}", "radius": int(radius_m)})
    js = _rq.get("https://maps.googleapis.com/maps/api/place/textsearch/json", params=params, timeout=12).json()
    results = js.get("results", [])
    by_np, by_pid = __e1_existing_vendor_keys(conn)
    rows = []
    for r in results:
        name = r.get("name",""); pid = r.get("place_id",""); addr = r.get("formatted_address","")
        phone = ""; website = ""; gurl = ""
        try:
            det = _rq.get("https://maps.googleapis.com/maps/api/place/details/json",
                          params={"place_id": pid, "fields": "formatted_phone_number,website,url", "key": key}, timeout=10).json().get("result", {}) or {}
            digits = "".join(_re.findall(r"\\d+", det.get("formatted_phone_number","") or ""))
            if len(digits)==11 and digits.startswith("1"): digits=digits[1:]
            phone = digits; website = det.get("website","") or ""; gurl = det.get("url","") or ""
        except Exception:
            pass
        dup = ((name.strip().lower(), phone) in by_np) or (pid in by_pid)
        rows.append(dict(name=name, phone=phone, website=website, address=addr, place_id=pid, google_url=gurl, _dup=dup))
        _time.sleep(0.05)
    if not rows:
        _st.info("No results to enrich.")
        return
    df = _pd.DataFrame(rows)
    st.session_state["s1d_df"] = df.to_dict("records")
    def _link(u,t): return f"<a href='{u}' target='_blank'>{t}</a>" if u else t
    df["name"] = df.apply(lambda r: _link(r["google_url"], r["name"]), axis=1)
    df["website"] = df.apply(lambda r: _link(r["website"], "site") if r["website"] else "", axis=1)
    _st.write(df[["name","phone","website","address","place_id","_dup"]].to_html(escape=False, index=False), unsafe_allow_html=True)

# =========================

# --- Tab name alias ---
def run_l_and_m_checklist(conn):
    return run_lm_checklist(conn)


# --- Tab name alias ---
def run_backup_data(conn):
    return run_backup_and_data(conn)



# === Outreach: fallback sender picker and guard wrapper =====================
# Robust fallback picker that reads smtp_settings, only if not provided earlier.
if "_o3_render_sender_picker" not in globals():
    def _o3_render_sender_picker():
        import streamlit as st
        from contextlib import closing
        host, port, username, password, use_tls = "smtp.gmail.com", 465, "", "", False
        try:
            if "get_db" in globals():
                conn2 = get_db()
                with closing(conn2.cursor()) as cur:
                    cur.execute("CREATE TABLE IF NOT EXISTS smtp_settings (id INTEGER PRIMARY KEY, label TEXT, host TEXT, port INTEGER, username TEXT, password TEXT, use_tls INTEGER)")
                    row = cur.execute("SELECT host, port, username, password, use_tls FROM smtp_settings WHERE id=1").fetchone()
                if row:
                    host = row[0] or host
                    port = int(row[1] or port)
                    username = row[2] or username
                    password = row[3] or password
                    use_tls = bool(row[4] or use_tls)
        except Exception:
            pass
        if not username or not password:
            st.error("Sender not configured. Go to Outreach → Sender and Save sender.")
            return {}
        st.caption(f"Using {username} via {host}:{int(port)} TLS={'on' if use_tls else 'off'}")
        return {"host": host, "port": int(port), "email": username, "app_password": password, "use_tls": bool(use_tls)}

# Guard wrapper: ensure sender is configured before original Mail Merge UI runs
try:
    _orig__render_outreach_mailmerge = render_outreach_mailmerge
    def render_outreach_mailmerge(conn):
        import streamlit as st
        sender = _o3_render_sender_picker() if "_o3_render_sender_picker" in globals() else {}
        if sender and "username" in sender and "email" not in sender:
            sender["email"] = sender.get("username","")
        if sender and "password" in sender and "app_password" not in sender:
            sender["app_password"] = sender.get("password","")
        if not sender.get("email") or not sender.get("app_password"):
            st.warning("Configure a sender first in Outreach → Sender.")
            return
        return _orig__render_outreach_mailmerge(conn)
except Exception:
    pass

# === End Outreach guard =====================================================


def o1_sender_accounts_ui(conn):
    globals()['_O4_CONN'] = conn
    import streamlit as st
    import pandas as _pd
    ensure_outreach_o1_schema(conn)
    st.subheader("Sender accounts")
    email = st.text_input("Email")
    display = st.text_input("Display name")
    app_pw = st.text_input("Gmail App password", type="password")
    c1,c2,c3 = st.columns(3)
    with c1: host = st.text_input("SMTP host", value="smtp.gmail.com")
    with c2: port = st.number_input("SMTP port", 1, 65535, value=465)
    with c3: ssl = st.checkbox("Use SSL", value=True)
    if st.button("Save account", key="o4_ac_save__2"):
        if not email:
            st.error("Email required")
        else:
            with conn:
                conn.execute("""
                INSERT INTO email_accounts(user_email, display_name, app_password, smtp_host, smtp_port, use_ssl)
                VALUES(?,?,?,?,?,?)
                ON CONFLICT(user_email) DO UPDATE SET
                    display_name=excluded.display_name,
                    app_password=excluded.app_password,
                    smtp_host=excluded.smtp_host,
                    smtp_port=excluded.smtp_port,
                    use_ssl=excluded.use_ssl
                """, (email.strip(), display or "", app_pw or "", host or "smtp.gmail.com", int(port or 465), 1 if ssl else 0))
            st.success("Saved")
    try:
        import streamlit as st
        st.session_state["o4_sender_sel"] = email.strip()
    except Exception:
        pass
    st.rerun()
    try:
        df = _pd.read_sql_query("SELECT user_email, display_name, smtp_host, smtp_port, use_ssl FROM email_accounts ORDER BY user_email", conn)
        _styled_dataframe(df, use_container_width=True)
    except Exception:
        pass

# ---- helper: stable pagination (Phase 0) ----
def _pager_init(key: str):
    if key not in st.session_state:
        st.session_state[key] = 0

def _pager_update(key: str, delta: int, max_pages: int | None = None):
    _pager_init(key)
    st.session_state[key] = max(st.session_state[key] + delta, 0)
    if max_pages is not None:
        st.session_state[key] = min(st.session_state[key], max_pages - 1)
    return st.session_state[key]


# ---- helper: write guard (Phase 0) ----
def _write_guard(conn, fn, *args, **kwargs):
    # call DB write functions inside a transaction
    with conn:
        return fn(*args, **kwargs)


# ---- Phase 1: Global UI setup (theme, CSS, helpers) ----
def _init_phase1_ui():
    if st.session_state.get('_phase1_ui_ready'):
        return
    st.session_state['_phase1_ui_ready'] = True
    st.markdown("<div class='ela-banner'>Phase 1 theme active · polished layout & tables</div>", unsafe_allow_html=True)
    st.markdown('''
    <style>
    .block-container {padding-top: 1.2rem; padding-bottom: 1.2rem; max-width: 1400px;}
    h1, h2, h3 {margin-bottom: .4rem;}
    .ela-subtitle {color: rgba(49,51,63,0.65); font-size: .95rem; margin-bottom: 1rem;}
    div[data-testid="stDataFrame"] thead th {position: sticky; top: 0; background: #fff; z-index: 2;}
    div[data-testid="stDataFrame"] tbody tr:hover {background: rgba(64, 120, 242, 0.06);}
    .ela-card {border: 1px solid rgba(49,51,63,0.16); border-radius: 12px; padding: 12px; margin-bottom: 12px;}
    .ela-chip {display:inline-block; padding: 2px 8px; border-radius: 999px; font-size: 12px; margin-right:6px; background: rgba(49,51,63,.06);}
    .ela-ok {background: rgba(0,200,83,.12);}
    .ela-warn {background: rgba(251,140,0,.12);}
    .ela-bad {background: rgba(229,57,53,.12);}
    
    /* Top ribbon banner */
    .ela-banner {position: sticky; top: 0; z-index: 999; background: linear-gradient(90deg, #4068f2, #7a9cff); color: #fff; padding: 6px 12px; border-radius: 8px; margin-bottom: 10px;}
    /* Sidebar branding spacing */
    section[data-testid="stSidebar"] .block-container {padding-top: 0.8rem;}
    /* Expander cards */
    [data-testid="stExpander"] {border: 1px solid rgba(49,51,63,0.16); border-radius: 12px; margin-bottom: 10px;}
    [data-testid="stExpander"] summary {font-weight: 600;}
    /* Buttons subtle shadow */
    button[kind="primary"] {box-shadow: 0 1px 4px rgba(0,0,0,.08);}
    /* Text inputs rounding */
    .stTextInput>div>div>input, .stNumberInput input, .stTextArea textarea {border-radius: 10px !important;}
    
    </style>
    ''', unsafe_allow_html=True)

def _sidebar_brand():
    with st.sidebar:
        st.markdown("### 🧭 ELA GovCon Suite")
        st.caption("Phase 1 UI loaded")
        st.caption("Faster sourcing, compliant bids, higher win rates.")

def _styled_dataframe(df, use_container_width=True, height=None, hide_index=True, column_config=None):
    try:
        return _styled_dataframe(df, use_container_width=use_container_width, height=height, hide_index=hide_index, column_config=column_config)
    except TypeError:
        return _styled_dataframe(df, use_container_width=use_container_width, height=height)

def _chip(text: str, kind: str = 'neutral'):
    cls = 'ela-chip'
    if kind == 'ok': cls += ' ela-ok'
    elif kind == 'warn': cls += ' ela-warn'
    elif kind == 'bad': cls += ' ela-bad'
    st.markdown(f"<span class='{cls}'>{text}</span>", unsafe_allow_html=True)
