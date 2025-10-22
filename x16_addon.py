
# x16_addon.py — X16 Capability Statement Builder
# Import in app.py:
#   import x16_addon
#
# Builds a one-page capability statement (Markdown, optional DOCX).

try:
    import streamlit as st
    import pandas as pd
except Exception as _e:
    raise

def _make_docx(md_text: str) -> bytes:
    # Best-effort DOCX; fall back if python-docx not present.
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        for line in md_text.splitlines():
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
            else:
                doc.add_paragraph(line)
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.read()
    except Exception:
        return None

def _x16_ui(conn):
    st.markdown("### X16 — Capability Statement Builder")
    st.caption("X16 active — one-page capability statement as Markdown, optional DOCX")
    c1, c2 = st.columns(2)
    with c1:
        company = st.text_input("Company name", value="", key="x16_c")
        cage = st.text_input("CAGE", value="", key="x16_cage")
        uei = st.text_input("UEI", value="", key="x16_uei")
        naics = st.text_input("Core NAICS", value="", key="x16_naics")
        setaside = st.text_input("Set-Aside statuses", value="", key="x16_set")
        contact = st.text_area("Contact block", value="", key="x16_contact", height=80)
    with c2:
        summary = st.text_area("Company summary", height=100, key="x16_sum")
        core = st.text_area("Core competencies (bullets)", height=120, key="x16_core")
        diff = st.text_area("Differentiators (bullets)", height=120, key="x16_diff")
        past = st.text_area("Past performance (bullets)", height=120, key="x16_pp")

    if st.button("Build statement", key="x16_go"):
        md = f"""# {company}
**CAGE:** {cage} | **UEI:** {uei} | **NAICS:** {naics} | **Set-Aside:** {setaside}

## Summary
{summary}

## Core Competencies
{core}

## Differentiators
{diff}

## Past Performance
{past}

## Contact
{contact}
"""
        st.text_area("Markdown", value=md, height=300, key="x16_md")
        st.download_button("Download .md", data=md.encode("utf-8"), file_name="capability_statement.md", mime="text/markdown", key="x16_dl_md")
        try:
            import io
            b = _make_docx(md)
            if b:
                st.download_button("Download .docx", data=b, file_name="capability_statement.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="x16_dl_docx")
            else:
                st.info("python-docx not available; DOCX export skipped.")
        except Exception as e:
            st.info(f"DOCX export unavailable: {e}")

# Hook
try:
    _orig_run_rfp_analyzer = run_rfp_analyzer
except NameError:
    _orig_run_rfp_analyzer = None

if _orig_run_rfp_analyzer:
    def run_rfp_analyzer(conn):
        _orig_run_rfp_analyzer(conn)
        try:
            _x16_ui(conn)
        except Exception as e:
            st.info(f"X16 panel unavailable: {e}")
